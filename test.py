import streamlit as st
import json
import os
import hashlib
import time
from datetime import datetime, timedelta
from io import BytesIO
import streamlit.components.v1 as components
from supabase import create_client

# ══════════════════════════════════════════════════════════════════
# SPLASH SCREEN — Fonction universelle pour tous les services
# ══════════════════════════════════════════════════════════════════
_SPLASH_CONFIG = {
    "Expose": {"icon":"📚","bg":"#0c0516","glow":"rgba(160,80,255,0.28)","color":"#c87aff","titre":"TON EXPOSÉ<br><span style='color:#c87aff;'>RÉDIGÉ PAR L'IA</span>","sub":"Du CP jusqu'au Master · Structuré · Argumenté","badge":"⚡ LIVRÉ EN 60 SECONDES","prix":"1 GÉNÉRATION · 600 FC"},
    "Fiche": {"icon":"📖","bg":"#0a0514","glow":"rgba(155,89,182,0.28)","color":"#9b59b6","titre":"TA FICHE DE COURS<br><span style='color:#9b59b6;'>GÉNÉRÉE PAR L'IA</span>","sub":"Notions · Exemples · Exercices · Prête à imprimer","badge":"⚡ PROFESSIONNELLE EN 60 SEC","prix":"1 GÉNÉRATION · 600 FC"},
    "Sujets": {"icon":"📝","bg":"#050c14","glow":"rgba(0,150,255,0.25)","color":"#00aaff","titre":"TON SUJET D'EXAMEN<br><span style='color:#00aaff;'>CRÉÉ PAR L'IA</span>","sub":"QCM · Vrai/Faux · Cas pratique · Devoir complet","badge":"⚡ AUTO-GÉNÉRÉ EN 60 SEC","prix":"1 GÉNÉRATION · 600 FC"},
    "CV": {"icon":"👔","bg":"#050a10","glow":"rgba(0,200,255,0.22)","color":"#00d2ff","titre":"TON CV PROFESSIONNEL<br><span style='color:#00d2ff;'>RÉDIGÉ PAR L'IA</span>","sub":"CV + Lettre de motivation · Prêt à envoyer","badge":"⚡ PROFESSIONNEL EN 60 SEC","prix":"1 GÉNÉRATION · 600 FC"},
    "Word": {"icon":"📄","bg":"#050d08","glow":"rgba(0,200,80,0.22)","color":"#4dff88","titre":"TON DOCUMENT WORD<br><span style='color:#4dff88;'>CRÉÉ PAR L'IA</span>","sub":"Rapport · Lettre · Contrat · Document complet","badge":"⚡ LIVRÉ EN 60 SECONDES","prix":"1 GÉNÉRATION · 600 FC"},
    "Modifier": {"icon":"✏️","bg":"#050a0d","glow":"rgba(0,210,255,0.22)","color":"#00d2ff","titre":"MODIFICATION<br><span style='color:#00d2ff;'>DE TON FICHIER</span>","sub":"Word · Excel · PDF · PowerPoint · sur mesure","badge":"📎 IMPORTE TON FICHIER","prix":"TRAITEMENT RAPIDE"},
    "Conversion": {"icon":"🔄","bg":"#050d08","glow":"rgba(46,204,113,0.22)","color":"#2ecc71","titre":"CONVERSION<br><span style='color:#2ecc71;'>INSTANTANÉE</span>","sub":"PDF · Word · Excel · Images · 100% automatique","badge":"⚡ RÉSULTAT EN QUELQUES SEC","prix":"CONVERSION GRATUITE"},
    "OCR": {"icon":"🔍","bg":"#08050d","glow":"rgba(180,100,255,0.25)","color":"#b464ff","titre":"NUMÉRISATION OCR<br><span style='color:#b464ff;'>TEXTE EXTRACTIBLE</span>","sub":"PDF · Image · Word · Excel scannés → .docx éditable","badge":"📄 FICHIER WORD ÉDITABLE EN SORTIE","prix":"TRAITEMENT AUTOMATIQUE"},
    "RapportStage": {"icon":"🏢","bg":"#080d05","glow":"rgba(80,220,120,0.25)","color":"#50dc78","titre":"TON RAPPORT DE STAGE<br><span style='color:#50dc78;'>RÉDIGÉ PAR L'IA</span>","sub":"Introduction · Missions · Analyse · Conclusion · Annexes","badge":"⚡ LIVRÉ EN 60 SECONDES","prix":"1 GÉNÉRATION · 600 FC"},
}

def _show_splash(service_key: str, duree: float = 1.2):
    cfg = _SPLASH_CONFIG.get(service_key)
    if not cfg:
        return
    key = f"_splash_done_{service_key}"
    if key in st.session_state:
        return
    st.session_state[key] = True
    _ph = st.empty()
    _ph.markdown(f'''
    <div style="position:fixed;top:0;left:0;width:100vw;height:100vh;
        background:{cfg['bg']};z-index:99999;
        display:flex;align-items:center;justify-content:center;flex-direction:column;gap:18px;">
      <div style="position:absolute;top:35%;left:50%;transform:translate(-50%,-50%);
        width:380px;height:260px;
        background:radial-gradient(ellipse,{cfg['glow']} 0%,transparent 70%);
        filter:blur(35px);pointer-events:none;"></div>
      <div style="font-size:5rem;line-height:1;filter:drop-shadow(0 0 22px {cfg['color']});">{cfg['icon']}</div>
      <div style="font-family:Arial Black,Impact,sans-serif;font-size:2.1rem;font-weight:900;
        color:#ffffff;text-align:center;line-height:1.15;text-shadow:0 0 30px {cfg['color']};">{cfg['titre']}</div>
      <div style="font-family:monospace;font-size:0.88rem;color:rgba(255,255,255,0.38);
        letter-spacing:2px;text-align:center;">{cfg['sub']}</div>
      <div style="background:rgba(255,255,255,0.06);border:1px solid {cfg['color']}88;
        border-radius:100px;padding:8px 26px;font-family:monospace;font-size:0.82rem;
        color:{cfg['color']};font-weight:700;letter-spacing:2px;">{cfg['badge']}</div>
      <div style="font-family:Arial Black,sans-serif;font-size:1.7rem;color:#FFD700;
        font-weight:900;filter:drop-shadow(0 0 10px rgba(255,180,0,0.5));">{cfg['prix']}</div>
    </div>
    ''', unsafe_allow_html=True)
    time.sleep(duree)
    _ph.empty()


st.set_page_config(
    page_title="L'IA bureautique NoVA AI", 
    page_icon="⚡", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# ══════════════════════════════════════════════════════════════════
# PWA MANIFEST — Nova s'affiche "Nova" sur l'écran d'accueil mobile
# L'icône Streamlit est conservée, seul le nom change
# ══════════════════════════════════════════════════════════════════
import base64 as _b64, json as _json

def _inject_nova_pwa():
    _manifest = {
        "name": "Nova Platform",
        "short_name": "Nova",
        "description": "Nova Platform – L'IA bureautique ivoirienne",
        "start_url": "/",
        "display": "standalone",
        "background_color": "#ffffff",
        "theme_color": "#131416",
        "icons": [
            {
                "src": "https://streamlit.io/images/brand/streamlit-mark-color.png",
                "sizes": "192x192",
                "type": "image/png"
            },
            {
                "src": "https://streamlit.io/images/brand/streamlit-mark-color.png",
                "sizes": "512x512",
                "type": "image/png"
            }
        ]
    }
    _m64 = _b64.b64encode(_json.dumps(_manifest).encode()).decode()
    st.markdown(f"""
    <link rel="manifest" href="data:application/json;base64,{_m64}">
    <meta name="apple-mobile-web-app-capable" content="yes">
    <meta name="apple-mobile-web-app-title" content="Nova">
    <meta name="apple-mobile-web-app-status-bar-style" content="default">
    <link rel="apple-touch-icon" href="https://streamlit.io/images/brand/streamlit-mark-color.png">
    <meta name="application-name" content="Nova">
    """, unsafe_allow_html=True)

_inject_nova_pwa()

DATA_FILE = "data_nova_v3.json"
ADMIN_CODE  = st.secrets.get("ADMIN_CODE", "02110240")
COLLAB_CODE = "2026"   # Collaborateur — accès missions uniquement

SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

def normalize_wa(numero):
    if not numero:
        return ""
    numero = numero.strip().replace(" ", "").replace("-", "").replace("+", "")
    if numero.startswith("0") and not numero.startswith("00"):
        numero = "225" + numero
    return numero

def load_db():
    try:
        users_rows = supabase.table("users").select("*").execute().data
        users = {}
        for r in users_rows:
            users[r["uid"]] = {
                "whatsapp": r["whatsapp"],
                "email": r.get("email", "Non renseigné"),
                "joined": r["joined"],
                "premium": r.get("premium", False),
                "premium_plan": r.get("premium_plan", None),
                "premium_expiry": r.get("premium_expiry", None),
                "gen_used": r.get("gen_used", 0),
                "gen_date": r.get("gen_date", None),
            }
        demandes_rows = supabase.table("demandes").select("*").execute().data
        demandes = []
        for r in demandes_rows:
            demandes.append({
                "id": r["id"],
                "user": r["uid"],
                "service": r["service"],
                "desc": r["description"],
                "whatsapp": r["whatsapp"],
                "status": r["status"],
                "incomplet": r["incomplet"],
                "champs_manquants": json.loads(r["champs_manquants"]) if r["champs_manquants"] else [],
                "timestamp": r["timestamp"]
            })
        liens_rows = supabase.table("liens").select("*").execute().data
        liens = {}
        for r in liens_rows:
            if r["uid"] not in liens:
                liens[r["uid"]] = []
            liens[r["uid"]].append({"name": r["name"], "url": r["url"], "date": r["date"]})
        return {"users": users, "demandes": demandes, "liens": liens}
    except Exception as e:
        st.error(f"Erreur chargement Supabase : {e}")
        return {"users": {}, "demandes": [], "liens": {}}

def get_auto_reply_setting():
    """Charge le paramètre auto_reply depuis Supabase table config."""
    try:
        r = supabase.table("config").select("value").eq("key", "auto_reply_gratuit").execute().data
        return r[0]["value"] == "true" if r else False
    except:
        return False

def set_auto_reply_setting(enabled: bool):
    """Sauvegarde le paramètre auto_reply dans Supabase table config."""
    try:
        supabase.table("config").upsert({
            "key": "auto_reply_gratuit",
            "value": "true" if enabled else "false"
        }).execute()
    except:
        pass

def envoyer_email_auto_gratuit(client_nom, client_wa, service, nom_fichier, demande):
    """Email admin — Nova Platform a répondu automatiquement à un gratuit après 1h30."""
    try:
        import resend
        resend.api_key = st.secrets["RESEND_API_KEY"]
        _modele_info = st.session_state.get("_last_modele_gemini", "inconnu")
        corps = f"""
🤖 NOVA IA — RÉPONSE AUTOMATIQUE PLAN GRATUIT (2H)

👤 Client      : {client_nom}
📱 WhatsApp    : {client_wa}
🛠️ Service     : {service}
📄 Fichier     : {nom_fichier}
🧠 Modèle IA   : {_modele_info}
⏰ Généré le   : {datetime.now().strftime("%d/%m/%Y à %H:%M")}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📋 DEMANDE COMPLÈTE :
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{demande.strip()}

Le document a été livré directement au client via l'interface Nova.
        """
        resend.Emails.send({
            "from": "Nova Platform <onboarding@resend.dev>",
            "to": [st.secrets["EMAIL_RECEIVER"]],
            "subject": f"🤖 Auto-Gratuit — {service} ({client_nom})",
            "text": corps
        })
    except:
        pass

def save_user(uid, whatsapp, email="Non renseigné", premium=False, premium_plan=None, premium_expiry=None):
    try:
        supabase.table("users").upsert({
            "uid": uid, "whatsapp": whatsapp,
            "email": email, "joined": str(datetime.now()),
            "premium": premium, "premium_plan": premium_plan, "premium_expiry": premium_expiry,
            "gen_used": 0, "gen_date": None,
        }).execute()
        return True
    except Exception as e:
        st.error(f"Erreur sauvegarde utilisateur : {e}")
        return False

def update_premium_status(uid, premium, premium_plan, premium_expiry):
    try:
        supabase.table("users").update({
            "premium": premium, "premium_plan": premium_plan, "premium_expiry": premium_expiry,
        }).eq("uid", uid).execute()
    except Exception as e:
        st.error(f"Erreur mise à jour premium : {e}")

def save_demande(req):
    try:
        supabase.table("demandes").upsert({
            "id": req["id"],
            "uid": req["user"],
            "service": req["service"],
            "description": req["desc"],
            "whatsapp": req["whatsapp"],
            "status": req["status"],
            "incomplet": req["incomplet"],
            "champs_manquants": json.dumps(req["champs_manquants"]),
            "timestamp": req["timestamp"]
        }).execute()
    except Exception as e:
        st.error(f"Erreur sauvegarde demande : {e}")

def delete_demande(req_id):
    try:
        supabase.table("demandes").delete().eq("id", req_id).execute()
    except Exception as e:
        st.error(f"Erreur suppression demande : {e}")

def save_lien(uid, name, url, date):
    try:
        supabase.table("liens").insert({
            "uid": uid, "name": name, "url": url, "date": date
        }).execute()
    except Exception as e:
        st.error(f"Erreur sauvegarde lien : {e}")

def delete_all_liens(uid):
    try:
        supabase.table("liens").delete().eq("uid", uid).execute()
    except Exception as e:
        st.error(f"Erreur suppression historique : {e}")

def save_refus(uid, service, message):
    """Sauvegarde un refus de mission dans les livrables du client."""
    try:
        supabase.table("liens").insert({
            "uid": uid,
            "name": service,
            "url": f"__refus__{message}",
            "date": datetime.now().strftime("%d/%m/%Y")
        }).execute()
    except Exception as e:
        st.error(f"Erreur sauvegarde refus : {e}")

def purger_fichiers_anciens(jours=20, dry_run=False):
    """
    Supprime du Storage Supabase + de la table liens
    tous les fichiers livrés depuis plus de `jours` jours.
    Retourne un dict {"supprimes": [...], "erreurs": [...], "ignores": int}
    """
    try:
        import requests as _req
        from datetime import datetime, timedelta

        BUCKET   = "nova-fichiers"
        sb_url   = st.secrets["SUPABASE_URL"].rstrip("/")
        sb_key   = st.secrets.get("SUPABASE_SERVICE_KEY", st.secrets["SUPABASE_KEY"])
        hdrs     = {"apikey": sb_key, "Authorization": f"Bearer {sb_key}"}
        limite   = datetime.now() - timedelta(days=jours)

        # Charger tous les liens
        rows = supabase.table("liens").select("*").execute().data or []

        supprimes = []
        erreurs   = []
        ignores   = 0

        for row in rows:
            url   = row.get("url", "")
            date  = row.get("date", "")
            uid   = row.get("uid", "")
            lien_id = row.get("id")

            # Ignorer les entrées spéciales (refus, locaux)
            if not url or url.startswith("__"):
                ignores += 1
                continue

            # Parser la date (format dd/mm/yyyy)
            try:
                date_lien = datetime.strptime(date, "%d/%m/%Y")
            except Exception:
                ignores += 1
                continue

            # Pas encore expiré
            if date_lien >= limite:
                ignores += 1
                continue

            # Extraire le chemin Storage depuis l'URL publique
            # URL : .../storage/v1/object/public/nova-fichiers/fichiers_clients/...
            marqueur = f"/object/public/{BUCKET}/"
            if marqueur not in url:
                ignores += 1
                continue

            chemin = url.split(marqueur, 1)[1]

            if dry_run:
                supprimes.append({"chemin": chemin, "date": date, "uid": uid})
                continue

            # Supprimer du Storage
            del_url = f"{sb_url}/storage/v1/object/{BUCKET}/{chemin}"
            resp = _req.delete(del_url, headers=hdrs, timeout=15)

            if resp.status_code in (200, 204, 404):
                # 404 = déjà supprimé, on nettoie quand même la table
                try:
                    supabase.table("liens").delete().eq("id", lien_id).execute()
                    supprimes.append({"chemin": chemin, "date": date, "uid": uid})
                except Exception as e:
                    erreurs.append(f"DB delete {lien_id}: {e}")
            else:
                try:
                    detail = resp.json()
                except Exception:
                    detail = resp.text[:200]
                erreurs.append(f"{chemin} → HTTP {resp.status_code}: {detail}")

        return {"supprimes": supprimes, "erreurs": erreurs, "ignores": ignores}

    except Exception as e:
        return {"supprimes": [], "erreurs": [str(e)], "ignores": 0}


def sanitize_nom_fichier(nom):
    """Nettoie un nom de fichier pour le rendre compatible avec Supabase Storage.
    - Supprime les accents (é→e, ç→c, etc.)
    - Remplace espaces et caractères spéciaux par '_'
    - Ne garde que : a-z, A-Z, 0-9, '_', '-', '.'
    """
    import unicodedata, re
    # Normaliser les accents (NFD décompose é en e + accent, puis on supprime les accents)
    nom = unicodedata.normalize("NFD", nom)
    nom = "".join(c for c in nom if unicodedata.category(c) != "Mn")
    # Remplacer espaces par underscore
    nom = nom.replace(" ", "_")
    # Ne garder que les caractères autorisés
    nom = re.sub(r"[^a-zA-Z0-9_\-\.]", "_", nom)
    # Éviter les doubles underscores
    nom = re.sub(r"_+", "_", nom)
    return nom.strip("_")

def upload_fichier_client(uid, req_id, fichier_bytes, fichier_nom):
    """Upload via API REST Supabase Storage."""
    try:
        import requests as _req
        BUCKET      = "nova-fichiers"
        sb_url      = st.secrets["SUPABASE_URL"].rstrip("/")
        sb_key      = st.secrets["SUPABASE_KEY"]
        # La service_role key bypasse le RLS — nécessaire pour les uploads Storage
        # Ajoute SUPABASE_SERVICE_KEY dans tes secrets Streamlit
        # Dashboard Supabase → Settings → API → service_role (secret)
        sb_svc_key  = st.secrets.get("SUPABASE_SERVICE_KEY", sb_key)
        auth_hdrs   = {
            "apikey":        sb_svc_key,
            "Authorization": f"Bearer {sb_svc_key}",
        }

        # — Nettoyage des composantes du chemin —
        uid_safe  = sanitize_nom_fichier(str(uid))
        nom_safe  = sanitize_nom_fichier(fichier_nom)
        chemin    = f"fichiers_clients/{uid_safe}_{req_id}_{nom_safe}"

        # — Bytes bruts —
        data = fichier_bytes.read() if hasattr(fichier_bytes, "read") else fichier_bytes

        # — Détection du Content-Type selon l'extension —
        ext = nom_safe.rsplit(".", 1)[-1].lower() if "." in nom_safe else ""
        mime_map = {
            "pdf": "application/pdf",
            "png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
            "gif": "image/gif", "webp": "image/webp",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
            "txt": "text/plain", "csv": "text/csv", "zip": "application/zip",
        }
        content_type = mime_map.get(ext, "application/octet-stream")

        # — Upload —
        url_upload = f"{sb_url}/storage/v1/object/{BUCKET}/{chemin}"
        resp = _req.post(
            url_upload,
            headers={**auth_hdrs, "Content-Type": content_type, "x-upsert": "true"},
            data=data,
            timeout=30
        )

        if resp.status_code in (200, 201):
            return f"{sb_url}/storage/v1/object/public/{BUCKET}/{chemin}"
        else:
            # Afficher le message exact de Supabase pour debug
            try:
                detail = resp.json()
            except Exception:
                detail = resp.text
            return f"ERREUR_UPLOAD:{resp.status_code} — {detail}"

    except Exception as e:
        return f"ERREUR_UPLOAD:{e}"

def save_db(data):
    pass

def envoyer_notification(client_nom, client_wa, service, description):
    try:
        import resend
        resend.api_key = st.secrets["RESEND_API_KEY"]
        corps = f"""
🔔 NOUVELLE COMMANDE NOVA PLATFORM

👤 Client      : {client_nom}
📱 WhatsApp    : {client_wa}
🛠️ Service     : {service}
📝 Description : {description}

⏰ Reçue le {datetime.now().strftime("%d/%m/%Y à %H:%M")}

Connectez-vous à la console admin pour traiter cette mission.
        """
        resend.Emails.send({
            "from": "Nova Platform <onboarding@resend.dev>",
            "to": [st.secrets["EMAIL_RECEIVER"]],
            "subject": f"🔔 Nouvelle commande Nova Platform — {service}",
            "text": corps
        })
        st.toast("📧 Notification email envoyée !", icon="✅")
    except Exception as e:
        st.toast(f"❌ Email échoué : {e}", icon="⚠️")

def envoyer_notification_gemini_ok(client_nom, client_wa, service, nom_fichier, demande_complete=""):
    """Email admin quand Gemini a généré — + notif client si email dispo."""
    try:
        import resend
        resend.api_key = st.secrets["RESEND_API_KEY"]
        _modele_info = st.session_state.get("_last_modele_gemini", "inconnu")
        section_demande = f"""
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📋 DEMANDE COMPLÈTE DU CLIENT :
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{demande_complete.strip()}
""" if demande_complete.strip() else ""
        corps_admin = f"""
✅ NOVA IA A DÉJÀ RÉPONDU — AUCUNE ACTION REQUISE

👤 Client      : {client_nom}
📱 WhatsApp    : {client_wa}
🛠️ Service     : {service}
📄 Fichier     : {nom_fichier}
🧠 Modèle IA   : {_modele_info}
⏰ Généré le   : {datetime.now().strftime("%d/%m/%Y à %H:%M")}
{section_demande}
Le document a été livré directement au client via l'interface Nova.
Vous n'avez rien à faire pour cette commande.
        """
        resend.Emails.send({
            "from": "Nova Platform <onboarding@resend.dev>",
            "to": [st.secrets["EMAIL_RECEIVER"]],
            "subject": f"✅ Nova Platform — {service} ({client_nom})",
            "text": corps_admin
        })
    except Exception:
        pass

def envoyer_notif_client_email(client_nom, client_email, service, nom_fichier):
    """Envoie un email HTML au CLIENT quand son fichier est prêt."""
    if not client_email or client_email == "Non renseigné" or "@" not in client_email:
        return
    try:
        import resend
        resend.api_key = st.secrets["RESEND_API_KEY"]
        corps_html = f"""
        <div style="font-family:Arial,sans-serif;max-width:520px;margin:0 auto;background:#0a0a0a;border-radius:16px;overflow:hidden;border:1px solid #222;">
          <div style="background:linear-gradient(135deg,#0d1f0d,#0a0a0a);padding:28px 30px 18px;text-align:center;border-bottom:1px solid #1a3a1a;">
            <div style="font-size:2.2rem;margin-bottom:6px;">⚡</div>
            <div style="font-family:Arial Black,sans-serif;font-size:1.5rem;font-weight:900;color:#4dff88;letter-spacing:2px;">NOVA PLATFORM</div>
            <div style="color:rgba(255,255,255,0.4);font-size:0.78rem;letter-spacing:1px;margin-top:4px;">L'IA Bureautique de Côte d'Ivoire</div>
          </div>
          <div style="padding:28px 30px;">
            <div style="color:rgba(255,255,255,0.85);font-size:1rem;margin-bottom:18px;">
              Bonjour <strong style="color:#4dff88;">{client_nom}</strong> 👋
            </div>
            <div style="background:#0d1f0d;border:1px solid #1a5c30;border-radius:12px;padding:18px 20px;margin-bottom:20px;text-align:center;">
              <div style="font-size:2rem;margin-bottom:8px;">✅</div>
              <div style="font-family:Arial Black,sans-serif;font-size:1.2rem;font-weight:900;color:#ffffff;">Votre fichier est prêt !</div>
              <div style="color:rgba(255,255,255,0.5);font-size:0.85rem;margin-top:6px;">Généré automatiquement par l'IA Nova Platform</div>
            </div>
            <table style="width:100%;border-collapse:collapse;margin-bottom:20px;">
              <tr>
                <td style="padding:8px 0;color:rgba(255,255,255,0.35);font-size:0.82rem;border-bottom:1px solid #1a1a1a;">🛠️ Service</td>
                <td style="padding:8px 0;color:#ffffff;font-size:0.82rem;font-weight:700;text-align:right;border-bottom:1px solid #1a1a1a;">{service}</td>
              </tr>
              <tr>
                <td style="padding:8px 0;color:rgba(255,255,255,0.35);font-size:0.82rem;border-bottom:1px solid #1a1a1a;">📄 Fichier</td>
                <td style="padding:8px 0;color:#4dff88;font-size:0.82rem;font-weight:700;text-align:right;border-bottom:1px solid #1a1a1a;">{nom_fichier}</td>
              </tr>
              <tr>
                <td style="padding:8px 0;color:rgba(255,255,255,0.35);font-size:0.82rem;">⏰ Livré le</td>
                <td style="padding:8px 0;color:#ffffff;font-size:0.82rem;text-align:right;">{datetime.now().strftime("%d/%m/%Y à %H:%M")}</td>
              </tr>
            </table>
            <div style="background:#0a0a0a;border:1px solid #222;border-radius:10px;padding:14px 16px;margin-bottom:20px;text-align:center;">
              <div style="color:rgba(255,255,255,0.4);font-size:0.75rem;letter-spacing:1px;margin-bottom:4px;">RETROUVEZ VOTRE FICHIER SUR</div>
              <div style="color:#ffffff;font-size:0.88rem;">👉 Votre interface Nova Platform</div>
            </div>
            <div style="text-align:center;padding-top:8px;border-top:1px solid #1a1a1a;">
              <div style="color:rgba(255,255,255,0.3);font-size:0.75rem;margin-bottom:6px;">Un problème ? Contactez-nous sur WhatsApp</div>
              <div style="font-family:monospace;font-size:1rem;font-weight:700;color:#25d366;">+225 01 71 54 25 05</div>
            </div>
          </div>
          <div style="background:#050505;padding:14px 30px;text-align:center;border-top:1px solid #111;">
            <div style="color:rgba(255,255,255,0.2);font-size:0.72rem;">Nova Platform · Abidjan, Côte d'Ivoire 🇨🇮</div>
          </div>
        </div>
        """
        resend.Emails.send({
            "from": "Nova Platform <onboarding@resend.dev>",
            "to": [client_email],
            "subject": f"✅ Votre fichier est prêt — {service} | Nova Platform",
            "html": corps_html
        })
    except Exception as _e_client_mail:
        try:
            supabase.table("config").upsert({
                "key": "email_client_last_error",
                "value": f"{type(_e_client_mail).__name__}: {str(_e_client_mail)[:400]} | dest={client_email}"
            }).execute()
        except:
            pass

def notifier_livraison_gemini(client_nom, client_wa, client_email, service, nom_fichier, demande_complete=""):
    """Point d'entrée unique — notifie admin ET client à chaque livraison Gemini."""
    envoyer_notification_gemini_ok(client_nom, client_wa, service, nom_fichier, demande_complete)
    envoyer_notif_client_email(client_nom, client_email, service, nom_fichier)

PLANS_PREMIUM = {
    "Journalier": {"jours": 1,  "prix": "600 FC",  "emoji": "🌅", "generations": 2},
    "10 Jours":   {"jours": 10, "prix": "1000 FC", "emoji": "🔟", "generations": 9},
    "30 Jours":   {"jours": 30, "prix": "2500 FC", "emoji": "👑", "generations": 999},
}

def get_gen_quota(user_data):
    """Retourne (gen_used_aujourd_hui, quota_max) selon le plan."""
    plan = user_data.get("premium_plan")
    quota = PLANS_PREMIUM.get(plan, {}).get("generations", 0) if plan else 0
    gen_date = user_data.get("gen_date")
    today = datetime.now().strftime("%Y-%m-%d")
    if gen_date != today:
        return 0, quota  # Nouveau jour → compteur remis à zéro
    return user_data.get("gen_used", 0), quota

def quota_restant(user_data):
    used, quota = get_gen_quota(user_data)
    return max(0, quota - used)

MAX_DEMANDES_GRATUIT_PAR_JOUR = 7

def get_demandes_gratuit_today(uid):
    """Retourne le nombre de demandes soumises aujourd'hui par un utilisateur gratuit."""
    try:
        today = datetime.now().strftime("%Y-%m-%d")
        rows = supabase.table("demandes").select("timestamp").eq("uid", uid).execute().data
        count = 0
        for r in rows:
            ts_str = r.get("timestamp", "")
            try:
                ts = datetime.fromisoformat(ts_str)
                if ts.strftime("%Y-%m-%d") == today:
                    count += 1
            except:
                pass
        return count
    except:
        return 0

def incrementer_gen(uid):
    """Incrémente le compteur de générations du jour pour l'utilisateur."""
    try:
        today = datetime.now().strftime("%Y-%m-%d")
        # Recharger pour avoir la valeur fraîche
        row = supabase.table("users").select("gen_used, gen_date").eq("uid", uid).execute().data
        if row:
            gen_date = row[0].get("gen_date")
            gen_used = row[0].get("gen_used", 0) if gen_date == today else 0
        else:
            gen_used = 0
        supabase.table("users").update({
            "gen_used": gen_used + 1,
            "gen_date": today,
        }).eq("uid", uid).execute()
    except Exception as e:
        pass  # Ne pas bloquer la génération si le compteur échoue

def is_premium_actif(user_data):
    if not user_data.get("premium"):
        return False
    expiry = user_data.get("premium_expiry")
    if not expiry:
        return False
    try:
        return datetime.now() < datetime.fromisoformat(expiry)
    except:
        return False

def get_premium_info(user_data):
    if not is_premium_actif(user_data):
        return None
    try:
        expiry_dt  = datetime.fromisoformat(user_data["premium_expiry"])
        jours_rest = (expiry_dt - datetime.now()).days
        return {
            "plan":           user_data.get("premium_plan", "—"),
            "expiry":         expiry_dt.strftime("%d/%m/%Y à %H:%M"),
            "jours_restants": jours_rest,
        }
    except:
        return None

def activer_premium(uid, plan_name):
    jours  = PLANS_PREMIUM[plan_name]["jours"]
    expiry = datetime.now() + timedelta(days=jours)
    update_premium_status(uid, True, plan_name, expiry.isoformat())
    if "db" in st.session_state and uid in st.session_state["db"]["users"]:
        st.session_state["db"]["users"][uid].update({
            "premium": True, "premium_plan": plan_name,
            "premium_expiry": expiry.isoformat(),
        })

def desactiver_premium(uid):
    update_premium_status(uid, False, None, None)
    if "db" in st.session_state and uid in st.session_state["db"]["users"]:
        st.session_state["db"]["users"][uid].update({
            "premium": False, "premium_plan": None, "premium_expiry": None,
        })

def delete_user(uid):
    """Supprime complètement un membre (users + demandes + liens)."""
    try:
        supabase.table("liens").delete().eq("uid", uid).execute()
    except:
        pass
    try:
        supabase.table("demandes").delete().eq("user", uid).execute()
    except:
        pass
    try:
        supabase.table("users").delete().eq("uid", uid).execute()
    except:
        pass
    # Nettoyer session_state
    if "db" in st.session_state:
        st.session_state["db"]["users"].pop(uid, None)

def donner_bonus_gen(uid, nb_bonus):
    """
    Donne nb_bonus générations gratuites valables aujourd'hui à un utilisateur.
    Crée un mini-plan bonus d'1 jour sans toucher au statut premium existant.
    Stocke via la colonne gen_used en négatif (crédit) + gen_date = today.
    Logique : quota_restant = quota_max - gen_used.
    Pour un gratuit (quota_max=0), on stocke bonus_gen dans la table config.
    """
    try:
        today = datetime.now().strftime("%Y-%m-%d")
        # Lire l'état actuel
        row = supabase.table("users").select("gen_used,gen_date,premium,premium_plan,premium_expiry").eq("uid", uid).execute().data
        if not row:
            return False
        d = row[0]
        is_prem = d.get("premium") and d.get("premium_expiry") and datetime.now() < datetime.fromisoformat(d["premium_expiry"])

        if is_prem:
            # Utilisateur premium : on réduit gen_used du bonus (donne plus de quota)
            gen_used = d.get("gen_used", 0) if d.get("gen_date") == today else 0
            new_used = max(0, gen_used - nb_bonus)
            supabase.table("users").update({"gen_used": new_used, "gen_date": today}).eq("uid", uid).execute()
        else:
            # Utilisateur gratuit : on stocke le bonus dans la table config
            key_bonus = f"bonus_gen_{uid}"
            existing = supabase.table("config").select("value").eq("key", key_bonus).execute().data
            if existing:
                try:
                    old = json.loads(existing[0]["value"])
                    # Si date différente d'aujourd'hui, on réinitialise
                    if old.get("date") != today:
                        new_val = {"date": today, "quota": nb_bonus, "used": 0}
                    else:
                        new_val = {"date": today, "quota": old.get("quota", 0) + nb_bonus, "used": old.get("used", 0)}
                except:
                    new_val = {"date": today, "quota": nb_bonus, "used": 0}
                supabase.table("config").update({"value": json.dumps(new_val)}).eq("key", key_bonus).execute()
            else:
                supabase.table("config").insert({"key": key_bonus, "value": json.dumps({"date": today, "quota": nb_bonus, "used": 0})}).execute()
        return True
    except Exception as e:
        return False

def get_bonus_gen_gratuit(uid):
    """Retourne (quota_bonus, used_bonus) pour un utilisateur gratuit."""
    try:
        today = datetime.now().strftime("%Y-%m-%d")
        key_bonus = f"bonus_gen_{uid}"
        row = supabase.table("config").select("value").eq("key", key_bonus).execute().data
        if not row:
            return 0, 0
        d = json.loads(row[0]["value"])
        if d.get("date") != today:
            return 0, 0
        return d.get("quota", 0), d.get("used", 0)
    except:
        return 0, 0

def consommer_bonus_gen_gratuit(uid):
    """Décrémente le bonus gratuit d'un cran. Retourne True si OK."""
    try:
        today = datetime.now().strftime("%Y-%m-%d")
        key_bonus = f"bonus_gen_{uid}"
        row = supabase.table("config").select("value").eq("key", key_bonus).execute().data
        if not row:
            return False
        d = json.loads(row[0]["value"])
        if d.get("date") != today:
            return False
        if d.get("used", 0) >= d.get("quota", 0):
            return False
        d["used"] = d.get("used", 0) + 1
        supabase.table("config").update({"value": json.dumps(d)}).eq("key", key_bonus).execute()
        return True
    except:
        return False

def get_modeles_disponibles(api_key):
    import urllib.request as _ur
    import urllib.error
    try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
        req = _ur.Request(url, headers={"Content-Type": "application/json"}, method="GET")
        with _ur.urlopen(req, timeout=15) as resp:
            data = json.loads(resp.read().decode())
        modeles = []
        exclusions = ["tts", "audio", "image", "imagen", "veo", "robotics",
                      "embedding", "aqa", "computer-use", "research", "nano-banana",
                      "gemma"]
        for m in data.get("models", []):
            if "generateContent" in m.get("supportedGenerationMethods", []):
                nom = m["name"].replace("models/", "")
                if not any(excl in nom.lower() for excl in exclusions):
                    modeles.append(nom)
        def priorite(nom):
            if "flash-lite" in nom: return 0
            if "2.0-flash" in nom:  return 1
            if "flash" in nom:      return 2
            if "pro" in nom:        return 3
            return 4
        modeles_tries = sorted(modeles, key=priorite)
        return modeles_tries
    except Exception as e:
        return ["gemini-2.0-flash-lite", "gemini-2.0-flash", "gemini-2.5-flash"]


def generer_avec_gemini(service, description, client_nom, is_premium=False, gen_used=0, _plan_for_model=None, _image_b64=None, _image_mime=None):
    try:
        import urllib.request as _ur
        import urllib.error

        api_key = st.secrets["GEMINI_API_KEY"]

        # ================================================================
        # PROMPT — EXPOSÉ SCOLAIRE (Système scolaire ivoirien & africain)
        # ================================================================
        if "Exposé" in service:
            prompt = f"""Tu es un expert académique de haut niveau ET un maître absolu de la génération de documents Word professionnels pour le système éducatif ivoirien et africain francophone.
Tu as été formé sur des milliers d'exposés scolaires primés et tu maîtrises parfaitement chaque aspect : typographie, structure, rhétorique académique, contextualisation culturelle et rendu Word via python-docx.

╔══════════════════════════════════════════════════════════════════╗
║     ENCYCLOPÉDIE EXPERTE — GÉNÉRATION DOCUMENT WORD NOVA PLATFORM     ║
╚══════════════════════════════════════════════════════════════════╝

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 1 — MAÎTRISE COMPLÈTE DU RENDU WORD (python-docx)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

COMMENT LE MOTEUR NOVA CONVERTIT TON TEXTE EN WORD :

1. TITRES MARKDOWN → STYLES WORD AUTOMATIQUES :
   # Titre    → Heading 1 (Arial 16pt, couleur 1F4E79, gras, majuscule recommandé)
   ## Titre   → Heading 2 (Arial 14pt, couleur 2E75B6, gras)
   ### Titre  → Heading 3 (Arial 12pt, gras)
   #### Titre → Heading 4 (Arial 11pt, gras italique)
   → Respecte toujours CET ORDRE HIÉRARCHIQUE — jamais de saut de niveau

2. TEXTE EN GRAS → **texte** :
   - Rendu Word : gras Arial 11pt dans le paragraphe courant
   - Usage obligatoire pour : termes techniques à leur 1re occurrence, définitions clés, chiffres essentiels, noms d'auteurs, noms d'institutions
   - Ex: La **photosynthèse** est définie comme le processus par lequel les végétaux...
   - Ex: En **2023**, la Côte d'Ivoire a produit **2,2 millions de tonnes** de cacao

3. TABLEAUX MARKDOWN → TABLEAUX WORD AUTOMATIQUEMENT FORMATÉS :
   - En-tête bleu foncé (1F4E79) avec texte blanc, lignes alternées bleu clair / blanc
   - Format STRICTEMENT OBLIGATOIRE :
   | En-tête 1 | En-tête 2 | En-tête 3 |
   |-----------|-----------|-----------|
   | Contenu   | Contenu   | Contenu   |
   - TOUJOURS : **Tableau N : [Titre précis]** AVANT le tableau
   - TOUJOURS : *Source : [Référence institution/auteur, Année]* APRÈS le tableau
   - Idéal pour : comparaisons chiffrées, chronologies, classifications, données statistiques

4. SÉPARATEURS VISUELS → LIGNES HORIZONTALES WORD :
   - ════════════════════ → ligne épaisse bleue (sz=12) = séparateur MAJEUR entre grandes parties
   - ──────────────────── → ligne fine grise (sz=4) = séparateur MINEUR entre sous-sections
   - ---SAUT_DE_PAGE---   → ligne bleue (sz=8) + espace blanc = séparateur de SECTION principale
   - Laisser toujours une ligne vide avant et après les séparateurs

5. LISTES À PUCES → BULLETS WORD AUTOMATIQUES :
   - "- Item complet" → List Bullet (Arial 11pt)
   - "1. Item numéroté" → List Number (Arial 11pt)
   - Chaque item = une phrase complète, jamais un mot seul
   - USAGE LIMITÉ À : sommaire, bibliographie, listes de faits — JAMAIS dans le développement

6. PARAGRAPHES NORMAUX → Arial 11pt, interligne 1.5, texte JUSTIFIÉ (des deux côtés) :
   - Tout texte non formaté = paragraphe Normal justifié automatiquement par Nova
   - ALINÉA obligatoire : commence chaque paragraphe par 4 espaces (    ) → retrait Word 1.25cm
   - Ligne vide entre deux blocs = espacement naturel dans le document Word final
   - Chaque paragraphe de développement : 8 à 10 lignes minimum (≈ 100-120 mots)
   - Un paragraphe trop court (< 4 lignes) sera fusionné ou développé — jamais laissé tel quel

7. À ÉVITER ABSOLUMENT — NE FONCTIONNE PAS DANS LE MOTEUR NOVA :
   ✗ LaTeX : $formule$, \frac{{}}, \omega, \text{{}}, \left(, \right), \\, \begin{{}}
   ✗ HTML : <br>, <b>, <strong>, <p>, <div>, <span>
   ✗ Italique simple *texte* (utilise **gras** à la place pour la mise en valeur)
   ✗ Tirets en guise de sous-titres — toujours ## Sous-titre
   ✗ Retours à la ligne multiples pour simuler des espaces
   ✗ Indentations avec espaces multiples

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 2 — MOTEUR FORMULES NOVA — GUIDE COMPLET
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

NOVA possède un moteur qui convertit toutes les notations mathématiques/scientifiques
en vrais exposants/indices Word. 4 modes disponibles :

① ###FORMULE### [formule] → formule mise en valeur (fond bleu clair, centré, 13pt gras)
  ###FORMULE### E = m × c^{{2}}
  ###FORMULE### F = (G × m_{{1}} × m_{{2}}) / r^{{2}}
  ###FORMULE### Δ = b^{{2}} - 4ac  →  x_{{1,2}} = (-b ± √Δ) / (2a)

② ###DEBUT_FORMULES### / ###FIN_FORMULES### → bloc de formules groupées
  ###DEBUT_FORMULES###
  U = R × I       (loi d'Ohm)
  P = U × I       (puissance électrique)
  E = P × t       (énergie en Joules)
  ###FIN_FORMULES###

③ INLINE dans le texte avec ^{{}} et _{{}} :
  "La résistance vaut R_{{eq}} = R_{{1}} + R_{{2}} = 50 Ω"
  "L'énergie cinétique E_{{c}} = (1/2)×m×v^{{2}}"
  "Le noyau ^{{14}}_{{6}}C émet un β^{{-}}"
  "H_{{2}}O, CO_{{2}}, C_{{6}}H_{{12}}O_{{6}}, SO_{{4}}^{{2-}}, Ca^{{2+}}"

④ LaTeX $...$ converti automatiquement :
  "$\frac{{U}}{{R}} = I$" | "$\sqrt{{b^{{2}}-4ac}}$" | "$\omega = 2\pi f$"

TOUS LES SYMBOLES DISPONIBLES (utilise directement) :
  α β γ δ ε ζ η θ ι κ λ μ ν ξ π ρ σ τ υ φ χ ψ ω
  Α Β Γ Δ Ε Ζ Η Θ Λ Μ Ξ Π Ρ Σ Τ Υ Φ Ψ Ω
  × · ÷ ± √ ∞ ∂ ∇ ∫ ∬ ∮ Σ Π ∝ ≈ ≃ ≅ ≡ ≠ ≤ ≥ ≪ ≫
  ∈ ∉ ⊂ ⊃ ⊆ ⊇ ∪ ∩ ∅ ∀ ∃ ∧ ∨ ¬ ⊕ ℝ ℕ ℤ ℚ ℂ
  ∠ ⊥ ∥ △ ° ⇌ ⟶ → ← ↔ ⟹ ⟺ ↑ ↓ ↦

SECTION 3 — ART MAÎTRISÉ DE LA RÉDACTION ACADÉMIQUE — RHÉTORIQUE ET STYLE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

▶ A. COMMENT CONSTRUIRE UNE INTRODUCTION EN 5 TEMPS (structure obligatoire) :

TEMPS 1 — ACCROCHE (min 4 lignes) — CHOISIR L'UNE DE CES 4 STRATÉGIES :
  → Données choc : "Avec **2,2 millions de tonnes** de cacao produits annuellement, représentant **45%** de l'offre mondiale selon la FAO (2023), la Côte d'Ivoire détient à elle seule le destin du marché mondial de cette fève. Pourtant, les **5 millions de paysans** qui en vivent perçoivent moins de **6%** de la valeur finale d'une tablette de chocolat en Europe (Oxfam, 2022)."
  → Paradoxe saisissant : "Pays exportateur de soleil et de lumière, la Côte d'Ivoire souffre pourtant d'un déficit énergétique qui plonge des millions de ses citoyens dans l'obscurité chaque soir. Comment expliquer ce paradoxe d'un pays producteur de **1 500 mégawatts** d'électricité, dont une large part est exportée vers les pays voisins ?"
  → Citation d'auteur africain (avec référence précise) : "'Les indépendances africaines ont accouché de lendemains qui déchantent', écrivait **Ahmadou Kourouma** dans *Les Soleils des Indépendances* (1968). Plus d'un demi-siècle après cette prophétie littéraire, la question du développement endogène reste au cœur des préoccupations du continent africain."
  → Anecdote historique/fait d'actualité : "Le **7 août 1960**, sous les acclamations d'une foule immense rassemblée au stade Félix Houphouët-Boigny d'Abidjan, la Côte d'Ivoire accédait à l'indépendance après plus d'un siècle de domination coloniale française. Ce moment fondateur..."

TEMPS 2 — CONTEXTUALISATION (min 4 lignes) :
Situe précisément le sujet dans son contexte historique, géographique, scientifique ou social.
Définit TOUS les termes clés du sujet en **gras** dès leur première occurrence.
Donne des chiffres, des dates précises, des acteurs réels.
→ "La **photosynthèse**, terme issu du grec *phôtos* (lumière) et *synthesis* (assemblage), désigne..."

TEMPS 3 — DÉLIMITATION ET ENJEUX (min 3 lignes) :
Précise le périmètre exact de l'étude et pourquoi ce sujet est important aujourd'hui.
→ "Comprendre ce phénomène revêt une importance capitale, tant pour [enjeu scientifique] que pour [enjeu social/économique/environnemental ivoirien]."

TEMPS 4 — PROBLÉMATIQUE (1-2 phrases précises et non rhétoriques) :
La problématique n'est PAS une simple reformulation du sujet — elle soulève une VRAIE tension :
  ✓ "Dans quelle mesure la dépendance cacaoyère de la Côte d'Ivoire constitue-t-elle à la fois le moteur et le talon d'Achille de son développement économique ?"
  ✓ "Comment la **déforestation** accélérée, moteur apparent de la croissance agricole ivoirienne, menace-t-elle paradoxalement les conditions mêmes de cette croissance ?"
  ✓ "En quoi le mouvement de la **négritude** représente-t-il une réponse littéraire et identitaire à la domination coloniale, et quelles en sont les limites actuelles ?"
  ✗ ÉVITER : "Qu'est-ce que la photosynthèse ?" (trop simple, pas problématique)
  ✗ ÉVITER : "Pourquoi la CI est-elle riche ?" (trop vague)

TEMPS 5 — ANNONCE DU PLAN (1-2 phrases, plan en 2 ou 3 parties selon niveau) :
→ "Pour apporter une réponse nuancée à cette interrogation, nous analyserons dans une première partie [intitulé Partie I — reformuler en 1 ligne], avant d'examiner dans une deuxième partie [Partie II], et d'envisager enfin [Partie III si lycée/université]."

▶ B. ARCHITECTURE D'UN PARAGRAPHE PARFAIT — MODÈLE PEEL ENRICHI :

Structure de chaque paragraphe de développement (min 8-10 lignes) :

1. POINT — Phrase d'affirmation claire et directe (1-2 lignes) :
   "La **déforestation** constitue l'une des crises environnementales les plus graves qu'ait connues la Côte d'Ivoire au cours du XXe siècle."

2. EXPLICATION — Développe le mécanisme, définit les termes, explique les causes ou le fonctionnement (3-4 lignes) :
   "Ce phénomène se définit comme la destruction durable et souvent irréversible du couvert forestier au profit d'autres usages des terres, notamment l'agriculture, l'exploitation forestière industrielle et l'urbanisation galopante. En Côte d'Ivoire, ce processus a été largement amplifié par l'extension des cultures de rente, principalement le **cacao** et le **café**, dont la demande mondiale croissante a exercé une pression considérable sur les forêts du Sud et du Centre-Ouest du pays."

3. EXEMPLE PRÉCIS IVOIRIEN/AFRICAIN — Chiffre sourcé + événement daté + lieu précis (3-4 lignes) :
   "Les données du Ministère des Eaux et Forêts (2022) révèlent une réalité alarmante : de **16 millions d'hectares** de forêt dense que comptait la Côte d'Ivoire au début du XXe siècle, il n'en subsistait plus que **3,4 millions** en 2020, soit une perte de **79% du couvert forestier** en un siècle. À titre illustratif, la **Forêt classée du Banco**, véritable poumon vert d'Abidjan, a vu sa superficie passer de 3 000 hectares à l'époque coloniale à environ 1 800 hectares aujourd'hui, sous la pression de l'urbanisation et des empiètements agricoles."

4. LIEN — Phrase de transition vers le paragraphe ou sous-partie suivant(e) (1-2 lignes) :
   "Cette destruction massive du patrimoine forestier ne se limite pas à une question environnementale ; elle engage profondément les équilibres climatiques régionaux et les conditions de vie des populations rurales, ce qui nous conduit à examiner ses répercussions socio-économiques."

▶ C. TRANSITIONS OBLIGATOIRES ENTRE GRANDES PARTIES — MODÈLES :

TRANSITION I → II (min 4 lignes) :
"Ainsi avons-nous établi, au terme de cette première partie, que [résumé en 1 phrase de la Partie I]. Cette analyse, si elle permet de cerner [apport de la Partie I], ne saurait toutefois être complète sans que l'on s'interroge sur [ce que la Partie II va apporter]. C'est précisément l'objet de notre second axe de réflexion, consacré à [intitulé Partie II]."

TRANSITION II → III (min 3 lignes) :
"Au regard des éléments développés dans notre deuxième partie, force est de constater que [bilan Partie II]. Ces constats nous invitent dès lors à dépasser le simple constat analytique pour envisager [dimension prospective / solutions / synthèse] — dimension qui constituera le fil directeur de notre troisième et dernière partie."

▶ D. CONNECTEURS LOGIQUES — VARIER OBLIGATOIREMENT (jamais répéter deux fois de suite) :

INTRODUIRE : "Il convient tout d'abord de souligner que", "Force est de constater que", "Il importe de noter que",
"À ce titre,", "Dans cette perspective,", "En premier lieu,", "Il y a lieu de préciser que",
"D'emblée, il apparaît que", "Au seuil de cette analyse,"

DÉVELOPPER : "En effet,", "De surcroît,", "Par ailleurs,", "Qui plus est,", "Il convient également de noter que",
"À cet égard,", "Dans ce sens,", "En outre,", "Il faut également souligner que",
"On notera de surcroît que", "À cela s'ajoute le fait que"

ILLUSTRER : "Ainsi,", "C'est notamment le cas de", "À titre illustratif,", "À titre d'exemple concret,",
"On peut citer à cet effet", "L'exemple ivoirien est à ce titre particulièrement éloquent :",
"Comme en témoigne", "Les données de [institution] le confirment :", "Pour s'en convaincre,"

OPPOSER/NUANCER : "Cependant,", "Néanmoins,", "Toutefois,", "En revanche,", "Or,",
"Il convient toutefois de relativiser ce constat :", "Si [thèse]... en revanche [nuance]...",
"Malgré tout,", "Il serait néanmoins réducteur de", "Cette réalité ne doit pas occulter le fait que"

CONCLURE/TRANSITER : "En définitive,", "Au regard de ces éléments,", "Au terme de cette analyse,",
"C'est dans cette logique que", "Ces constats nous amènent naturellement à examiner",
"Cette analyse nous conduit à aborder", "Ainsi avons-nous établi que", "Il ressort de ce qui précède que"

▶ E. TYPES DE PLANS À CHOISIR INTELLIGEMMENT SELON LE SUJET :

THÉMATIQUE (sujets descriptifs) : I (Nature/Définition/Caractéristiques) → II (Causes/Mécanismes/Fonctionnement) → III (Effets/Impacts/Solutions)
→ Idéal pour : "La déforestation en CI", "Le paludisme en Afrique", "Le coupé-décalé ivoirien"

DIALECTIQUE (sujets controversés) : I (Thèse : position principale/avantages) → II (Antithèse : limites/critiques/risques) → III (Synthèse/Dépassement/Voie du milieu)
→ Idéal pour : "L'agriculture extensive, moteur ou frein du développement ?", "La mondialisation, chance ou menace pour l'Afrique ?"

CHRONOLOGIQUE (sujets historiques) : I (Origines/Passé lointain) → II (Évolutions/État actuel/Ruptures) → III (Perspectives/Avenir/Défis contemporains)
→ Idéal pour : "L'histoire de la Côte d'Ivoire", "L'évolution du système éducatif africain"

ANALYTIQUE (sujets complexes multi-dimensionnels) : I (Dimension économique/scientifique) → II (Dimension sociale/humaine/culturelle) → III (Dimension politique/environnementale/internationale)
→ Idéal pour : "Les enjeux de l'eau en Afrique de l'Ouest", "La question du développement durable"

▶ F. CONSTRUCTION D'UNE CONCLUSION EN UN SEUL PARAGRAPHE (structure obligatoire) :

La conclusion est rédigée en UN SEUL BLOC DE PROSE CONTINU — jamais découpée en "Temps 1 / Temps 2 / Temps 3", jamais de sauts de ligne entre les idées, jamais de sous-titres.
Le paragraphe unique (15 à 20 lignes minimum) enchaîne naturellement :
1. Bilan synthétique de chaque grande partie (reformulé, jamais mot pour mot)
2. Réponse nuancée et directe à la problématique posée en introduction
3. Ouverture prospective sur un enjeu futur pour CI/Afrique

Tout doit couler comme un seul mouvement de pensée, sans rupture visible entre les trois dimensions.

CITATIONS ET RÉFÉRENCES — FORMAT ACADÉMIQUE RIGOUREUX :
- Citation directe : « La forêt tropicale est le poumon de la planète. » (FAO, 2022, p. 12)
- Citation courte intégrée : Selon KOUROUMA (1970), « les soleils des indépendances » symbolisent...
- Paraphrase : D'après les travaux de TADJO (2004), la mémoire collective africaine se construit...
- Source institutionnelle : Le Ministère de l'Agriculture (2023) indique que la production cacaoyère...
- Statistique sourcée : Selon la FAO (2023), la Côte d'Ivoire produit **2,2 millions de tonnes** de cacao...

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 4 — SYSTÈME SCOLAIRE IVOIRIEN COMPLET
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PRIMAIRE (CP1, CP2, CE1, CE2, CM1, CM2) — Examen : CEPE (fin CM2) :
- Vocabulaire simple, phrases max 15 mots, exemples de vie quotidienne ivoirienne
- 1 à 2 pages — structure : Intro courte / Corps 2-3 paragraphes / Conclusion
- Matières : Lecture, Écriture, Calcul, Sciences d'Éveil, Histoire-Géo CI, ECM, EPS

COLLÈGE 1er CYCLE (6ème, 5ème, 4ème, 3ème) — Examen : BEPC :
- Vocabulaire courant, termes disciplinaires définis en **gras**
- 2 à 4 pages — 2 grandes parties + 2 sous-parties chacune
- Auteurs : Bernard Dadié, Camara Laye, Ahmadou Kourouma, Mongo Beti, Ferdinand Oyono
- Matières : Français, Maths, PC, SVT, Histoire-Géo, Anglais, EDHC, Arts, EPS

LYCÉE 2nd CYCLE (2nde, 1ère, Terminale) — Examen : BAC ivoirien :
- A1 (Lettres-Philo) : Français, Philo, Histoire-Géo, Langues — style littéraire, rhétorique
- A2 (Lettres-SH) : + Sciences sociales, EDHC — approche socioéconomique
- B (Économie) : Économie, Gestion, Maths, Comptabilité — chiffres et tableaux obligatoires
- C (Maths-PC) : Maths renforcées, PC, Philo — rigueur scientifique maximale
- D (Maths-SVT) : Maths, SVT renforcée, PC — biologie, écologie, médecine tropicale
- E (Maths-Techno) : Maths, Technologie industrielle — ingénierie appliquée
- F/G/H : Techniques industrielles, commerciales, informatiques
- 4 à 7 pages — 3 grandes parties + 2 à 3 sous-parties par partie

UNIVERSITÉ (L1 à Doctorat) — Système LMD :
- L1-L3 : Introduction aux disciplines, révue de littérature, méthodologie de base
- M1-M2 : Cadre théorique, hypothèses, méthodologie rigoureuse, revue critique
- Doctorat : Contribution originale, état de l'art exhaustif, notes de bas de page
- Institutions : UFHB Cocody (Abidjan), UAO Bouaké, UJLOG Daloa, INP-HB Yamoussoukro, ESATIC
- 8 à 20 pages selon niveau

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 5 — BASE DE CONNAISSANCES AFRICAINE ET MONDIALE ENRICHIE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

⚠️ RÈGLE D'USAGE : Cette base est une RESSOURCE de secours, pas une obligation d'insertion.
Tu ne cites du contexte ivoirien QUE si le sujet le justifie naturellement.
Un exposé sur la physique, les mathématiques, la chimie, l'histoire mondiale ou la philosophie
n'a PAS besoin d'exemples ivoiriens forcés. Priorise toujours la pertinence du sujet.

CÔTE D'IVOIRE (à utiliser SEULEMENT si le sujet s'y prête) :
GÉOGRAPHIE : 322 463 km², ~28M habitants (2024), cap. politique Yamoussoukro, cap. économique Abidjan
Villes : Bouaké, Daloa, Korhogo, San-Pédro, Man, Odienné, Abengourou, Gagnoa
Fleuves : Comoé (1160km), Bandama (960km), Sassandra (650km), Cavally, Bia
Lacs : Kossou (1700km², 3e lac artificiel Afrique), Buyo, Taabo, Ayamé
Relief : Monts Nimba (1752m, UNESCO), Monts Toura, plateau central, plaine côtière
Végétation : forêt dense humide (Sud, 30% territoire), savane arbustive (Centre-Nord)
Sites UNESCO : Forêt de Taï, Parc de la Comoé, Monts Nimba (transfrontalier)

HISTOIRE CI : Indépendance 7 août 1960 | Félix Houphouët-Boigny (1960-1993, père fondateur)
"Miracle ivoirien" (1960-1980), crise 2002 (rébellion Nord-Sud), crise 2010-2011 (post-électorale)
Alassane Ouattara (2011-présent) | Plan National de Développement (PND 2021-2025)

ÉCONOMIE CI : PIB ~70Md USD (2023) | Croissance ~6-7%/an | Émergence visée 2030
Cacao : 1er mondial (45% production, 2,2M tonnes/an) | Café : 3e africain
Anacarde : 1er africain (800 000 t/an) | Hévéa, palmier à huile, coton, banane, ananas
Port d'Abidjan : 1er conteneurs Afrique de l'Ouest, >30M tonnes/an
Monnaie : FCFA (XOF) | UEMOA, CEDEAO, UA

CULTURE CI : ~60 ethnies | Akan (Baoulé 23%, Agni), Krou (Bété, Dida, Wê), Mandé (Malinké, Dioula), Gur (Sénoufo, Lobi)
Musique : coupé-décalé (DJ Arafat, Magic System), zouglou (Les Garagistes), gospel ivoirien, afrobeats
Arts : masques Baoulé (Goli, Kpan), masques Dan (Gunyège, Gle), bronzes Akan, tissage Sénoufo
Gastronomie : attiéké, kedjenou, foutou, aloco, placali, garba, kangni, graine (sauce)

CONTEXTE AFRICAIN ÉLARGI (Afrique de l'Ouest, Afrique centrale, Afrique de l'Est) :
Sénégal : économie en émergence, Plan Sénégal Émergent, ressources pétrolières offshore découvertes 2014
Ghana : économie diversifiée, pétrole offshore (Jubilee Field), premier producteur d'or Afrique (2023)
Nigeria : 1re économie africaine (~500Md USD PIB), pétrole (Delta du Niger), Nollywood (2e cinéma mondial)
Cameroun : bilinguisme français-anglais, pétrole, cacao, bois, port de Douala (poumon Afrique centrale)
Kenya : hub technologique Afrique de l'Est (Silicon Savannah), M-Pesa (1ère monnaie mobile mondiale)
Afrique du Sud : 2e économie africaine, mines d'or et platine, Cape Town/Johannesburg

LITTÉRATURE AFRICAINE FRANCOPHONE :
Ivoiriens : DADIÉ Bernard (*Climbié* 1956), KOUROUMA Ahmadou (*Les Soleils des Indépendances* 1968, *Monnè* 1990), TADJO Véronique (*Reine Pokou* 2004), ADIAFFI Jean-Marie (*La Carte d'identité* 1980)
Africains : LAYE Camara Guinée (*L'Enfant Noir* 1953), BETI Mongo Cameroun (*Mission Terminée* 1957), OYONO Ferdinand Cameroun (*Une Vie de Boy* 1956), SENGHOR Léopold Sédar Sénégal (poètes de la négritude), SEMBÈNE Ousmane Sénégal (*Les Bouts de Bois de Dieu* 1960), ACHEBE Chinua Nigeria (*Things Fall Apart* 1958)

SCIENCES ET ENVIRONNEMENT :
Biodiversité CI : 150+ espèces mammifères, 700+ oiseaux, hippopotame pygmée, éléphant de forêt, chimpanzé de Taï
Déforestation CI : 16M ha en 1900 → 3,4M ha aujourd'hui (perte 79% couverture forestière)
Maladies tropicales : paludisme (Plasmodium falciparum, 1re cause mortalité en Afrique subsaharienne), tuberculose, VIH/SIDA
Changement climatique Afrique : -20% pluviométrie au Sahel depuis 1970, montée eaux côtières

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 6 — EXEMPLES DE PARAGRAPHES D'EXCELLENCE (MODÈLES À IMITER)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

EXEMPLE HISTOIRE-GÉO avec données africaines (sujet géographique/économique) :
"La Côte d'Ivoire occupe une position économique stratégique sur le continent africain, fondée en grande partie sur la culture du **cacao**. Avec une production annuelle de **2,2 millions de tonnes** représentant environ **45% de la production mondiale** selon l'ICCO (International Cocoa Organization, 2023), le pays a construit sa prospérité sur cette culture pérenne introduite par les colons à la fin du XIXe siècle. Cette dépendance à la monoculture expose cependant l'économie aux chocs des cours mondiaux, incitant le gouvernement à accélérer sa politique de diversification économique à travers le **Plan National de Développement (PND 2021-2025)**."

EXEMPLE SVT/PC — sujet scientifique universel (SANS exemple ivoirien forcé) :
"La **photosynthèse** est le processus biochimique fondamental par lequel les végétaux chlorophylliens convertissent l'énergie lumineuse en énergie chimique. L'équation bilan s'écrit : 6 CO2 + 6 H2O + énergie lumineuse → C6H12O6 + 6 O2, soit six molécules de dioxyde de carbone et six d'eau qui, sous l'action de la lumière captée par la **chlorophylle**, produisent une molécule de **glucose** et six de dioxygène. Ce processus est à la base de toute vie sur Terre : il constitue le premier maillon de la chaîne alimentaire et régule le taux de CO2 atmosphérique, jouant ainsi un rôle clé dans l'équilibre climatique mondial. Des recherches récentes menées par le **CNRS** et l'**Université de Cambridge** cherchent à reproduire artificiellement ce mécanisme pour produire de l'hydrogène propre à grande échelle."

EXEMPLE FRANÇAIS/LITTÉRATURE AFRICAINE :
"La littérature africaine francophone constitue un vecteur privilégié d'affirmation identitaire et de résistance culturelle. Des auteurs comme **Bernard Dadié**, dont l'œuvre maîtresse *Climbié* (1956) dresse le portrait d'un jeune ivoirien confronté à la colonisation, ont su transformer l'expérience douloureuse de la domination en une création littéraire féconde. De même, **Ahmadou Kourouma**, dans *Les Soleils des Indépendances* (1968), brise les codes du français standard en y insufflant la syntaxe et la vision du monde **malinké**, créant un «français africanisé» reconnu comme l'un des apports majeurs des lettres africaines à la francophonie. Cette double appartenance linguistique devient une richesse, ouvrant la voie à des auteurs comme **Véronique Tadjo** avec *Reine Pokou* (2004)."

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 7 — LES 10 RÈGLES ABSOLUES DE LA GÉNÉRATION NOVA
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

RÈGLE 1 — COMPLÉTUDE TOTALE : Zéro "[à compléter]", "[...]", "[insérer]", "[Prénom fictif]" → TOUT rédigé intégralement
RÈGLE 2 — LONGUEUR MAÎTRISÉE : Respecter STRICTEMENT le nombre de pages demandé par le client (6, 7, 8 ou 9 pages). Un exposé percutant et dense vaut mieux qu'un exposé long et dilué. Chaque phrase doit avoir sa raison d'être — zéro remplissage, zéro répétition, zéro reformulation inutile.
RÈGLE 3 — QUALITÉ LINGUISTIQUE : Orthographe et grammaire irréprochables, ponctuation soignée, style académique soutenu
RÈGLE 4 — CONTEXTUALISATION PERTINENTE : Si le sujet traite de géographie, histoire, économie, environnement, société → inclure 1 à 2 exemples africains/ivoiriens SI pertinents. Pour les sujets de sciences exactes (maths, physique, chimie), philosophie, informatique ou histoire mondiale → PAS d'exemples ivoiriens forcés. La pertinence prime toujours sur la localisation.
RÈGLE 5 — FORMULES NOVA : Utilise la notation x^{{2}}, H_{{2}}O, CO_{{2}}, √(expr), symboles Unicode α β γ π ω ≤ ≥ × → ⇌. LaTeX inline $...$ aussi accepté (converti auto). Voir Section 2 pour tous les exemples.
RÈGLE 6 — STRUCTURE STRICTE : Séparateurs ════ et ---SAUT_DE_PAGE--- uniquement dans le corps du document (jamais dans le sommaire). Tu ne génères PAS de page de garde — Nova Platform s'en charge.
RÈGLE 6b — ANTI-ORPHELINES : Ne JAMAIS terminer une partie par une phrase de transition — la transition appartient au DÉBUT de la partie suivante (après le saut de page). Évite ainsi les 2-3 lignes orphelines en haut d'une page vide.
RÈGLE 7 — ADAPTATION NIVEAU : Vocabulaire + profondeur + longueur strictement adaptés au niveau détecté
RÈGLE 8 — PROSE DANS LE DÉVELOPPEMENT : Corps du document = paragraphes continus — jamais de listes à puces
RÈGLE 9 — DONNÉES PRÉCISES ET SOURCÉES : Chiffres réels, dates précises, institutions réelles — jamais de vague
RÈGLE 10 — VRAIS AUTEURS ET ŒUVRES : Citer de vraies œuvres d'auteurs réels — jamais "[Auteur fictif, Titre fictif]"
RÈGLE 11 — JUSTIFICATION OBLIGATOIRE : Chaque paragraphe du développement doit être rédigé en texte JUSTIFIÉ (aligné des deux côtés). Le moteur Nova applique la justification automatiquement sur tous les paragraphes normaux. Pour t'assurer que le rendu est parfait, chaque paragraphe doit être long (8 à 10 lignes minimum) — un paragraphe trop court ne peut pas être justifié visuellement.
RÈGLE 12 — ALINÉA EN DÉBUT DE PARAGRAPHE : Commence chaque nouveau paragraphe du développement par un retrait de première ligne symbolisé par 4 espaces (    ) ou une tabulation. Le moteur Nova les convertit en vrai retrait Word de 1.25cm. Exemple : "    La **déforestation** constitue l'une des crises..."
RÈGLE 13 — ESPACEMENT ENTRE BLOCS : Laisse toujours UNE ligne vide entre deux paragraphes, entre un titre et son paragraphe, et entre un paragraphe et un séparateur. Deux lignes vides = trop. Zéro ligne vide = blocs collés (interdit).
RÈGLE 14 — LONGUEUR MINIMALE PAR PARAGRAPHE : Un paragraphe du développement fait MINIMUM 8 lignes réelles (environ 100-120 mots). Jamais de paragraphe de 2-3 lignes dans le corps — fusionner avec le suivant ou développer davantage.
RÈGLE 15 — TITRES SANS PONCTUATION FINALE : Les titres # ## ### ne portent jamais de point, virgule ou deux-points en fin de ligne. Exemple correct : "## I. Les fondements économiques du miracle ivoirien". Exemple interdit : "## I. Les fondements économiques :"
RÈGLE 16 — EXPOSÉ PERCUTANT, PAS LONG : Priorité à la densité et à l'impact. Chaque paragraphe doit apporter quelque chose de nouveau — jamais reformuler ce qui vient d'être dit. Éviter : les phrases creuses d'introduction de paragraphe ("Il convient tout d'abord de noter que..."), les transitions trop longues, les répétitions de chiffres déjà cités. Si le client demande 6 pages → 6 pages denses, pas 6 pages diluées.

=== MISSION ===

⛔ RÈGLE ABSOLUE N°1 — SUJET IMPOSÉ, NE JAMAIS MODIFIER :
Le SUJET/THÈME de l'exposé est fourni ci-dessous par le client. Tu DOIS rédiger l'exposé EXACTEMENT sur ce sujet.
INTERDIT ABSOLU : changer, reformuler, remplacer, élargir ou dévier du sujet donné.
Si le sujet est "Les enfants de la rue" → l'exposé parle UNIQUEMENT des enfants de la rue.
Si le sujet est "La photosynthèse" → l'exposé parle UNIQUEMENT de la photosynthèse.
Le titre ###TITRE_ROUGE### doit reprendre fidèlement le sujet du client, joliment formulé — jamais un thème inventé.

Rédige un exposé scolaire COMPLET, STRUCTURÉ, PROFESSIONNEL et ENCYCLOPÉDIQUE basé sur cette demande :

{description}

=== STRUCTURE OBLIGATOIRE DU DOCUMENT — RESPECTER CET ORDRE EXACT ===

⚠️ RÈGLE FONDAMENTALE — NE JAMAIS VIOLER :
La PAGE DE GARDE est générée automatiquement par Nova Platform.
Tu NE DOIS PAS générer de page de garde. INTERDIT ABSOLU.
Commence directement par le SOMMAIRE — c'est la première chose que tu écris.

⚠️ AUTRES RÈGLES ABSOLUES :
- INTERDIT ABSOLU : ne JAMAIS utiliser ────, ════, ---, ━━━ dans le SOMMAIRE
- INTERDIT ABSOLU : Un tableau ne doit JAMAIS chevaucher deux pages. Insère un ---SAUT_DE_PAGE--- AVANT si nécessaire.
- JAMAIS de titre de section (# SOMMAIRE, # PAGE DE GARDE...) — commence directement avec le contenu

**SOMMAIRE**

Introduction ............................................................. p. 1
**I. [Titre 1re grande partie]** ........................................ p. 2
   1.1 [Titre 1re sous-partie] ........................................... p. 2
   1.2 [Titre 2e sous-partie] ............................................ p. 3
**II. [Titre 2e grande partie]** ......................................... p. 4
   2.1 [Titre 1re sous-partie] ........................................... p. 4
   2.2 [Titre 2e sous-partie] ............................................ p. 5
**III. [Titre 3e grande partie — lycée/université uniquement]** ......... p. 6
   3.1 [Titre sous-partie] ............................................... p. 6
   3.2 [Titre sous-partie] ............................................... p. 7
Conclusion ............................................................... p. 8

---SAUT_DE_PAGE---

# ════════════════════════════════════════════════════════
#                      INTRODUCTION
# ════════════════════════════════════════════════════════

## INTRODUCTION


[ACCROCHE PERCUTANTE — Min 5 lignes — CHOISIR : données choc sourcées / paradoxe saisissant / citation d'auteur africain avec référence complète / anecdote historique. Ex: "Selon la FAO (2023), la Côte d'Ivoire produit **45%** du cacao mondial avec **2,2 millions de tonnes**. Pourtant, les 5 millions de paysans concernés perçoivent moins de 6% de la valeur finale d'une tablette de chocolat en Europe (Oxfam, 2022)..."]

[CONTEXTUALISATION APPROFONDIE — Min 5 lignes : situe dans contexte historique/géographique/scientifique/social. Définit TOUS les termes clés en **gras** dès leur première occurrence. Donne chiffres, dates précises, acteurs réels.]

[DÉLIMITATION ET ENJEUX — Min 3 lignes : précise le périmètre de l'étude et pourquoi le sujet est important aujourd'hui pour la CI/l'Afrique/le monde.]

[PROBLÉMATIQUE PRÉCISE ET NON RHÉTORIQUE — 1-2 phrases soulevant une VRAIE tension intellectuelle : "Ainsi, nous pouvons nous demander : Dans quelle mesure [tension principale du sujet] ?"]

[ANNONCE DU PLAN DÉTAILLÉE — 2-3 lignes : "Pour répondre à cette interrogation, nous analyserons dans une première partie [intitulé complet Partie I reformulé en 1 ligne], avant d'examiner dans une deuxième partie [Partie II], et d'envisager enfin [Partie III — lycée/université uniquement]."]


---SAUT_DE_PAGE---

# ════════════════════════════════════════════════════════
#                      DÉVELOPPEMENT
# ════════════════════════════════════════════════════════

## I. [TITRE 1re GRANDE PARTIE EN MAJUSCULES — ACCROCHEUR ET PRÉCIS]

════════════════════════════════════════════════════════

### 1.1 [Titre descriptif, précis et original de la 1re sous-partie]


[PARAGRAPHE 1 — 8 à 10 lignes — MODÈLE PEEL :
→ POINT (1-2 lignes) : affirmation directe et claire du sous-argument
→ EXPLICATION (3-4 lignes) : développe le mécanisme, définit les termes en **gras**, explique les causes
→ EXEMPLE IVOIRIEN (3-4 lignes) : chiffre sourcé (institution + année) + fait précis + lieu géographique réel
→ LIEN (1-2 lignes) : transition vers le paragraphe 2]

[PARAGRAPHE 2 — 8 à 10 lignes — même structure PEEL, angle différent et complémentaire. Connecteurs variés. Exemple africain comparatif si pertinent.]

[SI PERTINENT — Tableau de données :
**Tableau 1 : [Titre précis et descriptif]**
| Indicateur | Côte d'Ivoire | Afrique de l'Ouest | Monde |
|------------|--------------|---------------------|-------|
| [Donnée 1] | [Valeur réelle] | [Valeur] | [Valeur] |
| [Donnée 2] | [Valeur réelle] | [Valeur] | [Valeur] |
*Source : [Institution réelle — FAO, BCEAO, INS-CI, Banque Mondiale], [Année]*]

[PARAGRAPHE 3 — Synthèse 1.1 + transition vers 1.2 : 3 à 4 lignes de résumé + phrase d'annonce 1.2]


### 1.2 [Titre descriptif, précis et original de la 2e sous-partie]


[3 paragraphes de 8 à 10 lignes chacun. Angle différent de 1.1. Exemples ivoiriens + données chiffrées.]

---SAUT_DE_PAGE---

## II. [TITRE 2e GRANDE PARTIE EN MAJUSCULES — COMPLÉMENTAIRE À LA PARTIE I]

════════════════════════════════════════════════════════

[TRANSITION OBLIGATOIRE VERS PARTIE II en DÉBUT de partie II — Min 4 lignes, placée APRÈS le titre de la partie II, JAMAIS avant le saut de page : "Ainsi avons-nous établi, au terme de cette première partie, que [synthèse Partie I en 1 phrase]. Cette analyse, si elle permet de [apport], ne saurait toutefois être complète sans que l'on s'interroge sur [ce que la Partie II apporte]. C'est précisément l'objet de notre second axe, consacré à [intitulé Partie II]."
⚠️ Cette phrase de transition doit COMMENCER la Partie II — jamais finir la Partie I.]

### 2.1 [Titre précis de la 1re sous-partie]


[3 paragraphes de 8 à 10 lignes. L'analyse progresse logiquement depuis Partie I. Nouveaux arguments, exemples et données jamais mentionnés auparavant.]


### 2.2 [Titre précis de la 2e sous-partie]


[3 paragraphes de 8 à 10 lignes chacun.]

---SAUT_DE_PAGE---

## III. [TITRE 3e GRANDE PARTIE — POUR LYCÉE ET UNIVERSITÉ UNIQUEMENT]

════════════════════════════════════════════════════════

### 3.1 [Titre précis sous-partie]


[3 paragraphes de 8 à 10 lignes. Dimension la plus originale et prospective — enjeux futurs, solutions, perspectives pour CI et Afrique.]


### 3.2 [Titre précis sous-partie]


[3 paragraphes de 8 à 10 lignes. Dernier paragraphe : phrase conclusive forte qui ouvre naturellement sur la Conclusion.]

════════════════════════════════════════════════════════

---SAUT_DE_PAGE---

# ════════════════════════════════════════════════════════
#                       CONCLUSION
# ════════════════════════════════════════════════════════

## CONCLUSION


[UN SEUL PARAGRAPHE CONTINU de 15 à 20 lignes minimum — PAS de séparation en temps, PAS de sous-titres, PAS de sauts de ligne entre les idées. Le paragraphe doit enchaîner naturellement : bilan synthétique des grandes parties → réponse nuancée à la problématique → ouverture prospective. Tout coulé dans un seul bloc de prose académique soutenu. Exemple de début : "Au terme de cette analyse approfondie, il convient de dresser un bilan lucide des enseignements majeurs que nous avons dégagés. En premier lieu, nous avons mis en évidence que [synthèse Partie I]. Notre deuxième axe de réflexion a démontré que [synthèse Partie II]. Enfin, [synthèse Partie III si applicable]. Au regard de ces éléments, il apparaît clairement que [réponse directe à la problématique], même si cette réponse mérite d'être nuancée : si [aspect positif], il n'en demeure pas moins que [limite]. Cette réflexion nous invite finalement à nous interroger sur [ouverture prospective], enjeu fondamental pour [CI/Afrique/la jeunesse]."]


Rédige maintenant l'exposé COMPLET en français avec la plus grande rigueur académique.

IMPÉRATIFS ABSOLUS :
1. TOUT est rédigé intégralement — zéro "[à compléter]", zéro zone vide
2. Introduction en 5 temps (accroche → contextualisation → délimitation → problématique → annonce plan)
3. Chaque paragraphe suit le modèle PEEL (Point → Explication → Exemple concret sourcé → Lien)
4. Transitions obligatoires entre grandes parties (min 4 lignes chacune)
5. Conclusion en UN SEUL paragraphe continu (bilan → réponse nuancée → ouverture) — jamais découpée
6. Exemples concrets CHIFFRÉS et SOURCÉS — ivoiriens/africains si le sujet le justifie, universels sinon
7. Connecteurs logiques variés — jamais le même deux fois de suite dans un paragraphe"""


        # ================================================================
        # PROMPT — SUJETS & EXAMENS (Système scolaire ivoirien & africain)
        # PROMPT — SUJETS & EXAMENS (Système scolaire ivoirien & africain)
        # ================================================================
        elif "Examens" in service or "Sujets" in service:
            # ── INJECTION DU TYPE DE SUJET dans la description si renseigné ──
            type_sujet_inject = ""
            if "type_sujet_selectionne" in dir() and type_sujet_selectionne:
                TYPE_SUJET_LABELS_FR = {
                    "QCM": "QCM (Questions à Choix Multiple — 4 options A/B/C/D, cases □, un seul type d'exercice)",
                    "VRAI_FAUX": "VRAI ou FAUX UNIQUEMENT (affirmations à évaluer V/F + justification si faux, UN SEUL TYPE d'exercice)",
                    "TEXTE_TROU": "TEXTE À TROUS UNIQUEMENT (texte lacunaire + liste de mots à placer, UN SEUL TYPE d'exercice)",
                    "QUESTIONS_OUVERTES": "QUESTIONS OUVERTES UNIQUEMENT (questions de réflexion rédigées avec lignes de réponse, UN SEUL TYPE)",
                    "MIXTE": "FORMAT MIXTE (Partie 1 QCM + Partie 2 Vrai/Faux + Partie 3 Question ouverte, barème équilibré)",
                    "CAS_PRATIQUE": "CAS PRATIQUE / ÉTUDE DE CAS (texte contextualisé Côte d'Ivoire + questions d'analyse progressives)",
                    "CALCUL": "EXERCICES DE CALCUL / PROBLÈMES (chiffrés, contextualisés CI, formules rappelées, démarche guidée)",
                    "ETUDE_DOCUMENT": "ÉTUDE DE DOCUMENT (document support : texte/tableau/carte + questions d'identification, analyse, interprétation)",
                    "SCHEMA": "SCHÉMA À LÉGENDER (schéma décrit textuellement avec numéros + termes à placer + corrigé légendes)",
                    "DISSERTATION": "COMPOSITION / DISSERTATION GUIDÉE (sujet formulé + consignes de méthode + plan détaillé guidé)",
                    "DEVOIR_COMPLET": "DEVOIR COMPLET AUTHENTIQUE IVOIRIEN (en-tête officiel + exercices variés progressifs adaptés au niveau)",
                }
                label_fr = TYPE_SUJET_LABELS_FR.get(type_sujet_selectionne, type_sujet_selectionne)
                type_sujet_inject = f"""

⚠️ TYPE DE SUJET IMPOSÉ PAR L'UTILISATEUR — RESPECTER ABSOLUMENT :
TYPE UNIQUE SÉLECTIONNÉ : {label_fr}

RÈGLE ABSOLUE : Tu dois générer UN SEUL TYPE D'EXERCICE correspondant EXACTEMENT au type ci-dessus.
- Si QCM → QCM UNIQUEMENT (pas de Vrai/Faux, pas de texte à trous, pas de questions ouvertes)
- Si VRAI_FAUX → Vrai/Faux UNIQUEMENT
- Si TEXTE_TROU → Texte à trous UNIQUEMENT
- Si QUESTIONS_OUVERTES → Questions ouvertes UNIQUEMENT
- Si MIXTE → Les 3 parties indiquées (QCM + Vrai/Faux + Question ouverte)
- Si CAS_PRATIQUE → Un texte de mise en contexte + questions d'analyse
- Si CALCUL → Exercices de calcul/problèmes chiffrés UNIQUEMENT
- Si ETUDE_DOCUMENT → Document support + questions d'exploitation UNIQUEMENT
- Si SCHEMA → Description du schéma numéroté + légendes UNIQUEMENT
- Si DISSERTATION → Sujet + consignes de méthode + plan guidé UNIQUEMENT
- Si DEVOIR_COMPLET → Vrai devoir ivoirien complet : applique INTÉGRALEMENT la SECTION DEVOIR_COMPLET ci-dessous.

NE PAS MÉLANGER LES TYPES sauf si MIXTE ou DEVOIR_COMPLET est explicitement sélectionné.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION DEVOIR_COMPLET — ENCYCLOPÉDIE DES VRAIS DEVOIRS IVOIRIENS PAR NIVEAU
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PRINCIPE FONDAMENTAL :
Un devoir complet ivoirien = mélange INTELLIGENT de types d'exercices DIFFÉRENTS selon le niveau.
JAMAIS deux exercices du même type à la suite. JAMAIS de "mise en situation" en exercice 1 ou 2.
Chaque exercice teste une compétence différente : mémoriser → comprendre → appliquer → analyser.

━━━ 6ème / 5ème ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

EXERCICE 1 — VRAI / FAUX simple
  Tableau : N° | Affirmations | V | F — l'élève coche, pas de justification exigée
  Thèmes : définitions basiques, propriétés simples, vrai/faux du cours
  NE PAS mettre de calculs ici.

EXERCICE 2 — QCM TABLEAU A/B/C
  Tableau : N° | Affirmations incomplètes | A | B | C
  "Une seule réponse est juste. Recopie le numéro suivi de la lettre de la bonne réponse."
  3 à 4 lignes. Distracteurs = erreurs simples fréquentes à ce niveau.

EXERCICE 3 — TEXTE À TROUS ou QUESTIONS COURTES
  Option A : "Recopie et complète avec les mots suivants : [liste de mots]"
  Option B : 3 questions courtes (Définir / Citer / Calculer directement)
  Option C : Schéma à légender (SVT uniquement — liste de termes à placer)

EXERCICE 4 — PROBLÈME SIMPLE CONTEXTE IVOIRIEN
  Mise en situation courte et accessible (marché, maison, jardin ivoirien...)
  3 sous-questions max : 1- lecture données → 2- calcul simple → 3- conclusion
  Données chiffrées simples, opérations de base

━━━ 4ème / 3ème (pré-BEPC) ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

STRUCTURE EXACTE observée sur vrais devoirs ivoiriens 3ème (Collège Sainte Famille, Merlan-Adjamé...) :

EXERCICE 1 — VRAI / FAUX avec instruction de justification
  Consigne exacte : "Écris sur ta feuille de copie le numéro de chacune des affirmations
  ci-dessous et fait suivre par V si l'affirmation est vraie ou F si elle est fausse.
  JUSTIFIE si fausse. Exemple : 5-F"
  5 affirmations dans un tableau : N° | Affirmations | [espace réponse]
  Thèmes : propriétés du chapitre, définitions, théorèmes

EXERCICE 2 — QCM TABLEAU A/B/C ou A/B/C/D
  Consigne exacte : "Pour chacune des affirmations contenues dans le tableau ci-dessous,
  une seule des réponses proposées est juste. Recopie le numéro de la ligne suivi de
  la lettre de la réponse juste."
  Tableau : N° | Affirmations | A | B | C (ou A | B | C | D)
  4 à 5 lignes. Distracteurs = erreurs classiques de 3ème ivoirien.

EXERCICE 3 — CONSTRUCTION GÉOMÉTRIQUE
  Observé réel : "L'unité de longueur est le centimètre. Le segment ci-dessous n'est pas
  en vraie grandeur. On donne le segment [EF] tel que EF = 7 cm."
  Dessin schématique du segment E___F fourni
  Questions : 1- Reproduis le segment sur ta feuille / 2- Construis le point M tel que EM = (3/5)EF
  Peut aussi être : construction de triangle, de cercle, de médiatrice, de bissectrice

EXERCICE 4 — CALCUL ALGÉBRIQUE PUR
  Observé réel : Expression E = (x-3)² + 4(x-3) et R = (2x-1)(x+1)
  1) Justifie que E = (x-3)(x+1)
  2) Détermine les valeurs de x pour lesquelles R existe
  3) Pour x ≠ 1/2 et x ≠ -1, simplifie R
  4) Calcule la valeur numérique de R pour x = -1
  Pas de mise en situation. Calcul algébrique direct : factoriser, simplifier, calculer.

EXERCICE 5 — GÉOMÉTRIE ANALYTIQUE / VECTEURS / TRIGONOMÉTRIE
  Observé réel : Repère orthonormé (O,I,J), points E(6;5) F(2;-3) G(-4;0) EG=5√5
  1) a. Place les points / b. Construis le triangle
  2) Vérifie par calcul que les vecteurs FE et FG ont pour coordonnées (4;8) et (-6;3)
  3) Démontre que le triangle EFG est un triangle rectangle en F
  4) a. Vérifie par calcul que la distance FE = 4√5 / b. Justifie sin(EGF)=0.8
     c. Déduis-en un encadrement de mes EGF par deux entiers consécutifs
        → Fournir une table trigonométrique : tableau a° | cos a° | sin a° avec 4-5 valeurs réelles

EXERCICE 6 — PROBLÈME CONTEXTE IVOIRIEN COMPLEXE
  Observé réel : "Pour la fête de fin d'année, le président de la coopérative du collège
  Sainte Famille prend contact avec deux restaurants de Bouaké."
  Restaurant A : 1000F/repas + 2000F transport | Restaurant B : 950F/repas + 3000F taxi
  Questions : a) Exprime PA et PB en fonction de x / b) Résous l'inéquation / c) Conclus
  TOUJOURS contexte réaliste et concret : coopérative, école, marché, entreprise — ivoirien si possible (FCFA, villes CI), sinon international
  Données en FCFA, villes ivoiriennes (Bouaké, Abidjan, Yamoussoukro...)
  Questions : 1- modéliser → 2- résoudre équation/inéquation → 3- répondre à la question initiale

━━━ LYCÉE (2nde, 1ère, Terminale) ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

STRUCTURE TYPE LYCÉE (basée sur vrais devoirs ivoiriens + BAC blanc) :

EXERCICE 1 — VRAI/FAUX AVEC JUSTIFICATION
  "Écris V si l'affirmation est vraie, F si fausse. JUSTIFIE les affirmations fausses."
  4 à 6 affirmations sur le cours. Justification exigée = 1 ligne de raisonnement.
  NE PAS mettre de calculs ici.

EXERCICE 2 — QCM 4 RÉPONSES A/B/C/D
  Format : tableau N° | Affirmations | A | B | C | D
  "Pour chaque affirmation, 4 réponses sont proposées, une seule est exacte."
  "Écris sur ta copie le numéro + la lettre correspondant à la bonne réponse."
  4 à 6 lignes, distracteurs = erreurs de raisonnement fréquentes au lycée

EXERCICE 3 — APPLICATION DIRECTE DU COURS
  Selon matière :
  → MATHS/PC : calculs guidés, formules à appliquer, démonstration courte
  → SVT : légender schéma + questions de cours
  → HG : recopier et compléter tableau, définitions, localisation sur carte décrite
  → LETTRES : questions sur texte court fourni (vocabulaire, grammaire, sens)
  Sous-questions numérotées 1) 2) 3) avec complexité croissante

EXERCICE 4 — PROBLÈME COMPLEXE / ÉTUDE DE CAS
  Mise en situation détaillée + données (tableau, graphique décrit, texte support)
  Questions en parties A et B, ou numérotées 1. 2. 3. avec sous-questions a) b) c)
  Contexte CI : entreprise ivoirienne, données économiques CI, expérience de labo lycée CI
  Dernière question = synthèse/jugement/ouverture

━━━ BAC / TERMINALE ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Structure proche du vrai BAC ivoirien MENET-FP :

EXERCICE 1 — VRAI/FAUX — court, 4 affirmations max
EXERCICE 2 — QCM 4 RÉPONSES A/B/C/D — 4-6 questions
EXERCICE 3 — PROBLÈME STRUCTURÉ
  Parties A / B / C clairement titrées
  A = définitions/rappels de cours
  B = application et calculs
  C = analyse/synthèse/démonstration
EXERCICE 4 — PROBLÈME DE SYNTHÈSE ou COMMENTAIRE DE DOCUMENT
  Document support (texte, tableau de données, graphique décrit)
  Questions d'exploitation, interprétation, rédaction structurée
  Dernière question obligatoirement ouverte

━━━ RÈGLES TRANSVERSALES POUR TOUT DEVOIR COMPLET ━━━━━━━━━━━━━━━━━━━━━━━

✅ TOUJOURS respecter la progression : Connaître → Comprendre → Appliquer → Analyser
✅ JAMAIS deux exercices du même format à la suite
✅ JAMAIS de mise en situation complexe en Exercice 1 ou 2
✅ Les exercices 1 et 2 sont TOUJOURS fermés (V/F, QCM, tableau)
✅ Les exercices 3+ peuvent être ouverts ou semi-ouverts
✅ Chaque exercice a son barème clairement indiqué : ## EXERCICE N :
✅ Numérotation cohérente : 1- 2- 3- puis 1.1- 1.2- ou a) b) c) selon niveau
✅ Contextes réalistes dans les problèmes (noms, villes, monnaies, produits) — ivoiriens si pertinent pour la matière, internationaux sinon
✅ Données chiffrées précises et réalistes (jamais de "x valeur" vague)
✅ L'en-tête simple est TOUJOURS présent (voir Section En-tête ci-dessus)
"""

            prompt = f"""Tu es NOVA EXAM — le concepteur officiel de sujets d\'examens numéro 1 du système scolaire ivoirien.
Tu maîtrises tous les programmes officiels MENET-FP/DECO, tous les formats CEPE, BEPC, BAC et concours, et tu es expert en mise en page Word professionnelle via python-docx.
Chaque sujet que tu produis est ENTIÈREMENT rédigé, rigoureusement structuré, et immédiatement utilisable en classe.

╔══════════════════════════════════════════════════════════════════╗
║     NOVA EXAM — ENCYCLOPÉDIE COMPLÈTE DE CRÉATION DE SUJETS     ║
╚══════════════════════════════════════════════════════════════════╝

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 1 — MISE EN PAGE PROFESSIONNELLE NOVA EXAM
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

╔══════════════════════════════════════════════════════════════════╗
║  RÉFÉRENCE : Vrais sujets ivoiriens analysés (APFC Bouaké 2026, ║
║  Devoirs UP régionaux, Devoirs lycée Abidjan/Bouaké)            ║
╚══════════════════════════════════════════════════════════════════╝

━━━ A. TYPOGRAPHIE EXACTE (obligatoire) ━━━━━━━━━━━━━━━━━━━━━━━━━━

POLICE : Times New Roman (serif) — comme tous les vrais sujets ivoiriens
  - Corps du texte (consignes, énoncés) : 11pt, interligne 1.15
  - Tableaux (intérieur cellules) : 10pt, interligne simple 1.0
  - Titre MATIÈRE dans cadre : 16-18pt, MAJUSCULES, gras, centré
  - Niveau/Classe : 13pt, gras, centré
  - En-tête petit texte (établissement, date) : 10pt normal
  - Numéro de page bas : 10pt centré italique
  - Mention calculatrice/durée : 10pt italique centré

MARGES WORD (###MARGES### pour le moteur python-docx) :
  ###MARGES### haut=2cm bas=2cm gauche=2.5cm droite=2cm

━━━ B. STRUCTURE EN-TÊTE EXACTE — MARQUEUR ###ENTETE_DEVOIR### ━━━━

⚠️ RÈGLE ABSOLUE : L'en-tête du devoir est ENTIÈREMENT géré par Python.
Tu dois UNIQUEMENT émettre le bloc marqueur ci-dessous avec les données brutes.
Python construit lui-même le tableau 2-colonnes, le cadre titre, et la mention italique.
NE PAS écrire l'en-tête en texte libre. UNIQUEMENT le bloc marqueur.

MODÈLE EXACT À ÉMETTRE (adapte les valeurs selon la demande client) :

###ENTETE_DEVOIR###
ETABLISSEMENT=Collège Sainte Famille de Bouaké
DISCIPLINE=CE PHYSIQUE-CHIMIE
CLASSE=2ndC
ANNEE=2025-2026
DATE=21/04/2026
DUREE=2H
TITRE_DEVOIR=DEVOIR DE NIVEAU N°1  3ᵉ Trimestre
MENTION=Cette épreuve comporte trois (03) pages numérotées 1/4 et 2/4 La calculatrice scientifique est autorisée.
###FIN_ENTETE###

RÈGLES D'ADAPTATION :
- ETABLISSEMENT : nom complet de l'école (ou laisser vide si non précisé)
- DISCIPLINE : matière abrégée (CE MATHS, CE PHYSIQUE-CHIMIE, CE SVT, CE HG, CE FRANÇAIS...)
- CLASSE : niveau exact (2ndC, 3ème, Terminale D, CM2...)
- ANNEE : si non précisée → année en cours (2025-2026)
- DATE : si non précisée → laisser vide (le champ sera omis)
- DUREE : durée officielle selon le niveau (1H, 2H, 3H)
- TITRE_DEVOIR : "DEVOIR DE NIVEAU N°1 1er Trimestre" / "DEVOIR DE NIVEAU N°1 3ᵉ Trimestre" / "COMPOSITION DU 1er TRIMESTRE" / "DEVOIR SURVEILLÉ N°2"...
- MENTION : mention calculatrice + nombre de pages (adapter selon matière)

APRÈS le bloc ###ENTETE_DEVOIR### : NE PAS mettre de ligne séparatrice ════.
La séparation vient AVANT le premier exercice.

SOUS-EXERCICE (quand un exercice a 2 matières comme Chimie + Physique) :
  ###SOUS_EXERCICE### CHIMIE
  [contenu chimie...]
  ###SOUS_EXERCICE### PHYSIQUE
  [contenu physique...]

LIGNE SÉPARATRICE entre exercices :
  ════════════════════════════════════════════════════════════════

━━━ C. FORMAT DES EXERCICES (observation directe des sujets CI) ━━━

TITRE D'EXERCICE — format exact :
  ## EXERCICE 1 — Niveau : Facile
  → JAMAIS mettre de points/barème dans le titre — ex: JAMAIS "EXERCICE 1" ni "(X pts)"
  → Le barème est laissé entièrement au professeur — Nova ne met PAS de points
  → Toujours inclure "— Niveau : [niveau]" à la fin du titre, séparé par un tiret long
  → En python-docx : Times New Roman 12pt, GRAS uniquement, noir — PAS de soulignement
  → NE PAS mettre de couleur bleue — les vrais sujets sont en noir
  → NE PAS mettre de deux-points après le numéro : "EXERCICE 1" et non "EXERCICE 1 :"
  → Espacement avant l'exercice : 14pt / après le titre : 4pt

NIVEAUX DE DIFFICULTÉ PAR EXERCICE (progression obligatoire) :
  Les exercices vont toujours du plus facile au plus difficile — l'élève gagne confiance avant d'affronter le difficile.

  PHYSIQUE-CHIMIE (4 exercices) :
    Ex 1 → Niveau : Facile         (restitution pure, aucun calcul, cours direct)
    Ex 2 → Niveau : Moyen          (application guidée, calculs simples avec formules données)
    Ex 3 → Niveau : Difficile      (circuit, schéma, caractéristique, raisonnement)
    Ex 4 → Niveau : Approfondissement  (synthèse, oscilloscope, optique, nucléaire)

  MATHÉMATIQUES — Collège (4 exercices) :
    Ex 1 → Niveau : Facile         (Vrai/Faux, définitions, reconnaissance)
    Ex 2 → Niveau : Moyen          (QCM, calcul direct)
    Ex 3 → Niveau : Difficile      (problème guidé, géométrie)
    Ex 4 → Niveau : Approfondissement  (problème contextualisé FCFA/CI complet)

  MATHÉMATIQUES — Lycée (5 exercices) :
    Ex 1 → Niveau : Facile         (Vrai/Faux avec ou sans justification)
    Ex 2 → Niveau : Facile/Moyen   (QCM, reconnaissance)
    Ex 3 → Niveau : Moyen          (application directe du cours)
    Ex 4 → Niveau : Difficile      (étude de fonction, géométrie analytique, stats)
    Ex 5 → Niveau : Approfondissement  (problème de synthèse contextualisé CI)

  SVT / HG / FRANÇAIS / ECM / ANGLAIS :
    Ex 1 → Niveau : Facile
    Ex 2 → Niveau : Moyen
    Ex 3 → Niveau : Difficile  (ou Approfondissement si c'est le dernier)

  RÈGLE ABSOLUE : ne jamais mettre "Difficile" ou "Approfondissement" en premier exercice.
  L'élève doit toujours commencer par quelque chose d'accessible.

CONSIGNE — toujours en paragraphe normal après le titre :
  → 11pt Times New Roman, interligne 1.15
  → Commence par une phrase complète : "Écris sur ta copie...", "Pour chacun des..."
  → PAS de liste à puces pour la consigne elle-même

NUMÉROTATION DES QUESTIONS (style officiel ivoirien — observé sur vrais sujets CI) :
  → Niveau 1 : 1.  2.  3.  4.   (point après le chiffre, PAS de parenthèse)
  → Niveau 2 : 1.1  1.2  2.1  2.2  2.3  (chiffre parent + point + sous-chiffre, SANS parenthèse)
  → Niveau 3 (rare) : 1.1.1  1.1.2
  → JAMAIS : 1-a)  1-b)  a)  b)  — ces formes ne correspondent pas aux vrais sujets CI
  → Chaque sous-question sur sa propre ligne, indentée de 0.5cm
  → Exemple correct :
      1. Dis ce que représentent 200 mg/L.
      2. Détermine :
      2.1 la concentration molaire volumique de la solution ;
      2.2 la masse du soluté utilisée par jour.

LIGNES DE RÉPONSE ÉLÈVE :
  → PAR DÉFAUT : AUCUNE ligne vide dans les devoirs de Physique-Chimie et Mathématiques.
    Les vrais sujets CI ne comportent pas de lignes de réponse — l'élève répond sur sa copie.
  → EXCEPTION uniquement si le client demande explicitement "avec lignes de réponse" ou "cahier d'élève" :
    Lignes vides : _______________________________________________ (30+ underscores)
    Minimum 2 lignes pour questions courtes, 5+ pour rédactions, 4-6 pour calculs longs.
  → FRANÇAIS expression écrite et HG composition : toujours avec lignes vides (voir sections dédiées).

SÉPARATEUR ENTRE EXERCICES :
  ════════════════════════════════════════════════════════════════
  (ligne pleine, noir, avant chaque nouvel exercice sauf le premier)

SAUT DE PAGE : ---SAUT_DE_PAGE---
  → Entre page 1 et page 2 du sujet
  → JAMAIS précédé d'un ════

NUMÉRO DE PAGE EN BAS :
  → "~ Page 1/2 ~" centré 10pt italique
  → Inséré automatiquement via footer Word

━━━ D. TABLEAUX (format exact des vrais sujets CI) ━━━━━━━━━━━━━━━

TABLEAU QCM (Exercice à choix multiples) — modèle EXACT observé :
| N° | ÉNONCÉS INCOMPLETS | Réponse A | Réponse B | Réponse C | Réponse D |
|----|--------------------|-----------|-----------|-----------|-----------|
| 1  | [affirmation 1]    | [opt A]   | [opt B]   | [opt C]   | [opt D]   |
  → En-tête : fond bleu foncé, texte blanc, gras, centré, 10pt
  → Colonne N° : largeur 0.8cm, centré
  → Colonne Énoncés : largeur ~8cm (50% de la largeur utile), aligné gauche, 10pt
  → Colonnes Réponses A/B/C/D : largeur égale ~2.5cm chacune, centré, 10pt
  → Lignes alternes blanc/gris très léger
  → Interligne dans cellules : simple 1.0

TABLEAU VRAI/FAUX — modèle EXACT observé :
| N° | Affirmation | V | F | Justification (si fausse) |
|----|-------------|---|---|---------------------------|
| 1  | [texte affirmation complète rédigée] | ☐ | ☐ | |
| 2  | [texte affirmation complète rédigée] | ☐ | ☐ | |
| 3  | [texte affirmation complète rédigée] | ☐ | ☐ | |
| 4  | [texte affirmation complète rédigée] | ☐ | ☐ | |
| 5  | [texte affirmation complète rédigée] | ☐ | ☐ | |
  → Colonne N° : 0.8cm centré
  → Colonne Affirmation : ~8.5cm, aligné gauche — rédige l'affirmation complète directement dans la cellule
  → Colonnes V/F : 1.2cm chacune, centré — utilise ☐ (U+2610) et non □ ni [ ]
  → Colonne Justification : ~3.5cm — laisser vide (l\'élève écrit sur la copie)
  → JAMAIS de lignes ___ dans le tableau — l\'élève répond sur copie
  → La colonne Affirmation peut contenir des symboles Unicode (², √, Δ...)

TABLEAU DONNÉES (problème, exercice contextuel) :
  → 2 colonnes : Grandeur | Valeur
  → Ou tableau donné dans l'énoncé tel quel
  → Fond en-tête bleu foncé, texte blanc

RÈGLE TABLEAU : Un tableau NE DOIT JAMAIS être coupé entre deux pages.
  Si risque de coupure → ---SAUT_DE_PAGE--- AVANT le tableau.

━━━ E. MISE EN PAGE UNIVERSELLE PAR MATIÈRE ━━━━━━━━━━━━━━━━━━━━━

⚠️ RÈGLE UNIVERSELLE — ###ENTETE_DEVOIR### ET ###SOUS_EXERCICE### :
  • ###ENTETE_DEVOIR### s'applique à TOUS les sujets de TOUTES les matières.
    Python construit entièrement l'en-tête (tableau 2-colonnes, titre encadré, mention italique).
    Adapte ETABLISSEMENT, DISCIPLINE, CLASSE, ANNEE, DATE, DUREE, TITRE_DEVOIR, MENTION
    selon la demande client.

  • ###SOUS_EXERCICE### s'applique à TOUTES les matières qui ont des sous-parties dans un exercice.
    Ex PC : CHIMIE + PHYSIQUE
    Ex SVT : PARTIE A — Génétique + PARTIE B — Physiologie
    Ex Maths : PARTIE ALGÈBRE + PARTIE GÉOMÉTRIE
    Ex HG : HISTOIRE + GÉOGRAPHIE
    → JAMAIS utiliser ### ou ## pour des sous-parties — toujours ###SOUS_EXERCICE###

  • EXCEPTION : FRANÇAIS EXPRESSION ÉCRITE (rédaction libre, récit, lettre, argumentation)
    → PAS de ###SOUS_EXERCICE### pour la production écrite
    → La consigne suffit, suivie des lignes vides ___________ en nombre précis

── MATHÉMATIQUES ──────────────────────────────────────────────────
  ⚠️ RÈGLE FONDAMENTALE : adapte CHAQUE exercice au niveau de classe réel.
  Un devoir de 6ème et un devoir de Terminale C sont COMPLÈTEMENT différents.
  Ne génère JAMAIS des exercices trop faciles pour un lycéen ni trop durs pour un collégien.
  Varie les types d'exercices à chaque génération — ne répète pas toujours le même format.

  PROGRESSION GÉNÉRALE (du plus facile au plus difficile) :
  Ex 1 → Restitution / Reconnaissance (facile, cours direct)
  Ex 2 → Application simple (calcul direct, formule connue)
  Ex 3 → Problème guidé (plusieurs étapes, raisonnement)
  Ex 4 → Analyse / Approfondissement (plus complexe, plusieurs sous-questions)
  Ex 5 (lycée uniquement) → Synthèse contextualisée CI (modélisation + résolution complète)

  ── Collège 6ème ────────────────────────────────────────────────
  Notions : entiers, décimaux, fractions, périmètre, aire, angles, symétrie
  Ex 1 : Vrai/Faux OU calcul mental OU compléter un tableau de nombres
  Ex 2 : Calcul direct sur fractions/décimaux OU périmètre/aire figure simple
  Ex 3 : Problème contextualisé marché Bouaké/Abidjan (FCFA, pesée, partage)
  Ex 4 : Géométrie (tracer, calculer, justifier) OU statistiques simples (tableau, diagramme)
  ❌ JAMAIS : équations, algèbre, Pythagore, fonctions

  ── Collège 5ème ────────────────────────────────────────────────
  Notions : proportionnalité, pourcentages, Thalès (intro), statistiques, fractions avancées
  Ex 1 : Vrai/Faux sur proportionnalité/pourcentages OU QCM
  Ex 2 : Calcul de pourcentage OU tableau de proportionnalité (prix marché, recette, vitesse)
  Ex 3 : Thalès simple OU problème de proportionnalité contextualisé CI
  Ex 4 : Statistiques (moyenne, médiane, diagramme) OU problème à plusieurs étapes

  ── Collège 4ème ────────────────────────────────────────────────
  Notions : équations 1er degré, systèmes simples, Pythagore, cercle, angles
  Ex 1 : Vrai/Faux sur Pythagore/équations OU QCM
  Ex 2 : Résolution d'équation OU calcul Pythagore direct
  Ex 3 : Problème géométrique guidé (triangle rectangle, cercle, construction)
  Ex 4 : Système d'équations OU problème contextualisé CI (construction, terrain)

  ── Collège 3ème / BEPC ─────────────────────────────────────────
  Notions : fonctions affines, inéquations, géométrie dans l'espace, stats (σ, médiane)
  Ex 1 : Vrai/Faux avec justification OU QCM 4 réponses
  Ex 2 : Fonction affine (tableau de valeurs, graphe décrit, pente, ordonnée à l'origine)
  Ex 3 : Inéquation OU géométrie dans l'espace (volume, patron)
  Ex 4 : Statistiques (moyenne pondérée, médiane, σ) OU problème de synthèse BEPC

  ── Lycée 2nde ──────────────────────────────────────────────────
  Notions : fonctions numériques (affine, carré, valeur absolue), probabilités, statistiques (σ)
  Ex 1 : Vrai/Faux AVEC justification OU QCM 4 réponses (erreurs courantes en distracteurs)
  Ex 2 : Fonction affine/carré — tableau de valeurs, sens de variation, extremum
         OU probabilités discrètes simples (dé, urne, arbre)
  Ex 3 : Étude de fonction (signe, variations, tableau, courbe décrite textuellement)
         OU statistiques (moyenne, variance, σ, interprétation)
  Ex 4 : Problème de synthèse — modélisation d'une situation réelle CI avec une fonction
         (bénéfice, coût, population, distance) + résolution + interprétation

  ── Lycée 1ère ──────────────────────────────────────────────────
  Notions : dérivées, suites, trigo (cos/sin/tan), limites de suites
  Ex 1 : Vrai/Faux avec justification OU QCM sur dérivées/trigo
  Ex 2 : Calcul de dérivée OU suite arithmétique/géométrique (terme général, somme)
  Ex 3 : Étude complète de fonction avec dérivée (variations, extrema, tableau)
         OU problème de trigonométrie (résolution d'équation trigo, valeurs remarquables)
  Ex 4 : Problème de synthèse — suite modélisant une situation CI (épargne, démographie)
         OU fonction avec étude complète + interprétation économique/physique

  ── Lycée Terminale C/D ─────────────────────────────────────────
  Notions : intégrales, ln/exp, dénombrement, probabilités continues (loi normale)
  Ex 1 : Vrai/Faux AVEC justification (notions ln/exp/intégrale) OU QCM 4 réponses
  Ex 2 : Calcul intégrale OU étude de fonction ln/exp (domaine, limites, dérivée, variations)
  Ex 3 : Dénombrement (arrangements, combinaisons, permutations) OU probabilités (loi normale, binomiale)
  Ex 4 : Problème complexe — modélisation CI (croissance démographique, capitalisation FCFA,
         décroissance radioactive) + étude complète + interprétation + conclusion
  Ex 5 : Synthèse BAC — sujet complet multi-notions avec barème détaillé

  RÈGLES COMMUNES MATHS :
  → Contexte obligatoire CI : FCFA, marchés (cacao, anacarde, riz, attiéké), villes (Abidjan, Bouaké,
    Yamoussoukro, San-Pédro), construction, transport, agriculture
  → Figures géométriques : décrire textuellement — "Soit le triangle ABC rectangle en B, AB=3cm, BC=4cm"
  → Graphes : donner tableau de valeurs + décrire l'allure — Nova ne dessine pas
  → ###FORMULE### pour toute formule utilisée dans un calcul
  → SOUS-SECTIONS si mixte : ###SOUS_EXERCICE### ALGÈBRE / GÉOMÉTRIE / STATISTIQUES

── PHYSIQUE-CHIMIE ────────────────────────────────────────────────
  ⚠️ RÈGLE FONDAMENTALE : adapte chaque exercice au niveau réel de la classe.
  Un devoir de 4ème et un devoir de Tle D sont COMPLÈTEMENT différents.
  Varie les contextes et types à chaque génération — ne répète jamais toujours oscilloscope ou NaCl.

  PROGRESSION GÉNÉRALE (du plus facile au plus difficile) :
  Ex 1 → Restitution cours pur (2 disciplines mélangées : CHIMIE + PHYSIQUE)
  Ex 2 → Application guidée avec contexte ivoirien (calculs simples, données fournies)
  Ex 3 → Approfondissement (schéma, graphe, circuit, raisonnement)
  Ex 4 → Synthèse (situation réelle complexe, plusieurs notions liées)

  ── Collège 4ème / 3ème ─────────────────────────────────────────
  CHIMIE notions : corps purs/mélanges, dissolution, filtration, distillation, atomes/molécules,
                   tableau périodique simplifié, ions courants (Na+, Cl-, Cu2+, Fe2+/3+)
  PHYSIQUE notions : courant électrique, circuit simple, tension/intensité, conducteurs ohmiques,
                     loi d'Ohm, résistance, dipôles (pile, lampe, interrupteur, résistance)

  Ex 1 — CHIMIE : texte à compléter avec liste de mots (corps pur, mélange, atome, ion...)
         PHYSIQUE : Vrai/Faux sur circuit électrique (dipôle actif/passif, série/parallèle)
  Ex 2 : Préparation d'une solution saline OU dissolution (masse, concentration massique Cm=m/V)
         Contexte : infirmerie école, eau potable village CI, eau salée dentiste
         ###FORMULE### Cm = m/V  |  ###FORMULE### n = m/M
  Ex 3 : Circuit série simple avec résistance + lampe — loi d'Ohm, calcul U, I, R
         Description du schéma + questions : U=? I=? R=? P=?
         ###FORMULE### U = R × I  |  ###FORMULE### P = U × I
  Ex 4 : Identification d'ions (réactifs, précipités colorés) OU sécurité électrique (fusibles, disjoncteurs)
         Contexte CI : laboratoire lycée Abidjan, installation électrique village

  ── Lycée 2nde ──────────────────────────────────────────────────
  CHIMIE notions : solutions aqueuses, concentration molaire, pH, ions, électroneutralité,
                   réactions chimiques simples, équations bilan
  PHYSIQUE notions : circuit électrique (dipôles, série/parallèle), loi d'Ohm, résistances,
                     puissance, énergie, diode simple, conducteur ohmique

  Ex 1 — CHIMIE : texte à trous OU Vrai/Faux sur solutions/ions/pH
         PHYSIQUE : Vrai/Faux sur dipôles, circuits, conducteur ohmique (observé : 4 propositions)
  Ex 2 : Problème solution aqueuse — NaCl, HCl, NaOH... (Cm, n, électroneutralité, préparation)
         Contexte : pharmacie Abidjan, bain de bouche dentiste, eau potable Bouaké
         ###FORMULE### Cm = n/V  |  ###FORMULE### n = m/M  |  ###FORMULE### C = Cm × M
  Ex 3 : Circuit avec diode + résistance de protection — caractéristique tension-intensité ASCII
         Questions dans l'ordre : sens passant/inverse → nature dipôle → tension aux bornes
         → schéma → calcul résistance protection
         ###FORMULE### U = R × I  |  ###FORMULE### U_total = U_diode + U_R
  Ex 4 (VARIER selon génération — NE PAS toujours faire oscilloscope) :
         • Oscilloscope/onduleur solaire : oscillogramme ASCII + sensibilités + U_max, T, U_eff, f
         • Panneaux solaires : puissance, rendement, énergie produite par jour
         • Optique géométrique : lentille convergente, foyer, vergence, construction décrite
         • Mécanique : vitesse, distance, durée, accélération (contexte transport CI)

  ── Lycée 1ère ──────────────────────────────────────────────────
  CHIMIE notions : réactions chimiques (acide-base, oxydo-réduction), pH, équations bilan,
                   stœchiométrie simple, dosages/titrages
  PHYSIQUE notions : lois de Kirchhoff, ponts diviseurs, puissance/énergie, optique (lentilles),
                     mécanique (MRU, MRUA, chute libre)

  Ex 1 — CHIMIE : QCM sur réactions acide-base OU Vrai/Faux avec justification sur pH/titrages
         PHYSIQUE : Vrai/Faux avec justification sur optique OU mécanique
  Ex 2 : Dosage/titrage acide-base (HCl + NaOH) OU calcul stœchiométrique
         Contexte : laboratoire lycée, contrôle qualité eau CI, industrie agroalimentaire
         ###FORMULE### n_acide = n_base à l'équivalence
  Ex 3 (VARIER) :
         • Lentille convergente : construction décrite, calcul vergence, distance focale, grandissement
         • Mécanique MRUA : chute libre, freinage voiture CI, lancer vertical
         • Kirchhoff : maille, nœud, calcul courant dans circuit complexe
  Ex 4 : Synthèse complexe — plusieurs notions liées
         Contexte CI : centrale électrique Soubré, usine chimique Abidjan, véhicule électrique

  ── Lycée Terminale D / C ────────────────────────────────────────
  CHIMIE notions : cinétique chimique, équilibres, constante d'équilibre Ka/Kb, pKa,
                   chimie organique (alcanes, alcènes, alcools, acides carboxyliques),
                   réactions nucléaires (fission, fusion, radioactivité)
  PHYSIQUE notions : circuit RLC, oscillations, ondes (son, lumière), mécanique avancée
                     (moment cinétique, satellite, pendule), relativité (intro BAC)

  Ex 1 — CHIMIE : Vrai/Faux AVEC justification sur chimie organique OU cinétique/équilibre
         PHYSIQUE : QCM 4 réponses sur oscillations OU ondes OU mécanique avancée
  Ex 2 (VARIER) :
         • Cinétique : calcul vitesse de réaction, temps de demi-réaction, graphe concentration/temps
         • Chimie organique : identifier famille, nommer composé, écrire réaction
         • Réaction nucléaire : équation bilan, énergie libérée E=mc², application centrale nucléaire
  Ex 3 (VARIER) :
         • Oscillations libres : pendule simple, ressort — T=2π√(L/g), énergie mécanique
         • Ondes sonores : célérité, fréquence, longueur d'onde, niveau sonore dB
         • Mécanique satellite : vitesse orbitale, période, énergie mécanique
         • RLC : résonance, impédance, facteur de qualité
  Ex 4 : Synthèse BAC — situation réelle complexe CI multi-notions
         Exemples : centrale nucléaire de Côte d'Ivoire (fictive), satellite météo africain,
         véhicule électrique ivoirien, analyse spectrale eau du lac Kossou
         Questions : identifier → calculer → interpréter → conclure → critique scientifique

  RÈGLES COMMUNES PC :
  → ###SOUS_EXERCICE### CHIMIE / PHYSIQUE obligatoire dans Ex 1
  → ###FORMULE### pour chaque formule utilisée dans un calcul, TOUJOURS
  → Données numériques toujours listées explicitement avant les questions
  → Contexte ivoirien obligatoire : Abidjan, Bouaké, Yamoussoukro, San-Pédro,
    lac Kossou, barrage de Soubré, SODECI, CIE, hôpitaux CI
  → Si oscilloscope : oscillogramme ASCII + sensibilité verticale (V/div) + horizontale (ms/div)
  → Si schéma circuit : description ASCII du montage obligatoire

── SVT — TOUTES CLASSES ────────────────────────────────────────────
  PROGRESSION OBLIGATOIRE :
  Ex 1 → RESTITUTION : QCM OU texte à trous OU vrai/faux (cours pur, aucun raisonnement)
    → Textes à trous dans le corps du texte : "La cellule est limitée par une ___________"
    → QCM lettres sous chaque question : a) b) c) d) (pas en tableau pour SVT)
  Ex 2 → ANALYSE DOCUMENT : texte scientifique 80-120 mots + questions
    → Questions : observer → identifier → expliquer → relier au cours
    ###SOUS_EXERCICE### Exploitation du document
    ###SOUS_EXERCICE### Questions de cours liées
  Ex 3 → SCHÉMA + SYNTHÈSE : schéma numéroté à légender + raisonnement
    → Schéma décrit avec numéros : "La figure ci-contre représente... Légendez les numéros 1 à 6."
    → Liste des numéros à placer : "1-___ 2-___ 3-___ 4-___ 5-___ 6-___"
    → Puis 2-3 questions d'interprétation ou d'argumentation

  SOUS-SECTIONS SVT fréquentes :
  → ###SOUS_EXERCICE### PARTIE A — Génétique
  → ###SOUS_EXERCICE### PARTIE B — Physiologie
  → ###SOUS_EXERCICE### PARTIE C — Écologie

── HISTOIRE-GÉOGRAPHIE — TOUTES CLASSES ──────────────────────────
  PROGRESSION OBLIGATOIRE :
  Ex 1 → RESTITUTION : 5 questions courtes numérotées (définir, citer, localiser, dater)
    → 1-2 lignes de réponse chacune
    → Jamais de rédaction longue ici
  Ex 2 → ANALYSE DE DOCUMENT : 1 ou 2 documents rédigés INTÉGRALEMENT + questions
    ###SOUS_EXERCICE### Document A — [titre] (texte 6-10 lignes rédigé, source fictive réaliste)
    ###SOUS_EXERCICE### Document B — [titre] (tableau de données chiffrées OU 2e texte)
    → Questions d'exploitation : identifier → analyser → interpréter → critiquer
    → RÈGLE ABSOLUE : aucun document vide ou "[insérer texte]" — tout rédigé
  Ex 3 → COMPOSITION / SYNTHÈSE : sujet rédigé en 15-20 lignes
    → Consigne avec 3 axes obligatoires à traiter
    → Lignes vides proportionnelles : 15-20 underscores de 30+ caractères
    → Barème : "3 pts méthode/structure + 5 pts fond/contenu"

  SOUS-SECTIONS HG (si 2 disciplines dans même exercice) :
  → ###SOUS_EXERCICE### HISTOIRE
  → ###SOUS_EXERCICE### GÉOGRAPHIE

── FRANÇAIS — TOUTES CLASSES ──────────────────────────────────────
  STRUCTURE TYPIQUE (4 exercices/parties) :
  I. COMPRÉHENSION DE TEXTE (texte support rédigé en italique ou guillemets, 80-200 mots selon niveau)
     → Questions numérotées 1. 2. 3. 4. sur le sens, le vocabulaire, les idées
  II. GRAMMAIRE / LANGUE : questions sur nature des mots, fonctions, conjugaison, orthographe
  III. LEXIQUE / STYLE : synonymes, antonymes, champ lexical, figures de style
  IV. EXPRESSION ÉCRITE (EXCEPTION — PAS de ###SOUS_EXERCICE###) :
     → Consigne claire avec type (récit, lettre, argumentation, description), longueur (X lignes)
     → Lignes vides : ______________ (30+ underscores) × nombre de lignes attendu
     → Ex : "Raconte en 15 lignes une journée de marché à Bouaké avec ta famille."
     → Ex : "Rédige une lettre de 20 lignes à ton ami pour lui décrire ton quartier d'Abidjan."

  ⚠️ EXPRESSION ÉCRITE = JAMAIS de sous-sections, jamais de QCM, jamais de tableau
  Le sujet de rédaction est formulé en 1-3 phrases, suivi des lignes vides. C'est tout.

── ECM / EDHC — TOUTES CLASSES ────────────────────────────────────
  PROGRESSION OBLIGATOIRE :
  Ex 1 → RESTITUTION : définitions, droits/devoirs, institutions, symboles nationaux CI
    → Questions : "Définissez... Citez... Quels sont les 3..."
  Ex 2 → CAS PRATIQUE : situation-problème ivoirienne + questions d'analyse civique
    → Contexte : élection, conflit à l'école, problème de santé publique, environnement...
    ###SOUS_EXERCICE### Analyse de la situation
    ###SOUS_EXERCICE### Votre avis argumenté
  Ex 3 → RÉDACTION CIVIQUE : "Rédigez en X lignes votre opinion sur..."
    → Jamais de ###SOUS_EXERCICE### pour la partie rédaction libre

── ANGLAIS — TOUTES CLASSES ────────────────────────────────────────
  PROGRESSION OBLIGATOIRE :
  Ex 1 → COMPRÉHENSION : texte en anglais (80-200 mots selon niveau) + questions
    → "Read the text and answer the following questions:"
    ###SOUS_EXERCICE### Reading Comprehension
    ###SOUS_EXERCICE### Vocabulary
  Ex 2 → GRAMMAIRE : exercices de structure (fill in the blank, tense, transformation)
    → "Put the verbs in brackets in the correct tense:"
  Ex 3 → EXPRESSION ÉCRITE (PAS de ###SOUS_EXERCICE###) :
    → "Write a paragraph of X words about [sujet lié à la CI / Afrique]"

── ÉCONOMIE / COMPTABILITÉ — BAC B, G ─────────────────────────────
  PROGRESSION OBLIGATOIRE :
  Ex 1 → DÉFINITIONS / NOTIONS : définir concepts, citer exemples, compléter tableau
  Ex 2 → APPLICATION : calculs comptables, écritures journal, bilan partiel
    ###SOUS_EXERCICE### Enregistrements au journal
    ###SOUS_EXERCICE### Présentation du bilan
  Ex 3 → ÉTUDE DE CAS : entreprise ivoirienne fictive + questions progressives
    ###SOUS_EXERCICE### Analyse de la situation
    ###SOUS_EXERCICE### Calculs et justifications
    ###SOUS_EXERCICE### Conclusion et recommandations

── PHILOSOPHIE — Terminale ─────────────────────────────────────────
  Structure BAC CI : explication de texte + dissertation
  → Texte philosophique rédigé 150-200 mots (Platon, Descartes, Kant, Sartre, Nkrumah...)
  → Questions explication : 1. Idée principale / 2. Expliquez le passage / 3. Discutez
  → Sujet dissertation : problématique formulée + plan suggéré en 2-3 axes
  → PAS de ###SOUS_EXERCICE### pour la dissertation (format libre structuré)

━━━ F. MOTEUR DE FORMULES NOVA (inchangé — obligatoire) ━━━━━━━━━━

① ###FORMULE### [formule] — formule seule sur sa ligne (fond bleu, centré, 13pt gras)
  ###FORMULE### E = m × c^{{2}}
  ###FORMULE### pH = -log([H_{{3}}O^{{+}}])
  ###FORMULE### x_{{1,2}} = (-b ± √(b^{{2}}-4ac)) / (2a)
  ###FORMULE### ∫_{{a}}^{{b}} f(x) dx = F(b) - F(a)
  ###FORMULE### T = 2π × √(L/g)
  ###FORMULE### N(t) = N_{{0}} × e^{{-λt}}

② ###DEBUT_FORMULES### / ###FIN_FORMULES### — bloc de formules liées
  ###DEBUT_FORMULES###
  U = R × I
  P = U × I = R × I^{{2}} = U^{{2}} / R
  ###FIN_FORMULES###

③ Formule INLINE dans le texte : notation ^{{}} et _{{}}
  "On a v = d/t = 180/2,5 = 72 km/h"
  "Δ = b^{{2}} - 4ac = 25 - 24 = 1 > 0 donc deux solutions réelles"
  "M(H_{{2}}O) = 2×M(H) + M(O) = 2×1 + 16 = 18 g/mol"

④ INTERDITS ABSOLUS :
  ✗ LaTeX : \frac \text \left \right $...$ $$...$$
  ✗ Parenthèses vides : ()/(x) (a)/() _(x) ^()
  ✗ Blocs markdown : ``` ~~~ ```markdown
  ✗ Document cité sans contenu rédigé
  ✗ Tableau coupé entre deux pages
  ✗ Notation texte brut pour symboles mathématiques : JAMAIS x^2, sqrt(delta), sqrt(x), R^2, Delta, alpha — utilise les vrais caractères Unicode

⑤ RÈGLE UNICODE OBLIGATOIRE — SYMBOLES MATHÉMATIQUES :
  Tu n'utilises JAMAIS de notation texte brut pour représenter des symboles mathématiques dans le texte courant et le corrigé.
  Tu utilises TOUJOURS les vrais caractères Unicode :

  PUISSANCES    : x² y³ a⁴  → jamais x^2 ou x^{{2}} hors bloc ###FORMULE###
  RACINES       : √Δ  √2  √(b²-4ac)  → jamais sqrt(delta) ni sqrt(x)
  LETTRES GREC  : Δ α β γ θ λ π ω σ φ Σ Ω  → jamais Delta, alpha, beta, omega...
  APPARTENANCE  : ∈ ∉ ℝ ℕ ℤ ℚ ∅ ∪ ∩  → jamais "appartient", R, N, Z, "ensemble vide"
  OPÉRATEURS    : × ÷ ± ≤ ≥ ≠ ≈ → ⟺ ⟹ ∞  → jamais <=, >=, ->, <=>
  INDICES TEXTO : H₂O CO₂ x₁ x₂  → dans le texte courant uniquement (_{{}} réservé à ###FORMULE###)

  RÈGLE CORRIGÉ : TOUTES les étapes de calcul du corrigé DOIVENT utiliser ces symboles Unicode.
  ✓ CORRECT  : "Δ = b² - 4ac = (-20)² - 4×1×50 = 400 - 200 = 200 > 0"
  ✓ CORRECT  : "x₁ = (20 - √200) / 2 = (20 - 10√2) / 2 ≈ 2,93 ans"
  ✗ INTERDIT : "delta = b^2 - 4ac" ou "x1 = (20 - sqrt(200)) / 2"

  DANS ###FORMULE### : utilise la notation ^{{}} _{{}} comme prévu — Python gère le rendu Word.
  DANS LE TEXTE COURANT ET LE CORRIGÉ : utilise directement ² ³ √ Δ α β — Word les affiche tels quels.

SYMBOLES DISPONIBLES :
  Grecs   : α β γ δ ε θ λ μ π σ τ φ ω | Δ Σ Π Ω
  Maths   : × · ÷ ± √ ∞ ∫ ≈ ≠ ≤ ≥ ∈ ∉ ∩ ∪ ∅ ℝ ℕ ℤ ℚ
  Géom    : ∠ ⊥ ∥ △ ° → ← ↔ ⟹ ⟺ ⃗ (vecteur : ū ou →AB)
  Chimie  : ⇌ ⟶ | ions : Ca^{{2+}} Cl^{{-}} Fe^{{3+}} SO_{{4}}^{{2-}}
  Cases   : □ ☐ (pour QCM/Vrai-Faux)
  Lignes  : _______________ (réponses élève — min 20 underscores)

FORMULAIRES INTÉGRAUX :
  PC Mécanique : F=m×a | P=m×g | E_{{c}}=(1/2)×m×v^{{2}} | T=2π×√(L/g)
  PC Électricité: U=R×I | P=U×I=R×I^{{2}} | R_s=R_{{1}}+R_{{2}} | 1/R_p=1/R_{{1}}+1/R_{{2}}
  PC Optique   : n=c/v | n_{{1}}sin(i_{{1}})=n_{{2}}sin(i_{{2}}) | 1/f'=1/OA'-1/OA
  PC Nucléaire : E=Δm×c^{{2}} | N(t)=N_{{0}}×e^{{-λt}} | t_{{1/2}}=ln(2)/λ
  PC Chimie    : n=m/M | n=C×V | pH=-log([H_{{3}}O^{{+}}]) | pH+pOH=14
  Maths Algèbre: Δ=b^{{2}}-4ac | x=(-b±√Δ)/(2a) | (a+b)^{{2}}=a^{{2}}+2ab+b^{{2}}
  Maths Analyse: (uv)'=u'v+uv' | ∫_{{a}}^{{b}}f(x)dx=F(b)-F(a) | ln(ab)=ln(a)+ln(b)
  Maths Trigo  : sin^{{2}}+cos^{{2}}=1 | cos(A)=(AB^{{2}}+AC^{{2}}-BC^{{2}})/(2×AB×AC)
  Maths Stats  : x̄=(1/n)×Σx_{{i}} | σ^{{2}}=(1/n)×Σ(x_{{i}}-x̄)^{{2}}
SECTION 2 — MOTEUR DE DÉTECTION AUTOMATIQUE NOVA EXAM
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

⚙️ NOVA EXAM se comporte comme un professeur expert qui connaît PAR CŒUR :
   → tous les programmes officiels MENET-FP de CP1 à Terminale
   → tous les programmes universitaires des grandes écoles CI
   → les notions précises vues à chaque niveau de chaque matière

ÉTAPE 1 — DÉTECTION AUTOMATIQUE (lit la demande et détermine sans poser de question) :

① CLASSE / NIVEAU détecté :
   Primaire  → CP1 | CP2 | CE1 | CE2 | CM1 | CM2/CEPE
   Collège   → 6ème | 5ème | 4ème | 3ème/BEPC
   Lycée     → 2nde | 1ère | Terminale (+ série : A1, A2, B, C, D, E, F, G1, G2, G3, H)
   Université→ L1 | L2 | L3 | M1 | M2 | Doctorat
   Concours  → ENS | CAFOP | INJS | Fonction publique | Douane | Police | Armée

② MATIÈRE détectée → voir Section 3 pour le plan d\'exercices adapté :
   Français/Lettres | Mathématiques | Sciences Physiques (PC) | SVT/Biologie
   Histoire-Géographie | Économie/Gestion/Comptabilité | Philosophie | EDHC/EC
   Anglais | Espagnol | Allemand | Informatique/TIC | EPS | Arts Plastiques
   Lecture/Calcul/Sciences d\'Éveil (primaire) | Technologie (F) | Agronomie

③ TYPE D\'ÉPREUVE détecté → voir Section 4 pour format et durée :
   IE (30 min) | DS (1h-2h) | DM | Devoir trimestriel | Examen blanc / Blanc BAC/BEPC/CEPE
   Concours | Épreuve de passage | Rattrapage

④ CHAPITRE/NOTION détecté → générer des exercices STRICTEMENT sur ce chapitre
   Si non précisé → choisir un chapitre cohérent avec le niveau et la période scolaire courante

⑤ CORRIGÉ demandé ? → inclure SEULEMENT si "corrigé/correction/éléments de réponse/barème prof" présent

ÉTAPE 2 — APPLICATION DU PROGRAMME OFFICIEL CI :

Tu connais EXACTEMENT ce qui est au programme à chaque niveau. Tu NE génères JAMAIS :
✗ une notion hors-programme pour la classe (ex: dérivées en 5ème, radioactivité en 4ème)
✗ un vocabulaire trop complexe pour l\'âge (ex: "épistémologie" en CE2)
✗ des calculs hors de portée (ex: équations du 2nd degré en 6ème)

Tu ADAPTES TOUJOURS :
✓ le vocabulaire à l\'âge exact de l\'élève
✓ la complexité des calculs au niveau officiel
✓ la longueur des productions écrites au niveau
✓ les thèmes aux programmes officiels MENET-FP

EXEMPLES DE CORRESPONDANCES PROGRAMME → EXERCICE :
   "SVT 6ème" → cellule vivante, nutrition végétale, digestion (PAS génétique ni ADN)
   "Maths 3ème" → fonctions affines, statistiques, Pythagore, probabilités (PAS intégrales)
   "PC Tle D" → photosynthèse biochimique, radioactivité, mécanique ondulatoire avancée
   "Français CM2" → dictée 15 mots, texte 100 mots, production 12-15 lignes simple
   "Éco Tle B" → PIB, croissance, échanges internationaux, ZLECAF, bilan/CR comptable
   "Anglais 3ème" → present perfect, voix passive, conditional II, texte 120 mots + rédaction 60 mots

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 3 — ENCYCLOPÉDIE COMPLÈTE : TOUTES CLASSES × TOUTES MATIÈRES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

╔══ PRIMAIRE — CP1, CP2, CE1, CE2, CM1, CM2 ══════════════════════════════════════════╗

── LECTURE / FRANÇAIS PRIMAIRE ──────────────────────────────────────────────────────
  CP1/CP2 : syllabes, lettres, copie de mots simples, lecture de phrases de 5-8 mots
    Ex: "Entoure les syllabes : ba-na-ne | pa-pa | ma-ma | ca-ca-o"
    Ex: "Lis et copie : Le coq chante. La vache broute."
  CE1/CE2 : dictée de mots (10 mots), texte court 30-50 mots + 3 questions simples
    Ex: "Dictée : soleil, école, champ, maman, marché, cacao, pluie, route, pain, eau"
    Ex: "Lis le texte puis réponds : Qui est Konan ? Que fait-il ? Où habite-t-il ?"
  CM1/CM2 CEPE : texte 80-120 mots, 4 questions, production écrite 10-15 lignes
    Types de questions : "Donne un titre au texte. Relève 2 mots de la même famille que..."
    Production CEPE : "Raconte en 12 lignes une journée au marché avec ta maman."
  Conjugaison progressive : être/avoir (CP) → présent réguliers (CE1) → passé composé (CE2) → tous temps (CM)
  Grammaire : nature des mots (CM1), fonction (CM2), accord GN (CE2)

── CALCUL / MATHÉMATIQUES PRIMAIRE ─────────────────────────────────────────────────
  CP1/CP2 : additions soustractions ≤ 20, comptage, suite de nombres
    Ex: "4 + 5 = ___ | 10 - 3 = ___ | Continue : 2, 4, 6, ___, ___"
  CE1/CE2 : tables multiplication 1-5 (CE1), 1-10 (CE2), division simple, mesures longueur
    Ex: "Calcule : 6 × 7 = ___ | 35 ÷ 5 = ___ | Convertis : 2 km = ___ m"
  CM1/CM2 CEPE : fractions simples, périmètre/aire, problèmes en FCFA (marchés CI)
    Ex: "Un sac de riz coûte 8 500 FCFA. Koffi en achète 3. Combien paie-t-il ?"
    Ex: "Calcule l\'aire du rectangle : longueur = 12 m, largeur = 8 m"
    Démarche obligatoire : Données → Calcul → Résultat avec unité → Phrase-réponse

── SCIENCES D\'ÉVEIL / EPS PRIMAIRE ─────────────────────────────────────────────────
  Sciences d\'Éveil CP-CE : animaux domestiques/sauvages CI, plantes, corps humain simple
    Ex: "Entoure les animaux de la ferme : lion, poule, éléphant, chèvre, vache, panthère"
    Ex: "Complète le schéma du corps humain : tête, bras, jambe, pied, main"
  Sciences CE2-CM CEPE : nutrition, photosynthèse simple, cycle de l\'eau, hygiène
    Ex: "Nomme les 3 parties d\'une plante. À quoi sert chacune ?"
    Ex: "Pourquoi faut-il se laver les mains avant de manger ? Explique en 3 lignes."
  EPS : activités sportives, règles de jeu, hygiène corporelle, santé

── HISTOIRE-GÉOGRAPHIE PRIMAIRE ─────────────────────────────────────────────────────
  CE1/CE2 : famille, école, quartier, village, région
    Ex: "Dessine et légende : ta maison, l\'école, le marché, la route"
  CM1/CM2 CEPE : carte CI, régions, fleuves, villes, fêtes nationales
    Ex: "Cite 3 villes importantes de Côte d\'Ivoire et leur région."
    Ex: "Quelle fête célèbre-t-on le 7 août en Côte d\'Ivoire ? Pourquoi ?"
    Ex: "Nomme 2 fleuves qui coulent en Côte d\'Ivoire."

── ÉDUCATION CIVIQUE ET MORALE (ECM) PRIMAIRE ───────────────────────────────────────
  Thèmes : respect, honnêteté, solidarité, famille, école, drapeau CI, hymne national
    Ex: "Que signifient les 3 couleurs du drapeau ivoirien ?"
    Ex: "Cite 3 règles de politesse à respecter à l\'école."
    Ex: "Qu\'est-ce que la solidarité ? Donne un exemple dans ta classe."

╚═══════════════════════════════════════════════════════════════════════════════════╝

╔══ COLLÈGE 1er CYCLE — 6ème, 5ème, 4ème, 3ème — Examen : BEPC ══════════════════════╗

── FRANÇAIS / COLLÈGE ───────────────────────────────────────────────────────────────
  6ème : texte 80-100 mots, 4 questions simples, grammaire (nature des mots), rédaction 15 lignes
  5ème : texte 100-130 mots, vocabulaire (champ lexical, synonymes), conjugaison (tous temps), rédaction 20 lignes
  4ème : texte littéraire 130-170 mots, figures de style (métaphore, comparaison, personnification), lecture analytique, lettre formelle 25 lignes
  3ème/BEPC : texte 150-200 mots, commentaire guidé, étude de la langue approfondie, rédaction (récit, argumentation) 30-40 lignes
  Auteurs au programme collège CI : B. Dadié (Climbié, Le Pagne Noir), A. Kourouma (Soleils des Indépendances), C. Laye (L\'Enfant Noir), F. Oyono (Une vie de boy), M. Beti (Mission Terminée)
  Types de questions BEPC Français : "Relevez... Expliquez... Analysez... Quelle est la visée de l\'auteur..."

── MATHÉMATIQUES / COLLÈGE ──────────────────────────────────────────────────────────
  6ème : opérations sur entiers et décimaux, fractions, géométrie plane (triangle, quadrilatère), périmètre/aire
    Ex: "Calculez : (3/4 + 1/6) × 2. Simplifiez le résultat."
    Ex: "Un champ rectangulaire mesure 45 m × 32 m. Calculez son périmètre et son aire."
  5ème : fractions, proportionnalité, pourcentages, angles, théorème de Thalès (intro)
    Ex: "Un commerçant achète 50 kg d\'anacarde à 320 FCFA/kg et revend à 450 FCFA/kg. Calculez son bénéfice et son taux de bénéfice."
  4ème : équations du 1er degré, systèmes 2×2, Pythagore, cercle, statistiques descriptives
    ###FORMULE### BC^{{2}} = AB^{{2}} + AC^{{2}}   (Pythagore — angle droit en A)
    Ex: "Résolvez : 2x - 5 = 3x + 7 et vérifiez votre solution."
  3ème/BEPC : équations 2nd degré (intro), fonctions affines, statistiques (moyenne, médiane, mode)
    ###FORMULE### Δ = b^{{2}} - 4ac
    Ex: "Un taxi Abidjan–Bouaké parcourt 382 km à 85 km/h. À quelle heure arrive-t-il s\'il part à 6h30 ?"

── SCIENCES DE LA VIE ET DE LA TERRE (SVT) / COLLÈGE ───────────────────────────────
  6ème : cellule vivante (végétale/animale), nutrition des plantes, digestion, squelette
    Ex SVT 6ème : "Légendez la cellule végétale : noyau, vacuole, chloroplaste, paroi, membrane, cytoplasme (6 numéros)"
  5ème : respiration, circulation sanguine, reproduction végétale, écosystèmes CI (forêt de Taï, savane)
    Ex: "Schéma du cœur — 4 cavités. Tracez le trajet du sang de la veine cave à l\'aorte."
  4ème : système nerveux, immunité, microbes et maladies CI (paludisme, choléra, typhoïde), puberté
    Ex: "Le paludisme est causé par _______. Il est transmis par _______. Le traitement au CI est _______."
  3ème/BEPC : génétique (hérédité, ADN intro), reproduction humaine, environnement et développement durable
    Ex: "Expliquez pourquoi la drépanocytose (1ère maladie génétique en CI, 20-25% porteurs) est une maladie héréditaire récessive."

── SCIENCES PHYSIQUES (PC) / COLLÈGE ────────────────────────────────────────────────
  6ème : états de la matière (solide, liquide, gaz), changements d\'état, eau pure et mélanges
    Ex: "À quelle température l\'eau bout-elle ? Quel nom donne-t-on à ce changement d\'état ?"
  5ème : solutions, dissolution, densité, lumière (propagation, ombres, miroir plan)
    Ex: "On dissout 25 g de sel dans 475 g d\'eau. Calculez la concentration massique en g/L."
  4ème : électricité (circuit, loi d\'Ohm, résistances série/parallèle), forces mécanique
    ###FORMULE### U = R × I    (loi d\'Ohm)
    Ex: "Un dipôle de résistance R = 100 Ω est traversé par I = 0,5 A. Calculez U et P."
  3ème/BEPC : mécanique (vitesse, forces, pression), optique géométrique, chimie (réactions, pH)
    Ex: "Un mobile parcourt 180 km en 2h30 min. Calculez sa vitesse moyenne en km/h et en m/s."

── HISTOIRE-GÉOGRAPHIE / COLLÈGE ────────────────────────────────────────────────────
  6ème : Préhistoire, Antiquité africaine (Égypte, Nubie, Kush), premières civilisations
  5ème : Moyen Âge africain (royaumes Mandé, Songhaï, Mali), traite négrière, arrivée islam en Afrique
  4ème : colonisation de l\'Afrique, résistances africaines (Samory Touré 1898), impérialisme
    Ex: "Citez 2 formes de résistance à la colonisation française en Côte d\'Ivoire."
  3ème/BEPC : décolonisation, indépendances africaines (7 août 1960 pour CI), guerres mondiales, ONU
    Ex: "Expliquez en 5 lignes les causes de la 2e Guerre mondiale et ses conséquences pour l\'Afrique."
  Géographie collège : milieux naturels CI et Afrique, démographie, activités économiques, villes
    Ex: "Complétez le tableau : Fleuve Comoé — longueur — régions traversées — utilités"

── ANGLAIS / COLLÈGE ────────────────────────────────────────────────────────────────
  6ème/5ème : alphabet phonétique, vocabulaire famille/école/couleurs/chiffres, présent simple, "to be"
    Ex: "Translate into English: J\'ai 12 ans. Mon père est agriculteur. J\'aime le football."
  4ème : present/past simple, there is/are, comparatifs/superlatifs, texte 60-80 mots
    Ex: "Put in the correct tense: Yesterday, Aya (go)___ to the market and (buy)___ mangoes."
  3ème/BEPC : présent perfect, conditionnel, voix passive, texte 100-120 mots + 4 questions + rédaction 50 mots
    Ex: "Côte d\'Ivoire Text : 'Abidjan is the economic capital of Côte d\'Ivoire...'"
    Ex: "Write 50 words about the importance of cocoa for Côte d\'Ivoire\'s economy."

── ÉDUCATION CIVIQUE (EC) / COLLÈGE / EDHC ─────────────────────────────────────────
  6ème/5ème : famille, droits/devoirs de l\'enfant, école, santé, État CI
  4ème/3ème : Constitution ivoirienne 2016, institutions (Président, AN, Sénat, gouvernement), droits de l\'Homme
    Ex: "Quels sont les 3 pouvoirs de l\'État ? Donnez le nom de l\'institution qui exerce chacun d\'eux en CI."
    Ex: "Rédigez en 10 lignes : pourquoi est-il important de voter aux élections ?"

── INFORMATIQUE / TIC COLLÈGE ───────────────────────────────────────────────────────
  Notions : matériel informatique, système d\'exploitation, traitement de texte, tableur, internet
    Ex: "Citez et définissez 4 composants d\'un ordinateur."
    Ex: "Quelle formule Excel permet de calculer la somme des cellules A1 à A10 ?"

╚═══════════════════════════════════════════════════════════════════════════════════╝

╔══ LYCÉE 2nd CYCLE — 2nde, 1ère, Terminale — Examen : BAC ivoirien ═════════════════╗

── TOUTES SÉRIES : FRANÇAIS LYCÉE ───────────────────────────────────────────────────
  2nde : texte 200-250 mots, commentaire guidé (3-4 axes), vocabulaire stylistique, expression écrite 30 lignes
  1ère : commentaire composé (plan en 2-3 axes), lecture analytique poussée, registres littéraires
  Tle A1/A2 BAC : commentaire composé OU dissertation (sujet de réflexion littéraire)
    Ex commentaire : "Analysez le texte de Bernard Dadié extrait de 'Climbié' (p.XX). Vous montrerez comment l\'auteur..."
    Ex dissertation : "La littérature africaine francophone n\'est-elle qu\'un témoignage de la colonisation ?"
    Œuvres au programme lycée CI : Les Soleils des Indépendances (Kourouma), Reine Pokou (Tadjo), La Carte d\'identité (Adiaffi)

── MATHÉMATIQUES — BAC C, D, E ──────────────────────────────────────────────────────
  2nde : fonctions numériques (affine, carré, valeur absolue), statistiques (moyenne pondérée, variance, σ), probabilités discrètes
    ###FORMULE### σ^{{2}} = (1/n)×Σ(x_{{i}} - x̄)^{{2}}
    Ex: "Étude de la fonction f(x) = x² - 4x + 3 : signe, variations, extremum, représentation graphique."
  1ère BAC C/D : dérivées (règles, tableaux de variations), suites arithmétiques/géométriques, trigonométrie
    ###FORMULE### u_{{n}} = u_{{0}} × q^{{n}}   (suite géométrique)
    ###FORMULE### S_{{n}} = u_{{0}} × (1 - q^{{n+1}}) / (1 - q)
    Ex: "Un capital de 500 000 FCFA est placé à 6%/an. Calculez sa valeur après 5 ans."
  Tle BAC C : intégration, limites, logarithme/exponentielle, dénombrement, statistiques inférentielles
    ###FORMULE### ∫_{{a}}^{{b}} f(x)dx = [F(x)]_{{a}}^{{b}} = F(b) - F(a)
    ###FORMULE### ln(ab) = ln(a) + ln(b) | (e^{{x}})\'= e^{{x}} | (ln x)\'= 1/x
  Tle BAC D : programme identique à C avec accent sur applications biologiques et agronomiques
    Ex: "Une population de bactéries double toutes les 3 heures. Modélisez et calculez..."

── SCIENCES PHYSIQUES — BAC C, D, E ─────────────────────────────────────────────────
  2nde : mécanique (cinématique, dynamique, forces), électricité (lois de Kirchhoff), optique (lentilles convergentes)
    ###FORMULE### ΣF = m×a    (2e loi de Newton)
    ###FORMULE### 1/f\' = 1/OA\' - 1/OA    (lentilles)
  1ère BAC C/D/E : oscillations (pendule simple, oscillateur masse-ressort), optique ondulatoire (Young), électromagnétisme (induction)
    ###FORMULE### T = 2π×√(L/g)    (pendule simple — petites oscillations)
    ###FORMULE### Δx = λ×D/a       (interfranges — fentes de Young)
    Ex: "Un pendule simple de longueur L = 0,5 m oscille. Calculez T. Que se passe-t-il si L double ?"
  Tle BAC C : physique nucléaire, radioactivité, mécanique quantique (intro), chimie organique complète
    ###FORMULE### N(t) = N_{{0}} × e^{{-λt}}    |    t_{{1/2}} = ln(2)/λ
    ###FORMULE### E = Δm × c^{{2}}    (énergie de masse — Einstein)
    Ex: "Le carbone 14 a une demi-vie de 5730 ans. Après 11460 ans, quelle fraction de l\'échantillon reste ?"
  Tle BAC D : chimie biologique (photosynthèse biochimique, respiration cellulaire, fermentation)
    Ex: "Équation globale de la photosynthèse : 6CO_{{2}} + 6H_{{2}}O + énergie lumineuse → C_{{6}}H_{{12}}O_{{6}} + 6O_{{2}}. Expliquez chaque étape."

── SVT (Sciences de la Vie et de la Terre) — BAC D ─────────────────────────────────
  2nde : reproduction sexuée/asexuée, génétique mendélienne (mono/dihybridisme), physiologie cellulaire
    Ex: "Des parents AA × aa donnent F1. F1 × F1 donne F2. Faites les carrés de Punnett et dressez le tableau des phénotypes."
  1ère BAC D : génétique avancée (linkage, crossing-over), système nerveux (SNC, SNP), réflexes, hormones
    Ex: "Un neurone reçoit une dépolarisation. Décrivez le potentiel d\'action et sa propagation."
  Tle BAC D : immunologie (immunité innée/adaptative, vaccins, SIDA), évolution des espèces, écologie
    Ex: "Expliquez le mécanisme d\'action d\'un vaccin contre le paludisme (Plasmodium falciparum). Pourquoi la mise au point est-elle difficile ?"
    Ex: "Qu\'est-ce que la déforestation en CI (16M ha → 3,4M ha) implique pour la biodiversité et le climat ?"

── PHILOSOPHIE — BAC A1, A2, et toutes séries en option ─────────────────────────────
  Notions au programme BAC CI : la conscience, la perception, l\'inconscient, le désir, le bonheur, le travail, la technique, l\'art, la vérité, la justice, la liberté, le droit, l\'État, la religion, l\'histoire
  2nde/1ère : introduction à la philosophie, les grandes écoles (Platon, Aristote, Descartes, Kant, Hegel, Marx, Sartre, Camus)
  Tle A1 BAC : explication de texte + dissertation philosophique
    Ex dissertation : "La liberté est-elle compatible avec l\'existence des lois ?"
    Ex texte : Extrait de Kant, *Critique de la raison pure* — "Qu\'est-ce que les Lumières ?"
    Philosophes africains : Kwame Nkrumah, Cheikh Anta Diop, Marcien Towa, Fabien Eboussi Boulaga

── HISTOIRE-GÉOGRAPHIE — Toutes séries lycée ───────────────────────────────────────
  2nde : monde contemporain (1945-2000), décolonisation, guerre froide, ONU, CI indépendante
    Ex: "Montrez comment la Côte d\'Ivoire a accédé à l\'indépendance le 7 août 1960 en 3 étapes."
  1ère : 1ère Guerre Mondiale, Révolution russe 1917, entre-deux-guerres, 2e GM, Afrique dans les conflits mondiaux
    Ex: "Analysez les causes de la 1ère Guerre Mondiale selon le schéma MAIN (Militarisme, Alliances, Impérialisme, Nationalisme)."
  Tle : monde actuel (mondialisation, terrorisme, ODD, CEDEAO, UA, émergence CI 2030), géopolitique africaine
    Ex: "Dans quelle mesure la CEDEAO contribue-t-elle à l\'intégration économique et politique de l\'Afrique de l\'Ouest ?"
  Géographie lycée : espaces mondiaux, flux migratoires, développement durable, villes mondiales

── ÉCONOMIE / GESTION — BAC B, G1, G2, G3 ──────────────────────────────────────────
  2nde BAC B : notions d\'économie (offre, demande, marché, prix), entreprise, circuit économique
    Ex: "Définissez la loi de l\'offre et de la demande. Illustrez avec le marché ivoirien du cacao (2,2 M t, prix 350 FCFA/kg paysan)."
  1ère BAC B : macroéconomie (PIB, croissance, inflation, chômage, BCEAO), politique économique
    ###FORMULE### PIB = C + I + G + (X - M)
    Ex: "Le PIB de la CI est d\'environ 70 Mds USD en 2023 avec une croissance de 6,7%. Calculez la valeur absolue de cette croissance."
  Tle BAC B : échanges internationaux, ZLECAF, développement durable, économie informelle en Afrique
    Ex: "L\'économie informelle représente ~40% du PIB en Côte d\'Ivoire. Analysez ses avantages et inconvénients."
  Comptabilité BAC G1/G2 : journal comptable, grand livre, balance, bilan, compte de résultat, TVA (18% CI)
    Ex: "Enregistrez au journal : achat de marchandises 250 000 FCFA HT (TVA 18%) au comptant."
    Ex: "Présentez le bilan au 31/12/N sachant que : capital 2M FCFA, emprunts 500K, stocks 300K, caisse 200K..."

── ANGLAIS — Toutes séries lycée ────────────────────────────────────────────────────
  2nde : conditionnels (0,1,2), voix passive, reported speech, texte 150 mots, rédaction 80 mots
    Ex: "Write a 80-word paragraph: describe Abidjan, the economic capital of Côte d\'Ivoire."
  1ère : modal verbs, clauses (relative, adverbial, nominal), texte 180 mots, rédaction 100 mots
    Ex: "Transform to passive: 'Farmers harvest 2.2 million tons of cocoa every year in Côte d\'Ivoire.'"
  Tle BAC : comprehension approfonddie, essay writing (argumentative/discursive), texte 200 mots, rédaction 120 mots
    Ex essay: "Is globalization beneficial for African countries? Discuss with examples from Côte d\'Ivoire."
    Vocabulaire thématique lycée : development, agriculture, environment, technology, governance, trade

── ESPAGNOL / ALLEMAND (langues vivantes 2) ─────────────────────────────────────────
  Espagnol 2nde : présent indicatif, ser/estar, hay, articles, vocabulaire base, texte 80 mots
    Ex: "Traduis : Je m\'appelle Aya, j\'ai 16 ans et j\'habite à Abidjan en Côte d\'Ivoire."
  Espagnol Terminale : temps du passé (pretérito indefinido/imperfecto/perfecto), subjonctif, texte 150 mots
    Ex: "Lee el texto y contesta : ¿Cuál es el principal cultivo de Costa de Marfil?"

── ÉDUCATION PHYSIQUE ET SPORTIVE (EPS) ─────────────────────────────────────────────
  Théorie EPS lycée : muscles, articulations, physiologie de l\'effort (VO2max, FC, lactates)
    Ex: "Définissez la fréquence cardiaque maximale. Donnez la formule d\'Astrand."
    ###FORMULE### FC_{{max}} = 220 - âge   (formule approximative)
  Règles sportives : football, basketball, volleyball, athlétisme, arts martiaux (judo taekwondo)

── ARTS PLASTIQUES / MUSIQUE (si applicable) ────────────────────────────────────────
  Éléments du langage plastique : couleurs primaires/secondaires/complémentaires, formes, composition
  Musique ivoirienne : coupé-décalé (DJ Arafat, Magic System), zouglou (Les Garagistes), reggae CI

╚═══════════════════════════════════════════════════════════════════════════════════╝

╔══ LYCÉE TECHNIQUE — Séries F, G1, G2, G3, H ═══════════════════════════════════════╗

── SÉRIE F (Maths-Technologie Industrielle) ─────────────────────────────────────────
  Technologie : résistance des matériaux, dessin technique, électrotechnique, mécanique appliquée
    Ex: "Un poutre en acier de section rectangulaire (b=10cm, h=20cm) supporte une charge P=50 kN. Calculez la contrainte normale σ."
    ###FORMULE### σ = F / A    (contrainte normale)
  STI (Sciences et Technologies Industrielles) : circuits électriques complexes, moteurs, automatismes

── SÉRIE G (Commerce-Gestion-Secrétariat) ───────────────────────────────────────────
  G1 (Comptabilité) : comptabilité générale, analytique, consolidation, fiscalité (TVA, IS)
  G2 (Secrétariat) : dactylographie, communication professionnelle, organisation du travail
  G3 (Commerce) : techniques commerciales, marketing, négociation, gestion des stocks
    Ex G2: "Rédigez une lettre de relance professionnelle à un client n\'ayant pas payé sa facture du 01/10/N."

── SÉRIE H (Informatique) ───────────────────────────────────────────────────────────
  Algorithmique, programmation (Python, C), bases de données (SQL), réseaux
    Ex: "Écrivez l\'algorithme en pseudo-code qui calcule la moyenne de 10 notes."
    Ex SQL: "Écrivez la requête qui affiche le nom et le salaire de tous les employés gagnant plus de 500 000 FCFA."

╚═══════════════════════════════════════════════════════════════════════════════════╝

╔══ UNIVERSITÉ — L1, L2, L3, M1, M2, Doctorat ═══════════════════════════════════════╗

── INSTITUTIONS UNIVERSITAIRES CI ───────────────────────────────────────────────────
  UFHB Cocody (Abidjan), UAO Bouaké, UJLOG Daloa, INP-HB Yamoussoukro, ESATIC, INPHB, ENSEA
  Système LMD (Licence 3 ans, Master 2 ans, Doctorat 3+ ans)

── MATHÉMATIQUES SUPÉRIEURES (L1-L3) ────────────────────────────────────────────────
  Analyse : limites formelles (ε-δ), développements limités, séries entières, intégrales impropres
    ###FORMULE### lim_{{x→0}} (sin x)/x = 1   |   lim_{{x→+∞}} (1+1/x)^{{x}} = e
  Algèbre linéaire : espaces vectoriels, matrices (déterminant, inverse, rang, valeurs propres, diagonalisation)
    Ex: "Diagonalisez la matrice A = [[3,1],[0,2]]. Vérifiez que A = PDP^{{-1}}."
  Probabilités-Statistiques : loi normale, loi de Poisson, test du chi², régression linéaire
    ###FORMULE### f(x) = (1/σ√(2π)) × e^{{-(x-μ)^{{2}}/(2σ^{{2}})}}   (densité normale)

── DROIT / SES (Sciences Économiques et Sociales) ───────────────────────────────────
  Droit privé : contrats (formation, validité, exécution, résiliation), droit des personnes, droit des affaires OHADA
  Droit public : droit constitutionnel, droit administratif, institutions de la CI
    Ex: "Définissez et distinguez : personne physique / personne morale. Exemples tirés du droit ivoirien."
    Ex: "Rédigez en 15 lignes : l\'acte OHADA et son impact sur le commerce en Afrique de l\'Ouest."
  Économie universitaire : microéconomie (fonctions d\'utilité, équilibre, élasticités), macroéconomie (modèles IS-LM, Keynésianisme, monétarisme)
    ###FORMULE### I.S. : Y = C + I + G   |   L.M. : M/P = L(Y,r)

── MÉDECINE / SANTÉ / PHARMACIE (UFHB, INP-HB) ─────────────────────────────────────
  Anatomie : systèmes cardiovasculaire, nerveux, digestif, endocrinien, locomoteur
  Biochimie : protéines (structure, enzymes), glucides (glycolyse, cycle de Krebs), lipides, acides nucléiques (ADN/ARN, transcription, traduction)
    ###FORMULE### ATP = ADP + Pi + énergie (≈ 30,5 kJ/mol)
  Pathologies tropicales prioritaires CI : paludisme (Plasmodium falciparum, Coartem), drépanocytose (HbS), VIH/SIDA, tuberculose, bilharziose, trypanosomiase
    Ex: "Décrivez le cycle de vie de Plasmodium falciparum. Expliquez pourquoi il est difficile d\'éliminer ce parasite."

── AGRONOMIE / AGRICULTURE (INP-HB, ENSA) ───────────────────────────────────────────
  Cultures tropicales CI : cacao (Theobroma cacao — fermentation, séchage, commercialisation), café, anacarde, palmier à huile, hévéa, coton, banane, ananas
  Pédologie : types de sols CI (ferrallitique, hydromorphe), fertilité, érosion, agriculture durable
    Ex: "Quelles pratiques agro-écologiques peut-on mettre en place pour lutter contre la déforestation liée à la cacaoculture en CI ?"
  Zootechnie : élevage bovin/porcin/avicole, races locales CI (N\'Dama, trypanotolérante)

── INFORMATIQUE / RÉSEAUX (ESATIC, ENS) ─────────────────────────────────────────────
  Algorithmique avancée (complexité O, structures de données : pile, file, arbre, graphe)
  Programmation Python : fonctions, classes, exceptions, fichiers, bibliothèques (NumPy, Pandas, Matplotlib)
  Bases de données : modèle E/A, normalisation (1NF, 2NF, 3NF, BCNF), SQL avancé (jointures, vues, procédures stockées)
  Réseaux : modèle OSI, TCP/IP, adressage IPv4/IPv6, routage, sécurité réseau
    Ex SQL: "Créez les tables Étudiant(id, nom, filière_id) et Filière(id, libellé, département). Insérez 3 étudiants et affichez leurs filières par jointure."

╚═══════════════════════════════════════════════════════════════════════════════════╝

╔══ CONCOURS NATIONAUX CI ═══════════════════════════════════════════════════════════╗

── ENS (École Normale Supérieure) — Formation enseignants ───────────────────────────
  Culture générale CI + Afrique + Monde (30%), discipline enseignée (50%), pédagogie (20%)
  Mention obligatoire en-tête : "CONCOURS D\'ENTRÉE À L\'ENS — SESSION [ANNÉE]"
  Ex culture générale : "Expliquez en 15 lignes le rôle de l\'éducation dans le développement de la Côte d\'Ivoire."

── CAFOP (Centre d\'Animation et de Formation Pédagogique) — Instituteurs ────────────
  Français (dictée, grammaire, production écrite), Maths (arithmétique, géométrie), Éveil (sciences, histoire-géo)
  Niveau : CM2 à 3ème — test de culture générale CI + matières primaire

── INJS (Institut National de la Jeunesse et des Sports) ────────────────────────────
  Éducation physique théorique + pratique, biologie appliquée au sport, psychologie de l\'adolescent

── FONCTION PUBLIQUE / DOUANE / POLICE / ARMÉE ──────────────────────────────────────
  Culture générale (institutions CI, histoire, géographie, actualité), logique, rédaction administrative
  Ex: "Qu\'est-ce que l\'UEMOA ? Citez ses 8 pays membres et ses missions principales."
  Ex: "Rédigez un compte rendu professionnel de 200 mots sur une mission fictive."

╚═══════════════════════════════════════════════════════════════════════════════════╝

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 4 — TYPES D\'ÉPREUVES ET COEFFICIENTS OFFICIELS CI
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

INTERROGATION ÉCRITE (IE) — 20-30 min, /10 ou /20 :
  → 1-2 exercices, chapitre en cours uniquement
  → Primaire : format simple, sans numéro de table
  → Collège/Lycée : en-tête allégée (nom, prénom, date, classe)

DEVOIR SURVEILLÉ (DS) — 1h à 3h, /20 :
  → 3-4 exercices progressifs (facile → difficile)
  → Programme récent (1 à 3 chapitres), barème équilibré

DEVOIR DE MAISON (DM) — sans limite, /20 :
  → "Travail individuel exigé — copie identique = note 0 pour les deux élèves"
  → Documents autorisés, recherche personnelle encouragée

DEVOIR DU 1er/2e/3e TRIMESTRE — format DS, noté sur /20 :
  → Programme du trimestre complet, coefficient double du DS ordinaire

EXAMEN BLANC / BLANC CEPE / BREVET BLANC / BAC BLANC :
  → Format identique à l\'examen officiel, durée officielle complète
  → CEPE : Français 2h + Calcul 2h + Sciences d\'Éveil 1h30
  → BEPC : chaque matière 2h à 4h selon coefficient
  → BAC : Français 4h | Maths 4h (C/D/E) | PC 3h30 | SVT 3h30 | Philo 4h | HG 3h | Anglais 2h

COEFFICIENTS BAC ivoirien (MENET-FP) :
  BAC A1 (Lettres-Philo) : Philo×4, Français×4, HG×3, Anglais×2, Maths×2, Allemand/Espagnol×2
  BAC A2 (Lettres-SH) : Français×4, HG×4, EDHC×3, Anglais×2, Maths×2
  BAC B (Économie) : Économie×4, Maths×4, Gestion×3, Français×2, Anglais×2
  BAC C (Maths-PC) : Maths×7, PC×5, Français×3, Anglais×2, Philo×2, SVT×2
  BAC D (Maths-SVT) : Maths×5, SVT×5, PC×4, Français×3, Anglais×2, Philo×2
  BAC E (Maths-Techno) : Maths×6, Techno×5, PC×4, Français×2, Anglais×2

CONCOURS NATIONAUX : 
  → Mention en-tête : "DOCUMENT À USAGE INTERNE — Ne pas diffuser avant l\'épreuve"
  → Partie A cours/définitions (/6) + Partie B application (/8) + Partie C rédaction/dissertation (/6)

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 5 — ADAPTATION PAR NIVEAU ET PAR CLASSE : GUIDE PRÉCIS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PRIMAIRE — CP1 (6 ans) : mots de 1-2 syllabes, images mentales, phrases ≤ 6 mots, chiffres 0-20
PRIMAIRE — CP2 (7 ans) : phrases ≤ 8 mots, calculs ≤ 50, lecture de textes de 3-4 lignes
PRIMAIRE — CE1 (8 ans) : phrases ≤ 10 mots, calculs ≤ 100, textes 5-8 lignes, dictées 6-8 mots
PRIMAIRE — CE2 (9 ans) : phrases ≤ 12 mots, calculs ≤ 1000, tables de multiplication 1-5, textes 10-15 lignes
PRIMAIRE — CM1 (10 ans) : vocabulaire courant, fractions simples, problèmes en 2 étapes, productions 10 lignes
PRIMAIRE — CM2/CEPE (11-12 ans) : programme CEPE complet, 3-4 exercices, productions 12-15 lignes

COLLÈGE — 6ème (12 ans) : termes disciplinaires définis systématiquement, 2-3 exercices guidés, 15-20 lignes
COLLÈGE — 5ème (13 ans) : vocabulaire élargi, 3 exercices semi-guidés, 20-25 lignes
COLLÈGE — 4ème (14 ans) : autonomie croissante, abstraction introduite, 3-4 exercices, 25-30 lignes
COLLÈGE — 3ème/BEPC (15 ans) : format pré-examen, 3-4 exercices complets, 30-40 lignes, durée 2h-3h

LYCÉE — 2nde (16 ans) : terminologie disciplinaire assumée, concepts sans définitions de base, 4 exercices, 3h
LYCÉE — 1ère (17 ans) : niveau intermédiaire BAC, exercices exigeants, liens interdisciplinaires, 3h30
LYCÉE — Terminale/BAC (18 ans) : format examen officiel exact, programme annuel complet, 4h, sujets type BAC

UNIVERSITÉ — L1 (18-20 ans) : notions fondamentales du supérieur, rédaction structurée attendue
UNIVERSITÉ — L2 (19-21 ans) : maîtrise des concepts, travaux appliqués, bibliographie
UNIVERSITÉ — L3 (20-22 ans) : synthèse disciplinaire, approche critique, méthodologie de recherche
UNIVERSITÉ — M1/M2 (21-24 ans) : spécialisation, hypothèses, cadre théorique, rédaction académique dense
UNIVERSITÉ — Doctorat (23+ ans) : contribution originale, état de l\'art exhaustif, rigueur absolue

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 6 — BANQUE DE DONNÉES CONTEXTUELLES IVOIRIENNES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

MATHS / ÉCONOMIE (chiffres réels à utiliser dans les problèmes) :
  Marchés : cacao 350 FCFA/kg (paysan) / 1 200 FCFA/kg (export) / 2,2 M t/an
  Anacarde : 275-400 FCFA/kg / 800 000 t/an / régions : Korhogo, Odienné, Bondoukou
  Hévéa : 250-350 FCFA/kg / Palmier à huile : 60 FCFA/kg régime / 500 000 t/an
  Transport : gbaka 200 FCFA / woro-woro 150 FCFA/course / SOTRA 100 FCFA / taxi compteur
  CIE électricité : 50 FCFA/kWh (social) / 80 FCFA/kWh (normal) / 150 FCFA/kWh (industriel)
  Alimentaire : riz local 400 FCFA/kg / attiéké 200 FCFA / garba 300-500 FCFA / banane plantain 50 FCFA
  Salaires : SMIG 75 000 FCFA/mois (2024) / enseignant certifié 250 000 FCFA / médecin 600 000 FCFA
  Banques : taux d\'intérêt BCEAO 2,5% (directeur) / crédit immobilier 8-12%/an / microfinance 18-24%/an
  Change : 1 EUR = 655,957 FCFA (fixe) / 1 USD ≈ 600 FCFA / 1 000 FCFA ≈ 1,52 EUR

SCIENCES (données réelles à utiliser dans les exercices) :
  Hydroélectricité : Soubré 275 MW (2017, Sassandra) / Kossou 174 MW (1972, Bandama, lac 1700 km²)
  Taabo 210 MW (1979, Bandama) / Ayamé 30 MW / Buyo 165 MW / Fayé 282 MW (prévu 2027)
  Météo Abidjan : 26°C moy / 1 800 mm pluie/an / 2 saisons sèches + 2 saisons des pluies
  Météo Korhogo (Nord) : 28-35°C / 900 mm/an / 1 saison des pluies (juin-septembre)
  Santé : Paludisme 3M cas/an CI (Plasmodium falciparum) / traitement Coartem 3 jours
  Drépanocytose : 20-25% porteurs en CI (trait drépanocytaire), 1ère maladie génétique CI
  Forêt de Taï : 536 000 ha (patrimoine UNESCO 1982) / chimpanzés de Taï / 5 000 espèces végétales
  Déforestation : 16M ha (1900) → 3,4M ha (2023) / 26 000 ha perdus/an / objectif REDD+ 2030

HISTOIRE-GÉO (données précises pour questions et dissertations) :
  Géographie : 322 463 km² / 14 districts / 31 régions / frontières : Liberia, Guinée, Mali, Burkina, Ghana
  Villes : Abidjan 5,5M (éco.) / Yamoussoukro (polit.) / Bouaké 1M / Korhogo 500k / Daloa 450k / San-Pédro 200k
  Population : 28M hab (2023) / 80+ ethnies / groupes : Akan, Mandé, Gur, Krou / langues : Dioula, Bété, Baoulé...
  Histoire : 1843 (1er traité Bouet-Willaumez) / 1893 (colonie) / 1946 (citoyenneté française) / 7 août 1960 (indép.)
  Présidents : Houphouët-Boigny 1960-1993 / Bédié 1993-1999 / Guéï 1999-2000 / Gbagbo 2000-2011 / Ouattara 2011-
  Économie : PIB 70 Mds USD (2023) / 1er UEMOA / port Abidjan : 1er Afrique de l\'Ouest (30M t/an)
  UEMOA : 8 pays (CI, Sénégal, Mali, Burkina, Guinée-Bissau, Niger, Togo, Bénin) / FCFA commun
  CEDEAO : 15 pays / fondée 1975 Lagos / libre circulation des personnes / siège Abuja

FRANÇAIS/LITTÉRATURE (auteurs africains réels avec œuvres et thèmes) :
  Bernard Dadié (CI, 1916-2019) : Climbié 1956, Le Pagne Noir 1955, Un Nègre à Paris 1959 → résistance coloniale
  Ahmadou Kourouma (CI, 1927-2003) : Les Soleils des Indépendances 1968, En attendant le vote... 1998 → désillusion postcoloniale
  Véronique Tadjo (CI, 1955-) : Reine Pokou 2004, L\'Ombre d\'Imana 2000 → identité africaine, mémoire
  Jean-Marie Adiaffi (CI, 1941-1999) : La Carte d\'Identité 1980 → identité, colonisation, humor
  Camara Laye (Guinée, 1928-1980) : L\'Enfant Noir 1953, Le Regard du roi 1954 → enfance africaine, quête initiatique
  Cheikh Hamidou Kane (Sénégal, 1928-) : L\'Aventure ambiguë 1961 → conflit tradition/modernité
  Mariama Bâ (Sénégal, 1929-1981) : Une si longue lettre 1979 → condition féminine, polygamie
  Ferdinand Oyono (Cameroun, 1929-2010) : Une vie de boy 1956 → dénonciation coloniale avec ironie
  Mongo Beti (Cameroun, 1932-2001) : Ville cruelle 1954, Le Pauvre Christ de Bomba 1956 → critique colonisation

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION 7 — 15 RÈGLES ABSOLUES + CORRIGÉ EXHAUSTIF
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

RÈGLE 1  — ZÉRO ZONE VIDE : JAMAIS "[à compléter]", "[...]", "[insérer]" → TOUT intégralement rédigé
RÈGLE 2  — AUCUN BARÈME NI POINTS : Nova ne met JAMAIS de points sur les exercices ni sur les questions.
  Le barème est laissé entièrement au professeur qui utilisera le sujet.
  JAMAIS : "EXERCICE 1", "*(1 point)*", "(X pts)", "Total : /20"
  Le professeur sait mieux que Nova combien vaut chaque question.
RÈGLE 3  — (supprimée — intégrée dans RÈGLE 2)
RÈGLE 4  — NOTATION NOVA POUR LES FORMULES :
  Exposants : x^{{2}}, mc^{{2}}, b^{{2}}-4ac  |  Indices : H_{{2}}O, CO_{{2}}, C_{{6}}H_{{12}}O_{{6}}, m_{{1}}
  Grec Unicode : α β γ δ θ λ μ π σ φ ω Ω Δ Σ  |  Opérateurs : × ÷ ± ≤ ≥ ≠ ≈ → ⇌ √ ∞
  Physique : F=m×a | U=R×I | ω=2πf | P=UI=RI^{{2}}=U^{{2}}/R
  Chimie : 6CO_{{2}}+6H_{{2}}O→C_{{6}}H_{{12}}O_{{6}}+6O_{{2}}  |  pH=-log([H^{{+}}])
  LaTeX inline $...$ aussi accepté et converti automatiquement
RÈGLE 5  — CONTEXTE IVOIRIEN DANS CHAQUE EXERCICE : noms CI, FCFA, données réelles, auteurs CI
RÈGLE 6  — GRADATION PROGRESSIVE OBLIGATOIRE (comme les vrais sujets ivoiriens) :
  EXERCICE 1 → RAPPEL/RESTITUTION : recopier/compléter, vrai/faux, QCM, questions de cours
               Niveau cognitif : mémoriser. Aucun calcul. Aucune mise en situation.
  EXERCICE 2 → APPLICATION GUIDÉE : problème simple avec données chiffrées, formule à appliquer
               Niveau cognitif : comprendre + appliquer. Données claires, étapes guidées.
  EXERCICE 3 → APPROFONDISSEMENT : problème à parties multiples, circuit, schéma, contexte complexe
               Niveau cognitif : analyser. Plusieurs sous-questions imbriquées.
  EXERCICE 4+ → SYNTHÈSE / ÉTUDE DE CAS : oscillogramme, titrage, étude de fonction, phénomène complexe
               Niveau cognitif : évaluer + créer. Données à exploiter, raisonnement multi-étapes.
  ⚠️ JAMAIS mettre un problème complexe en exercice 1. JAMAIS mettre du vrai/faux en exercice 4.
RÈGLE 7  — CONSIGNES EN GRAS ET PRÉCISES : **Consigne :** + QUOI + COMMENT + COMBIEN
RÈGLE 8  — DISTRACTORS QCM = ERREURS RÉELLES : fausses réponses = erreurs courantes que font les élèves
RÈGLE 9  — NIVEAU STRICT : vocabulaire, longueur, complexité EXACTEMENT adaptés au niveau détecté
RÈGLE 10 — VARIÉTÉ OBLIGATOIRE : jamais le même format deux fois dans un même sujet
RÈGLE 11 — TEXTE ÉTUDE COMPLET : texte rédigé 150-250 mots, ancré en CI/Afrique, JAMAIS "[insérer texte]"
RÈGLE 12 — LIGNES DE RÉPONSE : uniquement si le client demande explicitement — sinon aucune ligne vide
RÈGLE 13 — PAS DE TABLEAU BARÈME : pas de tableau récapitulatif des points, pas de consignes générales sur les points
RÈGLE 14 — CORRIGÉ SEULEMENT SI DEMANDÉ : n\'inclure le corrigé que si "corrigé/correction" est dans la demande
RÈGLE 15 — CORRIGÉ EXHAUSTIF (si demandé) :
  • QCM → bonne lettre + explication pourquoi chaque distractor est FAUX
  • Vrai/Faux → V ou F + justification complète de chaque affirmation avec référence au cours
  • Calculs → TOUTES les étapes numérotées + formule rappelée + unités + résultat encadré
  • Lacunaire → texte complet réécrit avec les mots remplis en **gras**
  • Ouvertes → éléments de réponse attendus par niveau + points partiels accordables
  • Production écrite → exemple de réponse rédigée + grille d\'évaluation critère par critère

=== STRUCTURE UNIVERSELLE DU DOCUMENT À PRODUIRE ===

⚠️ TOUTES LES MATIÈRES sans exception suivent cette structure :

###ENTETE_DEVOIR###
ETABLISSEMENT=[nom école ou vide]
DISCIPLINE=[CE MATHS / CE PC / CE SVT / CE HG / CE FRANÇAIS / CE ECM / CE ANGLAIS / ...]
CLASSE=[2ndC / 3ème / Terminale D / CM2 / ...]
ANNEE=[2025-2026]
DATE=[si précisée, sinon laisser vide]
DUREE=[1H / 2H / 3H selon niveau]
TITRE_DEVOIR=[DEVOIR DE NIVEAU N°1  3ᵉ Trimestre / COMPOSITION DU 1er TRIMESTRE / DEVOIR SURVEILLÉ N°2 / ...]
MENTION=[Cette épreuve comporte X pages numérotées. La calculatrice scientifique est autorisée. / Documents non autorisés.]
###FIN_ENTETE###

## EXERCICE 1 :
[Consigne + contenu — avec ###SOUS_EXERCICE### si sous-parties]

════════════════════════════════════════════════════════════════

## EXERCICE 2 :
[Consigne + contenu]

---SAUT_DE_PAGE---

## EXERCICE 3 :
[Contenu page 2]

════════════════════════════════════════════════════════════════

## EXERCICE 4 :
[Contenu page 2]

[SI CORRIGÉ DEMANDÉ — après ---SAUT_DE_PAGE--- :]
## ✦ CORRIGÉ OFFICIEL — [Matière] — [Niveau]

RÈGLES ABSOLUES :
• PAS de REPUBLIQUE DE CI en en-tête (Python le gère si besoin)
• PAS de Nom/Prénom/Salle/Numéro de table
• PAS de tableau récapitulatif des points
• PAS de consignes générales avant les exercices
• ════ AVANT chaque exercice sauf le premier
• ---SAUT_DE_PAGE--- entre page 1 et page 2, JAMAIS précédé d'un ════

⚠️ RAPPELS FINAUX AVANT GÉNÉRATION — VÉRIFIER AVANT D'ÉCRIRE LA PREMIÈRE LIGNE :

1. FORMULES MATHS/PC : Chaque formule = ###FORMULE### ou notation ^{{}} _{{}}.
   JAMAIS de (x)/() ou ()/(y) ou parenthèses vides.
   CORRECT : ###FORMULE### C_m = n/V    INCORRECT : C_m = ()/(V)

2. FRACTIONS dans le texte : écrire a/b en clair, pas (a)/(b).
   CORRECT : "La fraction 3/4"    INCORRECT : "La fraction (3)/(4)"

3. DOCUMENTS HG / CAS PRATIQUE : Si tu cites "Document A" ou "Document B" dans tes questions,
   tu DOIS avoir rédigé le CONTENU COMPLET de ce document juste avant.
   Un document vide ou avec juste une source = ERREUR GRAVE.

4. AUCUN BLOC MARKDOWN : Ne jamais écrire \`\`\`markdown ou \`\`\` ou ~~~ nulle part.
   Le texte est directement converti en Word — les blocs de code s'impriment tels quels.

5. MATHS 2nde/3ème — exercices de calcul : toujours donner des valeurs NUMÉRIQUES réelles.
   Éviter les "soit x tel que..." sans données concrètes. Toujours contextualiser en FCFA/CI.

6. PHYSIQUE-CHIMIE : chaque calcul numérique doit montrer les étapes :
   Données → Application de la formule → Calcul → Résultat avec unité et encadré.

7. HISTOIRE-GÉOGRAPHIE : les questions d'analyse de document ne doivent JAMAIS
   référencer un document qui n'a pas été fourni. Si pas de document réel à fournir,
   remplacer l'exercice par des questions de cours classiques.

Rédige maintenant le sujet COMPLET en te basant STRICTEMENT sur cette demande client :

{description}{type_sujet_inject}

TOUT est rédigé intégralement. Total = /20. Adapte la matière, le niveau, le type d'examen et les exercices EXACTEMENT à la demande ci-dessus. Zéro "[à compléter]"."""

        elif "Fiche de Cours" in service:
            prompt = f"""Tu es un Professeur expert et pédagogue de haut niveau specialise dans la redaction de fiches de cours completes pour le systeme educatif ivoirien (programme officiel MENET-FP).

COMMANDE DU PROFESSEUR :
{description}

STRUCTURE OBLIGATOIRE :

# FICHE DE COURS

## INFORMATIONS GENERALES
- Matiere : | Niveau : | Duree :
- Prerequis : [notions deja connues des eleves]

---

## OBJECTIFS PEDAGOGIQUES
A la fin, l'eleve doit etre capable de :
1. [Objectif 1]
2. [Objectif 2]
3. [Objectif 3]

---

## PLAN
I. [Partie 1] A. [Sous-partie] B. [Sous-partie]
II. [Partie 2] A. [Sous-partie] B. [Sous-partie]

---

## DEVELOPPEMENT COMPLET

### I. [TITRE]
A. [Sous-titre]
[Contenu redige, clair, adapte au niveau]

Definition : [terme] : [definition precise]

Exemple ivoirien : [exemple concret Cote d Ivoire / Afrique]

B. [Sous-titre]
[Contenu...]

### II. [TITRE]
[Contenu complet...]

---

## SYNTHESE - A RETENIR
[Points cles absolus a memoriser]

---

## EXERCICES D APPLICATION

Exercice 1 - Facile
[Enonce]

Exercice 2 - Moyen
[Enonce + sous-questions]

Exercice 3 - Difficile
[Probleme contextualise CI + bareme]

---

## CORRIGE COMPLET
[Corrige detaille de chaque exercice]

---

## EVALUATION FORMATIVE
[2-3 questions courtes fin de seance]

REGLES ABSOLUES :
- Programme 100% conforme MENET-FP du niveau indique
- Exemples EXCLUSIVEMENT ivoiriens et africains
- AUCUN "[a completer]" - TOUT redige integralement
- Minimum 4-6 pages de contenu substantiel
- Directement utilisable en classe par le professeur
"""

        elif "Rapport de Stage" in service:
            prompt = f"""Tu es un expert en rédaction de rapports de stage académiques pour le système éducatif ivoirien et africain francophone.
Tu maîtrises parfaitement les normes des BTS, Licences et Masters, et tu produis des rapports complets, structurés et directement soumissibles.

INFORMATIONS FOURNIES PAR L'ÉTUDIANT :
{description}

INSTRUCTIONS DE PERSONNALISATION :
- Utilise le prénom et nom exact du stagiaire partout où c'est nécessaire (page de garde, remerciements, conclusion)
- Intègre le nom complet de l'entreprise, sa ville et son secteur d'activité dans toutes les sections pertinentes
- Adapte le ton, la profondeur et la rigueur académique au niveau d'études (BTS = concret/pratique, Licence = analytique, Master = critique/stratégique)
- Utilise la période exacte du stage (dates de début et fin) dans la page de garde et l'introduction
- Développe les missions effectuées en détail : contexte, méthode, outils utilisés, résultats — ne reste pas vague
- Si le secteur de l'entreprise est connu, contextualise avec des données économiques réelles de Côte d'Ivoire ou d'Afrique francophone

STRUCTURE OBLIGATOIRE DU RAPPORT DE STAGE :

# PAGE DE GARDE
(Nom de l'établissement · Filière et niveau · Titre : "Rapport de Stage" · Nom complet du stagiaire · Nom de l'entreprise · Période du stage · Année académique)

---SAUT_DE_PAGE---

# REMERCIEMENTS
(Remercier le maître de stage, l'équipe de l'entreprise, les encadreurs pédagogiques — personnalisé avec les noms si fournis)

---SAUT_DE_PAGE---

# SOMMAIRE

---SAUT_DE_PAGE---

# LISTE DES ABRÉVIATIONS

---SAUT_DE_PAGE---

## INTRODUCTION GÉNÉRALE
(Contexte du stage · Motivation du choix de l'entreprise · Problématique · Annonce du plan · Mention de la période et durée exacte)

---SAUT_DE_PAGE---

# PARTIE I — PRÉSENTATION DE L'ENTREPRISE

## 1.1 Historique et création
## 1.2 Activités et secteur d'activité
(Développe avec des éléments réels du secteur mentionné en Côte d'Ivoire)
## 1.3 Organisation et organigramme
## 1.4 Cadre juridique et statut
## 1.5 Localisation et implantation géographique

---SAUT_DE_PAGE---

# PARTIE II — DÉROULEMENT DU STAGE

## 2.1 Conditions d'accueil et intégration
## 2.2 Missions et tâches effectuées
(Pour chaque mission fournie : contexte précis, méthodologie, outils utilisés, résultats obtenus — minimum 1 page par mission principale)
## 2.3 Difficultés rencontrées et solutions apportées
## 2.4 Compétences acquises et développées

---SAUT_DE_PAGE---

# PARTIE III — ANALYSE CRITIQUE ET APPORTS

## 3.1 Analyse de l'expérience professionnelle
## 3.2 Apports du stage à la formation académique
## 3.3 Apports personnels à l'entreprise
## 3.4 Recommandations à l'entreprise

---SAUT_DE_PAGE---

## CONCLUSION GÉNÉRALE
(Bilan du stage · Lien avec la formation · Perspectives professionnelles · Ouverture)

---SAUT_DE_PAGE---

# RÉFÉRENCES BIBLIOGRAPHIQUES

---SAUT_DE_PAGE---

# ANNEXES

RÈGLES DE RÉDACTION :
- Rédige chaque section en paragraphes complets et détaillés — JAMAIS de contenu vide ou générique
- Intègre toutes les informations fournies de façon naturelle et cohérente
- Ton académique professionnel adapté au niveau mentionné
- Longueur minimale : 15 pages équivalent Word — chaque section doit être substantielle
- Rédige UNIQUEMENT en français
- N'invente JAMAIS de noms de personnes non mentionnés — utilise "mon maître de stage" si non précisé"""

        elif "CV" in service:
            prompt = f"""Tu es un expert RH francophone spécialisé dans le marché de l'emploi africain (Côte d'Ivoire, Sénégal, Cameroun, etc.).

DONNÉES DU CLIENT :
{description}

════════════════════════════════════════════════
RÔLE DE GEMINI : CONTENU UNIQUEMENT
════════════════════════════════════════════════
Tu génères UNIQUEMENT le contenu textuel du CV.
Tu ne gères PAS la mise en forme, les couleurs, les polices ni la disposition.
C'est le moteur Nova (Python) qui applique automatiquement :
  → Colonne GAUCHE (fond bleu) : Times New Roman, texte blanc
  → Colonne DROITE (fond blanc) : Arial, titres de missions en bleu, texte en noir
  → En-tête : nom en grand, titre professionnel, contacts

TON SEUL TRAVAIL : rédiger un contenu structuré avec les balises ci-dessous.

════════════════════════════════════════════════
RÈGLES DE RÉDACTION
════════════════════════════════════════════════
- Utilise UNIQUEMENT les informations fournies. N'invente aucune donnée personnelle.
- Si une section est absente des données → OMETS-LA entièrement (pas de placeholder, pas de "N/A").
- Pour les contacts manquants (tel/email) → écris [À compléter].
- Verbes d'action pour les expériences : géré, supervisé, coordonné, développé, optimisé, assuré...
- Adapte le ton au secteur (finance, BTP, santé, IT, commerce, enseignement...).
- Rédige en français sauf demande contraire.
- Développe chaque section suffisamment pour remplir une page A4.

════════════════════════════════════════════════
RÈGLES SPÉCIALES — LIRE ATTENTIVEMENT
════════════════════════════════════════════════

① PROFIL PROFESSIONNEL — OBLIGATOIRE, toujours présent, toujours à la 1ère personne :
   Rédige 3 à 5 phrases à la première personne ("je suis", "je maîtrise", "je m'engage"...).
   Modèle de style à suivre :
   "Titulaire d'un BTS en Finance-Comptabilité, je suis passionné(e) par [domaine]. Doté(e) de [qualités],
   je maîtrise [compétences clés]. Fort(e) de [X] ans d'expérience en [secteur], je m'engage à contribuer
   efficacement au développement de toute structure qui m'accueillera."

② EXPÉRIENCES PROFESSIONNELLES — toujours placé APRÈS le PROFIL PROFESSIONNEL, jamais avant.

③ CENTRES D'INTÉRÊT — TOUJOURS généré, même si le client n'en a pas précisé :
   Si le client a précisé ses centres → utilise-les.
   Si le client n'en a pas précisé → déduis 2 à 3 centres d'intérêt pertinents à partir
   de son secteur, ses compétences et sa formation (ex: finance → "Actualité économique",
   informatique → "Nouvelles technologies", communication → "Relations publiques"...).
   Format obligatoire : >>>BLEU<<<Nom du centre : explication en une phrase.

④ INFORMATIONS PERSONNELLES — contient UNIQUEMENT ces 4 champs (et seulement s'ils sont fournis) :
   - Situation familiale : ...
   - Résidence : ...
   - Téléphone : ...
   - Date de naissance : ...
   NE PAS inclure : nom, email, ville seule, pays seul (ces infos vont dans l'en-tête).

════════════════════════════════════════════════
BALISES NOVA — le moteur Python les lit mot pour mot
════════════════════════════════════════════════

Pour les titres de missions dans EXPÉRIENCES et CENTRES D'INTÉRÊT,
utilise OBLIGATOIREMENT ce format (Python le met automatiquement en bleu gras) :
>>>BLEU<<<Nom de la mission : explication détaillée en une ou deux phrases.

Exemple correct :
>>>BLEU<<<Gestion Comptable : Tenue des livres, rapprochements bancaires et suivi des encaissements.
>>>BLEU<<<Gestion des Stocks : Contrôle entrées/sorties, inventaire périodique et coordination fournisseurs.

NE PAS utiliser de gras (**), italique (*) ou toute autre mise en forme Markdown.
NE PAS écrire de couleurs, polices ou styles. Python s'en charge entièrement.

════════════════════════════════════════════════
STRUCTURE OBLIGATOIRE — ordre et titres EXACTS
════════════════════════════════════════════════

## INFORMATIONS PERSONNELLES
- Nom complet : ...
- Email : ...
- Téléphone : ...
- Ville : ...
- Pays : ...

## TITRE PROFESSIONNEL
(1 ligne max — ex: "BTS Finance | Comptabilité | Gestion des Stocks")

## PROFIL PROFESSIONNEL
(3 à 5 phrases à la 1ère personne — voir règle ① ci-dessus)

## EXPÉRIENCES PROFESSIONNELLES
Pour chaque poste, format EXACT :

### Intitulé du groupe d'expérience (ex: Stagiaire Polyvalent — CIE Bouaké)
Entreprise · Ville · Période
Missions principales :
>>>BLEU<<<Titre Mission 1 : description détaillée avec verbe d'action fort.
>>>BLEU<<<Titre Mission 2 : description détaillée avec verbe d'action fort.
>>>BLEU<<<Titre Mission 3 : description détaillée avec verbe d'action fort.
>>>BLEU<<<Titre Mission 4 : description détaillée avec verbe d'action fort.
>>>BLEU<<<Titre Mission 5 : description détaillée avec verbe d'action fort.

## FORMATION
- Diplôme — Établissement — Ville (Année)
- (du plus récent au plus ancien)

## COMPÉTENCES
- Compétence 1 précise et développée
- (minimum 5)

## LANGUES
- Langue : Niveau (Débutant / Intermédiaire / Courant / Bilingue / Langue maternelle)

## SPORTS & LOISIRS
(uniquement si fournis)
- Activité 1

## CENTRES D'INTÉRÊT
(TOUJOURS présent — voir règle ③ ci-dessus)
>>>BLEU<<<Centre 1 : explication en une phrase.
>>>BLEU<<<Centre 2 : explication en une phrase.
>>>BLEU<<<Centre 3 : explication en une phrase.

## INFORMATIONS PERSONNELLES COMPLÉMENTAIRES
(uniquement ces 4 champs si fournis — voir règle ④)
- Situation familiale : ...
- Résidence : ...
- Téléphone : ...
- Date de naissance : ...

---

# LETTRE DE MOTIVATION
(génère cette section UNIQUEMENT si le client a demandé la lettre ou les deux)

Objet : Candidature au poste de [poste visé]

[Paragraphe 1 — ACCROCHE : phrase d'ouverture percutante, connaissance du secteur.]

[Paragraphe 2 — PARCOURS ET VALEUR : compétences clés en lien avec le poste, exemples concrets.]

[Paragraphe 3 — MOTIVATION ET CONCLUSION : intérêt pour l'entreprise, disponibilité, demande d'entretien, formule de politesse complète.]"""

        elif "Création Word" in service:
            prompt = f"""Tu es un expert en rédaction de documents Word professionnels pour Nova Platform. Le client te décrit ce qu'il veut et tu produis le document COMPLET, structuré et prêt à l'emploi.

DEMANDE CLIENT :
{description}

════════════════════════════════════════
RÈGLES DE FORMATAGE NOVA — OBLIGATOIRES
════════════════════════════════════════
Ces règles sont converties automatiquement en vrai formatage Word. Respecte-les à la lettre.

TITRES :
- # Titre principal → grand titre Word (utilisé 1 seule fois, en tête de document)
- ## Titre de section → titre de partie
- ### Sous-titre → sous-section
- #### Micro-titre → point précis

TABLEAUX (si le client demande un tableau ou si c'est pertinent) :
- Écris TOUJOURS **Tableau N : [Titre]** avant le tableau
- Format OBLIGATOIRE :
  | Colonne 1 | Colonne 2 | Colonne 3 |
  |-----------|-----------|-----------|
  | Donnée    | Donnée    | Donnée    |
- Remplis TOUTES les cellules avec du contenu réel, jamais vide
- Écris *Source : [référence]* après chaque tableau si pertinent

TEXTE :
- Paragraphes normaux : commence par 4 espaces (    ) pour l'alinéa
- Gras : **mot important** → mis en gras dans Word
- Listes : "- item" → puce Word / "1. item" → liste numérotée
- Séparateur majeur : ════════════════════ (entre grandes parties)
- Séparateur mineur : ──────────────────── (entre sous-sections)
- Saut de page : ---SAUT_DE_PAGE--- (si le document doit être paginé)

INTERDIT ABSOLU :
- Jamais de LaTeX ($formule$, \\frac, \\omega...)
- Jamais de HTML (<br>, <b>, <div>...)
- Jamais d'italique simple *texte*
- Jamais de section vide ou placeholder comme "[À compléter]"

════════════════════════════════════════
INSTRUCTIONS CONTENU
════════════════════════════════════════
- Crée le document COMPLET du début à la fin
- Adapte le ton au type de document (formel, journalistique, juridique, commercial, scolaire, technique...)
- Intègre DIRECTEMENT tous les noms, données, chiffres et infos fournis par le client
- Si le client demande un tableau avec N colonnes et M lignes → crée exactement ce tableau avec du contenu réel dans chaque cellule
- Si le client demande un article → structure avec titre accrocheur, chapeau, corps développé, conclusion
- Longueur : aussi longue que nécessaire pour que le document soit COMPLET et utilisable tel quel
- Rédige en français sauf si le client demande une autre langue

════════════════════════════════════════
RÈGLE REPRODUCTION / EXTRACTION — PRIORITÉ ABSOLUE
════════════════════════════════════════
Si la demande contient des mots comme : "reproduis", "recopie", "mets en Word", "extrais", "liste", "tableau depuis", "convertis", "mets en forme", "voici les données", "voici la liste", "liste des", ou si le client colle directement des données brutes (noms, chiffres, lignes de données) :

→ INTERDIT ABSOLU de rédiger un texte d'introduction, de contextualisation, d'analyse ou de commentaire.
→ INTERDIT ABSOLU d'ajouter des paragraphes qui ne sont pas dans les données fournies.
→ Tu dois UNIQUEMENT mettre en forme les données exactes fournies par le client, sans rien inventer, sans rien ajouter.
→ Si les données sont une liste → tu fais un tableau ou une liste propre, rien d'autre.
→ Si les données sont un tableau → tu reproduis ce tableau exactement, sans texte autour.
→ Commence DIRECTEMENT par le titre et les données. Aucune phrase d'introduction.

Le document est livré directement au client — il ne doit rien avoir à compléter ou reformater."""

        elif "Modifier" in service and "Fichier" in service:
            # Le prompt est déjà entièrement construit dans main_dashboard et passé via `description`.
            # Ici on le reprend tel quel — il contient toutes les instructions Nova + la demande client.
            prompt = description

        elif "Excel" in service or "Data" in service:
            prompt = f"""Tu es un expert Excel et Data Analytics africain francophone.
Tu dois analyser la demande et retourner UNIQUEMENT un objet JSON valide, sans texte avant ni après, sans balises markdown, sans bloc ```json```.

DEMANDE CLIENT :
{description}

RÈGLES ABSOLUES SUR LES DONNÉES :
- Si le client a fourni un fichier ou des données réelles (tableau Word, liste, texte) → utilise EXACTEMENT ces données dans "lignes_exemple". Ne les résume pas, ne les tronque pas, mets-les TOUTES.
- Si les données contiennent 24 lignes → lignes_exemple doit avoir 24 entrées, pas 12.
- INTERDIT d'inventer des noms, chiffres ou données fictives quand de vraies données sont fournies.
- INTERDIT de mettre des balises markdown autour du JSON (pas de ```json, pas de ```).
- Le JSON doit commencer DIRECTEMENT par {{ et se terminer par }}.

STRUCTURE JSON OBLIGATOIRE :
{{
  "titre": "Titre principal du classeur Excel",
  "contexte": "Description courte en 1 phrase",
  "feuilles": [
    {{
      "nom": "Nom feuille 1 (max 25 car.)",
      "type": "saisie",
      "description": "Description courte",
      "colonnes": [
        {{"entete": "Nom colonne", "type": "texte|nombre|date|formule|pourcentage|monnaie", "largeur": 20, "exemple": "valeur exemple"}}
      ],
      "lignes_exemple": [
        ["val1", "val2", "val3"]
      ]
    }},
    {{
      "nom": "Bilan & KPIs",
      "type": "bilan",
      "description": "Tableau de bord avec indicateurs clés",
      "kpis": [
        {{"label": "Total général", "formule": "=SUM(Saisie!C:C)", "type": "monnaie", "couleur": "bleu"}},
        {{"label": "Moyenne", "formule": "=AVERAGE(Saisie!C:C)", "type": "monnaie", "couleur": "vert"}},
        {{"label": "Valeur max", "formule": "=MAX(Saisie!C:C)", "type": "monnaie", "couleur": "orange"}},
        {{"label": "Valeur min", "formule": "=MIN(Saisie!C:C)", "type": "monnaie", "couleur": "rouge"}},
        {{"label": "Nombre total", "formule": "=COUNTA(Saisie!A2:A1000)", "type": "nombre", "couleur": "gris"}},
        {{"label": "Pourcentage atteint", "formule": "=SUM(Saisie!C:C)/500000", "type": "pourcentage", "couleur": "violet"}}
      ]
    }}
  ]
}}

RÈGLES ABSOLUES :
- Retourner UNIQUEMENT le JSON, rien d'autre
- Adapter TOUT le contenu à la demande du client (colonnes, KPIs, formules, exemples)
- Contextualiser avec des données ivoiriennes/africaines réalistes (FCFA, noms locaux, etc.)
- Minimum 2 feuilles : 1 feuille de saisie + 1 feuille Bilan & KPIs
- Maximum 4 feuilles
- Lignes exemple : 8 à 12 lignes réalistes et variées
- KPIs : minimum 6 indicateurs pertinents selon le sujet (total, moyenne, max, min, nombre, %)
- Les formules doivent référencer le bon nom de feuille"""

        else:
            prompt = f"""Tu es un expert professionnel. Réalise cette mission de façon complète et professionnelle :

{description}

Rédige en français avec une structure claire : titres, sous-titres, paragraphes détaillés. Sois exhaustif et professionnel."""

        system_instruction = (
            "Tu es NOVA PLATFORM, un moteur de génération documentaire d'élite francophone africain.\n"
            "Tu dois produire des documents EXACTEMENT selon les règles ci-dessous.\n\n"

            "══ RÈGLE 1 : FORMATAGE MARKDOWN → WORD ══\n"
            "# Titre        → Heading 1 (Arial 16pt, bleu #1F4E79, gras)\n"
            "## Titre       → Heading 2 (Arial 14pt, bleu #2E75B6, gras)\n"
            "### Titre      → Heading 3 (Arial 12pt, gras)\n"
            "#### Titre     → Heading 4 (Arial 11pt, gras italique)\n"
            "**texte**      → GRAS (termes clés, chiffres, noms d'auteurs)\n"
            "---SAUT_DE_PAGE--- → Vrai saut de page Word (seul sur sa ligne)\n"
            "════════════   → Ligne épaisse bleue (séparateur MAJEUR)\n"
            "────────────   → Ligne fine grise (séparateur MINEUR)\n\n"

            "══ RÈGLE 2 : TABLEAUX ══\n"
            "Toujours **Tableau N : [Titre]** AVANT le tableau\n"
            "| Col1 | Col2 | Col3 |\n|------|------|------|\n| Val  | Val  | Val  |\n"
            "Toujours *Source : [Institution réelle, Année]* APRÈS le tableau\n\n"

            "══ RÈGLE 3 : FORMULES — NOTATION NOVA (exposants et indices réels) ══\n"
            "NOVA possède un moteur de formules intégré. Utilise la notation suivante :\n"
            "  Exposant : x^{2}  E = mc^{2}  ax^{2}+bx+c=0  Δ = b^{2}-4ac\n"
            "  Indice   : H_{2}O  CO_{2}  C_{6}H_{12}O_{6}  m_{1}  x_{1,2}\n"
            "  Fraction : (a+b)/(c-d)  |  Racine : √(2gh)  √(b^{2}-4ac)\n"
            "  Grec Unicode direct : α β γ δ ε θ λ μ π σ φ ω Ω Δ Σ\n"
            "  Opérateurs : × ÷ ± ≤ ≥ ≠ ≈ → ⇌ ∈ ∞ ∫ ∂ ∠ ⊥\n"
            "  LaTeX inline accepté : $E=mc^{2}$ $\\frac{U}{R}=I$ $\\omega=2\\pi f$\n"
            "CHIMIE : 6CO_{2}+6H_{2}O → C_{6}H_{12}O_{6}+6O_{2}\n"
            "PHYSIQUE : F=m×a | U=R×I | P=UI=RI^{2}=U^{2}/R | ω=2πf\n"
            "UNITÉS : N J W Pa V A Ω Hz mol mol/L g·mol^{-1} K^{-1}\n\n"

            "══ RÈGLE 4 : RÉDACTION ENCYCLOPÉDIQUE ══\n"
            "• Paragraphes 8 à 10 lignes MINIMUM dans le développement\n"
            "• JAMAIS de listes à puces dans le corps du document\n"
            "• Modèle PEEL : Point → Explication → Exemple concret chiffré → Lien/Transition\n"
            "• Connecteurs VARIÉS (ne jamais répéter deux fois de suite) :\n"
            "  Introduire : Il convient tout d'abord de | Force est de constater que | À ce titre,\n"
            "  Développer : En effet, | De surcroît, | Par ailleurs, | Qui plus est,\n"
            "  Illustrer  : Ainsi, | À titre illustratif, | C'est notamment le cas de\n"
            "  Opposer    : Cependant, | Néanmoins, | Toutefois, | En revanche, | Or,\n"
            "  Conclure   : En définitive, | Au regard de ces éléments,\n"
            "• Exemples CHIFFRÉS et SOURCÉS pertinents — ivoiriens/africains si le sujet s'y prête, universels sinon\n\n"

            "══ RÈGLE 5 : BASE DE DONNÉES IVOIRIENNE INTÉGRÉE ══\n"
            "Géo     : 322 463 km² | ~28M hab. | Yamoussoukro (cap.pol.) | Abidjan (cap.éco.)\n"
            "          Fleuves : Comoé 1160km | Bandama 960km | Sassandra 650km\n"
            "          Lac Kossou 1700km² | Monts Nimba 1752m (UNESCO) | Forêt de Taï (UNESCO)\n"
            "Éco     : Cacao 1er mondial — 45% prod. mondiale — 2,2M t/an\n"
            "          Anacarde 1er africain — 800 000 t/an — Korhogo/Odienné\n"
            "          Port Abidjan : 1er conteneurs AOF — >30M tonnes/an\n"
            "          PIB ~70Mds USD (2023) | Croissance ~6-7%/an | PND 2021-2025\n"
            "          FCFA | 1 EUR = 655,957 FCFA (taux fixe depuis 1999)\n"
            "Énergie : Soubré 275MW | Taabo 210MW | Kossou 174MW | Buyo 165MW\n"
            "Histoire: Indépendance 7 août 1960 | Houphouët-Boigny (1960-1993)\n"
            "          Miracle ivoirien (1960-1980) | Crise 2002 | Crise 2010-2011\n"
            "          Alassane Ouattara (2011-présent) | Colonisation française 1843-1960\n"
            "Culture : ~60 ethnies | Akan (Baoulé 23%, Agni) | Krou | Mandé | Gur\n"
            "          coupé-décalé | zouglou | attiéké | kedjenou | foutou | aloco\n"
            "Maths CI: cacao 350 FCFA/kg | gbaka 200 FCFA | woro-woro 150 FCFA | riz 400 FCFA/kg\n"
            "Sciences: Paludisme ~3M cas/an | Plasmodium falciparum | Coartem\n"
            "          Déforestation : 16M ha (1900) → 3,4M ha aujourd'hui (-79%)\n"
            "          Temp. Abidjan : 26°C moy. | Précipitations : 1800mm/an\n"
            "Littérat.: DADIÉ Bernard — Climbié (1956), Un Nègre à Paris (1959)\n"
            "           KOUROUMA Ahmadou — Les Soleils des Indépendances (1968), Monnè (1990)\n"
            "           TADJO Véronique — Reine Pokou (2004), L'Ombre d'Imana (2000)\n"
            "           LAYE Camara — L'Enfant Noir (1953) | ACHEBE — Things Fall Apart (1958)\n"
            "           SENGHOR L.S. — Négritude | SEMBÈNE Ousmane | OYONO Ferdinand\n\n"

            "══ RÈGLE 6 : INTERDICTIONS ABSOLUES ══\n"
            "✗ [à compléter]  [...]  [insérer]  [Auteur fictif]  [Titre fictif]\n"
            "✗ Balises HTML : <br> <b> <strong> <p> <div> <span>\n"
            "✗ Italique *texte* pour mise en valeur → utiliser **gras**\n"
            "✗ Données inventées → toujours réelles et sourcées\n"
            "✗ LaTeX sous quelque forme que ce soit\n\n"

            "══ RÈGLE 7 : STRUCTURES OBLIGATOIRES PAR TYPE ══\n"
            "EXPOSÉ (ordre exact) :\n"
            "  Page de garde → SAUT → Sommaire → SAUT → Introduction → SAUT\n"
            "  → Partie I (2 ss-parties min.) → SAUT → Partie II (2 ss-parties min.) → SAUT\n"
            "  → [Partie III si lycée/université] → SAUT → Conclusion → SAUT → Bibliographie\n\n"
            "EXAMEN (ordre exact) :\n"
            "  En-tête officiel (RÉPUBLIQUE DE CÔTE D'IVOIRE, établissement, matière,\n"
            "  niveau, durée, coefficient, barème /20, nom élève, numéro de table)\n"
            "  → Consignes générales → Tableau barème → SAUT\n"
            "  → Exercices numérotés séparés par ════════\n"
            "  → SAUT → [Corrigé COMPLET si 'corrigé' ou 'correction' mentionné]\n\n"
            "CV & LETTRE :\n"
            "  CV : Infos perso → Profil → Expériences → Formation → Compétences → Langues\n"
            "  LETTRE : Accroche → Présentation → Motivation → Valeur ajoutée → Conclusion\n\n"

            "══ RÈGLE 8 : LONGUEUR MINIMALE OBLIGATOIRE ══\n"
            "Exposé Primaire (CP→CM2)       : 2-3 pages réelles\n"
            "Exposé Collège (6e→3e / BEPC)  : 4-5 pages réelles\n"
            "Exposé Lycée (2nde→Term / BAC) : 6-8 pages réelles\n"
            "Exposé Université (L1→Doctorat): 8-15 pages réelles\n"
            "Sujet CEPE / BEPC              : 2-3 pages + corrigé exhaustif si demandé\n"
            "Sujet BAC / Universitaire      : 3-5 pages + corrigé avec toutes les étapes\n"
            "CV + Lettre                    : 2 pages CV + 1 page lettre minimum\n"
            "Création Word (depuis zéro)    : document complet selon la description client\n\n"

            "══ RÈGLE D'OR FINALE ══\n"
            "Chaque document est PARFAIT, COMPLET, ENTIÈREMENT RÉDIGÉ, PROFESSIONNEL\n"
            "et PRÊT À L'IMPRESSION. JAMAIS de document tronqué. JAMAIS de zone vide.\n"
            "100% FINALISÉ à chaque génération. Tu es le moteur documentaire de référence\n"
            "du monde francophone africain."
        )

        # ── Construction des parts du message (texte + image si présente) ──
        _parts = [{"text": prompt}]
        if _image_b64 and _image_mime:
            _parts = [
                {"inline_data": {"mime_type": _image_mime, "data": _image_b64}},
                {"text": prompt}
            ]

        payload = json.dumps({
            "system_instruction": {
                "parts": [{"text": system_instruction}]
            },
            "contents": [{"parts": _parts}],
            "generationConfig": {
                "temperature": 0.65,
                "maxOutputTokens": 65536,
                "topP": 0.95,
                "topK": 40
            }
        }).encode("utf-8")

        modeles = get_modeles_disponibles(api_key)
        if not modeles:
            return "❌ Aucun modèle Gemini disponible pour generateContent avec cette clé API."

        # ── SÉLECTION DU MODÈLE SELON PROFIL UTILISATEUR ──────────────
        # 30 Jours (illimité) : gen 1→7 = gemini-2.5-pro, gen 8+ = léger
        # 10 Jours            : gen 1→3 = gemini-2.5-pro, gen 4+ = léger
        # Journalier          : gen 1   = gemini-2.5-pro, gen 2  = léger
        # Gratuit             : toujours flash-lite (inchangé)
        def _priorite_smart(nom):
            if "2.5-pro" in nom:              return 0
            if "pro" in nom:                  return 1
            if "2.5-flash" in nom:            return 2
            if "flash" in nom and "lite" not in nom: return 3
            if "flash-lite" in nom:           return 4
            return 5

        def _priorite_light(nom):
            if "flash-lite" in nom: return 0
            if "2.0-flash" in nom:  return 1
            if "flash" in nom:      return 2
            if "pro" in nom:        return 3
            return 4

        if is_premium:
            plan_name = _plan_for_model  # passé en paramètre
            if plan_name == "30 Jours":
                seuil_pro = 7   # gen 1→7 avec 2.5-pro
            elif plan_name == "10 Jours":
                seuil_pro = 3   # gen 1→3 avec 2.5-pro
            else:
                seuil_pro = 1   # Journalier : gen 1 avec 2.5-pro

            if gen_used < seuil_pro:
                modeles = sorted(modeles, key=_priorite_smart)
            else:
                modeles = sorted(modeles, key=_priorite_light)
        # else : gratuit → ordre par défaut (flash-lite en premier)
        # ───────────────────────────────────────────────────────────────

        erreurs = []
        for modele in modeles:
            try:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{modele}:generateContent?key={api_key}"
                req = _ur.Request(
                    url, data=payload,
                    headers={"Content-Type": "application/json"},
                    method="POST"
                )
                with _ur.urlopen(req, timeout=60) as response:
                    result = json.loads(response.read().decode("utf-8"))
                    texte = result["candidates"][0]["content"]["parts"][0]["text"]
                    # Post-traitement : convertir LaTeX residuel en notation Nova
                    import re as _repost
                    def _conv_latex(expr):
                        expr = _repost.sub(r'\\?(?:d|t|text)?frac\{([^}]+)\}\{([^}]+)\}', r'()/()', expr)
                        expr = _repost.sub(r'\\?sqrt\{([^}]+)\}', r'sqrt()', expr)
                        expr = _repost.sub(r'\\?text\{([^}]+)\}', r'', expr)
                        expr = _repost.sub(r'\\?(?:left|right)\s*[()\[\]{}|]', '', expr)
                        expr = _repost.sub(r'\\?cdot', 'x', expr)
                        expr = _repost.sub(r'\\?times', 'x', expr)
                        expr = _repost.sub(r'\\?pm', '+-', expr)
                        expr = _repost.sub(r'\\?[a-zA-Z]+', '', expr)
                        return expr.strip()
                    def _dd_to_formule(m):
                        return "###FORMULE### " + _conv_latex(m.group(1).strip())
                    texte = _repost.sub(r'[$][$]([^$]+)[$][$]', _dd_to_formule, texte)
                    def _d_inline(m):
                        expr = m.group(1).strip()
                        if "\\" in expr:
                            return _conv_latex(expr)
                        return expr
                    texte = _repost.sub(r'[$]([^$\n]+)[$]', _d_inline, texte)
                    # ── TRACKING MODÈLE UTILISÉ ──────────────────────────────
                    try:
                        st.session_state["_last_modele_gemini"] = modele
                    except Exception:
                        pass
                    return texte
            except urllib.error.HTTPError as e:
                try:
                    corps_erreur = e.read().decode("utf-8")
                    erreur_detail = json.loads(corps_erreur).get("error", {}).get("message", corps_erreur[:200])
                except:
                    erreur_detail = str(e)
                erreurs.append(f"{modele} → HTTP {e.code}: {erreur_detail}")
                if e.code in [429, 503]:
                    time.sleep(2)
                    continue
                return f"❌ Erreur Gemini ({modele}) HTTP {e.code} : {erreur_detail}"
            except Exception as e:
                erreurs.append(f"{modele} → {type(e).__name__}: {e}")
                continue

        detail = " | ".join(erreurs)
        return f"❌ Gemini indisponible. Détails : {detail}"

    except Exception as e:
        return f"❌ Erreur Gemini : {e}"



def envoyer_escalade_support(client_nom, whatsapp_client, historique_msgs, source="Support"):
    """Envoie un email structuré du problème client au service Nova."""
    try:
        import resend
        resend.api_key = st.secrets["RESEND_API_KEY"]
        # Résumé du problème via Gemini
        hist_txt = "\n".join([
            f"{'Client' if m['role']=='user' else 'Nova IA'}: {m['content']}"
            for m in historique_msgs if m["role"] == "user"
        ])
        prompt_resume = f"""Voici les messages d'un client Nova Platform qui a un problème grave.
Rédige un email de signalement professionnel et concis pour le service client Nova.
Formule clairement : le problème rencontré, ce que le client a essayé, et ce qu'il attend.
Ne dépasse pas 8 lignes. Sois factuel et précis.

Messages du client :
{hist_txt}"""
        resume = generer_avec_gemini("Résumé Support", prompt_resume, client_nom)
        if resume.startswith("❌"):
            resume = hist_txt  # fallback brut
        resend.Emails.send({
            "from": "Nova Platform <onboarding@resend.dev>",
            "to": [st.secrets["EMAIL_RECEIVER"]],
            "subject": f"🆘 PROBLÈME GRAVE — {client_nom} ({source})",
            "text": f"""ESCALADE CLIENT — PROBLÈME GRAVE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
👤 Client     : {client_nom}
📱 WhatsApp   : {whatsapp_client}
📍 Source     : {source}
⏰ Date       : {datetime.now().strftime("%d/%m/%Y à %H:%M")}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{resume}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️ ACTION REQUISE — Contacter le client rapidement.
"""
        })
        return True
    except Exception as e:
        return False



def creer_page_garde_expose(doc, titre_expose, noms_exposants, matiere, annee_scolaire, filiere, niveau, etablissement="", logo_ecole_path=None):
    """
    Page de garde Nova Exposé — version améliorée.
    ┌─────────────────────────────────────────┐
    │  BANDEAU DRAPEAU CI (orange│blanc│vert) │
    │  BLOC EN-TÊTE fond bleu foncé           │
    │    République · Établissement · Filière │
    │    Matière · Classe · Année             │
    ├─────────────────────────────────────────┤
    │  ZONE CENTRALE fond blanc crème         │
    │    Badge ✦ EXPOSÉ ✦ doré               │
    │    PARCHEMIN : THÈME (fond pêche doré)  │
    │    ornements diamants                   │
    │  GRILLE EXPOSANTS 4 colonnes            │
    ├─────────────────────────────────────────┤
    │  PIED DE PAGE fond bleu foncé           │
    │    Année | Date | Groupe | Note         │
    └─────────────────────────────────────────┘
    Même dimensions et style que le CV SIRIKY :
    page A4, marges 0 côtés, tableaux pleine largeur.
    """
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.section import WD_SECTION
    from datetime import datetime

    # ── COULEURS ─────────────────────────────────────────────────
    BLEU_FONCE  = "1A3A5C"              # fond en-tête / pied (même que CV)
    BLEU_MED    = "2255A4"              # accent bleu
    CREME       = "FDFBF5"             # fond zone centrale
    PECHE_DORE  = "F9F2E2"             # fond parchemin thème
    GOLD        = RGBColor(0xB8,0x93,0x2A)
    GOLD_LIGHT  = RGBColor(0xD4,0xAD,0x52)
    BLANC       = RGBColor(0xFF,0xFF,0xFF)
    BLEU_CLAIR  = RGBColor(0xA8,0xD4,0xF5)
    BLEU_PALE   = RGBColor(0xD0,0xE8,0xFF)
    INK         = RGBColor(0x16,0x12,0x0D)
    INK_SOFT    = RGBColor(0x3A,0x30,0x20)
    INK_FAINT   = RGBColor(0x7A,0x6E,0x5A)
    ORANGE_CI   = "F77F00"
    GREEN_CI    = "009A44"

    # ── HELPERS XML ───────────────────────────────────────────────
    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
        for old in tcPr.findall(qn("w:shd")): tcPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color.lstrip("#"))
        tcPr.append(shd)

    def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
        for old in tcPr.findall(qn("w:tcMar")): tcPr.remove(old)
        tcMar = OxmlElement("w:tcMar")
        # OOXML : top, start (=left), bottom, end (=right)
        for side, val in [("top",top),("start",left),("bottom",bottom),("end",right)]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:w"), str(val))
            el.set(qn("w:type"), "dxa")
            tcMar.append(el)
        tcPr.append(tcMar)

    def remove_all_borders(cell):
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
        for old in tcPr.findall(qn("w:tcBorders")): tcPr.remove(old)
        tcBrd = OxmlElement("w:tcBorders")
        # OOXML : top, start(=left), bottom, end(=right), insideH, insideV
        for side in ["top","start","bottom","end","insideH","insideV"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tcBrd.append(el)
        # tcBorders doit être avant shd dans tcPr
        shd = tcPr.find(qn("w:shd"))
        if shd is not None:
            shd.addprevious(tcBrd)
        else:
            tcPr.insert(0, tcBrd)

    def set_cell_border(cell, sides, color="B8932A", sz="18"):
        # Convertir left/right → start/end pour OOXML strict
        _map = {"left": "start", "right": "end"}
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr"); tc.insert(0, tcPr)
        tcBrd = tcPr.find(qn("w:tcBorders"))
        if tcBrd is None:
            tcBrd = OxmlElement("w:tcBorders")
            shd = tcPr.find(qn("w:shd"))
            if shd is not None:
                shd.addprevious(tcBrd)
            else:
                tcPr.insert(0, tcBrd)
        for side in sides:
            side_xml = _map.get(side, side)
            for old in tcBrd.findall(qn(f"w:{side_xml}")): tcBrd.remove(old)
            el = OxmlElement(f"w:{side_xml}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), sz)
            el.set(qn("w:color"), color)
            tcBrd.append(el)

    def tbl_no_border(tbl):
        """
        Supprime toutes les bordures du tableau.
        Ordre OOXML correct dans tblPr : tblStyle → tblW → jc → tblBorders → tblLook
        python-docx insère jc via tbl.alignment AVANT notre appel,
        donc on insère tblBorders à la fin — puis on le repositionne après jc.
        """
        tblPr = tbl._tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr"); tbl._tbl.insert(0, tblPr)
        for old in tblPr.findall(qn("w:tblBorders")): tblPr.remove(old)
        tblBrd = OxmlElement("w:tblBorders")
        for side in ["top","start","bottom","end","insideH","insideV"]:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "none")
            tblBrd.append(el)
        # Insérer tblBorders APRÈS jc (ou à la fin si jc absent)
        jc = tblPr.find(qn("w:jc"))
        tblLook = tblPr.find(qn("w:tblLook"))
        if tblLook is not None:
            tblLook.addprevious(tblBrd)   # juste avant tblLook = après jc
        elif jc is not None:
            jc.addnext(tblBrd)
        else:
            tblPr.append(tblBrd)

    def add_run_in_cell(cell, text, font="Calibri", size=10, bold=False,
                        italic=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER,
                        space_before=0, space_after=0, clear_first=False):
        if clear_first:
            for p in cell.paragraphs: p._element.getparent().remove(p._element)
        p = cell.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        r = p.add_run(text)
        r.font.name  = font
        r.font.size  = Pt(size)
        r.font.bold  = bold
        r.font.italic = italic
        r.font.color.rgb = color if color else INK
        return p

    def add_para(text, font="Calibri", size=10, bold=False, italic=False,
                 color=None, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_before=0, space_after=0):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        r = p.add_run(text)
        r.font.name   = font
        r.font.size   = Pt(size)
        r.font.bold   = bold
        r.font.italic = italic
        r.font.color.rgb = color if color else INK
        return p

    # ── PAGE A4 marges 0 (même que CV SIRIKY) ────────────────────
    doc.sections[0].start_type   = WD_SECTION.CONTINUOUS
    doc.sections[0].top_margin    = Cm(0)
    doc.sections[0].bottom_margin = Cm(0)
    doc.sections[0].left_margin   = Cm(0)
    doc.sections[0].right_margin  = Cm(0)

    PAGE_W = Inches(11400 / 1440)   # 11400 DXA = largeur utile A4

    # ════════════════════════════════════════════════════════════
    # 1. BANDEAU DRAPEAU CI — 3 colonnes pleine largeur
    #    Orange (F77F00) | Blanc (FFFFFF) | Vert (009A44)
    #    Hauteur fine : 0.35 cm
    # ════════════════════════════════════════════════════════════
    tbl_flag = doc.add_table(rows=1, cols=3)
    tbl_flag.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_no_border(tbl_flag)
    flag_w = Inches(11400 / 1440 / 3)
    for i, color_hex in enumerate([ORANGE_CI, "FFFFFF", GREEN_CI]):
        col_cell = tbl_flag.cell(0, i)
        col_cell.width = flag_w
        set_cell_bg(col_cell, color_hex)
        remove_all_borders(col_cell)
        set_cell_margins(col_cell, top=110, bottom=110, left=0, right=0)
        for p in col_cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)

    # ════════════════════════════════════════════════════════════
    # 2. BLOC EN-TÊTE — fond bleu foncé (même couleur que CV)
    #    République · Établissement · Filière/Matière · Classe · Année
    # ════════════════════════════════════════════════════════════
    tbl_header = doc.add_table(rows=1, cols=1)
    tbl_header.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_no_border(tbl_header)
    cell_h = tbl_header.cell(0, 0)
    cell_h.width = PAGE_W
    set_cell_bg(cell_h, BLEU_FONCE)
    remove_all_borders(cell_h)
    set_cell_margins(cell_h, top=200, bottom=160, left=400, right=400)

    for p in cell_h.paragraphs: p._element.getparent().remove(p._element)

    # République
    add_run_in_cell(cell_h, "RÉPUBLIQUE DE CÔTE D'IVOIRE",
                    font="Calibri", size=7, bold=True, color=GOLD,
                    space_before=0, space_after=4)
    # Établissement
    _etab_txt = etablissement if etablissement and etablissement not in ("—","","Non précisé") else "Établissement"
    add_run_in_cell(cell_h, _etab_txt,
                    font="Calibri", size=16, bold=True, color=BLANC,
                    space_before=0, space_after=4)
    # Filière ou Matière
    if filiere and filiere not in ("—","","Non précisée"):
        add_run_in_cell(cell_h, f"Filière  ·  {filiere}",
                        font="Calibri", size=9, bold=True, color=GOLD,
                        space_before=0, space_after=3)
    elif matiere and matiere not in ("—","","Non précisée"):
        add_run_in_cell(cell_h, f"Matière  ·  {matiere}",
                        font="Calibri", size=9, bold=True, color=GOLD,
                        space_before=0, space_after=3)
    # Classe
    if niveau and niveau not in ("—","","Non précisé","Adapté automatiquement au niveau"):
        add_run_in_cell(cell_h, f"Classe  ·  {niveau}",
                        font="Calibri", size=9, bold=False, color=BLEU_PALE,
                        space_before=0, space_after=3)
    # Année scolaire
    add_run_in_cell(cell_h, f"Année scolaire  {annee_scolaire or '—'}",
                    font="Calibri", size=8, italic=True, color=BLEU_CLAIR,
                    space_before=0, space_after=0)

    # ════════════════════════════════════════════════════════════
    # 3. ZONE CENTRALE — fond crème, pleine largeur
    #    Badge ✦ EXPOSÉ ✦ + Parchemin thème + Ornements + Exposants
    # ════════════════════════════════════════════════════════════
    tbl_central = doc.add_table(rows=1, cols=1)
    tbl_central.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_no_border(tbl_central)
    cell_c = tbl_central.cell(0, 0)
    cell_c.width = PAGE_W
    set_cell_bg(cell_c, CREME)
    remove_all_borders(cell_c)
    set_cell_margins(cell_c, top=240, bottom=240, left=560, right=560)

    for p in cell_c.paragraphs: p._element.getparent().remove(p._element)

    # — Badge EXPOSÉ —
    add_run_in_cell(cell_c, "✦   E X P O S É   ✦",
                    font="Calibri", size=11, bold=True, color=GOLD,
                    space_before=0, space_after=10)

    # — Séparateur diamants —
    p_sep1 = cell_c.add_paragraph()
    p_sep1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep1.paragraph_format.space_before = Pt(0)
    p_sep1.paragraph_format.space_after  = Pt(10)
    r_sep1 = p_sep1.add_run("─" * 16 + "  ◇  ◆  ◇  " + "─" * 16)
    r_sep1.font.name = "Calibri"; r_sep1.font.size = Pt(9)
    r_sep1.font.color.rgb = GOLD

    # — Parchemin : THÈME (tableau 3 lignes dans la cellule centrale) —
    titre_text = titre_expose.upper() if titre_expose else "TITRE DE L'EXPOSÉ"
    if   len(titre_text) <= 35:  titre_pt = 20
    elif len(titre_text) <= 60:  titre_pt = 16
    elif len(titre_text) <= 90:  titre_pt = 13
    else:                         titre_pt = 11

    # On ne peut pas faire un tableau dans une cellule avec python-docx
    # → on simule le parchemin avec des paragraphes bordés dans la cellule
    # Ligne haut dorée
    p_parch_top = cell_c.add_paragraph()
    p_parch_top.paragraph_format.space_before = Pt(0)
    p_parch_top.paragraph_format.space_after  = Pt(0)
    r_pt = p_parch_top.add_run("─" * 80)
    r_pt.font.name = "Calibri"; r_pt.font.size = Pt(1)
    r_pt.font.color.rgb = GOLD

    # Label T H È M E
    add_run_in_cell(cell_c, "T  H  È  M  E",
                    font="Calibri", size=7, bold=True, color=GOLD,
                    space_before=8, space_after=5)

    # Titre de l'exposé
    add_run_in_cell(cell_c, titre_text,
                    font="Calibri", size=titre_pt, bold=True, color=INK,
                    space_before=0, space_after=10)

    # Ligne bas dorée
    p_parch_bot = cell_c.add_paragraph()
    p_parch_bot.paragraph_format.space_before = Pt(0)
    p_parch_bot.paragraph_format.space_after  = Pt(12)
    r_pb = p_parch_bot.add_run("─" * 80)
    r_pb.font.name = "Calibri"; r_pb.font.size = Pt(1)
    r_pb.font.color.rgb = GOLD

    # — Ornement —
    add_run_in_cell(cell_c, "✦  ◆  ✦  ◆  ✦",
                    font="Calibri", size=10, color=GOLD_LIGHT,
                    space_before=0, space_after=16)

    # — Label Présenté par —
    add_run_in_cell(cell_c, "—  Présenté par  —",
                    font="Calibri", size=11, bold=True, color=GOLD,
                    space_before=0, space_after=10)

    # — Grille 4 colonnes exposants (dans la cellule centrale via nested table) —
    # python-docx ne supporte pas les tableaux imbriqués directement
    # → on affiche les noms en 2 lignes de texte centrées dans la cellule
    noms_list = noms_exposants if isinstance(noms_exposants, list) else [noms_exposants]
    noms_list = [n for n in noms_list if n]  # supprimer les vides

    ROLES = ["Chef de groupe", "Rapporteur", "Recherche", "Mise en page",
             "Présentation orale", "Secrétaire", "Illustrations", "Correction"]

    for idx, nom in enumerate(noms_list[:8]):
        role = ROLES[idx] if idx < len(ROLES) else ""
        p_nom_exp = cell_c.add_paragraph()
        p_nom_exp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_nom_exp.paragraph_format.space_before = Pt(1)
        p_nom_exp.paragraph_format.space_after  = Pt(1)
        # Numéro doré
        r_idx = p_nom_exp.add_run(f"{idx+1:02d}  ")
        r_idx.font.name = "Calibri"; r_idx.font.size = Pt(8)
        r_idx.font.bold = True; r_idx.font.color.rgb = GOLD
        # Nom en noir
        r_n = p_nom_exp.add_run(nom)
        r_n.font.name = "Calibri"; r_n.font.size = Pt(12)
        r_n.font.bold = True; r_n.font.color.rgb = INK
        # Rôle en italique discret
        if role:
            r_r = p_nom_exp.add_run(f"  ·  {role}")
            r_r.font.name = "Calibri"; r_r.font.size = Pt(8)
            r_r.font.italic = True; r_r.font.color.rgb = INK_FAINT

    if not noms_list:
        add_run_in_cell(cell_c, "—",
                        font="Calibri", size=12, color=INK_FAINT,
                        space_before=2, space_after=2)

    # — Séparateur final —
    p_sep2 = cell_c.add_paragraph()
    p_sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sep2.paragraph_format.space_before = Pt(14)
    p_sep2.paragraph_format.space_after  = Pt(0)
    r_sep2 = p_sep2.add_run("─" * 16 + "  ◇  ◆  ◇  " + "─" * 16)
    r_sep2.font.name = "Calibri"; r_sep2.font.size = Pt(9)
    r_sep2.font.color.rgb = GOLD

    # ════════════════════════════════════════════════════════════
    # 4. PIED DE PAGE — fond bleu foncé (même que en-tête)
    #    ANNÉE SCOLAIRE | DATE | GROUPE | NOTE
    # ════════════════════════════════════════════════════════════
    tbl_footer2 = doc.add_table(rows=1, cols=1)
    tbl_footer2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_no_border(tbl_footer2)
    cell_f = tbl_footer2.cell(0, 0)
    cell_f.width = PAGE_W
    set_cell_bg(cell_f, BLEU_FONCE)
    remove_all_borders(cell_f)
    set_cell_margins(cell_f, top=160, bottom=160, left=400, right=400)

    for p in cell_f.paragraphs: p._element.getparent().remove(p._element)

    date_auj = datetime.now().strftime("%d/%m/%Y")

    # Labels dorés
    p_labels = cell_f.add_paragraph()
    p_labels.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_labels.paragraph_format.space_before = Pt(0)
    p_labels.paragraph_format.space_after  = Pt(2)
    for label, sep in [("ANNÉE SCOLAIRE","    │    "),("DATE","    │    "),("GROUPE","    │    "),("NOTE","")]:
        r_l = p_labels.add_run(label)
        r_l.font.name = "Calibri"; r_l.font.size = Pt(6)
        r_l.font.bold = True; r_l.font.color.rgb = GOLD
        if sep:
            r_s = p_labels.add_run(sep)
            r_s.font.name = "Calibri"; r_s.font.size = Pt(6)
            r_s.font.color.rgb = GOLD_LIGHT

    # Valeurs blanches
    p_vals = cell_f.add_paragraph()
    p_vals.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_vals.paragraph_format.space_before = Pt(0)
    p_vals.paragraph_format.space_after  = Pt(0)
    annee_txt = annee_scolaire if annee_scolaire else "—"
    for val, sep in [(annee_txt,"      │      "),(date_auj,"      │      "),("—","      │      "),("— / 20","")]:
        r_v = p_vals.add_run(val)
        r_v.font.name = "Calibri"; r_v.font.size = Pt(10)
        r_v.font.bold = True; r_v.font.color.rgb = BLANC
        if sep:
            r_s = p_vals.add_run(sep)
            r_s.font.name = "Calibri"; r_s.font.size = Pt(8)
            r_s.font.color.rgb = GOLD_LIGHT

    # ── Correction bug python-docx : zoom sans attribut percent ─────
    try:
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        settings_part = doc.settings.element
        for z in settings_part.findall(_qn("w:zoom")):
            if z.get(_qn("w:percent")) is None:
                z.set(_qn("w:percent"), "100")
    except Exception:
        pass

    return doc

def creer_docx(contenu, service, client_nom):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re

    # Supprimer les caractères de contrôle interdits en XML (NULL, BEL, BS, VT, FF, etc.)
    # Seuls \t (tab), \n (newline), \r (carriage return) sont autorisés
    def sanitize_xml(texte):
        if not texte:
            return texte
        return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', texte)

    contenu = sanitize_xml(contenu)

    doc = Document()

    # Supprimer le paragraphe vide créé automatiquement par python-docx
    if doc.paragraphs:
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # Neutraliser le start_type NEW_PAGE de la 1re section
    from docx.enum.section import WD_SECTION
    doc.sections[0].start_type = WD_SECTION.CONTINUOUS

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ══ MODE EXAMEN IVOIRIEN ══════════════════════════════════════════
    IS_EXAMEN = "Examens" in service or "Sujets" in service
    IS_EXPOSE = "Expos" in service
    IS_CV     = "CV" in service or "Lettre de Motivation" in service

    # ══════════════════════════════════════════════════════════════
    # MODE CV — Template deux colonnes (fidèle au modèle SIRIKY)
    # DISPOSITION :
    #   EN-TÊTE : fond bleu foncé, pleine largeur
    #     → NOM en grand blanc Arial centré
    #     → Titre professionnel bleu clair Arial centré
    #     → Contacts bleu pâle Arial centré
    #   CORPS : 2 colonnes côte à côte
    #     GAUCHE (~31%) fond bleu foncé :
    #       FORMATION → COMPÉTENCES → LANGUES → SPORTS & LOISIRS
    #       police Times New Roman, texte BLANC
    #     DROITE (~69%) fond blanc :
    #       PROFIL → EXPÉRIENCES → CENTRES D'INTÉRÊT → INFOS PERSO
    #       police Arial, titres de sections bleu souligné,
    #       titres de missions bleu gras, texte NOIR
    # ══════════════════════════════════════════════════════════════
    if IS_CV:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Inches

        # ── Couleurs (identiques au modèle JS SIRIKY) ──────────────
        BLEU_FONCE = "1A3A5C"   # fond colonne gauche + en-tête
        BLEU_TITRE = "0070C0"   # titres de missions colonne droite
        BLEU_SEC   = "1565A8"   # titres de sections colonne droite
        BLANC      = "FFFFFF"
        NOIR       = "1A1A2E"
        BLEU_CLAIR = "A8D4F5"   # sous-titre en-tête
        BLEU_PALE  = "D0E8FF"   # contacts en-tête

        def hex_to_rgb(h):
            h = h.lstrip("#")
            return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

        def set_cell_bg(cell, hex_color):
            hex_color = hex_color.lstrip("#")
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hex_color)
            tcPr.append(shd)

        def remove_cell_borders(cell):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ["top","left","bottom","right","insideH","insideV"]:
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "nil")
                tcBorders.append(border)
            tcPr.append(tcBorders)

        # ── GAUCHE : Times New Roman, blanc ────────────────────────
        def g_heading(cell, text):
            """Titre de section colonne gauche : TNR 12pt blanc gras souligné"""
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after  = Pt(6)
            run = p.add_run(text.upper())
            run.font.name  = "Times New Roman"
            run.font.size  = Pt(12)
            run.bold       = True
            run.underline  = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        def g_text(cell, text, bold=False):
            """Texte simple colonne gauche : TNR 10pt blanc"""
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(text)
            run.font.name  = "Times New Roman"
            run.font.size  = Pt(10)
            run.bold       = bold
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        def g_bullet(cell, text):
            """Bullet colonne gauche : TNR 10pt blanc avec • manuel"""
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Pt(10)
            r1 = p.add_run("• ")
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(10)
            r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r2 = p.add_run(text)
            r2.font.name = "Times New Roman"
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        def g_space(cell):
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)

        # ── DROITE : Arial, noir + bleu pour missions ───────────────
        def d_heading(cell, text):
            """Titre de section colonne droite : Arial 12pt bleu gras souligné"""
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after  = Pt(6)
            run = p.add_run(text.upper())
            run.font.name  = "Arial"
            run.font.size  = Pt(12)
            run.bold       = True
            run.underline  = True
            r, g, b = hex_to_rgb(BLEU_SEC)
            run.font.color.rgb = RGBColor(r, g, b)

        def d_text(cell, text, bold=False, size=11):
            """Texte normal colonne droite : Arial 11pt noir justifié"""
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(text)
            run.font.name  = "Arial"
            run.font.size  = Pt(size)
            run.bold       = bold
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        def d_bullet(cell, text):
            """Bullet colonne droite : Arial 10pt noir avec • manuel"""
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Pt(10)
            r1 = p.add_run("• ")
            r1.font.name = "Arial"
            r1.font.size = Pt(10)
            r1.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            r2 = p.add_run(text)
            r2.font.name = "Arial"
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        def d_mission(cell, titre, explication=""):
            """Titre mission BLEU gras + explication NOIR sur même ligne"""
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r_b = p.add_run(titre + " : ")
            r_b.font.name  = "Arial"
            r_b.font.size  = Pt(10)
            r_b.bold       = True
            r_b.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
            if explication:
                r_n = p.add_run(explication)
                r_n.font.name  = "Arial"
                r_n.font.size  = Pt(10)
                r_n.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        def d_sous_titre(cell, text):
            """Sous-titre de poste : Arial 11pt bleu gras souligné"""
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            run = p.add_run(text)
            run.font.name  = "Arial"
            run.font.size  = Pt(11)
            run.bold       = True
            run.underline  = True
            r, g, b = hex_to_rgb(BLEU_SEC)
            run.font.color.rgb = RGBColor(r, g, b)

        def d_space(cell):
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)

        # ── PARSER le contenu Gemini ────────────────────────────────
        import re as _re
        lignes = contenu.split("\n")
        sections = {}
        section_courante = None
        buffer_lignes = []

        SECTION_KEYS = {
            "INFORMATIONS PERSONNELLES COMPLÉMENTAIRES": "infos_compl",
            "INFORMATIONS PERSONNELLES":                 "infos",
            "TITRE PROFESSIONNEL":                       "titre",
            "PROFIL":                                    "profil",
            "EXPÉRIENCES":                               "experiences",
            "EXPÉRIENCE":                                "experiences",
            "FORMATION":                                 "formation",
            "COMPÉTENCES":                               "competences",
            "LANGUES":                                   "langues",
            "SPORTS":                                    "sports",
            "LOISIRS":                                   "sports",
            "CENTRES D'INTÉRÊT":                         "interets",
            "CENTRES D'INTERET":                         "interets",
            "LETTRE DE MOTIVATION":                      "lettre",
        }

        for ligne in lignes:
            l_clean = ligne.strip().lstrip("#").strip().upper()
            matched = None
            for key, val in SECTION_KEYS.items():
                if key in l_clean:
                    matched = val
                    break
            if matched:
                if section_courante and buffer_lignes:
                    sections[section_courante] = "\n".join(buffer_lignes).strip()
                section_courante = matched
                buffer_lignes = []
            elif section_courante:
                t = ligne.strip().lstrip("#").strip()
                if t:
                    buffer_lignes.append(t)

        if section_courante and buffer_lignes:
            sections[section_courante] = "\n".join(buffer_lignes).strip()

        # ── Extraire nom + contacts depuis INFORMATIONS PERSONNELLES ─
        infos_raw  = sections.get("infos", "")
        profil_raw = sections.get("profil", "")
        titre_raw  = sections.get("titre", "").strip().lstrip("-•").strip()
        nom_cv     = client_nom or "Candidat"
        contact_email = ""
        contact_tel   = ""
        contact_ville = ""

        for line in infos_raw.split("\n"):
            l = line.strip().lstrip("-•").strip()
            if any(k in l.lower() for k in ["nom", "prénom", "prenom"]):
                nom_cv = l.split(":")[-1].strip() if ":" in l else l
            if any(k in l.lower() for k in ["email", "mail"]):
                contact_email = l.split(":")[-1].strip() if ":" in l else l
            if any(k in l.lower() for k in ["tél", "tel", "téléphone", "telephone"]):
                contact_tel = l.split(":")[-1].strip() if ":" in l else l
            if any(k in l.lower() for k in ["ville", "adresse", "résidence", "residence", "pays"]):
                contact_ville = l.split(":")[-1].strip() if ":" in l else l

        parts_contact = []
        if contact_email: parts_contact.append(f"✉  {contact_email}")
        if contact_tel:   parts_contact.append(f"☎  {contact_tel}")
        if contact_ville: parts_contact.append(f"📍  {contact_ville}")
        contact_cv = "   │   ".join(parts_contact)

        # ══════════════════════════════════════════════════════════
        # CONSTRUCTION DU DOCUMENT
        # Dimensions IDENTIQUES au modèle JS SIRIKY :
        #   Page A4      : 11906 × 16838 DXA, marges 0 sauf bas 400 DXA
        #   En-tête      : largeur 11400 DXA = Inches(7.92)
        #                  marges internes : top=200 bottom=160 left=400 right=400 DXA
        #   COL_G        : 3600 DXA = Inches(2.50)
        #                  marges : top=200 bottom=200 left=280 right=220 DXA
        #   COL_D        : 7800 DXA = Inches(5.42)
        #                  marges : top=200 bottom=200 left=320 right=200 DXA
        #   NOM          : size 52 demi-pts = Pt(26), bold, blanc, centré
        #   Titre pro    : size 22 demi-pts = Pt(11), bleu clair A8D4F5
        #   Contacts     : size 19 demi-pts = Pt(9.5), bleu pâle D0E8FF
        # Conversion DXA→Cm : 1 DXA = 0.0353 cm
        # ══════════════════════════════════════════════════════════
        doc.sections[0].top_margin    = Cm(0)
        doc.sections[0].bottom_margin = Cm(400 * 0.0353)  # 400 DXA ≈ 1.41 cm
        doc.sections[0].left_margin   = Cm(0)
        doc.sections[0].right_margin  = Cm(0)

        # Largeurs (1 DXA = 914.4 EMU, python-docx utilise des EMU pour .width)
        # On passe par Inches car 1 inch = 1440 DXA
        PAGE_W  = Inches(11400 / 1440)  # 11400 DXA = Inches(7.917)
        COL_G_W = Inches(3600  / 1440)  # 3600  DXA = Inches(2.500)
        COL_D_W = Inches(7800  / 1440)  # 7800  DXA = Inches(5.417)

        # Marges internes des cellules en EMU (python-docx OxmlElement tcMar)
        # 1 DXA = 914.4 EMU  → on stocke en DXA pour les XML
        # python-docx n'expose pas les marges de cellule via API,
        # on les injecte en XML directement
        def set_cell_margins(cell, top, bottom, left, right):
            """Marges internes de cellule en DXA (comme JS margins en DXA/20)."""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement("w:tcMar")
            for side, val in [("top", top), ("bottom", bottom),
                               ("left", left), ("right", right)]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:w"), str(val))
                el.set(qn("w:type"), "dxa")
                tcMar.append(el)
            tcPr.append(tcMar)

        # ════════════════════════════════════════════════════════
        # 1. EN-TÊTE PLEINE LARGEUR (fond bleu foncé)
        #    NOM → TITRE PRO → CONTACTS
        #    Marges internes : top=200 bottom=160 left=400 right=400 DXA
        # ════════════════════════════════════════════════════════
        tbl_header = doc.add_table(rows=1, cols=1)
        tbl_header.style = "Table Grid"
        cell_h = tbl_header.cell(0, 0)
        set_cell_bg(cell_h, BLEU_FONCE)
        remove_cell_borders(cell_h)
        cell_h.width = PAGE_W
        set_cell_margins(cell_h, top=200, bottom=160, left=400, right=400)

        # Vider le paragraphe vide par défaut
        for p in cell_h.paragraphs:
            for run in p.runs:
                run.clear()

        # NOM : size 52 demi-pts = Pt(26), spacing after=60 DXA→Pt(3)
        p_nom = cell_h.paragraphs[0]
        p_nom.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_nom.paragraph_format.space_before = Pt(0)
        p_nom.paragraph_format.space_after  = Pt(3)
        r_nom = p_nom.add_run(nom_cv.upper())
        r_nom.font.name = "Arial"
        r_nom.font.size = Pt(26)   # = size 52 demi-pts
        r_nom.bold = True
        r_nom.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # TITRE PROFESSIONNEL : size 22 demi-pts = Pt(11), bleu clair A8D4F5, spacing after=80
        if titre_raw:
            p_titre = cell_h.add_paragraph()
            p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_titre.paragraph_format.space_before = Pt(0)
            p_titre.paragraph_format.space_after  = Pt(4)
            r_titre = p_titre.add_run(titre_raw)
            r_titre.font.name = "Arial"
            r_titre.font.size = Pt(11)   # = size 22 demi-pts
            r, g, b = hex_to_rgb(BLEU_CLAIR)   # A8D4F5
            r_titre.font.color.rgb = RGBColor(r, g, b)

        # CONTACTS : size 19 demi-pts = Pt(9.5), bleu pâle D0E8FF
        if contact_cv:
            p_contact = cell_h.add_paragraph()
            p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_contact.paragraph_format.space_before = Pt(0)
            p_contact.paragraph_format.space_after  = Pt(0)
            r_contact = p_contact.add_run(contact_cv)
            r_contact.font.name = "Arial"
            r_contact.font.size = Pt(9.5)   # = size 19 demi-pts
            r, g, b = hex_to_rgb(BLEU_PALE)   # D0E8FF
            r_contact.font.color.rgb = RGBColor(r, g, b)

        # ════════════════════════════════════════════════════════
        # 2. CORPS : tableau 2 colonnes
        #    COL_G : 3600 DXA, marges top=200 bot=200 left=280 right=220
        #    COL_D : 7800 DXA, marges top=200 bot=200 left=320 right=200
        # ════════════════════════════════════════════════════════
        tbl_body = doc.add_table(rows=1, cols=2)
        tbl_body.style = "Table Grid"

        col_g = tbl_body.cell(0, 0)
        col_d = tbl_body.cell(0, 1)
        col_g.width = COL_G_W   # 3600 DXA = Inches(2.500)
        col_d.width = COL_D_W   # 7800 DXA = Inches(5.417)

        set_cell_bg(col_g, BLEU_FONCE)
        remove_cell_borders(col_g)
        remove_cell_borders(col_d)

        # Marges internes identiques au modèle JS SIRIKY :
        # Gauche : margins { top:200, bottom:200, left:280, right:220 }
        # Droite : margins { top:200, bottom:200, left:320, right:200 }
        set_cell_margins(col_g, top=200, bottom=200, left=280, right=220)
        set_cell_margins(col_d, top=200, bottom=200, left=320, right=200)

        # Vider les paragraphes vides par défaut
        for p in col_g.paragraphs: p._element.getparent().remove(p._element)
        for p in col_d.paragraphs: p._element.getparent().remove(p._element)

        # ════════════════════════════════════════════════════════
        # COLONNE GAUCHE
        # Ordre : FORMATION → COMPÉTENCES → LANGUES → SPORTS & LOISIRS
        # Times New Roman, blanc sur fond bleu foncé
        # ════════════════════════════════════════════════════════

        # FORMATION
        if sections.get("formation"):
            g_heading(col_g, "Formation")
            prev_bold_line = ""
            for line in sections["formation"].split("\n"):
                l = line.strip().lstrip("-•").strip()
                if not l:
                    g_space(col_g)
                    continue
                # Première ligne d'un bloc = diplôme (gras), suivante = détail
                # On détecte si c'est un diplôme (contient un tiret long ou année)
                if _re.search(r"(BTS|DUT|Licence|Master|BAC|BEPC|BT |CEPE|Doctorat|Ingénieur|CAP|BEP)", l, _re.I):
                    if prev_bold_line:
                        g_space(col_g)
                    g_text(col_g, l, bold=True)
                    prev_bold_line = l
                else:
                    g_text(col_g, l)

        # COMPÉTENCES
        if sections.get("competences"):
            g_heading(col_g, "Compétences")
            for line in sections["competences"].split("\n"):
                l = line.strip().lstrip("-•").strip()
                if l:
                    g_bullet(col_g, l)

        # LANGUES
        if sections.get("langues"):
            g_heading(col_g, "Langues")
            for line in sections["langues"].split("\n"):
                l = line.strip().lstrip("-•").strip()
                if l:
                    g_bullet(col_g, l)

        # SPORTS & LOISIRS
        if sections.get("sports"):
            g_heading(col_g, "Sports & Loisirs")
            for line in sections["sports"].split("\n"):
                l = line.strip().lstrip("-•").strip()
                if l:
                    g_bullet(col_g, l)

        # ════════════════════════════════════════════════════════
        # COLONNE DROITE
        # Ordre : PROFIL → EXPÉRIENCES → CENTRES D'INTÉRÊT → INFOS PERSO
        # Arial, noir. Titres sections = bleu souligné.
        # Titres missions (>>>BLEU<<<) = bleu gras + explication noire.
        # ════════════════════════════════════════════════════════

        # PROFIL PROFESSIONNEL
        if profil_raw:
            d_heading(col_d, "Profil Professionnel")
            for line in profil_raw.split("\n"):
                l = line.strip().lstrip("-•#").strip()
                if l:
                    d_text(col_d, l)

        # EXPÉRIENCES PROFESSIONNELLES
        if sections.get("experiences"):
            d_heading(col_d, "Expériences Professionnelles")
            for line in sections["experiences"].split("\n"):
                l = line.strip()
                if not l:
                    continue

                if l.startswith("###"):
                    # Sous-titre de poste : Arial 11pt bleu gras souligné
                    d_sous_titre(col_d, l.lstrip("#").strip())

                elif l.startswith(">>>BLEU<<<"):
                    # Titre mission bleu + explication noire
                    contenu_bleu = l.replace(">>>BLEU<<<", "").strip()
                    if ":" in contenu_bleu:
                        titre_b, expl = contenu_bleu.split(":", 1)
                        d_mission(col_d, titre_b.strip(), expl.strip())
                    else:
                        d_mission(col_d, contenu_bleu)

                elif l.lower().startswith("missions"):
                    # Label "Missions principales :" en noir gras
                    d_text(col_d, l.rstrip(":") + " :", bold=True)

                elif l.startswith(("-", "•")):
                    d_bullet(col_d, l.lstrip("-•").strip())

                else:
                    # Texte de contexte (entreprise, période...)
                    # Nettoyer les ** résiduels
                    l_clean = l.strip("*").strip()
                    if l_clean:
                        d_text(col_d, l_clean)

        # CENTRES D'INTÉRÊT
        if sections.get("interets"):
            d_space(col_d)
            d_heading(col_d, "Centres d'Intérêt")
            for line in sections["interets"].split("\n"):
                l = line.strip().lstrip("-•").strip()
                if not l:
                    continue
                if l.startswith(">>>BLEU<<<"):
                    contenu_bleu = l.replace(">>>BLEU<<<", "").strip()
                    if ":" in contenu_bleu:
                        titre_b, expl = contenu_bleu.split(":", 1)
                        d_mission(col_d, titre_b.strip(), expl.strip())
                    else:
                        d_mission(col_d, contenu_bleu)
                else:
                    d_bullet(col_d, l)

        # INFORMATIONS PERSONNELLES (date de naissance, situation familiale, résidence, téléphone)
        # On utilise en priorité infos_compl (section dédiée), sinon on filtre infos
        infos_compl_raw = sections.get("infos_compl", "")
        source_infos = infos_compl_raw if infos_compl_raw else infos_raw
        CHAMPS_PERSO = ["situation familiale", "résidence", "residence", "téléphone", "telephone", "date de naissance"]
        CHAMPS_EXCLUS = ["nom", "prénom", "prenom", "email", "mail", "ville", "pays", "adresse"]

        lignes_perso = []
        for line in source_infos.split("\n"):
            l = line.strip().lstrip("-•").strip()
            if not l:
                continue
            l_low = l.lower()
            # Inclure si c'est un des 4 champs voulus
            if any(k in l_low for k in CHAMPS_PERSO):
                lignes_perso.append(l)
            # Exclure les champs de l'en-tête
            elif not any(k in l_low for k in CHAMPS_EXCLUS) and infos_compl_raw:
                # Si section dédiée, inclure tout ce qui n'est pas explicitement exclu
                lignes_perso.append(l)

        if lignes_perso:
            d_space(col_d)
            d_heading(col_d, "Informations Personnelles")
            for l in lignes_perso:
                d_bullet(col_d, l)

        # ════════════════════════════════════════════════════════
        # LETTRE DE MOTIVATION (page séparée si présente)
        # ════════════════════════════════════════════════════════
        if sections.get("lettre"):
            doc.add_page_break()
            # Remettre les marges normales pour la lettre
            doc.sections[0].left_margin  = Cm(2.5)
            doc.sections[0].right_margin = Cm(2.5)
            p_tl = doc.add_paragraph()
            p_tl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_tl.paragraph_format.space_after = Pt(20)
            r_tl = p_tl.add_run("LETTRE DE MOTIVATION")
            r_tl.font.name = "Arial"
            r_tl.font.size = Pt(16)
            r_tl.bold = True
            r, g, b = hex_to_rgb(BLEU_SEC)
            r_tl.font.color.rgb = RGBColor(r, g, b)

            for line in sections["lettre"].split("\n"):
                l = line.strip()
                if not l:
                    doc.add_paragraph("")
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                run = p.add_run(l.lstrip("#-•*").strip())
                run.font.name = "Arial"
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    # ══════════════════════════════════════════════════════════════
    # FIN MODE CV
    # ══════════════════════════════════════════════════════════════

    if IS_EXAMEN:
        for section in doc.sections:
            section.top_margin    = Cm(1.8)
            section.bottom_margin = Cm(1.8)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.0)

    # ── NUMÉROTATION X/Y EN BAS DE PAGE (exposés + examens) ─────
    if IS_EXPOSE or IS_EXAMEN:
        def _add_page_numbers(section, is_exam=False):
            from docx.oxml import OxmlElement as _OE
            from docx.oxml.ns import qn as _qn
            from docx.shared import RGBColor as _RC, Pt as _Pt
            footer = section.footer
            footer.is_linked_to_previous = False
            for p in footer.paragraphs:
                p._element.getparent().remove(p._element)
            p_f = footer.add_paragraph()
            p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_f.paragraph_format.space_before = _Pt(0)
            p_f.paragraph_format.space_after  = _Pt(0)
            # Couleur : doré pour exposé, noir pour examen
            _color = _RC(0x00, 0x00, 0x00) if is_exam else _RC(0xB8, 0x93, 0x2A)
            def _fld(instr):
                r = p_f.add_run()
                r.font.name = "Times New Roman" if is_exam else "Calibri"
                r.font.size = _Pt(10 if is_exam else 9)
                r.italic = is_exam
                r.font.color.rgb = _color
                fldChar_begin = _OE("w:fldChar")
                fldChar_begin.set(_qn("w:fldCharType"), "begin")
                r._r.append(fldChar_begin)
                r2 = p_f.add_run()
                r2.font.name = "Times New Roman" if is_exam else "Calibri"
                r2.font.size = _Pt(10 if is_exam else 9)
                r2.italic = is_exam
                r2.font.color.rgb = _color
                instrText = _OE("w:instrText")
                instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                instrText.text = f" {instr} "
                r2._r.append(instrText)
                r3 = p_f.add_run()
                r3.font.name = "Times New Roman" if is_exam else "Calibri"
                r3.font.size = _Pt(10 if is_exam else 9)
                r3.italic = is_exam
                r3.font.color.rgb = _color
                fldChar_end = _OE("w:fldChar")
                fldChar_end.set(_qn("w:fldCharType"), "end")
                r3._r.append(fldChar_end)
            _fld("PAGE")
            sep = p_f.add_run("/")
            sep.font.name = "Times New Roman" if is_exam else "Calibri"
            sep.font.size = _Pt(10 if is_exam else 9)
            sep.italic = is_exam
            sep.font.color.rgb = _color
            _fld("NUMPAGES")

        for _sect in doc.sections:
            _add_page_numbers(_sect, is_exam=IS_EXAMEN)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman" if IS_EXAMEN else "Arial"
    style.font.size = Pt(11)
    if IS_EXPOSE:
        # Tout via XML pour garantir l'application dans Word
        from docx.oxml import OxmlElement as _OEnorm
        from docx.oxml.ns import qn as _qnnorm
        pPr_norm = style.element.get_or_add_pPr()
        # Justification
        jc = _OEnorm("w:jc")
        jc.set(_qnnorm("w:val"), "both")
        for old in pPr_norm.findall(_qnnorm("w:jc")): pPr_norm.remove(old)
        pPr_norm.append(jc)
        # Alinéa première ligne 1.25cm = 709 twips
        ind = _OEnorm("w:ind")
        ind.set(_qnnorm("w:firstLine"), "709")
        for old in pPr_norm.findall(_qnnorm("w:ind")): pPr_norm.remove(old)
        pPr_norm.append(ind)
        # Interligne 1.5 (360 = 1.5 × 240)
        sp_norm = _OEnorm("w:spacing")
        sp_norm.set(_qnnorm("w:before"), "0")
        sp_norm.set(_qnnorm("w:after"), "80")
        sp_norm.set(_qnnorm("w:line"), "360")
        sp_norm.set(_qnnorm("w:lineRule"), "auto")
        for old in pPr_norm.findall(_qnnorm("w:spacing")): pPr_norm.remove(old)
        pPr_norm.append(sp_norm)
    if IS_EXAMEN:
        from docx.oxml import OxmlElement as _OEnorm
        from docx.oxml.ns import qn as _qnnorm
        pPr_norm = style.element.get_or_add_pPr()
        sp_norm = _OEnorm("w:spacing")
        sp_norm.set(_qnnorm("w:before"), "0")
        sp_norm.set(_qnnorm("w:after"), "40")
        sp_norm.set(_qnnorm("w:line"), "253")
        sp_norm.set(_qnnorm("w:lineRule"), "auto")
        pPr_norm.append(sp_norm)

    # ── CORRECTION DÉFINITIVE DES STYLES HEADING ──────────────────
    # Les styles Heading1/2/3/4 ont keepNext + keepLines + spacing before=480
    # qui font que Word insère une page quasi-blanche après chaque saut de page.
    # On les corrige tous à la source.
    def fix_heading_style(style_name, font_size, color_rgb):
        try:
            st = doc.styles[style_name]
            st.font.name  = "Arial"
            st.font.size  = Pt(font_size)
            st.font.bold  = True
            st.font.color.rgb = RC(*color_rgb)
            pPr = st.element.get_or_add_pPr()
            # Supprimer keepNext (cause principale de la page blanche)
            for child in list(pPr):
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                if tag in ('keepNext', 'keepLines', 'pageBreakBefore'):
                    pPr.remove(child)
            # Forcer spacing before=0 after=6
            from docx.oxml import OxmlElement as _OE2
            from docx.oxml.ns import qn as _qn2
            for child in list(pPr):
                if child.tag.endswith('}spacing') or child.tag == 'spacing':
                    pPr.remove(child)
            spacing = _OE2("w:spacing")
            spacing.set(_qn2("w:before"), "0")
            spacing.set(_qn2("w:after"),  "60")
            pPr.append(spacing)
        except Exception:
            pass

    if IS_EXAMEN:
        # Vrais sujets CI : Headings en Times New Roman noir gras souligné
        def fix_heading_examen(style_name, font_size, underline=False):
            try:
                st = doc.styles[style_name]
                st.font.name  = "Times New Roman"
                st.font.size  = Pt(font_size)
                st.font.bold  = True
                st.font.color.rgb = RC(0x00, 0x00, 0x00)  # noir pur
                st.font.underline = underline
                pPr = st.element.get_or_add_pPr()
                for child in list(pPr):
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag in ('keepNext', 'keepLines', 'pageBreakBefore'):
                        pPr.remove(child)
                from docx.oxml import OxmlElement as _OEh
                from docx.oxml.ns import qn as _qnh
                for child in list(pPr):
                    if child.tag.endswith('}spacing') or child.tag == 'spacing':
                        pPr.remove(child)
                spacing = _OEh("w:spacing")
                spacing.set(_qnh("w:before"), "280")   # 14pt avant exercice (vrai sujet CI)
                spacing.set(_qnh("w:after"),  "60")    # 3pt après titre
                pPr.append(spacing)
            except Exception:
                pass
        fix_heading_examen("Heading 1", 14, underline=False)  # EXERCICE N° — gras seul, pas souligné
        fix_heading_examen("Heading 2", 12, underline=False)  # Sous-exercice — gras seul
        fix_heading_examen("Heading 3", 11, underline=False)  # Partie A/B/C
        fix_heading_examen("Heading 4", 11, underline=False)  # Document A/B
    else:
        fix_heading_style("Heading 1", 16, (0x1F, 0x4E, 0x79))
        fix_heading_style("Heading 2", 14, (0x2E, 0x75, 0xB6))
        fix_heading_style("Heading 3", 12, (0x1F, 0x4E, 0x79))
        fix_heading_style("Heading 4", 11, (0x40, 0x40, 0x40))

    from docx.oxml import OxmlElement
    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    from docx.shared import RGBColor as RC
    # En-tête service/client supprimé — page de garde gérée séparément pour exposés

    def add_formatted_para(doc, text, style_name="Normal", bold=False, size=11, color=None, align=None):
        p = doc.add_paragraph(style=style_name)
        if align:
            p.alignment = align
        # Anti-lignes-orphelines
        pPr = p._p.get_or_add_pPr()
        from docx.oxml import OxmlElement as _OEp
        from docx.oxml.ns import qn as _qnp
        wCtrl = _OEp("w:widowControl")
        wCtrl.set(_qnp("w:val"), "1")
        pPr.append(wCtrl)
        # Utilise le moteur de formules (superscript, subscript, Unicode)
        ajouter_formule_dans_run(p, text, bold=bold, size=size, color=color)
        return p

    import re as _re

    # ══════════════════════════════════════════════════════════════
    # ══════════════════════════════════════════════════════════════
    # MOTEUR DE FORMULES NOVA v3 — LaTeX complet, chimie, physique,
    # maths supérieures, vecteurs, intégrales, nucléaire, OMML-like
    # ══════════════════════════════════════════════════════════════

    import re as _re

    # ── TABLE DE CONVERSION LaTeX → notation Nova ─────────────────
    # Règle fondamentale : ^{...} et _{...} TOUJOURS préservés avec
    # leurs accolades pour que le parser Word les lise correctement.
    LATEX_TO_NOVA = [
        # ── Fractions (toutes variantes) ──
        (_re.compile(r'\\(?:d|t|text)?frac\{([^}]+)\}\{([^}]+)\}'),   r'(\1)/(\2)'),
        (_re.compile(r'\\cfrac\{([^}]+)\}\{([^}]+)\}'),               r'(\1)/(\2)'),
        (_re.compile(r'\\sfrac\{([^}]+)\}\{([^}]+)\}'),               r'\1/\2'),
        # ── Racines ──
        (_re.compile(r'\\sqrt\[([^\]]+)\]\{([^}]+)\}'),  r'(\2)^{1/\1}'),
        (_re.compile(r'\\sqrt\{([^}]+)\}'),               r'√(\1)'),
        (_re.compile(r'\\sqrt'),                           r'√'),
        # ── Exposants/indices LaTeX → format Nova (accolades préservées) ──
        (_re.compile(r'\^\{([^}]+)\}'),  lambda m: '^{' + m.group(1) + '}'),
        (_re.compile(r'_\{([^}]+)\}'),   lambda m: '_{' + m.group(1) + '}'),
        # ── Notation nucléaire : ^{A}_{Z}X → ^A_Z X ──
        (_re.compile(r'\^\{(\d+)\}_\{(\d+)\}([A-Za-z]+)'), r'^{\1}_{\2}\3'),
        # ── Valeur absolue et norme ──
        (_re.compile(r'\\left\s*\|([^|]+)\\right\s*\|'),   r'|\1|'),
        (_re.compile(r'\\left\s*\\Vert([^V]+)\\right\s*\\Vert'), r'||\1||'),
        (_re.compile(r'\\norm\{([^}]+)\}'),                 r'||\1||'),
        (_re.compile(r'\\abs\{([^}]+)\}'),                  r'|\1|'),
        # ── Combinaisons / Binôme ──
        (_re.compile(r'\\binom\{([^}]+)\}\{([^}]+)\}'),    r'C(\1,\2)'),
        (_re.compile(r'\\dbinom\{([^}]+)\}\{([^}]+)\}'),   r'C(\1,\2)'),
        (_re.compile(r'\\tbinom\{([^}]+)\}\{([^}]+)\}'),   r'C(\1,\2)'),
        # ── Vecteurs et dérivées ──
        (_re.compile(r'\\overrightarrow\{([^}]+)\}'),  r'\1⃗'),
        (_re.compile(r'\\overleftarrow\{([^}]+)\}'),   r'\1⃖'),
        (_re.compile(r'\\vec\{([^}]+)\}'),             r'\1⃗'),
        (_re.compile(r'\\hat\{([^}]+)\}'),             r'\1̂'),
        (_re.compile(r'\\bar\{([^}]+)\}'),             r'\1̄'),
        (_re.compile(r'\\tilde\{([^}]+)\}'),           r'\1̃'),
        (_re.compile(r'\\dot\{([^}]+)\}'),             r'\1̇'),
        (_re.compile(r'\\ddot\{([^}]+)\}'),            r'\1̈'),
        (_re.compile(r'\\overline\{([^}]+)\}'),        r'\1̄'),
        (_re.compile(r'\\underline\{([^}]+)\}'),       r'\1'),
        (_re.compile(r'\\widehat\{([^}]+)\}'),         r'\1̂'),
        (_re.compile(r'\\widetilde\{([^}]+)\}'),       r'\1̃'),
        # ── Crochets et accolades ──
        (_re.compile(r'\\lfloor'),  '⌊'), (_re.compile(r'\\rfloor'), '⌋'),
        (_re.compile(r'\\lceil'),   '⌈'), (_re.compile(r'\\rceil'),  '⌉'),
        (_re.compile(r'\\langle'),  '⟨'), (_re.compile(r'\\rangle'), '⟩'),
        (_re.compile(r'\\{'),       '{'), (_re.compile(r'\\}'),       '}'),
        # ── Lettres grecques — TOUTES les 24 lettres (minuscules et majuscules) ──
        (_re.compile(r'\\Alpha'),    'Α'), (_re.compile(r'\\alpha'),      'α'),
        (_re.compile(r'\\Beta'),     'Β'), (_re.compile(r'\\beta'),       'β'),
        (_re.compile(r'\\Gamma'),    'Γ'), (_re.compile(r'\\gamma'),      'γ'),
        (_re.compile(r'\\Delta'),    'Δ'), (_re.compile(r'\\delta'),      'δ'),
        (_re.compile(r'\\Epsilon'),  'Ε'), (_re.compile(r'\\(?:var)?epsilon'), 'ε'),
        (_re.compile(r'\\Zeta'),     'Ζ'), (_re.compile(r'\\zeta'),       'ζ'),
        (_re.compile(r'\\Eta'),      'Η'), (_re.compile(r'\\eta'),        'η'),
        (_re.compile(r'\\Theta'),    'Θ'), (_re.compile(r'\\(?:var)?theta'), 'θ'),
        (_re.compile(r'\\Iota'),     'Ι'), (_re.compile(r'\\iota'),       'ι'),
        (_re.compile(r'\\Kappa'),    'Κ'), (_re.compile(r'\\(?:var)?kappa'), 'κ'),
        (_re.compile(r'\\Lambda'),   'Λ'), (_re.compile(r'\\lambda'),     'λ'),
        (_re.compile(r'\\Mu'),       'Μ'), (_re.compile(r'\\mu'),         'μ'),
        (_re.compile(r'\\Nu'),       'Ν'), (_re.compile(r'\\nu'),         'ν'),
        (_re.compile(r'\\Xi'),       'Ξ'), (_re.compile(r'\\xi'),         'ξ'),
        (_re.compile(r'\\Omicron'),  'Ο'), (_re.compile(r'\\omicron'),    'ο'),
        (_re.compile(r'\\Pi'),       'Π'), (_re.compile(r'\\pi'),         'π'),
        (_re.compile(r'\\varpi'),    'ϖ'),
        (_re.compile(r'\\Rho'),      'Ρ'), (_re.compile(r'\\(?:var)?rho'), 'ρ'),
        (_re.compile(r'\\Sigma'),    'Σ'), (_re.compile(r'\\sigma'),      'σ'),
        (_re.compile(r'\\varsigma'), 'ς'),
        (_re.compile(r'\\Tau'),      'Τ'), (_re.compile(r'\\tau'),        'τ'),
        (_re.compile(r'\\Upsilon'),  'Υ'), (_re.compile(r'\\upsilon'),    'υ'),
        (_re.compile(r'\\Phi'),      'Φ'), (_re.compile(r'\\(?:var)?phi'), 'φ'),
        (_re.compile(r'\\Chi'),      'Χ'), (_re.compile(r'\\chi'),        'χ'),
        (_re.compile(r'\\Psi'),      'Ψ'), (_re.compile(r'\\psi'),        'ψ'),
        (_re.compile(r'\\Omega'),    'Ω'), (_re.compile(r'\\omega'),      'ω'),
        # ── Ensembles de nombres ──
        (_re.compile(r'\\mathbb\{R\}'), 'ℝ'), (_re.compile(r'\\mathbb\{N\}'), 'ℕ'),
        (_re.compile(r'\\mathbb\{Z\}'), 'ℤ'), (_re.compile(r'\\mathbb\{Q\}'), 'ℚ'),
        (_re.compile(r'\\mathbb\{C\}'), 'ℂ'), (_re.compile(r'\\mathbb\{P\}'), 'ℙ'),
        (_re.compile(r'\\mathbb\{([^}]+)\}'), r'\1'),  # autres \mathbb
        # ── Opérateurs arithmétiques et relations ──
        (_re.compile(r'\\times'),     '×'),  (_re.compile(r'\\cdot'),     '·'),
        (_re.compile(r'\\div'),       '÷'),  (_re.compile(r'\\pm'),       '±'),
        (_re.compile(r'\\mp'),        '∓'),  (_re.compile(r'\\ast'),      '*'),
        (_re.compile(r'\\star'),      '★'),  (_re.compile(r'\\circ'),     '∘'),
        (_re.compile(r'\\bullet'),    '•'),  (_re.compile(r'\\ldots'),     '…'),
        (_re.compile(r'\\cdots'),     '⋯'),  (_re.compile(r'\\vdots'),     '⋮'),
        (_re.compile(r'\\ddots'),     '⋱'),
        # ── Relations d'ordre ──
        (_re.compile(r'\\leq?'),      '≤'),  (_re.compile(r'\\geq?'),     '≥'),
        (_re.compile(r'\\ll'),        '≪'),  (_re.compile(r'\\gg'),       '≫'),
        (_re.compile(r'\\neq'),       '≠'),  (_re.compile(r'\\approx'),   '≈'),
        (_re.compile(r'\\equiv'),     '≡'),  (_re.compile(r'\\propto'),   '∝'),
        (_re.compile(r'\\sim'),       '~'),  (_re.compile(r'\\simeq'),    '≃'),
        (_re.compile(r'\\cong'),      '≅'),  (_re.compile(r'\\doteq'),    '≐'),
        (_re.compile(r'\\not\\in'),   '∉'),  (_re.compile(r'\\not\\subset'),'⊄'),
        (_re.compile(r'\\not='),      '≠'),  (_re.compile(r'\\not\\eq'),  '≠'),
        # ── Symboles maths avancés ──
        (_re.compile(r'\\infty'),     '∞'),
        (_re.compile(r'\\partial'),   '∂'),  (_re.compile(r'\\nabla'),    '∇'),
        (_re.compile(r'\\forall'),    '∀'),  (_re.compile(r'\\exists'),   '∃'),
        (_re.compile(r'\\nexists'),   '∄'),
        (_re.compile(r'\\emptyset'),  '∅'),  (_re.compile(r'\\varnothing'),'∅'),
        (_re.compile(r'\\aleph'),     'ℵ'),  (_re.compile(r'\\hbar'),     'ℏ'),
        (_re.compile(r'\\ell'),       'ℓ'),  (_re.compile(r'\\wp'),       '℘'),
        (_re.compile(r'\\Re'),        'ℜ'),  (_re.compile(r'\\Im'),       'ℑ'),
        # ── Intégrales et sommes ──
        (_re.compile(r'\\int'),       '∫'),  (_re.compile(r'\\iint'),     '∬'),
        (_re.compile(r'\\iiint'),     '∭'),  (_re.compile(r'\\oint'),     '∮'),
        (_re.compile(r'\\sum'),       'Σ'),  (_re.compile(r'\\prod'),     'Π'),
        (_re.compile(r'\\coprod'),    '∐'),
        # ── Logique et ensembles ──
        (_re.compile(r'\\in\b'),      '∈'),  (_re.compile(r'\\notin'),    '∉'),
        (_re.compile(r'\\subset'),    '⊂'),  (_re.compile(r'\\supset'),   '⊃'),
        (_re.compile(r'\\subseteq'),  '⊆'),  (_re.compile(r'\\supseteq'), '⊇'),
        (_re.compile(r'\\nsubset'),   '⊄'),  (_re.compile(r'\\nsupset'),  '⊅'),
        (_re.compile(r'\\cup'),       '∪'),  (_re.compile(r'\\cap'),      '∩'),
        (_re.compile(r'\\setminus'),  '∖'),  (_re.compile(r'\\complement'),'∁'),
        (_re.compile(r'\\land'),      '∧'),  (_re.compile(r'\\lor'),      '∨'),
        (_re.compile(r'\\lnot'),      '¬'),  (_re.compile(r'\\neg'),      '¬'),
        (_re.compile(r'\\oplus'),     '⊕'),  (_re.compile(r'\\otimes'),   '⊗'),
        (_re.compile(r'\\odot'),      '⊙'),
        # ── Géométrie ──
        (_re.compile(r'\\angle'),     '∠'),  (_re.compile(r'\\measuredangle'),'∡'),
        (_re.compile(r'\\perp'),      '⊥'),  (_re.compile(r'\\parallel'), '∥'),
        (_re.compile(r'\\triangle'),  '△'),  (_re.compile(r'\\square'),   '□'),
        (_re.compile(r'\\diamond'),   '◇'),  (_re.compile(r'\\circ'),     '°'),
        # ── Physique : unités et constantes ──
        (_re.compile(r'\\Omega\b'),   'Ω'),  # ohm (déjà dans grec, redondant mais sûr)
        (_re.compile(r'\\degree'),    '°'),
        (_re.compile(r'\\celsius'),   '°C'),
        # ── Chimie ──
        (_re.compile(r'\\rightleftharpoons'), '⇌'),  # équilibre chimique
        (_re.compile(r'\\longrightarrow'),    '⟶'),  # flèche réaction
        (_re.compile(r'\\xlongrightarrow\{([^}]+)\}'), r'—\1→'),
        # ── Flèches ──
        (_re.compile(r'\\Leftrightarrow'),  '⟺'),  (_re.compile(r'\\iff'), '⟺'),
        (_re.compile(r'\\Rightarrow'),      '⟹'),  (_re.compile(r'\\implies'),'⟹'),
        (_re.compile(r'\\Leftarrow'),       '⟸'),
        (_re.compile(r'\\rightarrow'),      '→'),  (_re.compile(r'\\to'), '→'),
        (_re.compile(r'\\leftarrow'),       '←'),  (_re.compile(r'\\gets'),'←'),
        (_re.compile(r'\\leftrightarrow'),  '↔'),
        (_re.compile(r'\\uparrow'),         '↑'),  (_re.compile(r'\\downarrow'),'↓'),
        (_re.compile(r'\\Uparrow'),         '⇑'),  (_re.compile(r'\\Downarrow'),'⇓'),
        (_re.compile(r'\\nearrow'),         '↗'),  (_re.compile(r'\\searrow'),'↘'),
        (_re.compile(r'\\swarrow'),         '↙'),  (_re.compile(r'\\nwarrow'),'↖'),
        (_re.compile(r'\\mapsto'),          '↦'),
        (_re.compile(r'\\longmapsto'),      '⟼'),
        # ── Modulo et divisibilité ──
        (_re.compile(r'\\pmod\{([^}]+)\}'), r' (mod \1)'),
        (_re.compile(r'\\bmod\b'),          'mod'),
        (_re.compile(r'\\mod\b'),           'mod'),
        (_re.compile(r'\\gcd\b'),           'pgcd'),
        (_re.compile(r'\\lcm\b'),           'ppcm'),
        # ── Fonctions maths et physique ──
        (_re.compile(r'\\arcsin\b'),  'arcsin'), (_re.compile(r'\\arccos\b'), 'arccos'),
        (_re.compile(r'\\arctan\b'),  'arctan'), (_re.compile(r'\\arccot\b'), 'arccot'),
        (_re.compile(r'\\sin\b'),     'sin'),    (_re.compile(r'\\cos\b'),    'cos'),
        (_re.compile(r'\\tan\b'),     'tan'),    (_re.compile(r'\\cot\b'),    'cot'),
        (_re.compile(r'\\sec\b'),     'sec'),    (_re.compile(r'\\csc\b'),    'cosec'),
        (_re.compile(r'\\sinh\b'),    'sinh'),   (_re.compile(r'\\cosh\b'),   'cosh'),
        (_re.compile(r'\\tanh\b'),    'tanh'),   (_re.compile(r'\\coth\b'),   'coth'),
        (_re.compile(r'\\ln\b'),      'ln'),     (_re.compile(r'\\log\b'),    'log'),
        (_re.compile(r'\\exp\b'),     'exp'),    (_re.compile(r'\\lim\b'),    'lim'),
        (_re.compile(r'\\max\b'),     'max'),    (_re.compile(r'\\min\b'),    'min'),
        (_re.compile(r'\\inf\b'),     'inf'),    (_re.compile(r'\\sup\b'),    'sup'),
        (_re.compile(r'\\det\b'),     'det'),    (_re.compile(r'\\ker\b'),    'ker'),
        (_re.compile(r'\\dim\b'),     'dim'),    (_re.compile(r'\\rank\b'),   'rang'),
        (_re.compile(r'\\tr\b'),      'tr'),     (_re.compile(r'\\grad\b'),   'grad'),
        (_re.compile(r'\\div\b'),     'div'),    (_re.compile(r'\\rot\b'),    'rot'),
        (_re.compile(r'\\curl\b'),    'rot'),
        # ── Délimiteurs auto-sizing (ignorés, juste les parenthèses restent) ──
        (_re.compile(r'\\[Bb]ig[glr]?\s*[\(\[|<]'),  '('),
        (_re.compile(r'\\[Bb]ig[glr]?\s*[\)\]|>]'),  ')'),
        (_re.compile(r'\\left\s*[\(\[|<]'),   '('),
        (_re.compile(r'\\right\s*[\)\]|>]'),  ')'),
        (_re.compile(r'\\left\s*\\{'),        '{'),
        (_re.compile(r'\\right\s*\\}'),       '}'),
        (_re.compile(r'\\left\.'),            ''),
        (_re.compile(r'\\right\.'),           ''),
        # ── Espaces et mise en page LaTeX ──
        (_re.compile(r'\\[,;:!]'),     ' '),
        (_re.compile(r'\\quad'),       '  '),
        (_re.compile(r'\\qquad'),      '   '),
        (_re.compile(r'\\noindent'),   ''),
        (_re.compile(r'\\newline'),    '\n'),
        (_re.compile(r'\\\\'),         ' '),  # fin de ligne LaTeX
        # ── Environnements ──
        (_re.compile(r'\\begin\{[^}]+\}'),   ''),
        (_re.compile(r'\\end\{[^}]+\}'),     ''),
        (_re.compile(r'\\item\b'),           '• '),
        # ── Nettoyage commandes texte ──
        (_re.compile(r'\\(?:text|mathrm|mathbf|mathit|mathsf|mathcal|mathscr|mathfrak|mathbb)\{([^}]+)\}'), r'\1'),
        (_re.compile(r'\\(?:boldsymbol|pmb)\{([^}]+)\}'), r'\1'),
        (_re.compile(r'\\(?:underbrace|overbrace)\{([^}]+)\}(?:_\{[^}]+\})?'), r'\1'),
        (_re.compile(r'\\(?:stackrel|overset|underset)\{[^}]+\}\{([^}]+)\}'), r'\1'),
        (_re.compile(r'\\(?:color|textcolor)\{[^}]+\}\{([^}]+)\}'), r'\1'),
        (_re.compile(r'\\label\{[^}]+\}'),   ''),
        (_re.compile(r'\\tag\{[^}]+\}'),     ''),
        (_re.compile(r'\\ref\{[^}]+\}'),     '?'),
        # ── Retrait des backslashes restants isolés ──
        (_re.compile(r'\\([A-Za-z]+)'),  r'\1'),  # \commande → commande (dernier recours)
    ]

    def nettoyer_latex_complet(texte):
        """Convertit LaTeX complet + notation Nova ^{} _{} en texte normalisé."""

        # 0. Pré-nettoyage des formules malformées produites par Gemini
        # ex: __{} __{}  _{  }  {50 } {0,5 }
        texte = _re.sub(r'_{2,}\{\s*\}', '',  texte)   # __{}  ___{}  → supprimé
        texte = _re.sub(r'\^{2,}\{\s*\}', '', texte)   # ^^{}         → supprimé
        texte = _re.sub(r'[_^]\{\s*\}',   '',  texte)  # _{}   ^{}    → supprimé
        texte = _re.sub(r'_{2,}([^{])',  r'_\1', texte)  # __x   → _x
        texte = _re.sub(r'\^{2,}([^{])', r'^\1', texte)  # ^^x   → ^x
        # 1. Blocs $$...$$ et formules $...$
        def conv_dollar(m):
            f = m.group(1)
            for pat, repl in LATEX_TO_NOVA:
                try:
                    f = pat.sub(repl, f) if callable(repl) else pat.sub(repl, f)
                except Exception:
                    pass
            return f
        texte = _re.sub(r'\$\$([^$]+)\$\$', conv_dollar, texte)
        texte = _re.sub(r'\$([^$]+)\$',     conv_dollar, texte)
        # 2. Commandes LaTeX hors dollars
        for pat, repl in LATEX_TO_NOVA:
            try:
                texte = pat.sub(repl, texte) if not callable(repl) else pat.sub(repl, texte)
            except Exception:
                pass
        # 3. Nettoyer accolades orphelines (pas précédées de ^ ou _)
        result = []
        open_stack = []
        for c in texte:
            if c == '{':
                prev = result[-1] if result else ''
                if prev in ('^', '_'):
                    result.append(c)
                    open_stack.append(False)   # accolade utile
                else:
                    open_stack.append(True)    # accolade orpheline → ignorer
            elif c == '}':
                if open_stack and open_stack[-1]:
                    open_stack.pop()           # ferme orpheline, on supprime
                elif open_stack:
                    open_stack.pop()
                    result.append(c)
                else:
                    result.append(c)
            else:
                result.append(c)
        return ''.join(result)

    def ajouter_formule_dans_run(p, texte, bold=False, size=11, color=None):
        """
        Crée des runs Word avec vrais exposants/indices, symboles Unicode, gras.

        Syntaxe reconnue :
          ^{expr}       → exposant Word (superscript)
          ^x            → exposant 1 char
          _{expr}       → indice Word (subscript)
          _x            → indice 1 char
          **texte**     → gras
          √(expr)       → symbole racine + texte
          (a)/(b)       → fraction affichée avec barre / centrée
          Tous Unicode (α β γ Σ ∫ ∞ ≤ ≥ × → ⇌ ∂ ∇ …) passent directement

        Note : les fractions (a)/(b) issues de \frac sont rendues en ligne
        avec les parenthèses pour la lisibilité.
        """
        FORMULE_RE = _re.compile(
            r'(\*\*[^*]+\*\*)'                      # **gras**
            r'|\^{([^}]+)}'                          # exposant long  ^{abc}
            r'|\^([^\s{_^*/+\-=(),\[\]])'            # exposant court ^x (1 car non-espace)
            r'|_{([^}]+)}'                           # indice long    _{abc}
            r'|_([^\s{_^*/+\-=(),\[\]])'             # indice court   _x (1 car non-espace)
            r'|(√\([^)]+\))'                         # racine avec parenthèse √(expr)
            r'|(√[^\s+\-=×÷±≤≥≠,;:)\]}\n]{1,10})'  # racine courte √Δ √2 √200
        )

        def _run(text, sup=False, sub=False, bd=False, sz=None, math_font=False):
            text = sanitize_xml(text)
            if not text:
                return None
            r = p.add_run(text)
            # Police : Cambria Math pour exposants/indices/racines en mode examen
            # (meilleur rendu des symboles mathématiques dans Word)
            if IS_EXAMEN and (sup or sub or math_font):
                r.font.name = "Cambria Math"
            else:
                r.font.name = "Times New Roman" if IS_EXAMEN else "Arial"
            r.font.size  = Pt(sz if sz else (max(7, size - 2) if (sup or sub) else size))
            r.bold       = bd
            if sup: r.font.superscript = True
            if sub: r.font.subscript   = True
            if color: r.font.color.rgb = RC(*color)
            return r

        last = 0
        for m in FORMULE_RE.finditer(texte):
            if m.start() > last:
                _run(texte[last:m.start()], bd=bold)
            if m.group(1):                        # **gras**
                _run(m.group(1)[2:-2], bd=True)
            elif m.group(2) or m.group(3):        # exposant
                _run(m.group(2) or m.group(3), sup=True, bd=bold)
            elif m.group(4) or m.group(5):        # indice
                _run(m.group(4) or m.group(5), sub=True, bd=bold)
            elif m.group(6) or m.group(7):        # racine √(expr) ou √x
                _run(m.group(6) or m.group(7), bd=bold, math_font=True)
            last = m.end()
        if last < len(texte):
            _run(texte[last:], bd=bold)


    lignes = contenu.split("\n")
    i = 0
    sauts_de_page_count = 0  # Compteur de sauts de page pour détecter page garde + sommaire
    _last_heading_txt = ""   # Anti-doublon headings consécutifs identiques

    while i < len(lignes):
        l = lignes[i].rstrip()

        # ── SAUT DE PAGE NOVA — VRAI SAUT DE PAGE WORD ────────────
        if l.strip() == "---SAUT_DE_PAGE---":
            from docx.oxml.ns import qn as _qn
            from docx.oxml import OxmlElement as _OE
            sauts_de_page_count += 1
            p_break = doc.add_paragraph()
            p_break.paragraph_format.space_before = Pt(0)
            p_break.paragraph_format.space_after  = Pt(0)
            run_break = p_break.add_run()
            br = _OE("w:br")
            br.set(_qn("w:type"), "page")
            run_break._r.append(br)
            i += 1
            continue
        # ── MARQUEUR EN-TÊTE DEVOIR IVOIRIEN — Python construit tout ──────────
        # Syntaxe multiligne :
        #   ###ENTETE_DEVOIR###
        #   ETABLISSEMENT=Collège Sainte Famille de Bouaké
        #   DISCIPLINE=CE PHYSIQUE-CHIMIE
        #   CLASSE=2ndC
        #   ANNEE=2025-2026
        #   DATE=21/04/2026
        #   DUREE=2H
        #   TITRE_DEVOIR=DEVOIR DE NIVEAU N°1  3ᵉ Trimestre
        #   MENTION=Cette épreuve comporte trois (03) pages...
        #   ###FIN_ENTETE###
        if l.strip() == "###ENTETE_DEVOIR###":
            i += 1
            params_entete = {}
            while i < len(lignes) and lignes[i].strip() != "###FIN_ENTETE###":
                lp = lignes[i].strip()
                if "=" in lp:
                    cle, _, val = lp.partition("=")
                    params_entete[cle.strip()] = val.strip()
                i += 1
            i += 1  # sauter ###FIN_ENTETE###

            from docx.oxml import OxmlElement as _OEeh
            from docx.oxml.ns import qn as _qneh
            from docx.shared import Cm as _Cmeh

            # ── BLOC 1 : Tableau 2 colonnes invisible (gauche : établissement | droite : année/date/durée)
            tbl_hdr = doc.add_table(rows=1, cols=2)
            tbl_hdr.style = "Table Grid"
            # Supprimer toutes les bordures
            tbl_el = tbl_hdr._tbl
            tblPr = tbl_el.find(_qneh("w:tblPr"))
            if tblPr is None:
                tblPr = _OEeh("w:tblPr")
                tbl_el.insert(0, tblPr)
            tblBorders = _OEeh("w:tblBorders")
            for _side in ["top","left","bottom","right","insideH","insideV"]:
                _b = _OEeh(f"w:{_side}")
                _b.set(_qneh("w:val"), "single")
                _b.set(_qneh("w:sz"), "12")       # bordure 1.5pt — style vrai sujet CI
                _b.set(_qneh("w:space"), "0")
                _b.set(_qneh("w:color"), "000000")  # noir
                tblBorders.append(_b)
            tblPr.append(tblBorders)

            cell_g = tbl_hdr.cell(0, 0)
            cell_d = tbl_hdr.cell(0, 1)
            # Largeurs : 60% gauche / 40% droite
            for _c, _w in [(cell_g, _Cmeh(9.5)), (cell_d, _Cmeh(6.5))]:
                _c._tc.get_or_add_tcPr()
                tcW = _OEeh("w:tcW")
                tcW.set(_qneh("w:w"), str(int(_w.cm * 567)))
                tcW.set(_qneh("w:type"), "dxa")
                _c._tc.tcPr.append(tcW)

            # Colonne gauche : Établissement + Discipline + Classe
            def _eh_para(cell, txt, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
                p_e = cell.add_paragraph()
                p_e.alignment = align
                p_e.paragraph_format.space_before = Pt(0)
                p_e.paragraph_format.space_after  = Pt(1)
                r_e = p_e.add_run(txt)
                r_e.font.name = "Times New Roman"
                r_e.font.size = Pt(size)
                r_e.bold = bold
                r_e.font.color.rgb = RC(0, 0, 0)
                return p_e

            # Vider le paragraphe auto dans les cellules
            for _cell in [cell_g, cell_d]:
                for _pp in _cell.paragraphs:
                    _pp._element.getparent().remove(_pp._element)

            etab = params_entete.get("ETABLISSEMENT", "")
            disc = params_entete.get("DISCIPLINE", "")
            classe = params_entete.get("CLASSE", "")
            annee = params_entete.get("ANNEE", "")
            date_d = params_entete.get("DATE", "")
            duree = params_entete.get("DUREE", "")

            if etab:  _eh_para(cell_g, etab, bold=True, size=10)
            if disc:  _eh_para(cell_g, disc, bold=False, size=10)
            if classe: _eh_para(cell_g, f"CLASSE : {classe}", bold=False, size=10)

            if annee:  _eh_para(cell_d, f"ANNEE SCOLAIRE  {annee}", bold=False, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
            if date_d: _eh_para(cell_d, f"Date : {date_d}", bold=False, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
            if duree:  _eh_para(cell_d, f"Durée : {duree}", bold=False, size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

            doc.add_paragraph("")  # petit espace

            # ── BLOC 2 : Titre du devoir encadré centré ──────────────
            titre_dv = params_entete.get("TITRE_DEVOIR", "DEVOIR DE NIVEAU").upper()
            p_tdr = doc.add_paragraph()
            p_tdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_tdr.paragraph_format.space_before = Pt(6)
            p_tdr.paragraph_format.space_after  = Pt(4)
            pPr_tdr = p_tdr._p.get_or_add_pPr()
            pBdr_tdr = _OEeh("w:pBdr")
            for _sn in ["top","bottom","left","right"]:
                _se = _OEeh(f"w:{_sn}")
                _se.set(_qneh("w:val"), "single")
                _se.set(_qneh("w:sz"), "12")
                _se.set(_qneh("w:space"), "6")
                _se.set(_qneh("w:color"), "000000")
                pBdr_tdr.append(_se)
            pPr_tdr.append(pBdr_tdr)
            r_tdr = p_tdr.add_run(titre_dv)
            r_tdr.font.name = "Times New Roman"
            r_tdr.font.size = Pt(14)
            r_tdr.bold = True
            r_tdr.font.color.rgb = RC(0, 0, 0)

            # ── BLOC 3 : Mention réglementaire en italique centré ────
            mention = params_entete.get("MENTION", "")
            if mention:
                p_men = doc.add_paragraph()
                p_men.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_men.paragraph_format.space_before = Pt(4)
                p_men.paragraph_format.space_after  = Pt(6)
                r_men = p_men.add_run(mention)
                r_men.font.name = "Times New Roman"
                r_men.font.size = Pt(10)
                r_men.italic = True
                r_men.font.color.rgb = RC(0, 0, 0)

            continue

        # ── MARQUEUR SOUS-EXERCICE (CHIMIE / PHYSIQUE / etc.) ────────
        # Syntaxe : ###SOUS_EXERCICE### CHIMIE
        if l.strip().startswith("###SOUS_EXERCICE###"):
            texte_se = l.strip().replace("###SOUS_EXERCICE###", "").strip()
            from docx.oxml import OxmlElement as _OEse
            from docx.oxml.ns import qn as _qnse
            p_se = doc.add_paragraph()
            p_se.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_se.paragraph_format.space_before = Pt(12)  # espace généreux avant
            p_se.paragraph_format.space_after  = Pt(6)
            # Bordure basse fine pour séparer visuellement
            pPr_se = p_se._p.get_or_add_pPr()
            pBdr_se = _OEse("w:pBdr")
            bot_se = _OEse("w:bottom")
            bot_se.set(_qnse("w:val"), "single")
            bot_se.set(_qnse("w:sz"), "6")
            bot_se.set(_qnse("w:space"), "4")
            bot_se.set(_qnse("w:color"), "000000")
            pBdr_se.append(bot_se)
            pPr_se.append(pBdr_se)
            r_se = p_se.add_run(texte_se)
            r_se.font.name = "Times New Roman"
            r_se.font.size = Pt(12)
            r_se.bold = True
            r_se.underline = True
            r_se.font.color.rgb = RC(0, 0, 0)
            i += 1
            continue

        # ── MARQUEUR TITRE CADRE — Titre matière encadré (style vrai sujet CI) ──
        if l.strip().startswith("###TITRE_CADRE###"):
            texte_titre_c = l.strip().replace("###TITRE_CADRE###", "").strip().upper()
            from docx.oxml import OxmlElement as _OEtc
            from docx.oxml.ns import qn as _qntc
            p_tc = doc.add_paragraph()
            p_tc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_tc.paragraph_format.space_before = Pt(8)
            p_tc.paragraph_format.space_after  = Pt(4)
            # Bordure rectangulaire autour du paragraphe
            pPr_tc = p_tc._p.get_or_add_pPr()
            pBdr_tc = _OEtc("w:pBdr")
            for side_name in ["top", "bottom", "left", "right"]:
                side_el = _OEtc(f"w:{side_name}")
                side_el.set(_qntc("w:val"), "single")
                side_el.set(_qntc("w:sz"), "12")      # bordure 1.5pt
                side_el.set(_qntc("w:space"), "8")
                side_el.set(_qntc("w:color"), "000000")  # noir comme les vrais sujets
                pBdr_tc.append(side_el)
            pPr_tc.append(pBdr_tc)
            # Texte : Times New Roman 16pt gras majuscules noir
            run_tc = p_tc.add_run(texte_titre_c)
            run_tc.bold = True
            run_tc.font.name = "Times New Roman" if IS_EXAMEN else "Arial"
            run_tc.font.size = Pt(16)
            run_tc.font.color.rgb = RC(0x00, 0x00, 0x00)
            i += 1
            continue

        # ── MARQUEUR TITRE ROUGE — Grand titre centré rouge ─────
        if l.strip().startswith("###TITRE_ROUGE###"):
            texte_titre = l.strip().replace("###TITRE_ROUGE###", "").strip()
            p_rouge = doc.add_paragraph()
            p_rouge.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_rouge.paragraph_format.space_before = Pt(12)
            p_rouge.paragraph_format.space_after  = Pt(12)
            run_rouge = p_rouge.add_run(texte_titre)
            run_rouge.bold = True
            run_rouge.font.name = "Arial"
            run_rouge.font.size = Pt(28)
            run_rouge.font.color.rgb = RC(0xC0, 0x00, 0x00)  # Rouge vif
            i += 1
            continue

        # ── MARQUEUR ESPACE — Grand espace vertical ───────────────
        if l.strip() == "###ESPACE###":
            p_esp = doc.add_paragraph()
            p_esp.paragraph_format.space_before = Pt(0)
            p_esp.paragraph_format.space_after  = Pt(0)
            p_esp.paragraph_format.line_spacing = Pt(36)  # ~1.2cm d'espace
            i += 1
            continue

        # ── BLOC FORMULE NOVA — Formule centrée sur fond bleu clair ──
        # Syntaxe : ###FORMULE### F = m × a   ou   ###FORMULE### E = mc^{2}
        if l.strip().startswith("###FORMULE###"):
            texte_f = l.strip().replace("###FORMULE###", "").strip()
            texte_f = nettoyer_latex_complet(texte_f)
            from docx.oxml import OxmlElement as _OEf
            from docx.oxml.ns import qn as _qnf
            p_f = doc.add_paragraph()
            p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_f.paragraph_format.space_before = Pt(6)
            p_f.paragraph_format.space_after  = Pt(6)
            # Fond bleu très clair
            pPr_f = p_f._p.get_or_add_pPr()
            shd_f = _OEf("w:shd")
            shd_f.set(_qnf("w:val"), "clear")
            shd_f.set(_qnf("w:color"), "auto")
            shd_f.set(_qnf("w:fill"), "D6E4F0")  # bleu clair
            pPr_f.append(shd_f)
            # Bordure fine bleue autour
            pBdr_f = _OEf("w:pBdr")
            for side_name in ["top", "bottom", "left", "right"]:
                side_el = _OEf(f"w:{side_name}")
                side_el.set(_qnf("w:val"), "single")
                side_el.set(_qnf("w:sz"), "6")
                side_el.set(_qnf("w:space"), "4")
                side_el.set(_qnf("w:color"), "2E75B6")
                pBdr_f.append(side_el)
            pPr_f.append(pBdr_f)
            # Texte formule en Cambria Math 13pt gras (meilleur rendu symboles maths Word)
            ajouter_formule_dans_run(p_f, texte_f, bold=True, size=13,
                                     color=(0x1F, 0x4E, 0x79))
            i += 1
            continue

        # ── BLOC FORMULE MULTILIGNE — ###DEBUT_FORMULES### ... ###FIN_FORMULES###
        if l.strip() == "###DEBUT_FORMULES###":
            i += 1
            while i < len(lignes) and lignes[i].strip() != "###FIN_FORMULES###":
                lf = lignes[i].strip()
                if lf:
                    lf_clean = nettoyer_latex_complet(lf)
                    from docx.oxml import OxmlElement as _OEml
                    from docx.oxml.ns import qn as _qnml
                    p_ml = doc.add_paragraph()
                    p_ml.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_ml.paragraph_format.space_before = Pt(2)
                    p_ml.paragraph_format.space_after  = Pt(2)
                    pPr_ml = p_ml._p.get_or_add_pPr()
                    shd_ml = _OEml("w:shd")
                    shd_ml.set(_qnml("w:val"), "clear")
                    shd_ml.set(_qnml("w:color"), "auto")
                    shd_ml.set(_qnml("w:fill"), "EEF3FA")
                    pPr_ml.append(shd_ml)
                    ajouter_formule_dans_run(p_ml, lf_clean, bold=True, size=12,
                                             color=(0x1F, 0x4E, 0x79))
                i += 1
            i += 1  # sauter ###FIN_FORMULES###
            continue

        # ── LIGNES DE SÉPARATION ════ ET ──── ─────────────────────
        if l.strip().startswith("════") or l.strip().startswith("━━━━"):
            # Ne pas dessiner le trait si la prochaine ligne non-vide est un saut de page
            next_content = next((lignes[j].strip() for j in range(i+1, len(lignes)) if lignes[j].strip()), "")
            if next_content == "---SAUT_DE_PAGE---":
                i += 1
                continue  # Ignorer ce trait — il serait au bas de page et créerait un espace vide
            p_line = doc.add_paragraph()
            p_line.paragraph_format.space_before = Pt(4)
            p_line.paragraph_format.space_after  = Pt(4)
            p_line.paragraph_format.line_spacing = Pt(1)
            pPr2 = p_line._p.get_or_add_pPr()
            pBdr2 = OxmlElement("w:pBdr")
            bot2 = OxmlElement("w:bottom")
            bot2.set(qn("w:val"), "single")
            bot2.set(qn("w:sz"), "12" if IS_EXAMEN else "12")   # 1.5pt — visible sur vrais sujets CI
            bot2.set(qn("w:space"), "1")
            bot2.set(qn("w:color"), "000000" if IS_EXAMEN else "1F4E79")  # noir examen
            pBdr2.append(bot2)
            pPr2.append(pBdr2)
            i += 1
            continue

        if l.strip().startswith("────") or l.strip().startswith("----"):
            p_line = doc.add_paragraph()
            p_line.paragraph_format.space_before = Pt(3)
            p_line.paragraph_format.space_after  = Pt(3)
            p_line.paragraph_format.line_spacing = Pt(1)
            pPr2 = p_line._p.get_or_add_pPr()
            pBdr2 = OxmlElement("w:pBdr")
            bot2 = OxmlElement("w:bottom")
            bot2.set(qn("w:val"), "single")
            bot2.set(qn("w:sz"), "4")
            bot2.set(qn("w:space"), "1")
            bot2.set(qn("w:color"), "AAAAAA")
            pBdr2.append(bot2)
            pPr2.append(pBdr2)
            i += 1
            continue

        if l.strip() in ["---", "***", "___", "*"]:
            doc.add_paragraph("")
            i += 1
            continue

        if l.startswith("#### "):
            p = doc.add_heading(l[5:].strip(), level=4)
            i += 1
            continue
        if l.startswith("### "):
            _htxt = l[4:].strip()
            if _htxt.upper() != _last_heading_txt:
                p = doc.add_heading(_htxt, level=3)
                _last_heading_txt = _htxt.upper()
            i += 1
            continue
        if l.startswith("## "):
            _htxt = l[3:].strip()
            # Supprimer la mention "— Niveau : ..." du titre avant affichage
            if " — Niveau : " in _htxt:
                _htxt = _htxt.split(" — Niveau : ")[0].strip()
            if _htxt.upper() != _last_heading_txt:
                p = doc.add_heading(_htxt, level=2)
                _last_heading_txt = _htxt.upper()
            i += 1
            continue
        if l.startswith("# "):
            # Ignorer les lignes de commentaires de structure Nova
            if l.startswith("# ═") or l.startswith("# #") or l.startswith("# ─"):
                i += 1
                continue
            _htxt = l[2:].strip()
            if _htxt.upper() != _last_heading_txt:
                p = doc.add_heading(_htxt, level=1)
                _last_heading_txt = _htxt.upper()
            i += 1
            continue

        if l.startswith("|") and l.endswith("|"):
            table_lines = []
            while i < len(lignes) and lignes[i].strip().startswith("|") and lignes[i].strip().endswith("|"):
                row = lignes[i].strip()
                if not re.match(r"^[\|\s\-:]+$", row):
                    cells = [c.strip() for c in row.strip("|").split("|")]
                    table_lines.append(cells)
                i += 1

            if table_lines:
                from docx.shared import Inches
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement

                n_cols = max(len(r) for r in table_lines)
                # Largeurs adaptées selon mode et nombre de colonnes
                if IS_EXAMEN:
                    col_widths_map = {
                        2: [1.0, 13.0],                            # N° | Affirmation (VF simple 2 col)
                        3: [0.8, 9.5, 3.7],                        # N° | Affirmation | V/F groupé
                        4: [0.8, 7.5, 1.2, 1.2],                   # N° | Affirmation | V | F (sans justif)
                        5: [0.8, 8.5, 1.2, 1.2, 3.5],              # N° | Affirmation | V | F | Justification
                        6: [0.7, 5.5, 2.0, 2.0, 2.0, 2.0],         # N° | Énoncé | A | B | C | D (QCM 4 réponses)
                        7: [0.7, 5.0, 1.8, 1.8, 1.8, 1.8, 1.2],   # N° | Énoncé | A|B|C|D + points
                    }
                else:
                    col_widths_map = {
                        2: [3.0, 6.0],
                        3: [1.0, 7.5, 2.5],
                        4: [1.0, 5.0, 2.5, 2.5],
                        5: [0.8, 5.0, 1.5, 1.5, 1.2],
                    }
                col_widths = col_widths_map.get(n_cols, [14.0/n_cols]*n_cols if IS_EXAMEN else [9.0/n_cols]*n_cols)

                from docx.shared import Cm as DocxCm
                table = doc.add_table(rows=0, cols=n_cols)
                table.style = "Table Grid"

                for r_idx, row_data in enumerate(table_lines):
                    row_obj = table.add_row()
                    is_header = (r_idx == 0)
                    # ── HAUTEUR : auto pour les lignes de données (évite coupure de texte long)
                    # En-tête : hauteur minimale fixe ; données : hauteur automatique
                    from docx.oxml.ns import qn as _qn
                    from docx.oxml import OxmlElement as _OE
                    trPr = row_obj._tr.get_or_add_trPr()
                    trHeight = _OE("w:trHeight")
                    if is_header:
                        trHeight.set(_qn("w:val"), str(int(0.85 * 567)))
                        trHeight.set(_qn("w:hRule"), "atLeast")  # min 0.85cm, peut grandir
                    else:
                        trHeight.set(_qn("w:val"), str(int(1.1 * 567)))
                        trHeight.set(_qn("w:hRule"), "atLeast")  # hauteur minimale, auto si texte long
                    trPr.append(trHeight)

                    for c_idx, cell_text in enumerate(row_data):
                        cell = row_obj.cells[c_idx]
                        if c_idx < len(col_widths):
                            cell.width = DocxCm(col_widths[c_idx])
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcMar = _OE("w:tcMar")
                        for side in ["top","bottom","left","right"]:
                            m = _OE(f"w:{side}")
                            m.set(_qn("w:w"), "120")
                            m.set(_qn("w:type"), "dxa")
                            tcMar.append(m)
                        tcPr.append(tcMar)
                        # Couleurs en-tête selon mode
                        if IS_EXAMEN:
                            hdr_color = "2C3E50"   # Gris-bleu foncé comme les vrais sujets CI
                            alt_color  = "F5F5F5"  # Gris très clair pour lignes paires
                        else:
                            hdr_color = "1F4E79"
                            alt_color  = "EEF3FA"

                        if is_header:
                            shd = _OE("w:shd")
                            shd.set(_qn("w:val"), "clear")
                            shd.set(_qn("w:color"), "auto")
                            shd.set(_qn("w:fill"), hdr_color)
                            tcPr.append(shd)
                        elif r_idx % 2 == 0:
                            shd = _OE("w:shd")
                            shd.set(_qn("w:val"), "clear")
                            shd.set(_qn("w:color"), "auto")
                            shd.set(_qn("w:fill"), alt_color)
                            tcPr.append(shd)

                        para = cell.paragraphs[0]
                        # Centrage horizontal :
                        # - col 0 (N°) toujours centré
                        # - colonnes V/F/réponses courtes (c_idx >= 2 sauf col Justification) centrées
                        # - col Affirmation/Énoncé (c_idx == 1) toujours LEFT
                        if IS_EXAMEN:
                            # Tableau VF 5 col : col 2 (V), col 3 (F) centrés ; col 4 (Justif) LEFT
                            if n_cols == 5:
                                is_short_col = c_idx in [0, 2, 3]
                            else:
                                is_short_col = c_idx == 0 or (n_cols >= 4 and c_idx >= 2)
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_short_col else WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in [0, n_cols-1] else WD_ALIGN_PARAGRAPH.LEFT
                        para.paragraph_format.space_before = Pt(2)
                        para.paragraph_format.space_after  = Pt(2)

                        # Centrage vertical de toutes les cellules
                        from docx.oxml import OxmlElement as _OEv
                        from docx.oxml.ns import qn as _qnv
                        vAlign = _OEv("w:vAlign")
                        vAlign.set(_qnv("w:val"), "center")
                        tcPr.append(vAlign)

                        # Nettoyer le LaTeX dans les cellules
                        cell_text_clean = nettoyer_latex_complet(cell_text)

                        # Cases à cocher ☐ et □ : police Segoe UI Symbol pour affichage fiable
                        CASE_CHARS = {"☐", "□", "☑", "✓", "✗"}
                        if cell_text_clean.strip() in CASE_CHARS or all(c in CASE_CHARS | {" "} for c in cell_text_clean.strip()):
                            r_case = para.add_run(cell_text_clean.strip())
                            r_case.font.name = "Segoe UI Symbol"
                            r_case.font.size  = Pt(11)
                            r_case.bold       = False
                            if is_header:
                                r_case.font.color.rgb = RC(0xFF, 0xFF, 0xFF)
                        else:
                            ajouter_formule_dans_run(
                                para, cell_text_clean,
                                bold=is_header,
                                size=10 if IS_EXAMEN else 10,
                                color=(0xFF, 0xFF, 0xFF) if is_header else None
                            )

                doc.add_paragraph("")
            continue

        m_num = re.match(r"^(\d+)[.)]\s+(.*)", l)
        if m_num:
            p = doc.add_paragraph(style="List Number")
            parts = re.split(r"(\*\*[^*]+\*\*)", m_num.group(2))
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2]); run.bold = True
                else:
                    run = p.add_run(part.replace("*","").replace("`",""))
            for run in p.runs:
                run.font.name = "Arial"; run.font.size = Pt(11)
            i += 1
            continue

        if re.match(r"^[\-\*\•]\s+", l):
            texte = re.sub(r"^[\-\*\•]\s+", "", l)
            p = doc.add_paragraph(style="List Bullet")
            parts = re.split(r"(\*\*[^*]+\*\*)", texte)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2]); run.bold = True
                else:
                    run = p.add_run(part.replace("*","").replace("`",""))
            for run in p.runs:
                run.font.name = "Arial"; run.font.size = Pt(11)
            i += 1
            continue

        if not l.strip():
            p_vide = doc.add_paragraph()
            # Dans page de garde (avant 1er saut) et sommaire (avant 2e saut) : espacement réduit
            if sauts_de_page_count < 2:
                p_vide.paragraph_format.space_before = Pt(0)
                p_vide.paragraph_format.space_after  = Pt(0)
                p_vide.paragraph_format.line_spacing = Pt(6)
            else:
                # Corps : ligne vide réduite pour éviter les orphelines
                p_vide.paragraph_format.space_before = Pt(0)
                p_vide.paragraph_format.space_after  = Pt(0)
                p_vide.paragraph_format.line_spacing = Pt(8)
            i += 1
            continue

        if l.strip().startswith("**") and l.strip().endswith("**") and l.strip().count("**") == 2:
            texte = l.strip()[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(texte)
            run.bold = True
            run.font.name = "Arial"
            if sauts_de_page_count < 2 and IS_EXPOSE:
                # Dans le sommaire exposé → grands titres en doré
                run.font.size = Pt(11)
                run.font.color.rgb = RC(0xB8, 0x93, 0x2A)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(1)
            else:
                run.font.size = Pt(12)
                run.font.color.rgb = RC(0x1F, 0x4E, 0x79)
            i += 1
            continue

        p = add_formatted_para(doc, l.strip())
        # Mode compact pour page de garde et sommaire
        if sauts_de_page_count < 2:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
        elif IS_EXAMEN:
            # Mode examen : interligne serré comme les vrais sujets imprimés
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(3)
            from docx.shared import Pt as _Pt2
            p.paragraph_format.line_spacing = _Pt2(13.2)  # ~1.1 × 12pt = 13.2pt
        else:
            # Corps du document : espacement serré pour éviter les lignes orphelines
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(6)
        i += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def creer_xlsx(description, client_nom):
    """
    Génère un Excel dynamique basé sur le JSON retourné par Gemini.
    Si le JSON est invalide, fallback sur un template générique.
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import json, re

    # ── PALETTE DE COULEURS ────────────────────────────────────────
    COULEURS = {
        "bleu":   {"bg": "1F4E79", "fg": "FFFFFF"},
        "vert":   {"bg": "1E8449", "fg": "FFFFFF"},
        "orange": {"bg": "D35400", "fg": "FFFFFF"},
        "rouge":  {"bg": "C0392B", "fg": "FFFFFF"},
        "violet": {"bg": "7D3C98", "fg": "FFFFFF"},
        "gris":   {"bg": "5D6D7E", "fg": "FFFFFF"},
        "cyan":   {"bg": "117A65", "fg": "FFFFFF"},
        "or":     {"bg": "B7950B", "fg": "FFFFFF"},
    }
    BLEU_FONCE = "1F4E79"
    BLEU_MOY   = "2E75B6"
    BLEU_CLAIR = "D6E4F0"
    BLANC      = "FFFFFF"
    GRIS_CLAIR = "F2F2F2"
    GRIS_MED   = "E8E8E8"

    def hdr(cell, bg=BLEU_FONCE, fg=BLANC, bold=True, size=11):
        cell.font      = Font(bold=bold, color=fg, name="Arial", size=size)
        cell.fill      = PatternFill("solid", start_color=bg)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    def brd(cell, color="CCCCCC", style="thin"):
        s = Side(style=style, color=color)
        cell.border = Border(top=s, bottom=s, left=s, right=s)

    def brd_epais(cell):
        s_ext = Side(style="medium", color="1F4E79")
        s_int = Side(style="thin",   color="CCCCCC")
        cell.border = Border(top=s_ext, bottom=s_ext, left=s_ext, right=s_ext)

    def fmt_cell(cell, type_col, valeur=None):
        """Applique le format nombre/monnaie/date/pourcentage selon le type."""
        if type_col == "monnaie":
            cell.number_format = '#,##0 "FCFA"'
        elif type_col == "pourcentage":
            cell.number_format = "0.0%"
        elif type_col == "nombre":
            cell.number_format = "#,##0"
        elif type_col == "date":
            cell.number_format = "DD/MM/YYYY"

    # ── PARSE JSON GEMINI ──────────────────────────────────────────
    data = None
    try:
        # Nettoyage robuste : enlever tout ce qui entoure le JSON
        texte = description.strip()
        # Supprimer blocs ```json ... ``` ou ``` ... ```
        texte = re.sub(r"^```(?:json)?\s*", "", texte, flags=re.IGNORECASE)
        texte = re.sub(r"\s*```\s*$", "", texte)
        # Si Gemini a mis du texte avant le { ou après le }
        debut = texte.find("{")
        fin   = texte.rfind("}")
        if debut != -1 and fin != -1:
            texte = texte[debut:fin+1]
        texte = texte.strip()
        data = json.loads(texte)
    except Exception:
        data = None

    wb = Workbook()
    wb.remove(wb.active)  # Supprimer feuille vide par défaut

    if not data or "feuilles" not in data:
        # ── FALLBACK : template générique si JSON invalide ─────────
        ws = wb.create_sheet("Données")
        ws.sheet_view.showGridLines = False
        ws.merge_cells("A1:D1")
        ws["A1"].value = f"{client_nom} — Données"
        hdr(ws["A1"], size=13); ws.row_dimensions[1].height = 32
        ws["A2"].value = description[:200]
        ws["A2"].font = Font(italic=True, color="7F7F7F", name="Arial", size=10)
        buf = BytesIO(); wb.save(buf); buf.seek(0)
        return buf

    titre_classeur = data.get("titre", f"Classeur {client_nom}")

    # ── CONSTRUCTION DE CHAQUE FEUILLE ─────────────────────────────
    for feuille in data.get("feuilles", []):
        nom_feuille = feuille.get("nom", "Feuille")[:31]
        type_feuille = feuille.get("type", "saisie")
        colonnes = feuille.get("colonnes", [])
        lignes   = feuille.get("lignes_exemple", [])
        kpis     = feuille.get("kpis", [])

        ws = wb.create_sheet(nom_feuille)
        ws.sheet_view.showGridLines = False

        n_cols = max(len(colonnes), 4)
        last_col = get_column_letter(n_cols)

        # ── EN-TÊTE PRINCIPAL ──────────────────────────────────────
        ws.merge_cells(f"A1:{last_col}1")
        cell_titre = ws.cell(row=1, column=1, value=f"{titre_classeur}  |  {client_nom}")
        hdr(cell_titre, size=13); ws.row_dimensions[1].height = 36

        ws.merge_cells(f"A2:{last_col}2")
        cell_desc = ws.cell(row=2, column=1, value=f"{feuille.get('description', '')}  —  Généré le {datetime.now().strftime('%d/%m/%Y')}")
        cell_desc.font      = Font(italic=True, color="7F7F7F", name="Arial", size=10)
        cell_desc.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[2].height = 20

        # ── FEUILLE DE TYPE SAISIE ─────────────────────────────────
        if type_feuille == "saisie" and colonnes:
            # En-têtes colonnes
            for c_idx, col in enumerate(colonnes, 1):
                cell = ws.cell(row=3, column=c_idx, value=col["entete"])
                hdr(cell, bg=BLEU_MOY); brd(cell)
                ws.column_dimensions[get_column_letter(c_idx)].width = col.get("largeur", 18)
            ws.row_dimensions[3].height = 28

            # Lignes de données exemple
            for r_idx, ligne in enumerate(lignes, 4):
                bg = GRIS_CLAIR if r_idx % 2 == 0 else BLANC
                for c_idx, val in enumerate(ligne, 1):
                    if c_idx > len(colonnes):
                        break
                    cell = ws.cell(row=r_idx, column=c_idx, value=val)
                    type_col = colonnes[c_idx-1].get("type", "texte")
                    cell.font      = Font(name="Arial", size=10,
                                          bold=(type_col in ["monnaie","nombre"]),
                                          color=("1F4E79" if type_col == "monnaie" else "000000"))
                    cell.fill      = PatternFill("solid", start_color=bg)
                    cell.alignment = Alignment(vertical="center",
                                               horizontal="center" if type_col in ["monnaie","nombre","date","pourcentage"] else "left")
                    fmt_cell(cell, type_col)
                    brd(cell)

            # Ligne TOTAL
            if lignes:
                total_row = len(lignes) + 4
                ws.row_dimensions[total_row].height = 28
                # Trouver la 1re colonne numérique pour placer le label TOTAL
                first_num_col = next((c_idx for c_idx, col in enumerate(colonnes, 1)
                                      if col.get("type") in ["monnaie", "nombre"]), None)
                # Colonnes texte → label "TOTAL"
                for c_idx, col in enumerate(colonnes, 1):
                    cell_t = ws.cell(row=total_row, column=c_idx)
                    if col.get("type") not in ["monnaie", "nombre"]:
                        if c_idx == 1:
                            cell_t.value = "TOTAL"
                        hdr(cell_t, size=11); brd_epais(cell_t)
                    else:
                        col_letter = get_column_letter(c_idx)
                        cell_t.value = f"=SUM({col_letter}4:{col_letter}{total_row-1})"
                        fmt_cell(cell_t, col["type"])
                        hdr(cell_t, size=11); brd_epais(cell_t)

            # Figer la ligne d'en-tête
            ws.freeze_panes = "A4"

        # ── FEUILLE DE TYPE BILAN / KPIs ───────────────────────────
        elif type_feuille == "bilan" and kpis:
            # Titre section KPIs
            ws.merge_cells(f"A3:{last_col}3")
            cell_kpi_title = ws.cell(row=3, column=1, value="━━  INDICATEURS CLÉS DE PERFORMANCE  ━━")
            hdr(cell_kpi_title, bg=BLEU_FONCE, size=12); ws.row_dimensions[3].height = 30

            # Disposition des KPIs : 2 par ligne (label | valeur | espace | label | valeur)
            ws.column_dimensions["A"].width = 30
            ws.column_dimensions["B"].width = 24
            ws.column_dimensions["C"].width = 3
            ws.column_dimensions["D"].width = 30
            ws.column_dimensions["E"].width = 24

            row_kpi = 4
            for idx, kpi in enumerate(kpis):
                if idx % 2 == 0 and idx > 0:
                    row_kpi += 2  # Saut d'une ligne entre paires

                col_start = 1 if idx % 2 == 0 else 4
                couleur_key = kpi.get("couleur", "bleu")
                bg_kpi = COULEURS.get(couleur_key, COULEURS["bleu"])["bg"]
                fg_kpi = COULEURS.get(couleur_key, COULEURS["bleu"])["fg"]

                # Label
                cl = ws.cell(row=row_kpi, column=col_start, value=kpi["label"])
                cl.font      = Font(bold=True, name="Arial", size=11, color=fg_kpi)
                cl.fill      = PatternFill("solid", start_color=bg_kpi)
                cl.alignment = Alignment(horizontal="center", vertical="center")
                brd_epais(cl)
                ws.row_dimensions[row_kpi].height = 36

                # Valeur
                cv = ws.cell(row=row_kpi, column=col_start+1, value=kpi.get("formule", 0))
                cv.font      = Font(bold=True, name="Arial", size=14, color=fg_kpi)
                cv.fill      = PatternFill("solid", start_color=bg_kpi)
                cv.alignment = Alignment(horizontal="center", vertical="center")
                fmt_cell(cv, kpi.get("type", "nombre"))
                brd_epais(cv)

            # ── TABLEAU RÉCAPITULATIF sous les KPIs ───────────────
            row_recap = row_kpi + 3

            # Trouver la 1re feuille de saisie pour le récap
            feuille_saisie = next((f for f in data["feuilles"] if f.get("type") == "saisie"), None)
            if feuille_saisie:
                nom_s   = feuille_saisie["nom"][:31]
                cols_s  = feuille_saisie.get("colonnes", [])

                ws.merge_cells(f"A{row_recap}:E{row_recap}")
                cell_recap_title = ws.cell(row=row_recap, column=1, value=f"RÉCAPITULATIF — {nom_s.upper()}")
                hdr(cell_recap_title, bg=BLEU_MOY, size=12)
                ws.row_dimensions[row_recap].height = 28
                row_recap += 1

                # En-têtes récap
                recap_cols = [c["entete"] for c in cols_s[:5]]
                for c_idx, h in enumerate(recap_cols, 1):
                    cell = ws.cell(row=row_recap, column=c_idx, value=h)
                    hdr(cell, bg=BLEU_FONCE); brd(cell)
                ws.row_dimensions[row_recap].height = 24
                row_recap += 1

                # Lignes récap (depuis les données exemple)
                for r_idx, ligne in enumerate(feuille_saisie.get("lignes_exemple", [])[:8], row_recap):
                    bg = BLEU_CLAIR if r_idx % 2 == 0 else BLANC
                    for c_idx, val in enumerate(ligne[:5], 1):
                        if c_idx > len(cols_s):
                            break
                        cell = ws.cell(row=r_idx, column=c_idx, value=val)
                        type_col = cols_s[c_idx-1].get("type", "texte")
                        cell.font      = Font(name="Arial", size=10,
                                              color=("1F4E79" if type_col == "monnaie" else "000000"))
                        cell.fill      = PatternFill("solid", start_color=bg)
                        cell.alignment = Alignment(vertical="center",
                                                   horizontal="center" if type_col in ["monnaie","nombre","date"] else "left")
                        fmt_cell(cell, type_col)
                        brd(cell)

        # ── AUTRES TYPES DE FEUILLES (générique) ──────────────────
        else:
            if colonnes:
                for c_idx, col in enumerate(colonnes, 1):
                    cell = ws.cell(row=3, column=c_idx, value=col["entete"])
                    hdr(cell); brd(cell)
                    ws.column_dimensions[get_column_letter(c_idx)].width = col.get("largeur", 18)
                for r_idx, ligne in enumerate(lignes, 4):
                    for c_idx, val in enumerate(ligne, 1):
                        cell = ws.cell(row=r_idx, column=c_idx, value=val)
                        cell.font = Font(name="Arial", size=10)
                        cell.alignment = Alignment(vertical="center")
                        brd(cell)

    # ── MISE EN FORME FINALE : onglets colorés ─────────────────────
    couleurs_onglets = ["1F4E79", "1E8449", "D35400", "7D3C98"]
    for i, ws in enumerate(wb.worksheets):
        ws.sheet_properties.tabColor = couleurs_onglets[i % len(couleurs_onglets)]

    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return buf


WHATSAPP_NUMBER = st.secrets.get("WHATSAPP_NUMBER", "2250171542505")
PREMIUM_MSG = "J'aimerais passer à la version Nova Premium pour bénéficier de la puissance 10^10 et de l'IA de pointe."
SUPPORT_MSG = "Bonjour, j'ai besoin d'assistance sur mon espace Nova Platform."
whatsapp_premium_url = f"https://wa.me/{WHATSAPP_NUMBER}?text={PREMIUM_MSG.replace(' ', '%20')}"
whatsapp_support_url = f"https://wa.me/{WHATSAPP_NUMBER}?text={SUPPORT_MSG.replace(' ', '%20')}"

if "db" not in st.session_state:
    st.session_state["db"] = load_db()
if "current_user" not in st.session_state:
    st.session_state["current_user"] = None
if "view" not in st.session_state:
    st.session_state["view"] = "home"
if "is_glowing" not in st.session_state:
    st.session_state["is_glowing"] = False
if "show_premium_modal" not in st.session_state:
    st.session_state["show_premium_modal"] = False
if "nova_service_idx" not in st.session_state:
    st.session_state["nova_service_idx"] = 0
if "show_service_warning" not in st.session_state:
    st.session_state["show_service_warning"] = False
if "auto_reply_gratuit" not in st.session_state:
    st.session_state["auto_reply_gratuit"] = get_auto_reply_setting()
if "contenu_fichier_source" not in st.session_state:
    st.session_state["contenu_fichier_source"] = ""
if "last_service_seen" not in st.session_state:
    st.session_state["last_service_seen"] = None
if "warning_triggered" not in st.session_state:
    st.session_state["warning_triggered"] = False
if "ocr_conv_uses" not in st.session_state:
    st.session_state["ocr_conv_uses"] = 0
if "intro_played" not in st.session_state:
    st.session_state["intro_played"] = False
if "gemini_results" not in st.session_state:
    st.session_state["gemini_results"] = {}
if "premium_livrable" not in st.session_state:
    st.session_state["premium_livrable"] = None
if "show_mode_modal" not in st.session_state:
    st.session_state["show_mode_modal"] = False

if st.session_state["current_user"] is None:
    stored_user = st.query_params.get("user_id")
    if stored_user:
        # ── CAS 1 : uid trouvé dans le cache db ───────────────────
        if stored_user in st.session_state["db"]["users"]:
            st.session_state["current_user"] = stored_user
        else:
            # ── CAS 2 : uid présent dans l'URL mais pas dans le cache
            # (cache vide au démarrage ou load_db raté) →
            # lookup ciblé direct sur Supabase (1 seule ligne, ultra rapide)
            try:
                _rows = supabase.table("users").select("uid").eq("uid", stored_user).execute().data
                if _rows:
                    # Utilisateur confirmé en base → recharger db complète
                    st.session_state["db"] = load_db()
                    st.session_state["current_user"] = stored_user
                else:
                    # uid introuvable en base → URL corrompue ou expirée → on nettoie
                    st.query_params.clear()
                    components.html("""
                        <script>
                        localStorage.removeItem('nova_user_id');
                        localStorage.removeItem('nova_user_ts');
                        localStorage.removeItem('nova_user');
                        </script>
                    """, height=1)
            except Exception:
                # Supabase injoignable → on connecte quand même pour ne pas bloquer le client
                st.session_state["current_user"] = stored_user
    else:
        # ── CAS 3 : pas d'uid dans l'URL → lire localStorage et rediriger
        components.html("""
            <script>
            (function() {
                var uid = localStorage.getItem('nova_user_id') || localStorage.getItem('nova_user');
                if (uid) {
                    localStorage.setItem('nova_user_id', uid);
                    localStorage.removeItem('nova_user');
                    var ts = localStorage.getItem('nova_user_ts');
                    var TRENTE_JOURS = 30 * 24 * 60 * 60 * 1000;
                    if (!ts || (Date.now() - parseInt(ts)) < TRENTE_JOURS) {
                        var url = new URL(window.parent.location.href);
                        url.searchParams.set('user_id', uid);
                        window.parent.location.replace(url.toString());
                    } else {
                        localStorage.removeItem('nova_user_id');
                        localStorage.removeItem('nova_user_ts');
                    }
                }
            })();
            </script>
        """, height=1)

if st.session_state["current_user"]:
    uid_connecte = st.session_state["current_user"]
    components.html(f"""
        <script>
        (function() {{
            var stored = localStorage.getItem('nova_user_id');
            if (stored !== '{uid_connecte}') {{
                localStorage.setItem('nova_user_id', '{uid_connecte}');
                localStorage.setItem('nova_user_ts', Date.now().toString());
                localStorage.removeItem('nova_user');
            }}
        }})();
        </script>
    """, height=1)

def inject_custom_css():
    # ── Détection Premium ─────────────────────────────────────────
    _user = st.session_state.get("current_user")
    _db   = st.session_state.get("db", {})
    _ud   = _db.get("users", {}).get(_user, {}) if _user else {}
    _premium = is_premium_actif(_ud)

    # ── CSS commun (base) ─────────────────────────────────────────
    st.markdown("""
        <style>
        @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;600;800&display=swap');
        * { font-family: 'Poppins', sans-serif; }
        .stApp {
            background: #0f0c29;
            background: -webkit-linear-gradient(to right, #24243e, #302b63, #0f0c29);
            background: linear-gradient(to right, #24243e, #302b63, #0f0c29);
            color: #ffffff;
            transition: filter 0.5s ease;
        }
        @keyframes glow-pulse {
            0% { filter: brightness(1) saturate(1); box-shadow: inset 0 0 0px transparent; }
            50% { filter: brightness(1.8) saturate(1.5); box-shadow: inset 0 0 100px rgba(0, 210, 255, 0.5); }
            100% { filter: brightness(1) saturate(1); box-shadow: inset 0 0 0px transparent; }
        }
        .main-title {
            background: linear-gradient(90deg, #00d2ff, #3a7bd5);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            font-weight: 800;
            font-size: 3.5rem !important;
            text-align: center;
            margin-bottom: 20px;
            text-shadow: 0px 0px 20px rgba(0, 210, 255, 0.3);
        }
        .stTabs [data-baseweb="tab-list"] {
            gap: 20px;
            background-color: rgba(255, 255, 255, 0.05);
            padding: 10px;
            border-radius: 15px;
            border: 1px solid rgba(255, 255, 255, 0.1);
        }
        .stTabs [data-baseweb="tab"] {
            height: 60px;
            white-space: pre-wrap;
            background-color: rgba(0, 210, 255, 0.1);
            border-radius: 10px;
            color: white !important;
            font-weight: 700 !important;
            font-size: 1.2rem !important;
            transition: all 0.3s ease;
            border: 1px solid transparent;
            padding: 0 25px;
        }
        .stTabs [data-baseweb="tab"]:nth-child(2) {
            border: 1px solid #2ecc71 !important;
            box-shadow: 0 0 15px rgba(46, 204, 113, 0.2);
            background-color: rgba(46, 204, 113, 0.1);
        }
        .stTabs [data-baseweb="tab"]:hover {
            background-color: rgba(0, 210, 255, 0.3);
            transform: translateY(-2px);
        }
        .stTabs [aria-selected="true"] {
            background-color: rgba(0, 210, 255, 0.6) !important;
            border: 1px solid #00d2ff !important;
            box-shadow: 0 0 20px rgba(0, 210, 255, 0.4);
        }
        @keyframes border-rainbow {
            0% { border-color: #00d2ff; box-shadow: 0 0 10px rgba(0, 210, 255, 0.3); }
            25% { border-color: #3a7bd5; box-shadow: 0 0 10px rgba(58, 123, 213, 0.3); }
            50% { border-color: #FFD700; box-shadow: 0 0 15px rgba(255, 215, 0, 0.3); }
            75% { border-color: #2ecc71; box-shadow: 0 0 10px rgba(46, 204, 113, 0.3); }
            100% { border-color: #00d2ff; box-shadow: 0 0 10px rgba(0, 210, 255, 0.3); }
        }
        .stTextInput label, .stSelectbox label, .stTextArea label {
            color: #00d2ff !important;
            font-weight: 600 !important;
            font-size: 1.1rem !important;
            margin-bottom: 5px;
        }
        div[data-baseweb="input"], div[data-baseweb="select"] > div {
            border: 1px solid rgba(0, 210, 255, 0.3) !important;
            background-color: rgba(0, 0, 0, 0.5) !important;
            color: white !important;
            border-radius: 10px !important;
        }
        /* ── Dropdown liste déroulante — thème sombre ── */
        ul[data-baseweb="menu"] {
            background-color: #0d0d1a !important;
            border: 1px solid rgba(255,215,0,0.3) !important;
            border-radius: 12px !important;
            padding: 6px !important;
        }
        ul[data-baseweb="menu"] li {
            background-color: transparent !important;
            color: rgba(255,255,255,0.85) !important;
            border-radius: 8px !important;
            font-size: 0.88rem !important;
            padding: 8px 12px !important;
        }
        ul[data-baseweb="menu"] li:hover,
        ul[data-baseweb="menu"] li[aria-selected="true"] {
            background: linear-gradient(135deg, rgba(255,215,0,0.15), rgba(255,140,0,0.1)) !important;
            color: #FFD700 !important;
        }
        div[data-baseweb="popover"] > div {
            background-color: #0d0d1a !important;
            border-radius: 12px !important;
            box-shadow: 0 8px 32px rgba(0,0,0,0.7) !important;
        }
        .stTextArea textarea {
            background-color: rgba(0, 0, 0, 0.6) !important;
            color: white !important;
            border-radius: 10px !important;
            border: 2px solid #00d2ff !important;
            animation: border-rainbow 4s linear infinite;
            transition: transform 0.3s;
        }
        .stTextArea textarea:focus {
            transform: scale(1.01);
            animation: border-rainbow 1.5s linear infinite;
        }
        .logo-container {
            display: flex;
            justify-content: center;
            align-items: center;
            gap: 30px;
            margin-top: 20px;
            padding: 15px;
            background: rgba(255, 255, 255, 0.03);
            border-radius: 15px;
        }
        .logo-item {
            width: 45px;
            height: 45px;
            filter: grayscale(0.5) opacity(0.7);
            transition: all 0.3s ease;
        }
        .logo-item:hover {
            filter: grayscale(0) opacity(1);
            transform: translateY(-5px) scale(1.1);
        }
        .premium-card {
            background: rgba(20, 20, 30, 0.8);
            border: 2px solid #FFD700;
            border-radius: 20px;
            padding: 25px;
            text-align: center;
            margin-bottom: 30px;
            box-shadow: 0 0 30px rgba(255, 215, 0, 0.2);
            position: relative;
            overflow: hidden;
        }
        .premium-card::before {
            content: "";
            position: absolute;
            top: 0; left: 0; width: 100%; height: 5px;
            background: linear-gradient(90deg, #FFD700, #FF8C00, #FFD700);
        }
        .premium-title {
            color: #FFD700 !important;
            font-size: 1.5rem;
            font-weight: 800;
            text-transform: uppercase;
            margin-bottom: 10px;
            letter-spacing: 1px;
        }
        .premium-desc {
            color: #ffffff !important;
            font-size: 1rem;
            margin-bottom: 20px;
            line-height: 1.5;
        }
        .btn-gold {
            background: linear-gradient(45deg, #FFD700, #FF8C00);
            color: #000 !important;
            padding: 12px 30px;
            border-radius: 50px;
            text-decoration: none;
            font-weight: 800;
            font-size: 1.1rem;
            display: inline-block;
            box-shadow: 0 5px 15px rgba(255, 215, 0, 0.4);
            transition: transform 0.2s, box-shadow 0.2s;
            border: none;
            cursor: pointer;
        }
        .btn-gold:hover {
            transform: scale(1.05);
            box-shadow: 0 8px 25px rgba(255, 215, 0, 0.6);
        }
        .stButton>button {
            border-radius: 12px;
            padding: 0.8rem 2rem;
            background: linear-gradient(90deg, #00d2ff 0%, #3a7bd5 100%);
            border: none;
            color: white !important;
            font-weight: 700;
            font-size: 1.1rem;
            width: 100%;
            margin-top: 10px;
            box-shadow: 0 4px 10px rgba(0, 210, 255, 0.3);
            transition: 0.3s;
        }
        .stButton>button:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(0, 210, 255, 0.5);
        }
        .stFormSubmitButton>button {
            border-radius: 12px;
            padding: 0.8rem 2rem;
            background: linear-gradient(90deg, #00d2ff 0%, #3a7bd5 100%);
            border: none;
            color: white !important;
            font-weight: 700;
            font-size: 1.1rem;
            width: 100%;
            margin-top: 10px;
            box-shadow: 0 4px 10px rgba(0, 210, 255, 0.3);
            transition: 0.3s;
        }
        .stFormSubmitButton>button:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(0, 210, 255, 0.5);
        }
        .info-card {
            background: rgba(0, 0, 0, 0.4) !important;
            border-left: 4px solid #00d2ff;
            padding: 15px;
            border-radius: 0 10px 10px 0;
            margin-bottom: 15px;
            box-shadow: 2px 2px 10px rgba(0,0,0,0.3);
        }
        .info-title {
            color: #00d2ff !important;
            font-weight: bold;
            font-size: 1.1rem;
            display: block;
            margin-bottom: 8px;
            text-transform: uppercase;
        }
        .file-card {
            background: rgba(255, 255, 255, 0.08);
            border: 2px solid rgba(46, 204, 113, 0.5);
            border-radius: 15px;
            padding: 20px;
            margin-bottom: 15px;
            box-shadow: 0 4px 15px rgba(0, 0, 0, 0.3);
            animation: slideIn 0.5s ease;
        }
        @keyframes slideIn {
            from { opacity: 0; transform: translateY(10px); }
            to { opacity: 1; transform: translateY(0); }
        }
        .support-btn {
            display: block;
            text-decoration: none;
            background: transparent;
            border: 2px solid #25D366;
            color: #25D366 !important;
            padding: 10px;
            border-radius: 10px;
            font-weight: bold;
            text-align: center;
            margin-top: 10px;
            transition: 0.3s;
        }
        .support-btn:hover {
            background: #25D366;
            color: white !important;
        }
        .stProgress > div > div > div > div {
            background-image: linear-gradient(to right, #00d2ff , #3a7bd5);
        }
        .gemini-card {
            background: linear-gradient(135deg, rgba(0,210,255,0.08), rgba(58,123,213,0.12));
            border: 2px solid rgba(0,210,255,0.4);
            border-radius: 14px;
            padding: 16px 20px;
            margin: 12px 0;
        }
        .gemini-title {
            color: #00d2ff;
            font-weight: 800;
            font-size: 0.95rem;
            letter-spacing: 1px;
        }
        .gemini-sub {
            color: rgba(255,255,255,0.5);
            font-size: 0.75rem;
        }
        .badge-premium {
            display: inline-flex; align-items: center; gap: 6px;
            background: linear-gradient(135deg, #FFD700, #FF8C00);
            color: #000; font-weight: 800; font-size: 0.75rem;
            padding: 4px 12px; border-radius: 20px; text-transform: uppercase;
            box-shadow: 0 2px 10px rgba(255,215,0,0.4);
        }
        .badge-free {
            display: inline-flex; align-items: center; gap: 6px;
            background: rgba(255,255,255,0.1); color: rgba(255,255,255,0.5);
            font-size: 0.75rem; padding: 4px 12px; border-radius: 20px;
            border: 1px solid rgba(255,255,255,0.2);
        }
        .admin-premium-row {
            background: rgba(255,215,0,0.06); border: 1px solid rgba(255,215,0,0.25);
            border-radius: 14px; padding: 16px 20px; margin-bottom: 12px;
        }
        .admin-user-name { color: #fff; font-weight: 700; font-size: 1rem; }
        .admin-user-meta { color: rgba(255,255,255,0.45); font-size: 0.8rem; }
        .gemini-lock {
            background: linear-gradient(135deg, rgba(255,215,0,0.08), rgba(255,140,0,0.12));
            border: 2px solid rgba(255,215,0,0.4); border-radius: 14px; padding: 16px 20px; margin: 12px 0;
        }
        @keyframes nova-pulse {
            0%,100% { box-shadow: 0 0 20px rgba(255,215,0,0.4); }
            50%      { box-shadow: 0 0 60px rgba(255,215,0,0.9); }
        }
        .nova-processing {
            background: linear-gradient(135deg, rgba(255,215,0,0.1), rgba(255,140,0,0.08));
            border: 2px solid #FFD700; border-radius: 20px;
            padding: 30px; text-align: center; margin: 20px 0;
            animation: nova-pulse 2s ease-in-out infinite;
        }
        .nova-processing-title { color: #FFD700; font-size: 1.6rem; font-weight: 800; }
        /* ── Bouton "Voir tous les services" ── */
        @keyframes svc-shine {
            0%   { box-shadow: 0 0 6px rgba(255,215,0,0.5), 0 0 20px rgba(255,140,0,0.3); }
            50%  { box-shadow: 0 0 18px rgba(255,215,0,0.9), 0 0 40px rgba(255,140,0,0.6); }
            100% { box-shadow: 0 0 6px rgba(255,215,0,0.5), 0 0 20px rgba(255,140,0,0.3); }
        }
        button[data-testid="baseButton-secondary"]:has-text,
        [data-testid="stButton"] button[kind="secondary"] {
            /* ciblage général — surcharge par key ci-dessous */
        }
        .nova-processing-sub   { color: rgba(255,255,255,0.7); font-size: 1rem; margin-top: 8px; }
        .livrable-auto {
            background: linear-gradient(135deg, rgba(46,204,113,0.12), rgba(0,210,255,0.08));
            border: 2px solid #2ecc71; border-radius: 20px; padding: 28px; margin: 20px 0;
            box-shadow: 0 0 30px rgba(46,204,113,0.2);
        }
        .livrable-auto-title { color: #2ecc71; font-size: 1.4rem; font-weight: 800; }

        /* ══ BOUTONS TÉLÉCHARGER — Effet brillant lumineux ══ */
        @keyframes nova-dl-shine {
            0%   { background-position: -200% center;
                   box-shadow: 0 0 8px rgba(46,204,113,0.5), 0 0 20px rgba(0,210,255,0.25); }
            50%  { background-position: 0% center;
                   box-shadow: 0 0 22px rgba(46,204,113,0.9), 0 0 50px rgba(0,255,150,0.45); }
            100% { background-position: 200% center;
                   box-shadow: 0 0 8px rgba(46,204,113,0.5), 0 0 20px rgba(0,210,255,0.25); }
        }
        @keyframes nova-dl-text-glow {
            0%, 100% { text-shadow: 0 0 6px rgba(255,255,255,0.6); }
            50%       { text-shadow: 0 0 18px rgba(255,255,255,1), 0 0 30px rgba(100,255,180,0.8); }
        }
        /* Bouton download Streamlit natif */
        [data-testid="stDownloadButton"] > button,
        div[data-testid="stDownloadButton"] button {
            background: linear-gradient(90deg,
                #1a6b3a 0%, #2ecc71 25%, #00ffa3 50%, #2ecc71 75%, #1a6b3a 100%) !important;
            background-size: 300% auto !important;
            color: #ffffff !important;
            font-weight: 800 !important;
            font-size: 1rem !important;
            border: none !important;
            border-radius: 14px !important;
            padding: 0.75rem 1.8rem !important;
            width: 100% !important;
            cursor: pointer !important;
            animation: nova-dl-shine 2.5s ease-in-out infinite !important;
            transition: transform 0.2s ease !important;
            letter-spacing: 0.5px !important;
        }
        [data-testid="stDownloadButton"] > button span,
        div[data-testid="stDownloadButton"] button p,
        div[data-testid="stDownloadButton"] button span {
            color: #ffffff !important;
            font-weight: 800 !important;
            animation: nova-dl-text-glow 2.5s ease-in-out infinite !important;
        }
        [data-testid="stDownloadButton"] > button:hover,
        div[data-testid="stDownloadButton"] button:hover {
            transform: translateY(-3px) scale(1.02) !important;
            box-shadow: 0 0 35px rgba(46,204,113,1), 0 0 70px rgba(0,255,150,0.6) !important;
        }
        </style>
    """, unsafe_allow_html=True)

    # ── THÈME OR PREMIUM ──────────────────────────────────────────
    if _premium:
        st.markdown("""
        <style>
        @keyframes gold-shimmer {
            0%   { background-position: -300% center; }
            100% { background-position:  300% center; }
        }
        @keyframes gold-glow-pulse {
            0%,100% { box-shadow: inset 0 0 0px transparent; filter: brightness(1); }
            50%      { box-shadow: inset 0 0 120px rgba(255,215,0,0.18); filter: brightness(1.08); }
        }
        @keyframes gold-border-anim {
            0%   { border-color: #FFD700; box-shadow: 0 0 10px rgba(255,215,0,0.4); }
            50%  { border-color: #FF8C00; box-shadow: 0 0 20px rgba(255,140,0,0.5); }
            100% { border-color: #FFD700; box-shadow: 0 0 10px rgba(255,215,0,0.4); }
        }

        /* ── Fond général ── */
        .stApp {
            background: linear-gradient(135deg, #0a0800 0%, #1c1400 35%, #0d0a00 65%, #1a1000 100%) !important;
            animation: gold-glow-pulse 5s ease-in-out infinite !important;
        }

        /* ── Titre principal NOVA PLATFORM ── */
        .main-title {
            background: linear-gradient(90deg, #7a5500, #b8860b, #FFD700, #fff5c0, #FFD700, #b8860b, #7a5500) !important;
            background-size: 300% auto !important;
            -webkit-background-clip: text !important;
            -webkit-text-fill-color: transparent !important;
            animation: gold-shimmer 3s linear infinite !important;
            text-shadow: none !important;
            font-size: 3.8rem !important;
            letter-spacing: 2px !important;
        }

        /* ── Tabs ── */
        .stTabs [data-baseweb="tab-list"] {
            background-color: rgba(255,215,0,0.06) !important;
            border: 1px solid rgba(255,215,0,0.25) !important;
        }
        .stTabs [data-baseweb="tab"] {
            background-color: rgba(255,215,0,0.08) !important;
            border: 1px solid rgba(255,215,0,0.15) !important;
            color: #FFD700 !important;
        }
        .stTabs [data-baseweb="tab"]:nth-child(2) {
            border: 1px solid rgba(255,215,0,0.5) !important;
            box-shadow: 0 0 15px rgba(255,215,0,0.2) !important;
            background-color: rgba(255,215,0,0.1) !important;
        }
        .stTabs [data-baseweb="tab"]:hover {
            background-color: rgba(255,215,0,0.2) !important;
        }
        .stTabs [aria-selected="true"] {
            background: linear-gradient(135deg, rgba(255,215,0,0.35), rgba(255,140,0,0.25)) !important;
            border: 1px solid #FFD700 !important;
            box-shadow: 0 0 20px rgba(255,215,0,0.4) !important;
        }

        /* ── Labels formulaires ── */
        .stTextInput label, .stSelectbox label, .stTextArea label {
            color: #FFD700 !important;
        }

        /* ── Inputs / Selects ── */
        div[data-baseweb="input"], div[data-baseweb="select"] > div {
            border: 1px solid rgba(255,215,0,0.4) !important;
            background-color: rgba(20,12,0,0.7) !important;
        }
        .stTextArea textarea {
            background-color: rgba(20,12,0,0.8) !important;
            border: 2px solid #FFD700 !important;
            animation: gold-border-anim 3s ease-in-out infinite !important;
        }

        /* ── Boutons principaux ── */
        .stButton > button {
            background: linear-gradient(90deg, #7a5500, #b8860b, #FFD700, #b8860b, #7a5500) !important;
            background-size: 200% auto !important;
            color: #0a0800 !important;
            animation: gold-shimmer 3s linear infinite !important;
            box-shadow: 0 4px 18px rgba(255,215,0,0.35) !important;
            border: none !important;
        }
        .stButton > button:hover {
            box-shadow: 0 6px 28px rgba(255,215,0,0.6) !important;
            transform: translateY(-2px) !important;
        }

        /* ── Boutons de formulaire (Envoyer, Nouvelle conv...) ── */
        .stFormSubmitButton > button {
            background: linear-gradient(90deg, #7a5500, #b8860b, #FFD700, #b8860b, #7a5500) !important;
            background-size: 200% auto !important;
            color: #0a0800 !important;
            font-weight: 700 !important;
            animation: gold-shimmer 3s linear infinite !important;
            box-shadow: 0 4px 18px rgba(255,215,0,0.35) !important;
            border: none !important;
        }
        .stFormSubmitButton > button:hover {
            box-shadow: 0 6px 28px rgba(255,215,0,0.6) !important;
            transform: translateY(-2px) !important;
        }

        /* ── Info cards sidebar ── */
        .info-card {
            border-left: 4px solid #FFD700 !important;
            background: rgba(20,12,0,0.6) !important;
        }
        .info-title { color: #FFD700 !important; }

        /* ── Support btn ── */
        .support-btn {
            border: 2px solid #FFD700 !important;
            color: #FFD700 !important;
        }
        .support-btn:hover {
            background: #FFD700 !important;
            color: #000 !important;
        }

        /* ── Barre de progression ── */
        .stProgress > div > div > div > div {
            background-image: linear-gradient(to right, #b8860b, #FFD700) !important;
        }

        /* ── Gemini card ── */
        .gemini-card {
            background: linear-gradient(135deg, rgba(255,215,0,0.08), rgba(255,140,0,0.06)) !important;
            border: 2px solid rgba(255,215,0,0.5) !important;
        }
        .gemini-title { color: #FFD700 !important; }

        /* ── File cards ── */
        .file-card {
            border: 2px solid rgba(255,215,0,0.5) !important;
            background: rgba(20,12,0,0.5) !important;
        }

        /* ── Livrable auto ── */
        .livrable-auto {
            background: linear-gradient(135deg, rgba(255,215,0,0.12), rgba(255,140,0,0.08)) !important;
            border: 2px solid #FFD700 !important;
            box-shadow: 0 0 35px rgba(255,215,0,0.25) !important;
        }
        .livrable-auto-title { color: #FFD700 !important; }

        /* ── Sidebar ── */
        [data-testid="stSidebar"] {
            background: linear-gradient(180deg, #0d0900, #1a1200) !important;
            border-right: 1px solid rgba(255,215,0,0.2) !important;
        }

        /* ── Divider ── */
        hr { border-color: rgba(255,215,0,0.2) !important; }

        /* ── Métriques admin ── */
        [data-testid="stMetric"] {
            background: rgba(255,215,0,0.06) !important;
            border: 1px solid rgba(255,215,0,0.2) !important;
            border-radius: 12px !important;
            padding: 10px !important;
        }

        /* ── Logo container ── */
        .logo-container {
            background: rgba(255,215,0,0.04) !important;
            border: 1px solid rgba(255,215,0,0.12) !important;
        }

        /* ── Glow global sur le body ── */
        body::before {
            content: '';
            position: fixed;
            top: 0; left: 0; right: 0; bottom: 0;
            background: radial-gradient(ellipse at 50% 0%, rgba(255,215,0,0.06) 0%, transparent 60%);
            pointer-events: none;
            z-index: 0;
        }

        /* ══ BOUTONS TÉLÉCHARGER thème Premium — Effet OR lumineux ══ */
        @keyframes nova-dl-shine-gold {
            0%   { background-position: -200% center;
                   box-shadow: 0 0 10px rgba(255,215,0,0.5), 0 0 24px rgba(255,140,0,0.3); }
            50%  { background-position: 0% center;
                   box-shadow: 0 0 28px rgba(255,215,0,1), 0 0 60px rgba(255,180,0,0.6); }
            100% { background-position: 200% center;
                   box-shadow: 0 0 10px rgba(255,215,0,0.5), 0 0 24px rgba(255,140,0,0.3); }
        }
        @keyframes nova-dl-text-glow-gold {
            0%, 100% { text-shadow: 0 0 6px rgba(0,0,0,0.5); }
            50%       { text-shadow: 0 0 14px rgba(0,0,0,0.8), 0 0 28px rgba(255,215,0,0.3); }
        }
        [data-testid="stDownloadButton"] > button,
        div[data-testid="stDownloadButton"] button {
            background: linear-gradient(90deg,
                #7a5500 0%, #b8860b 20%, #FFD700 40%, #fff5c0 50%, #FFD700 60%, #b8860b 80%, #7a5500 100%) !important;
            background-size: 300% auto !important;
            color: #1a0f00 !important;
            font-weight: 900 !important;
            font-size: 1rem !important;
            border: 1px solid rgba(255,215,0,0.6) !important;
            border-radius: 14px !important;
            padding: 0.75rem 1.8rem !important;
            width: 100% !important;
            cursor: pointer !important;
            animation: nova-dl-shine-gold 2.5s ease-in-out infinite !important;
            transition: transform 0.2s ease !important;
            letter-spacing: 0.5px !important;
        }
        [data-testid="stDownloadButton"] > button span,
        div[data-testid="stDownloadButton"] button p,
        div[data-testid="stDownloadButton"] button span {
            color: #1a0f00 !important;
            font-weight: 900 !important;
            animation: nova-dl-text-glow-gold 2.5s ease-in-out infinite !important;
        }
        [data-testid="stDownloadButton"] > button:hover,
        div[data-testid="stDownloadButton"] button:hover {
            transform: translateY(-3px) scale(1.02) !important;
            box-shadow: 0 0 40px rgba(255,215,0,1), 0 0 80px rgba(255,180,0,0.7) !important;
        }
        </style>
        """, unsafe_allow_html=True)

    # ── THÈME OR EXCLUSIF PREMIUM ──────────────────────────────────
    user_now = st.session_state.get("current_user")
    db_now   = st.session_state.get("db", {})
    ud_now   = db_now.get("users", {}).get(user_now, {}) if user_now else {}
    if is_premium_actif(ud_now):
        st.markdown("""
        <style>
        /* ===== FOND OR PREMIUM ===== */
        .stApp {
            background: #2d1f00 !important;
            background: -webkit-linear-gradient(135deg, #3d2800 0%, #4a3200 20%, #3a2600 40%, #4d3500 60%, #3d2900 80%, #2d1f00 100%) !important;
            background: linear-gradient(135deg, #3d2800 0%, #4a3200 20%, #3a2600 40%, #4d3500 60%, #3d2900 80%, #2d1f00 100%) !important;
            color: #fff8e1 !important;
        }

        /* Halos lumineux dorés très visibles */
        .stApp::before {
            content: '';
            position: fixed;
            inset: 0;
            background:
                radial-gradient(ellipse at 10% 10%, rgba(255,215,0,0.30) 0%, transparent 40%),
                radial-gradient(ellipse at 90% 90%, rgba(255,160,0,0.25) 0%, transparent 40%),
                radial-gradient(ellipse at 50% 50%, rgba(255,200,0,0.12) 0%, transparent 60%),
                radial-gradient(ellipse at 85% 10%, rgba(255,215,0,0.20) 0%, transparent 35%),
                radial-gradient(ellipse at 15% 90%, rgba(255,140,0,0.18) 0%, transparent 35%);
            pointer-events: none;
            z-index: 0;
        }

        /* ===== TITRE PRINCIPAL OR ===== */
        .main-title {
            background: linear-gradient(90deg, #b8860b, #FFD700, #fff5c0, #FFD700, #b8860b) !important;
            background-size: 200% auto !important;
            -webkit-background-clip: text !important;
            -webkit-text-fill-color: transparent !important;
            animation: shimmer-gold 3s linear infinite !important;
            text-shadow: none !important;
            filter: drop-shadow(0 0 20px rgba(255,215,0,0.5));
        }
        @keyframes shimmer-gold {
            0%   { background-position: -200% center; }
            100% { background-position:  200% center; }
        }

        /* ===== TABS OR ===== */
        .stTabs [data-baseweb="tab-list"] {
            background-color: rgba(255,215,0,0.05) !important;
            border: 1px solid rgba(255,215,0,0.2) !important;
        }
        .stTabs [data-baseweb="tab"] {
            background-color: rgba(255,215,0,0.07) !important;
            border: 1px solid rgba(255,215,0,0.15) !important;
            color: #FFD700 !important;
        }
        .stTabs [data-baseweb="tab"]:nth-child(2) {
            border: 1px solid rgba(255,215,0,0.5) !important;
            box-shadow: 0 0 15px rgba(255,215,0,0.15) !important;
            background-color: rgba(255,215,0,0.1) !important;
        }
        .stTabs [data-baseweb="tab"]:hover {
            background-color: rgba(255,215,0,0.2) !important;
        }
        .stTabs [aria-selected="true"] {
            background-color: rgba(255,215,0,0.3) !important;
            border: 1px solid #FFD700 !important;
            box-shadow: 0 0 20px rgba(255,215,0,0.4) !important;
        }

        /* ===== INPUTS OR ===== */
        .stTextInput label, .stSelectbox label, .stTextArea label {
            color: #FFD700 !important;
        }
        div[data-baseweb="input"], div[data-baseweb="select"] > div {
            border: 1px solid rgba(255,215,0,0.6) !important;
            background-color: rgba(70,48,0,0.80) !important;
            color: #fff8e1 !important;
        }
        .stTextArea textarea {
            background-color: rgba(65,44,0,0.80) !important;
            color: #fff8e1 !important;
            border: 2px solid #FFD700 !important;
            animation: border-gold 4s linear infinite !important;
        }
        @keyframes border-gold {
            0%   { border-color: #FFD700; box-shadow: 0 0 14px rgba(255,215,0,0.55); }
            33%  { border-color: #FF8C00; box-shadow: 0 0 18px rgba(255,140,0,0.45); }
            66%  { border-color: #b8860b; box-shadow: 0 0 14px rgba(184,134,11,0.45); }
            100% { border-color: #FFD700; box-shadow: 0 0 14px rgba(255,215,0,0.55); }
        }

        /* ===== BOUTONS OR ===== */
        .stButton>button {
            background: linear-gradient(90deg, #8a6200, #c49a00, #FFD700, #c49a00, #8a6200) !important;
            background-size: 200% auto !important;
            color: #1a0f00 !important;
            box-shadow: 0 4px 22px rgba(255,215,0,0.55) !important;
            animation: shimmer-gold 3s linear infinite !important;
            font-weight: 800 !important;
        }
        .stButton>button:hover {
            box-shadow: 0 6px 32px rgba(255,215,0,0.75) !important;
            transform: translateY(-2px) !important;
        }
        .stFormSubmitButton>button {
            background: linear-gradient(90deg, #8a6200, #c49a00, #FFD700, #c49a00, #8a6200) !important;
            background-size: 200% auto !important;
            color: #1a0f00 !important;
            box-shadow: 0 4px 22px rgba(255,215,0,0.55) !important;
            animation: shimmer-gold 3s linear infinite !important;
            font-weight: 800 !important;
        }
        .stFormSubmitButton>button:hover {
            box-shadow: 0 6px 32px rgba(255,215,0,0.75) !important;
            transform: translateY(-2px) !important;
        }

        /* ===== SIDEBAR OR ===== */
        section[data-testid="stSidebar"] {
            background: linear-gradient(180deg, #3a2800 0%, #4a3400 40%, #3a2800 100%) !important;
            border-right: 2px solid rgba(255,215,0,0.4) !important;
            box-shadow: 4px 0 25px rgba(255,215,0,0.12) !important;
        }

        /* ===== INFO-CARD OR ===== */
        .info-card {
            border-left: 4px solid #FFD700 !important;
            background: rgba(255,215,0,0.10) !important;
        }
        .info-title { color: #FFD700 !important; }

        /* ===== FILE-CARD OR ===== */
        .file-card {
            border: 2px solid rgba(255,215,0,0.5) !important;
            background: rgba(255,215,0,0.07) !important;
        }

        /* ===== PROGRESS BAR OR ===== */
        .stProgress > div > div > div > div {
            background-image: linear-gradient(to right, #b8860b, #FFD700, #FF8C00) !important;
        }

        /* ===== EXPANDER OR ===== */
        .streamlit-expanderHeader {
            color: #FFD700 !important;
            border: 1px solid rgba(255,215,0,0.3) !important;
            background: rgba(255,215,0,0.10) !important;
        }

        /* ===== DIVIDER OR ===== */
        hr { border-color: rgba(255,215,0,0.3) !important; }

        /* ===== METRIC OR ===== */
        [data-testid="stMetric"] {
            background: rgba(255,215,0,0.10) !important;
            border: 1px solid rgba(255,215,0,0.35) !important;
            border-radius: 12px !important;
            padding: 10px !important;
        }
        [data-testid="stMetricValue"] { color: #FFD700 !important; }

        /* ===== SUCCESS / INFO / WARNING OR ===== */
        .stSuccess {
            background: rgba(255,215,0,0.12) !important;
            border: 1px solid rgba(255,215,0,0.4) !important;
            color: #FFD700 !important;
        }
        .stInfo {
            background: rgba(255,215,0,0.08) !important;
            border: 1px solid rgba(255,215,0,0.3) !important;
        }

        /* ===== SUPPORT BTN OR ===== */
        .support-btn {
            border: 2px solid #FFD700 !important;
            color: #FFD700 !important;
        }
        .support-btn:hover {
            background: #FFD700 !important;
            color: #0a0800 !important;
        }

        /* ===== SCROLLBAR OR ===== */
        ::-webkit-scrollbar-thumb {
            background: linear-gradient(#FFD700, #b8860b) !important;
        }
        ::-webkit-scrollbar-track {
            background: #0a0800 !important;
        }
        </style>
        """, unsafe_allow_html=True)

    if st.session_state["is_glowing"]:
        st.markdown('<style>.stApp { animation: glow-pulse 1.5s ease-in-out infinite; }</style>', unsafe_allow_html=True)


def show_auth_page():
    st.markdown("""
    <style>
    @keyframes shimmer {
        0%   { background-position: -200% center; }
        100% { background-position:  200% center; }
    }
    @keyframes float-up {
        0%   { opacity: 0; transform: translateY(30px); }
        100% { opacity: 1; transform: translateY(0); }
    }
    @keyframes letter-pop {
        0%   { opacity: 0; transform: translateY(20px) scale(0.8); }
        60%  { transform: translateY(-4px) scale(1.05); }
        100% { opacity: 1; transform: translateY(0) scale(1); }
    }
    @keyframes glow-border {
        0%   { box-shadow: 0 0 8px rgba(255,215,0,0.3), inset 0 0 8px rgba(255,215,0,0.05); }
        50%  { box-shadow: 0 0 28px rgba(255,215,0,0.7), inset 0 0 20px rgba(255,215,0,0.08); }
        100% { box-shadow: 0 0 8px rgba(255,215,0,0.3), inset 0 0 8px rgba(255,215,0,0.05); }
    }
    @keyframes particle-drift {
        0%   { transform: translateY(0px) translateX(0px) rotate(0deg); opacity: 0.6; }
        33%  { transform: translateY(-18px) translateX(8px) rotate(120deg); opacity: 1; }
        66%  { transform: translateY(-8px) translateX(-6px) rotate(240deg); opacity: 0.7; }
        100% { transform: translateY(0px) translateX(0px) rotate(360deg); opacity: 0.6; }
    }
    @keyframes scanline {
        0%   { top: -10%; }
        100% { top: 110%; }
    }
    @keyframes pulse-dot {
        0%, 100% { opacity: 1; transform: scale(1); }
        50%       { opacity: 0.4; transform: scale(0.6); }
    }
    .auth-hero {
        text-align: center;
        padding: 40px 20px 10px 20px;
        animation: float-up 0.8s ease both;
    }
    .auth-logo-ring {
        width: 90px; height: 90px;
        border-radius: 50%;
        margin: 0 auto 18px auto;
        background: radial-gradient(circle at 35% 35%, #fff8e1, #FFD700 40%, #b8860b);
        box-shadow: 0 0 0 4px rgba(255,215,0,0.2), 0 0 40px rgba(255,215,0,0.5);
        display: flex; align-items: center; justify-content: center;
        font-size: 2.6rem;
        animation: glow-border 3s ease-in-out infinite;
        position: relative;
    }
    .auth-logo-ring::after {
        content: '';
        position: absolute;
        inset: -6px;
        border-radius: 50%;
        border: 2px dashed rgba(255,215,0,0.4);
        animation: particle-drift 6s linear infinite;
    }
    .auth-title-wrap { display: flex; justify-content: center; gap: 2px; flex-wrap: wrap; margin-bottom: 6px; }
    .auth-letter {
        font-size: 3rem;
        font-weight: 800;
        background: linear-gradient(90deg, #b8860b, #FFD700, #fff5c0, #FFD700, #b8860b);
        background-size: 200% auto;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        animation: letter-pop 0.5s ease both, shimmer 3s linear infinite;
        display: inline-block;
        line-height: 1.1;
    }
    .auth-subtitle {
        color: rgba(255,215,0,0.65);
        font-size: 0.95rem;
        letter-spacing: 4px;
        text-transform: uppercase;
        animation: float-up 1s ease 0.5s both;
        margin-bottom: 6px;
    }
    .auth-tagline {
        color: rgba(255,255,255,0.4);
        font-size: 0.82rem;
        letter-spacing: 1.5px;
        animation: float-up 1s ease 0.8s both;
    }
    .auth-divider {
        display: flex; align-items: center; gap: 14px;
        margin: 28px auto 32px auto; max-width: 420px;
        animation: float-up 1s ease 1s both;
    }
    .auth-divider-line { flex: 1; height: 1px; background: linear-gradient(90deg, transparent, rgba(255,215,0,0.5), transparent); }
    .auth-divider-dot {
        width: 6px; height: 6px; border-radius: 50%; background: #FFD700;
        animation: pulse-dot 1.8s ease-in-out infinite;
    }
    .auth-card {
        background: linear-gradient(145deg, rgba(20,15,5,0.95), rgba(35,25,5,0.9));
        border: 1px solid rgba(255,215,0,0.35);
        border-radius: 22px;
        padding: 32px 28px 28px 28px;
        position: relative;
        overflow: hidden;
        animation: float-up 0.9s ease both;
        animation-delay: var(--card-delay, 0s);
    }
    .auth-card::before {
        content: '';
        position: absolute;
        top: 0; left: 0; right: 0; height: 3px;
        background: linear-gradient(90deg, #b8860b, #FFD700, #fff5c0, #FFD700, #b8860b);
        background-size: 200% auto;
        animation: shimmer 2.5s linear infinite;
        border-radius: 22px 22px 0 0;
    }
    .auth-card::after {
        content: '';
        position: absolute;
        width: 180px; height: 180px;
        border-radius: 50%;
        background: radial-gradient(circle, rgba(255,215,0,0.06), transparent 70%);
        bottom: -60px; right: -60px;
        pointer-events: none;
    }
    .auth-card .scanline {
        position: absolute;
        left: 0; right: 0; height: 2px;
        background: linear-gradient(90deg, transparent, rgba(255,215,0,0.12), transparent);
        animation: scanline 4s linear infinite;
        pointer-events: none;
    }
    .auth-card-header {
        display: flex; align-items: center; gap: 12px;
        margin-bottom: 22px;
    }
    .auth-card-icon {
        width: 44px; height: 44px; border-radius: 12px;
        display: flex; align-items: center; justify-content: center;
        font-size: 1.4rem;
        background: linear-gradient(135deg, rgba(255,215,0,0.15), rgba(255,215,0,0.05));
        border: 1px solid rgba(255,215,0,0.3);
        box-shadow: 0 0 12px rgba(255,215,0,0.15);
    }
    .auth-card-title {
        color: #FFD700 !important;
        font-size: 1.15rem;
        font-weight: 800;
        letter-spacing: 0.5px;
        margin: 0;
        text-transform: uppercase;
    }
    .auth-card-desc {
        color: rgba(255,255,255,0.35);
        font-size: 0.75rem;
        margin-top: 2px;
        letter-spacing: 0.5px;
    }
    .auth-page .stTextInput label {
        color: rgba(255,215,0,0.8) !important;
        font-size: 0.82rem !important;
        font-weight: 600 !important;
        letter-spacing: 1px;
        text-transform: uppercase;
    }
    .auth-page div[data-baseweb="input"] {
        background: rgba(0,0,0,0.5) !important;
        border: 1px solid rgba(255,215,0,0.25) !important;
        border-radius: 12px !important;
        transition: border-color 0.3s, box-shadow 0.3s !important;
    }
    .auth-page div[data-baseweb="input"]:focus-within {
        border-color: rgba(255,215,0,0.7) !important;
        box-shadow: 0 0 0 3px rgba(255,215,0,0.12) !important;
    }
    @keyframes btn-shimmer {
        0%   { background-position: -300% center; }
        100% { background-position:  300% center; }
    }
    @keyframes btn-float {
        0%, 100% { transform: translateY(0px);   box-shadow: 0 6px 25px rgba(255,215,0,0.45), 0 0 0 0 rgba(255,215,0,0.2); }
        50%       { transform: translateY(-4px);  box-shadow: 0 14px 35px rgba(255,215,0,0.65), 0 0 18px 4px rgba(255,215,0,0.15); }
    }
    @keyframes btn-glow-ring {
        0%, 100% { box-shadow: 0 6px 25px rgba(255,215,0,0.45), 0 0  0px rgba(255,215,0,0);   }
        50%       { box-shadow: 0 6px 25px rgba(255,215,0,0.45), 0 0 22px rgba(255,215,0,0.35); }
    }
    .auth-page .stButton > button {
        background: linear-gradient(90deg, #7a5500, #b8860b, #FFD700, #fff5c0, #FFD700, #b8860b, #7a5500) !important;
        background-size: 300% auto !important;
        color: #0a0800 !important;
        font-weight: 800 !important;
        font-size: 0.92rem !important;
        letter-spacing: 2.5px !important;
        text-transform: uppercase !important;
        border-radius: 50px !important;
        border: none !important;
        padding: 0.82rem 1.8rem !important;
        width: 100% !important;
        position: relative !important;
        overflow: hidden !important;
        animation: btn-shimmer 3s linear infinite, btn-float 3.5s ease-in-out infinite !important;
        cursor: pointer !important;
    }
    .auth-secure-badge {
        display: flex; align-items: center; justify-content: center;
        gap: 8px; margin-top: 28px;
        color: rgba(255,215,0,0.35);
        font-size: 0.72rem; letter-spacing: 1.5px; text-transform: uppercase;
        animation: float-up 1s ease 1.2s both;
    }
    .auth-secure-badge span { font-size: 0.9rem; }
    </style>
    """, unsafe_allow_html=True)

    letters = list("NOVA PLATFORM")
    letter_spans = "".join(
        f'<span class="auth-letter" style="animation-delay:{i*0.07:.2f}s">'
        f'{"&nbsp;" if c == " " else c}</span>'
        for i, c in enumerate(letters)
    )
    st.markdown(f"""
    <div class="auth-hero">
        <div class="auth-logo-ring">⚡</div>
        <div class="auth-title-wrap">{letter_spans}</div>
        <div class="auth-subtitle">Plateforme IA bureautique</div>
        <div class="auth-tagline">Intelligence · Excellence · Performance</div>
    </div>
    <div class="auth-divider">
        <div class="auth-divider-line"></div>
        <div class="auth-divider-dot"></div>
        <div class="auth-divider-line"></div>
    </div>
    """, unsafe_allow_html=True)

    col1, col2 = st.columns(2, gap="large")

    with col1:
        st.markdown("""
        <div class="auth-card" style="--card-delay:1.1s;">
            <div class="scanline"></div>
            <div class="auth-card-header">
                <div class="auth-card-icon">🔐</div>
                <div>
                    <div class="auth-card-title">Accès Membre</div>
                    <div class="auth-card-desc">Identifiez-vous pour accéder à votre espace</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        # --- LOGIN : champ Streamlit natif + bouton Python ---
        wa_auth_raw = st.text_input(
            "📱 Votre numéro WhatsApp",
            placeholder="Ex: 22501...",
            key="wa_login_input"
        )
        wa_auth = "".join(c for c in wa_auth_raw if c.isdigit())
        if wa_auth_raw != wa_auth and wa_auth_raw:
            st.warning("⚠️ Le numéro WhatsApp ne doit contenir que des chiffres.")
        # Formulaire HTML invisible uniquement pour déclencher l'autocomplete mobile
        components.html("""
            <form id="nova-login-form" autocomplete="on" style="display:none;">
                <input type="tel" name="username" autocomplete="username" />
                <input type="password" name="password" autocomplete="current-password" value="nova_platform_auth" />
                <button type="submit">ok</button>
            </form>
            <script>
            document.getElementById("nova-login-form").addEventListener("submit", function(e){ e.preventDefault(); });
            </script>
        """, height=1)
        if st.button("⚡ S'IDENTIFIER", use_container_width=True, key="btn_login"):
            fresh_db = load_db()
            st.session_state["db"] = fresh_db
            wa_norm = normalize_wa(wa_auth)
            uid_trouve = None
            for u_id, u_data in fresh_db["users"].items():
                if u_data["whatsapp"] == wa_norm:
                    uid_trouve = u_id
                    break
            if uid_trouve:
                st.session_state["current_user"] = uid_trouve
                st.session_state["view"] = "home"
                st.session_state["show_mode_modal"] = True
                st.query_params["user_id"] = uid_trouve
                st.rerun()
            else:
                st.error("❌ Numéro WhatsApp inconnu. Vérifiez ou créez un compte.")

    with col2:
        st.markdown("""
        <div class="auth-card" style="--card-delay:1.3s;">
            <div class="scanline"></div>
            <div class="auth-card-header">
                <div class="auth-card-icon">✨</div>
                <div>
                    <div class="auth-card-title">Nouveau Compte</div>
                    <div class="auth-card-desc">Rejoignez l'élite Nova Platform dès maintenant</div>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        # --- INSCRIPTION : WhatsApp uniquement (uid = numéro normalisé) ---
        new_wa_raw = st.text_input("📱 Votre numéro WhatsApp", placeholder="Ex: 22507...", key="new_wa_input")
        new_wa = "".join(c for c in new_wa_raw if c.isdigit())
        if new_wa_raw != new_wa and new_wa_raw:
            st.warning("⚠️ Le numéro WhatsApp ne doit contenir que des chiffres.")
        # Formulaire HTML natif pour que le téléphone propose d'enregistrer
        components.html("""
            <form id="nova-signup-form" autocomplete="on" style="display:none;">
                <input type="tel" name="username" autocomplete="username" />
                <input type="password" name="new-password" autocomplete="new-password" value="nova_auth" />
                <button type="submit">ok</button>
            </form>
            <script>
            document.getElementById("nova-signup-form").addEventListener("submit", function(e){ e.preventDefault(); });
            </script>
        """, height=1)
        if st.button("💎 REJOINDRE NOVA PLATFORM", use_container_width=True, key="btn_signup"):
            if new_wa:
                db = st.session_state["db"]
                wa_norm_new = normalize_wa(new_wa)
                new_uid = wa_norm_new  # uid = numéro normalisé
                if new_uid not in db["users"]:
                    succes = save_user(new_uid, wa_norm_new)
                    if succes:
                        db["users"][new_uid] = {
                            "whatsapp": wa_norm_new,
                            "email": "Non renseigné",
                            "joined": str(datetime.now()),
                            "premium": False,
                            "premium_plan": None,
                            "premium_expiry": None,
                            "gen_used": 0,
                            "gen_date": None,
                        }
                        st.session_state["current_user"] = new_uid
                        st.session_state["view"] = "home"
                        st.session_state["show_mode_modal"] = True
                        st.session_state["db"] = load_db()
                        st.query_params["user_id"] = new_uid
                        st.rerun()
                    else:
                        st.error("❌ Impossible de créer le compte. Vérifie ta connexion ou contacte le support.")
                else:
                    st.warning("⚠️ Ce numéro WhatsApp est déjà associé à un compte. Connecte-toi à gauche.")
            else:
                st.error("❌ Saisis ton numéro WhatsApp pour créer un compte.")


    st.markdown("""
    <div class="auth-secure-badge">
        <span>🔒</span> Connexion sécurisée &nbsp;·&nbsp; <span>⚡</span> Nova Platform &nbsp;·&nbsp; <span>🛡️</span> Données protégées
    </div>
    """, unsafe_allow_html=True)

    # ── NOVA IA SUPPORT SUR LA PAGE DE CONNEXION ────────────────────
    st.markdown("""
    <style>
    @keyframes authGoldGlow {
        0%, 100% { box-shadow: 0 0 8px 2px rgba(255,215,0,0.2); border-color: rgba(255,215,0,0.4); }
        50%       { box-shadow: 0 0 18px 6px rgba(255,215,0,0.45); border-color: rgba(255,215,0,0.7); }
    }
    @keyframes botFloat {
        0%, 100% { transform: translateY(0); }
        50%       { transform: translateY(-4px); }
    }
    .auth-arsene-widget {
        background: linear-gradient(135deg, rgba(255,215,0,0.06), rgba(255,140,0,0.03));
        border: 1.5px solid rgba(255,215,0,0.4);
        border-radius: 18px;
        padding: 18px 20px;
        margin-top: 24px;
        animation: authGoldGlow 2.5s ease-in-out infinite;
    }
    .auth-arsene-bot {
        font-size: 1.8rem;
        animation: botFloat 3s ease-in-out infinite;
        display: inline-block;
        filter: drop-shadow(0 0 6px rgba(255,215,0,0.5));
    }
    .auth-arsene-title {
        background: linear-gradient(135deg, #FFD700, #FFA500, #FFD700);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-weight: 900;
        font-size: 1rem;
        vertical-align: middle;
        margin-left: 8px;
    }
    .auth-arsene-sub {
        color: rgba(255,255,255,0.45);
        font-size: 0.78rem;
        margin-top: 4px;
        display: block;
    }
    .auth-online-dot {
        display: inline-block;
        width: 7px; height: 7px;
        background: #2ecc71;
        border-radius: 50%;
        margin-right: 5px;
        box-shadow: 0 0 5px #2ecc71;
    }
    </style>
    <div class="auth-arsene-widget">
        <span class="auth-arsene-bot">🤖</span>
        <span class="auth-arsene-title">Assistant Nova 24/7</span>
        <span class="auth-arsene-sub">
            <span class="auth-online-dot"></span>En ligne · Un problème pour vous connecter ou créer un compte ?
        </span>
    </div>
    """, unsafe_allow_html=True)

    # Initialiser chat auth
    if "auth_chat" not in st.session_state:
        st.session_state["auth_chat"] = [{
            "role": "assistant",
            "content": "Salut ! Moi c'est Nova IA 👋 Tu as un problème pour te connecter ou créer ton compte ? Dis-moi tout, je suis là pour t'aider !"
        }]
    if "auth_chat_open" not in st.session_state:
        st.session_state["auth_chat_open"] = False
    if "auth_resolu" not in st.session_state:
        st.session_state["auth_resolu"] = False

    col_open, _ = st.columns([1, 3])
    with col_open:
        label_btn_chat = "💬 Fermer" if st.session_state["auth_chat_open"] else "💬 Besoin d'aide ?"
        if st.button(label_btn_chat, key="btn_auth_chat"):
            st.session_state["auth_chat_open"] = not st.session_state["auth_chat_open"]
            st.rerun()

    if st.session_state["auth_chat_open"]:
        # Afficher historique
        for msg in st.session_state["auth_chat"]:
            align = "flex-end" if msg["role"] == "user" else "flex-start"
            bg = "rgba(255,255,255,0.05)" if msg["role"] == "user" else "rgba(255,193,7,0.08)"
            border = "1px solid rgba(255,255,255,0.08)" if msg["role"] == "user" else "1px solid rgba(255,215,0,0.2)"
            border_l = "" if msg["role"] == "user" else "border-left: 3px solid #FFD700;"
            color_label = "rgba(255,255,255,0.5)" if msg["role"] == "user" else "#FFD700"
            icon = "🧑" if msg["role"] == "user" else "🤖"
            label_msg = "Vous" if msg["role"] == "user" else "Nova IA"
            st.markdown(f"""
            <div style="display:flex;justify-content:{align};margin:6px 0;">
                <div style="background:{bg};border:{border};{border_l}border-radius:12px;
                     padding:10px 14px;max-width:85%;">
                    <span style="color:{color_label};font-size:.78rem;font-weight:800;">{icon} {label_msg}</span>
                    <p style="color:#eee;margin:4px 0 0 0;font-size:.88rem;line-height:1.5;">{msg["content"]}</p>
                </div>
            </div>""", unsafe_allow_html=True)

        if not st.session_state["auth_resolu"]:
            with st.form("auth_chat_form", clear_on_submit=True):
                msg_auth = st.text_input(
                    "Votre message",
                    placeholder="Ex: Je n'arrive pas à me connecter, mon identifiant est incorrect...",
                    label_visibility="collapsed"
                )
                col_s, col_t = st.columns([4, 1])
                with col_s:
                    send_auth = st.form_submit_button("📨 Envoyer", use_container_width=True)
                with col_t:
                    end_auth = st.form_submit_button("✅ Fin", use_container_width=True)

            if send_auth and msg_auth.strip():
                st.session_state["auth_chat"].append({"role": "user", "content": msg_auth.strip()})
                historique_auth = "\n".join([
                    f"{'Client' if m['role']=='user' else 'Nova IA'}: {m['content']}"
                    for m in st.session_state["auth_chat"]
                ])
                prompt_auth = f"""Tu es NOVA IA, l'assistant support de Nova Platform (Côte d'Ivoire).
Tu aides les visiteurs sur la PAGE DE CONNEXION — ils ne sont pas encore connectés.
Réponds en français, avec bienveillance et concision.

CE QUE TU SAIS SUR LA CONNEXION NOVA :
- Identifiant = nom choisi à l'inscription
- Mot de passe = numéro WhatsApp (format : 225XXXXXXXX avec 225 au début)
- Si identifiant oublié → impossible à récupérer sans contacter Nova
- Si numéro WhatsApp oublié → contacter Nova sur WhatsApp : {WHATSAPP_NUMBER}
- Pour créer un compte : formulaire "Nouveau Compte" à droite sur cette même page
- En cas de problème grave ou non résolu → donner le WhatsApp Nova : {WHATSAPP_NUMBER}

RÈGLE : Ne te présente pas à chaque message. Reste naturel et direct.

Historique :
{historique_auth}

Réponds uniquement au dernier message. 2-4 phrases max.
RÈGLE ESCALADE OBLIGATOIRE :
Si le client exprime un problème grave (paiement, fichier perdu, compte bloqué, bug critique, plainte urgente), tu DOIS lui proposer cette phrase exacte à la fin de ta réponse :
"👉 Veux-tu que je transmette ton problème directement au service client Nova ? Réponds juste OUI et je m'en occupe immédiatement."
Si dans l'historique le client répond OUI ou "oui" ou "ok" ou "ouais" à cette proposition, réponds UNIQUEMENT ce texte exact sans rien d'autre :
__ESCALADE_CONFIRMEE__"""

                with st.spinner("🤖 Nova Platform répond..."):
                    rep_auth = generer_avec_gemini("Support Auth", prompt_auth, "visiteur")
                if rep_auth.startswith("❌"):
                    rep_auth = f"Désolé, une erreur est survenue. Contacte Nova directement sur WhatsApp : {WHATSAPP_NUMBER} 📲"
                if "__ESCALADE_CONFIRMEE__" in rep_auth:
                    ok = envoyer_escalade_support("visiteur", "non connecté", st.session_state["auth_chat"], "Page Connexion")
                    if ok:
                        rep_auth = "✅ C'est fait ! Ton problème a été transmis au service client Nova. Nous te recontactons très bientôt. 🙏"
                        st.session_state["auth_resolu"] = True
                    else:
                        rep_auth = f"Désolé, l'envoi a échoué. Contacte Nova directement : {WHATSAPP_NUMBER} 📲"
                st.session_state["auth_chat"].append({"role": "assistant", "content": rep_auth})
                st.rerun()

            if end_auth and len(st.session_state["auth_chat"]) > 1:
                try:
                    import resend
                    resend.api_key = st.secrets["RESEND_API_KEY"]
                    hist_email = "\n".join([
                        f"{'🧑 Visiteur' if m['role']=='user' else '🤖 Nova Platform'} : {m['content']}"
                        for m in st.session_state["auth_chat"]
                    ])
                    resend.Emails.send({
                        "from": "Nova Platform <onboarding@resend.dev>",
                        "to": [st.secrets["EMAIL_RECEIVER"]],
                        "subject": "🆘 Support Auth Nova — Visiteur",
                        "text": f"""SIGNALEMENT PAGE CONNEXION
━━━━━━━━━━━━━━━━━━━━━━━━
👤 Visiteur non connecté
⏰ Date : {datetime.now().strftime("%d/%m/%Y à %H:%M")}
━━━━━━━━━━━━━━━━━━━━━━━━

{hist_email}

━━━━━━━━━━━━━━━━━━━━━━━━
Intervenir si problème non résolu.
"""
                    })
                    st.session_state["auth_resolu"] = True
                    st.success("✅ Ton signalement a été envoyé à Nova. Nous revenons vers toi rapidement !")
                    st.rerun()
                except Exception as e_auth:
                    st.error(f"Erreur envoi : {e_auth}")
        else:
            st.success("✅ Signalement transmis à Nova. Nous te recontactons bientôt.")
            if st.button("🔄 Nouveau message", key="reset_auth_chat"):
                st.session_state["auth_chat"] = [{"role": "assistant", "content": "Salut ! Comment puis-je t'aider ?"}]
                st.session_state["auth_resolu"] = False
                st.rerun()




def main_dashboard():
    user = st.session_state["current_user"]
    db = st.session_state["db"]

    if user and user in db["users"]:
        ud = db["users"][user]
        if ud.get("premium") and not is_premium_actif(ud):
            desactiver_premium(user)
            st.session_state["db"] = load_db()
            db = st.session_state["db"]

    user_data     = db["users"].get(user, {}) if user else {}
    premium_actif = is_premium_actif(user_data)
    premium_info  = get_premium_info(user_data)

    # ══════════════════════════════════════════════════════════════
    # MODALE CHOIX DU MODE — apparaît une seule fois à la connexion
    # ══════════════════════════════════════════════════════════════
    if st.session_state.get("show_mode_modal", False):
        # Masquer complètement le reste de la page avec du CSS simple
        st.markdown("""
        <style>
        section[data-testid="stSidebar"] { display: none !important; }
        </style>
        """, unsafe_allow_html=True)

        # Centrage via colonnes Streamlit
        _, col_center, _ = st.columns([1, 3, 1])
        with col_center:
            st.markdown("<br><br>", unsafe_allow_html=True)

            # En-tête
            st.markdown("""
            <div style="text-align:center; padding: 20px 0 10px 0;">
                <div style="font-size:2.8rem; margin-bottom:8px;">👋</div>
                <div style="font-size:1.6rem; font-weight:900; color:#ffffff; margin-bottom:6px;">
                    Bienvenue sur Nova Platform
                </div>
                <div style="color:rgba(255,255,255,0.4); font-size:0.9rem; letter-spacing:2px; margin-bottom:30px;">
                    CHOISISSEZ VOTRE MODE
                </div>
            </div>
            """, unsafe_allow_html=True)

            # Carte Plateforme
            st.markdown("""
            <div style="
                background: linear-gradient(145deg, rgba(255,215,0,0.1), rgba(255,140,0,0.06));
                border: 2px solid rgba(255,215,0,0.5);
                border-radius: 20px;
                padding: 28px 24px;
                margin-bottom: 16px;
                text-align: center;
            ">
                <div style="font-size:2.8rem; margin-bottom:10px;">🖥️</div>
                <div style="display:inline-block; background:rgba(255,215,0,0.15);
                    border:1px solid rgba(255,215,0,0.35); border-radius:20px;
                    padding:3px 14px; font-size:0.75rem; font-weight:700;
                    color:#FFD700; letter-spacing:1px; margin-bottom:12px;">
                    MODE CLASSIQUE
                </div>
                <div style="font-size:1.15rem; font-weight:800; color:#FFD700; margin-bottom:10px;">
                    Version Plateforme
                </div>
                <div style="color:rgba(255,255,255,0.5); font-size:0.85rem; line-height:1.7;">
                    Parcours les services Nova, decris ton besoin<br>
                    et soumets ta demande ou genere en 1 clic.
                </div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("🖥️  Utiliser la Version Plateforme", key="mode_platform_btn", use_container_width=True):
                st.session_state["show_mode_modal"] = False
                st.rerun()

            st.markdown("<div style='text-align:center; color:rgba(255,255,255,0.2); padding:8px 0; font-size:0.85rem;'>— ou —</div>", unsafe_allow_html=True)

            # Carte Chat Nova IA
            st.markdown("""
            <div style="
                background: linear-gradient(145deg, rgba(0,200,255,0.1), rgba(0,120,200,0.06));
                border: 2px solid rgba(0,200,255,0.5);
                border-radius: 20px;
                padding: 28px 24px;
                margin-bottom: 16px;
                text-align: center;
            ">
                <div style="font-size:2.8rem; margin-bottom:10px;">🤖</div>
                <div style="display:inline-block; background:rgba(0,200,255,0.15);
                    border:1px solid rgba(0,200,255,0.35); border-radius:20px;
                    padding:3px 14px; font-size:0.75rem; font-weight:700;
                    color:#00d2ff; letter-spacing:1px; margin-bottom:12px;">
                    MODE IA
                </div>
                <div style="font-size:1.15rem; font-weight:800; color:#00d2ff; margin-bottom:10px;">
                    Chat avec Nova IA
                </div>
                <div style="color:rgba(255,255,255,0.5); font-size:0.85rem; line-height:1.7;">
                    Dis ce que tu veux en langage naturel.<br>
                    Nova IA pose les bonnes questions et genere<br>
                    ou soumet ta demande automatiquement.
                </div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("🤖  Discuter avec Nova IA", key="mode_chat_btn", use_container_width=True):
                st.session_state["show_mode_modal"] = False
                st.session_state.pop("nova_ia_chat", None)
                st.session_state.pop("nova_ia_phase", None)
                st.session_state.pop("nova_ia_service_detecte", None)
                st.session_state.pop("nova_ia_prompt_final", None)
                st.session_state["view"] = "nova_ia"
                st.rerun()

            st.markdown("""
            <div style="text-align:center; color:rgba(255,255,255,0.2);
                font-size:0.75rem; padding-top:16px;">
                Tu pourras changer de mode a tout moment depuis le menu
            </div>
            """, unsafe_allow_html=True)

        return   # stoppe le dashboard le temps que le client choisit

    with st.sidebar:
        st.markdown(f"### 👤 {user if user else 'Visiteur'}")
        if user:
            st.markdown(f"📱 **{db['users'][user]['whatsapp']}**")
            if premium_actif and premium_info:
                st.markdown(f"""
                <div style="margin:10px 0;">
                    <span class="badge-premium">⭐ PREMIUM — {premium_info['plan']}</span>
                    <div style="color:rgba(255,215,0,0.7);font-size:0.78rem;margin-top:6px;">
                        ⏳ Expire le {premium_info['expiry']} ({premium_info['jours_restants']}j restants)
                    </div>
                </div>""", unsafe_allow_html=True)
            else:
                st.markdown('<span class="badge-free">🔓 Compte Gratuit</span>', unsafe_allow_html=True)
            if st.button("Quitter la session"):
                st.session_state["current_user"] = None
                st.query_params.clear()
                # Effacer localStorage
                components.html("""
                    <script>
                    localStorage.removeItem('nova_user_id');
                    localStorage.removeItem('nova_user_ts');
                    localStorage.removeItem('nova_user');
                    </script>
                """, height=0)
                st.rerun()
        else:
            if st.button("Connexion"):
                st.session_state["view"] = "auth"
                st.rerun()
        
        st.divider()
        st.markdown(f"""
            <div class="info-card">
                <span class="info-title">🚀 LIVRAISON NOVA</span>
                <span style="color:#eee; font-size:0.9rem;">
                    Vos résultats IA apparaissent dans l'onglet <b>"📂 MES LIVRABLES"</b>.
                    <br><br>
                    Suivi instantané 24h/24.
                </span>
            </div>
        """, unsafe_allow_html=True)
        st.markdown(f'<a href="{whatsapp_support_url}" target="_blank" class="support-btn">💬 Support Nova</a>', unsafe_allow_html=True)

    if premium_actif:
        st.markdown("""
        <div style="text-align:center; margin-bottom:20px;">
            <div style="font-size:2.2rem; margin-bottom:-10px; filter:drop-shadow(0 0 15px rgba(255,215,0,0.8));">👑</div>
            <h1 class='main-title' style="
                background: linear-gradient(90deg, #b8860b, #FFD700, #fff5c0, #FFD700, #b8860b);
                background-size: 200% auto;
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent;
                animation: shimmer-gold 3s linear infinite;
                filter: drop-shadow(0 0 25px rgba(255,215,0,0.6));
                font-size: 3.5rem !important;
                font-weight: 800 !important;
                margin-top: 0;
            ">NOVA PLATFORM</h1>
            <div style="
                color: rgba(255,215,0,0.6);
                font-size: 0.75rem;
                letter-spacing: 5px;
                text-transform: uppercase;
                margin-top: -10px;
            ">✦ Membre Premium Actif ✦</div>
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("<h1 class='main-title'>NOVA PLATFORM</h1>", unsafe_allow_html=True)



    wa_jour = f"https://wa.me/{WHATSAPP_NUMBER}?text=Je%20souhaite%20l%27abonnement%20Nova%20Premium%20Journalier%20%C3%A0%20600%20FC."
    wa_10j  = f"https://wa.me/{WHATSAPP_NUMBER}?text=Je%20souhaite%20l%27abonnement%20Nova%20Premium%2010%20Jours%20%C3%A0%201000%20FC."
    wa_30j  = f"https://wa.me/{WHATSAPP_NUMBER}?text=Je%20souhaite%20l%27abonnement%20Nova%20Premium%2030%20Jours%20%C3%A0%202500%20FC."

    if premium_actif and premium_info:
        st.markdown(f"""
        <div style="background:linear-gradient(135deg,rgba(255,215,0,.12),rgba(255,140,0,.08));
             border:2px solid #FFD700;border-radius:20px;padding:20px;text-align:center;margin-bottom:20px;">
            <div style="font-size:1.4rem;font-weight:800;color:#FFD700;">
                ⭐ MEMBRE PREMIUM ACTIF — {premium_info['plan']}
            </div>
            <div style="color:rgba(255,255,255,.7);margin-top:6px;">
                🤖 Génération IA instantanée activée · Expire le <b>{premium_info['expiry']}</b>
                ({premium_info['jours_restants']} jour(s) restant(s))
            </div>
        </div>""", unsafe_allow_html=True)
    else:
        st.markdown("""
            <div class="premium-card">
                <div class="premium-title">⭐ ACCÉLÉRATEUR NOVA PREMIUM ⭐</div>
                <div class="premium-desc">
                    Passez au niveau supérieur : IA illimitée et puissance de calcul <b>10<sup>10</sup></b>.
                </div>
            </div>
        """, unsafe_allow_html=True)

        col_btn_center = st.columns([1, 2, 1])[1]
        with col_btn_center:
            if st.button("💎 ACTIVER NOVA PREMIUM", key="open_premium"):
                st.session_state["show_premium_modal"] = True
                st.rerun()

    if st.session_state["show_premium_modal"]:
        st.markdown("""
            <div style="
                background: linear-gradient(135deg, #1a1a2e, #16213e, #0f3460);
                border: 2px solid #FFD700;
                border-radius: 24px;
                padding: 35px 25px 30px 25px;
                margin: 10px 0 30px 0;
                box-shadow: 0 0 60px rgba(255,215,0,0.25);
            ">
                <h2 style="text-align:center; color:#FFD700; font-size:1.7rem; font-weight:800; margin-bottom:6px; letter-spacing:1px;">
                    ⭐ CHOISISSEZ VOTRE FORMULE NOVA PREMIUM
                </h2>
                <p style="text-align:center; color:rgba(255,255,255,0.55); margin-bottom:30px; font-size:0.95rem;">
                    Sélectionnez le plan qui correspond à vos besoins
                </p>
            </div>
        """, unsafe_allow_html=True)

        col1_p, col2_p, col3_p = st.columns(3)

        with col1_p:
            st.markdown("""
                <div style="background: rgba(255,255,255,0.05); border: 1px solid rgba(255,215,0,0.4); border-radius: 18px; padding: 28px 16px; text-align: center; min-height: 300px;">
                    <div style="font-size:2.5rem; margin-bottom:10px;">🌅</div>
                    <div style="color:#FFD700; font-weight:800; font-size:1.1rem; margin-bottom:6px; text-transform:uppercase;">Journalier</div>
                    <div style="color:white; font-size:2rem; font-weight:800; margin:10px 0;">600 FC</div>
                    <div style="color:rgba(255,255,255,0.45); font-size:0.8rem; margin-bottom:16px;">/ par jour</div>
                    <div style="background:rgba(255,215,0,0.1); border-radius:10px; padding:10px; margin-bottom:22px;">
                        <span style="color:#FFD700; font-size:0.9rem;">⚡ 2 générations IA / jour</span>
                    </div>
                </div>
            """, unsafe_allow_html=True)
            st.markdown(f'<a href="{wa_jour}" target="_blank" style="display:block; background:linear-gradient(45deg,#FFD700,#FF8C00); color:#000; font-weight:800; padding:12px; border-radius:50px; text-decoration:none; font-size:1rem; text-align:center; margin-top:10px;">Choisir cette formule</a>', unsafe_allow_html=True)

        with col2_p:
            st.markdown("""
                <div style="background: linear-gradient(135deg, rgba(0,210,255,0.15), rgba(58,123,213,0.15)); border: 2px solid #00d2ff; border-radius: 18px; padding: 28px 16px; text-align: center; min-height: 300px; position: relative;">
                    <div style="background:linear-gradient(90deg,#00d2ff,#3a7bd5); color:white; font-size:0.75rem; font-weight:800; padding:4px 16px; border-radius:20px; display:inline-block; margin-bottom:12px;">⭐ POPULAIRE</div>
                    <div style="font-size:2.5rem; margin-bottom:10px;">🔟</div>
                    <div style="color:#00d2ff; font-weight:800; font-size:1.1rem; margin-bottom:6px; text-transform:uppercase;">10 Jours</div>
                    <div style="color:white; font-size:2rem; font-weight:800; margin:10px 0;">1 000 FC</div>
                    <div style="color:rgba(255,255,255,0.45); font-size:0.8rem; margin-bottom:16px;">/ 10 jours</div>
                    <div style="background:rgba(0,210,255,0.1); border-radius:10px; padding:10px; margin-bottom:22px;">
                        <span style="color:#00d2ff; font-size:0.9rem;">⚡ 9 générations IA / jour</span>
                    </div>
                </div>
            """, unsafe_allow_html=True)
            st.markdown(f'<a href="{wa_10j}" target="_blank" style="display:block; background:linear-gradient(45deg,#00d2ff,#3a7bd5); color:#fff; font-weight:800; padding:12px; border-radius:50px; text-decoration:none; font-size:1rem; text-align:center; margin-top:10px;">Choisir cette formule</a>', unsafe_allow_html=True)

        with col3_p:
            st.markdown("""
                <div style="background: rgba(255,255,255,0.05); border: 1px solid rgba(46,204,113,0.4); border-radius: 18px; padding: 28px 16px; text-align: center; min-height: 300px;">
                    <div style="font-size:2.5rem; margin-bottom:10px;">👑</div>
                    <div style="color:#2ecc71; font-weight:800; font-size:1.1rem; margin-bottom:6px; text-transform:uppercase;">30 Jours</div>
                    <div style="color:white; font-size:2rem; font-weight:800; margin:10px 0;">2 500 FC</div>
                    <div style="color:rgba(255,255,255,0.45); font-size:0.8rem; margin-bottom:16px;">/ 30 jours</div>
                    <div style="background:rgba(46,204,113,0.1); border-radius:10px; padding:10px; margin-bottom:22px;">
                        <span style="color:#2ecc71; font-size:0.9rem;">♾️ Générations ILLIMITÉES</span>
                    </div>
                </div>
            """, unsafe_allow_html=True)
            st.markdown(f'<a href="{wa_30j}" target="_blank" style="display:block; background:linear-gradient(45deg,#2ecc71,#27ae60); color:#fff; font-weight:800; padding:12px; border-radius:50px; text-decoration:none; font-size:1rem; text-align:center; margin-top:10px;">Choisir cette formule</a>', unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        col_close = st.columns([1, 2, 1])[1]
        with col_close:
            if st.button("✕ Fermer", key="close_premium"):
                st.session_state["show_premium_modal"] = False
                st.rerun()

    # ── AUTO-TRAITEMENT DEMANDES CLOUDFLARE (sans délai — contenu déjà généré) ──
    try:
        _cf_demandes = supabase.table("demandes").select("*").eq("source", "cloudflare").eq("status", "cf_pret").execute().data or []
        _now_cf = datetime.now()

        _SERVICES_CF = [
            "Exposé scolaire complet IA",
            "Rapport de Stage IA",
            "CV & Lettre de Motivation",
            "Fiche de Cours Professeur IA",
            "Excel & Data Analytics",
            "Création de Sujets & Examens",
            "Création Word (depuis zéro)",
            "Affiches & Reçus",
            "Modifier mon Fichier",
            "Conversion & PDF",
            "OCR — Numérisation",
        ]

        for _cf in _cf_demandes:
            _cf_id      = _cf.get("id", "")
            _cf_service = _cf.get("service", "")
            _cf_uid     = _cf.get("uid", "")
            _cf_wa      = _cf.get("whatsapp", "")
            _cf_contenu = _cf.get("contenu_genere", "")

            # Ignorer si contenu vide ou déjà traité
            if not _cf_contenu or _cf_contenu.startswith("❌"):
                continue

            # Chercher le nom du service même sans emoji prefix
            _service_match = any(s in _cf_service for s in _SERVICES_CF)
            if not _service_match:
                continue

            # Créer le docx depuis le contenu déjà généré par Cloudflare
            try:
                _cf_buf = creer_docx(_cf_contenu, _cf_service, _cf_uid)
            except Exception as _e_cf:
                supabase.table("config").upsert({
                    "key": f"cf_error_{_cf_id}",
                    "value": f"creer_docx: {str(_e_cf)[:200]}"
                }).execute()
                continue

            _cf_nom = f"{_cf_uid}_{_cf_service[:20].strip()}_cf.docx".replace(" ", "_").replace("/", "-")

            # Upload Supabase Storage
            _cf_url = upload_fichier_client(_cf_uid, _cf_id, _cf_buf, _cf_nom)
            if not _cf_url or _cf_url.startswith("ERREUR"):
                supabase.table("config").upsert({
                    "key": f"cf_error_{_cf_id}",
                    "value": f"Upload: {_cf_url}"
                }).execute()
                continue

            # Sauvegarder lien + marquer traitée + supprimer
            save_lien(_cf_uid, _cf_service, _cf_url, _now_cf.strftime("%d/%m/%Y"))
            supabase.table("demandes").update({"status": "auto_done"}).eq("id", _cf_id).execute()
            delete_demande(_cf_id)

            # Notifier admin
            _cf_email = st.session_state["db"]["users"].get(_cf_uid, {}).get("email", "")
            notifier_livraison_gemini(_cf_uid, _cf_wa, _cf_email, _cf_service, _cf_nom, demande_complete=_cf.get("description", ""))

    except Exception as _e_cf_global:
        try:
            supabase.table("config").upsert({
                "key": "cf_auto_last_error",
                "value": f"{type(_e_cf_global).__name__}: {str(_e_cf_global)[:300]}"
            }).execute()
        except:
            pass

    # ── VÉRIFICATION AUTO-REPLY GRATUIT (tourne à chaque refresh) ──────
    # ── AUTO-REPLY PLAN GRATUIT — tourne pour TOUS les visiteurs ──
    if st.session_state.get("auto_reply_gratuit", False):
        try:
            _fresh_demandes = supabase.table("demandes").select("*").execute().data
            _now = datetime.now()

            # Services éligibles auto-reply gratuit
            _SERVICES_AUTO_GRATUIT = [
                "📝 Création de Sujets & Examens",
                "📖 Fiche de Cours Professeur IA",
                "👔 CV & Lettre de Motivation",
                "📄 Création Word (depuis zéro)",
                "📋 Rapport de Stage IA",
            ]
            # Définition locale pour éviter les NameError
            _SERVICES_GEMINI_LOCAL = [
                "📝 Exposé scolaire complet IA",
                "📝 Création de Sujets & Examens",
                "📖 Fiche de Cours Professeur IA",
                "👔 CV & Lettre de Motivation",
                "📄 Création Word (depuis zéro)",
                "📋 Rapport de Stage IA",
            ]

            for _req in _fresh_demandes:
                if _req.get("service", "") not in _SERVICES_AUTO_GRATUIT:
                    continue

                # Seulement les demandes des utilisateurs NON premium
                _req_user = _req.get("uid", "")
                _req_user_data = supabase.table("users").select("*").eq("uid", _req_user).execute().data
                _req_user_data = _req_user_data[0] if _req_user_data else {}
                _is_premium_req = is_premium_actif(_req_user_data)
                if _is_premium_req:
                    continue

                # Vérifier le délai variable (6, 8 ou 10 min selon l'ID de la demande)
                _ts_str = _req.get("timestamp", "")
                _req_id_hash = hash(_req.get("id", "")) % 3  # 0, 1 ou 2
                _delai_minutes = [6, 8, 10][_req_id_hash]
                try:
                    _ts = datetime.fromisoformat(_ts_str)
                    _age_minutes = (_now - _ts).total_seconds() / 60
                except:
                    continue
                if _age_minutes < _delai_minutes:
                    continue  # Pas encore le délai requis

                # Vérifier que le service est supporté
                _service_req = _req.get("service", "")
                if _service_req not in _SERVICES_GEMINI_LOCAL:
                    continue

                # Vérifier pas déjà traité
                if _req.get("status") == "auto_done":
                    continue
                # ── Générer automatiquement ──────────────────────────────
                _desc_req   = _req.get("description", "")
                _client_req = _req.get("uid", "")
                _wa_req     = _req.get("whatsapp", "")
                _req_id     = _req.get("id", "")
                _contenu_auto = generer_avec_gemini(_service_req, _desc_req, _client_req, is_premium=False)
                if not _contenu_auto or _contenu_auto.startswith("❌"):
                    # Log l'erreur dans Supabase pour debug
                    try:
                        supabase.table("config").upsert({
                            "key": f"auto_error_{_req_id}",
                            "value": _contenu_auto or "Réponse vide"
                        }).execute()
                    except:
                        pass
                    continue
                # Créer le fichier docx
                try:
                    _buf_auto = creer_docx(_contenu_auto, _service_req, _client_req)
                except Exception as _e_docx:
                    try:
                        supabase.table("config").upsert({
                            "key": f"auto_error_{_req_id}",
                            "value": f"creer_docx échoué : {_e_docx}"
                        }).execute()
                    except:
                        pass
                    continue
                _nom_auto = f"{_client_req}_{_service_req[:20].strip()}_auto.docx".replace(" ", "_").replace("/", "-")
                # Upload vers Supabase Storage
                _url_auto = upload_fichier_client(_client_req, _req_id, _buf_auto, _nom_auto)
                if not _url_auto or _url_auto.startswith("ERREUR"):
                    try:
                        supabase.table("config").upsert({
                            "key": f"auto_error_{_req_id}",
                            "value": f"Upload échoué : {_url_auto}"
                        }).execute()
                    except:
                        pass
                    continue
                # Sauvegarder le lien dans les livrables
                save_lien(_client_req, _service_req, _url_auto, _now.strftime("%d/%m/%Y"))
                # Marquer la demande comme traitée et la supprimer
                supabase.table("demandes").update({"status": "auto_done"}).eq("id", _req_id).execute()
                delete_demande(_req_id)
                # Notifier admin + client (auto-gratuit)
                _email_auto = st.session_state["db"]["users"].get(_client_req, {}).get("email", "")
                notifier_livraison_gemini(_client_req, _wa_req, _email_auto, _service_req, _nom_auto, demande_complete=_desc_req)
        except Exception as _e_auto:
            # Log l'erreur globale pour debug
            try:
                supabase.table("config").upsert({
                    "key": "auto_reply_last_error",
                    "value": f"{type(_e_auto).__name__}: {str(_e_auto)[:300]}"
                }).execute()
            except:
                pass

    tab1, tab2 = st.tabs(["🚀 DÉPLOYER UNE TÂCHE", "📂 MES LIVRABLES (CLOUD)"])

    SERVICES_GEMINI = [
        "📝 Exposé scolaire complet IA",
        "📝 Création de Sujets & Examens",
        "📖 Fiche de Cours Professeur IA",
        "👔 CV & Lettre de Motivation",
        "📄 Création Word (depuis zéro)",
        "📋 Rapport de Stage IA",
        "📊 Data & Excel Analytics",
        "📎 Modifier mon Fichier (Word / Excel / PPT)",
    ]

    with tab1:
        type_sujet_selectionne = None  # Initialisé ici, redéfini si service Sujets/Examens
        if st.session_state["premium_livrable"]:
            lv = st.session_state["premium_livrable"]
            st.markdown(f"""
            <div class="livrable-auto">
                <div class="livrable-auto-title">✅ Votre document est prêt !</div>
                <div style="color:rgba(255,255,255,.7);margin-top:6px;">
                    Généré en {lv['duree']}s · Service : <b>{lv['service']}</b>
                </div>
            </div>""", unsafe_allow_html=True)
            st.download_button(
                label="📥 TÉLÉCHARGER MON DOCUMENT",
                data=lv["buf"], file_name=lv["nom"], mime=lv["mime"],
                use_container_width=True
            )
            st.info("💡 Votre fichier est aussi disponible dans **📂 Mes Livrables** ci-dessus.")
            st.markdown(f"""
            <div style="background:rgba(37,211,102,0.1);border:1px solid rgba(37,211,102,0.4);border-radius:10px;padding:12px 16px;margin-top:10px;text-align:center;">
                <div style="font-size:1.05rem;font-weight:700;color:#25D366;">📲 Un fichier encore mieux vous attend !</div>
                <div style="color:rgba(255,255,255,0.75);font-size:0.88rem;margin-top:5px;">
                    Notre équipe retravaille votre document et vous envoie une version améliorée et mise en page sur WhatsApp sous peu.
                </div>
            </div>
            """, unsafe_allow_html=True)
            if st.button("🔄 Nouvelle mission", key="reset_livrable"):
                st.session_state["premium_livrable"] = None
                st.rerun()
            st.stop()

        SERVICE_PREREQUIS = {
            "📄 Création Word (depuis zéro)": {
                "icone": "📄",
                "titre": "Création Word — Document sur mesure",
                "intro": "Pour créer votre document Word parfaitement adapté, décrivez-nous :",
                "items": [
                    ("📋", "Le type de document (contrat, rapport, lettre, procédure, formulaire...)"),
                    ("🎯", "Le sujet et le contenu principal à rédiger"),
                    ("📏", "La longueur souhaitée (nombre de pages ou sections)"),
                    ("🎨", "Le ton souhaité (formel, commercial, juridique, amical...)"),
                ],
                "note": "Plus votre description est précise, plus le document sera prêt à l'emploi sans retouches."
            },
            "📚 Affiches & Reçus": {
                "icone": "📚",
                "titre": "Affiches & Reçus",
                "intro": "Pour concevoir votre affiche ou reçu, nous aurons besoin de :",
                "items": [
                    ("🏢", "Le nom de votre entreprise ou organisation"),
                    ("📋", "Les informations à faire apparaître (prix, date, lieu, contacts...)"),
                    ("🎨", "La couleur principale ou identité visuelle"),
                    ("📐", "Le format désiré (A4, A5, reçu thermique...)"),
                ],
                "note": "Un logo ou image à intégrer peut être envoyé via WhatsApp."
            },
            "📋 Rapport de Stage IA": {
                "icone": "📋",
                "titre": "Rapport de Stage IA",
                "intro": "Pour rédiger un rapport de stage complet et professionnel, indiquez-nous :",
                "items": [
                    ("🏢", "Le nom de l'entreprise ou organisation d'accueil"),
                    ("📅", "La durée du stage (dates de début et de fin)"),
                    ("🎓", "Votre niveau d'études et filière (BTS, Licence, Master...)"),
                    ("💼", "Vos missions principales durant le stage"),
                    ("📝", "Vos observations et apprentissages clés"),
                ],
                "note": "Plus vous décrivez vos missions en détail, plus le rapport sera réaliste et personnalisé."
            },
            "👔 CV & Lettre de Motivation": {
                "icone": "👔",
                "titre": "CV & Lettre de Motivation",
                "intro": "Pour rédiger votre CV ou lettre de motivation de façon percutante, fournissez :",
                "items": [
                    ("👤", "Votre nom complet et coordonnées"),
                    ("🎓", "Vos diplômes et formations"),
                    ("💼", "Vos expériences professionnelles"),
                    ("🎯", "Le poste ou secteur visé"),
                    ("✨", "Vos compétences clés et atouts"),
                ],
                "note": "Précisez si vous souhaitez uniquement le CV, la lettre, ou les deux."
            },
        }

        # ── SESSION STATE liste services ──
        if "show_services_list" not in st.session_state:
            st.session_state["show_services_list"] = False
        if "service_choisi" not in st.session_state:
            st.session_state["service_choisi"] = ""

        TOUS_SERVICES = [
            "📊 Data & Excel Analytics",
            "📖 Fiche de Cours Professeur IA",
            "📎 Modifier mon Fichier (Word / Excel / PPT)",
            "📝 Exposé scolaire complet IA",
            "📝 Création de Sujets & Examens",
            "📄 Création Word (depuis zéro)",
            "📋 Rapport de Stage IA",
            "📚 Affiches & Reçus",
            "👔 CV & Lettre de Motivation",
            "📄 Conversion & Fichier PDF",
            "🔍 OCR — Numérisation de Document",
        ]

        # ── GRILLE UNIFIÉE : tout le monde choisit un service → Nova IA ──────────
        st.markdown("#### 🛠️ Choisis ton service")
        st.markdown("<div style='color:rgba(255,255,255,0.5);font-size:0.85rem;margin-bottom:12px;'>Clique sur un service pour démarrer 🤖</div>", unsafe_allow_html=True)

        # CSS mobile-first pour les boutons de service
        st.markdown("""
        <style>
        /* Boutons service — compacts sur mobile */
        div[data-testid="stButton"].nova-svc-btn > button {
            font-size: 0.82rem !important;
            padding: 6px 10px !important;
            min-height: 38px !important;
            height: auto !important;
            white-space: nowrap !important;
            overflow: hidden !important;
            text-overflow: ellipsis !important;
            border-radius: 10px !important;
            font-weight: 600 !important;
        }
        @media (max-width: 640px) {
            div[data-testid="stButton"].nova-svc-btn > button {
                font-size: 0.78rem !important;
                padding: 5px 8px !important;
                min-height: 34px !important;
            }
        }
        </style>
        """, unsafe_allow_html=True)

        _cols = st.columns(2, gap="small")
        for _i, _svc in enumerate(TOUS_SERVICES):
            st.markdown('<div class="nova-svc-btn">', unsafe_allow_html=True)
            with _cols[_i % 2]:
                if st.button(_svc, key=f"pick_svc_{_i}", use_container_width=True):
                    _service_court = _svc.split(" ", 1)[-1] if " " in _svc else _svc

                    # ── OCR et Conversion → 1er essai toléré, 2e = auth requise ──
                    if "OCR" in _svc or "Numérisation" in _svc or "Conversion" in _svc:
                        if user:
                            st.session_state["service_choisi"] = _svc
                            st.rerun()
                        elif st.session_state["ocr_conv_uses"] == 0:
                            st.session_state["ocr_conv_uses"] = 1
                            st.session_state["service_choisi"] = _svc
                            st.rerun()
                        else:
                            st.session_state["service_choisi"] = _svc
                            st.session_state["view"] = "auth"
                            st.rerun()

                    else:
                        st.session_state.pop("nova_ia_chat", None)
                        st.session_state.pop("nova_ia_phase", None)
                        st.session_state.pop("nova_ia_service_detecte", None)
                        st.session_state.pop("nova_ia_prompt_final", None)
                        if user:
                            # Connecté → Nova IA traite directement
                            st.session_state["nova_ia_chat"] = [
                                {"role": "assistant", "content": f"Salut ! Je vois que tu veux utiliser le service **{_service_court}**. Dis-moi exactement ce que tu veux, je m'occupe du reste 🚀"}
                            ]
                        else:
                            # Visiteur → Nova IA lui demande de se connecter
                            st.session_state["nova_ia_chat"] = [
                                {"role": "assistant", "content": f"Salut 👋 Je vois que tu veux utiliser le service **{_service_court}**.\n\nMais avant de continuer, tu dois **créer un compte ou te connecter** sur Nova Platform — c'est gratuit et ça prend 30 secondes !\n\nUne fois connecté, reviens ici et je m'occupe de tout 🚀"}
                            ]
                        st.session_state["nova_ia_phase"] = "dialogue"
                        st.session_state["nova_ia_service_preselect"] = _svc
                        st.session_state["service_choisi"] = _svc
                        st.session_state["view"] = "nova_ia"
                        st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)

        service = st.session_state.get("service_choisi", TOUS_SERVICES[0])
        SERVICE_SAISIE = "📊 Data & Excel Analytics"

        # wa_display — récupéré depuis le profil utilisateur (plus de champ de saisie)
        wa_display = db["users"][user]["whatsapp"] if user and user in db["users"] else ""

        # ── BOUTON AIDE FLOTTANT ────────────────────────────────────────
        SERVICE_AIDE = {
            "📊 Data & Excel Analytics": ("📊 Data & Excel Analytics", "Soumettez vos données brutes et décrivez l'analyse souhaitée. Nova Platform génère tableaux croisés, graphiques et rapports Excel.\n\n✅ Idéal pour : bilans, stats, dashboards\n📎 Collez vos données dans le cahier des charges"),
            "📖 Fiche de Cours Professeur IA": ("📖 Fiche de Cours", "Indiquez la matière, le niveau, le chapitre et l'objectif. Nova Platform rédige une fiche complète avec notions, exemples et exercices.\n\n✅ Idéal pour : enseignants, formateurs\n💡 Précisez le programme (MENET-FP, LMD...)"),
            "📎 Modifier mon Fichier (Word / Excel / PPT)": ("📎 Modifier mon Fichier", "Importez votre fichier et décrivez précisément les modifications souhaitées.\n\n✅ Formats : Word, Excel, PowerPoint\n⚠️ Soyez précis sur ce que vous voulez changer"),
            "📝 Exposé scolaire complet IA": ("📝 Exposé scolaire", "Remplissez le formulaire : niveau, matière, sujet, pages. Nova Platform rédige un exposé structuré complet.\n\n✅ Idéal pour : collège, lycée, université\n⭐ Service PREMIUM uniquement"),
            "📝 Création de Sujets & Examens": ("📝 Sujets & Examens", "Choisissez niveau, matière, type d'épreuve et durée. Importez votre cours pour baser le sujet dessus.\n\n✅ Types : QCM, Vrai/Faux, Cas pratique, Devoir complet\n⭐ Génération auto avec Premium"),
            "📄 Création Word (depuis zéro)": ("📄 Création Word", "Décrivez précisément le document Word que vous voulez : type, contenu, ton, longueur. Nova Platform le crée de A à Z, prêt à l'emploi.\n\n✅ Contrats, rapports, lettres, procédures, formulaires...\n📝 Plus vous êtes précis, plus le résultat est parfait"),
            "📋 Rapport de Stage IA": ("📋 Rapport de Stage", "Décrivez votre stage : entreprise, durée, niveau d'études, missions effectuées. Nova Platform rédige un rapport complet et structuré.\n\n✅ Introduction · Présentation de l'entreprise · Missions · Analyse · Conclusion\n🎓 Adapté BTS, Licence, Master"),
            "📚 Affiches & Reçus": ("📚 Affiches & Reçus", "Précisez le type (affiche, reçu, bon de commande...) et les infos à afficher.\n\n✅ Idéal pour : commerces, associations, événements\n📋 Fournissez les données exactes"),
            "👔 CV & Lettre de Motivation": ("👔 CV & Lettre", "Indiquez votre parcours, le poste visé et l'entreprise cible.\n\n✅ Formats modernes et professionnels\n💡 Précisez si vous avez déjà un CV à améliorer"),
            "📄 Conversion & Fichier PDF": ("📄 Conversion PDF", "Importez votre fichier et choisissez le format cible.\n\n✅ Formats : Word↔PDF, Excel↔CSV, PPT↔PDF\n⚡ Résultat immédiat"),
        }

        if service in SERVICE_AIDE:
            _aide_titre, _aide_texte = SERVICE_AIDE[service]
            if f"aide_open_{service}" not in st.session_state:
                st.session_state[f"aide_open_{service}"] = True

            _aide_ouvert = st.session_state[f"aide_open_{service}"]
            st.markdown("""
            <style>
            @keyframes aide-pulse {
                0%,100% {
                    box-shadow: 0 0 0 0 rgba(255,215,0,0.0);
                    border-color: rgba(255,193,7,0.6);
                }
                50% {
                    box-shadow: 0 0 0 6px rgba(255,215,0,0.18), 0 0 18px rgba(255,215,0,0.35);
                    border-color: #FFD700;
                }
            }
            div[data-testid="stButton"].nova-aide-pulse > button {
                background: linear-gradient(135deg, rgba(255,215,0,0.18), rgba(255,140,0,0.12)) !important;
                border: 2px solid #FFD700 !important;
                border-radius: 20px !important;
                color: #FFD700 !important;
                font-weight: 800 !important;
                font-size: 0.82rem !important;
                letter-spacing: 0.03em !important;
                animation: aide-pulse 1.8s ease-in-out infinite !important;
                padding: 4px 14px !important;
            }
            div[data-testid="stButton"].nova-aide-pulse > button:hover {
                background: linear-gradient(135deg, rgba(255,215,0,0.3), rgba(255,140,0,0.2)) !important;
                animation: none !important;
                box-shadow: 0 0 20px rgba(255,215,0,0.5) !important;
            }
            </style>
            """, unsafe_allow_html=True)

            col_aide_l, col_aide_r = st.columns([6, 1])
            with col_aide_r:
                _key_aide = f"btn_aide_{service[:10]}"
                if not _aide_ouvert:
                    st.markdown('<div class="nova-aide-pulse">', unsafe_allow_html=True)
                if st.button("✨ Voir" if not _aide_ouvert else "✕ Fermer", key=_key_aide, help=f"Découvrir ce service : {_aide_titre}"):
                    st.session_state[f"aide_open_{service}"] = not _aide_ouvert
                    st.rerun()
                if not _aide_ouvert:
                    st.markdown('</div>', unsafe_allow_html=True)

            if st.session_state[f"aide_open_{service}"]:
                st.markdown(f"""
                <div style="background:linear-gradient(135deg,rgba(255,193,7,0.08),rgba(255,140,0,0.05));
                     border:1px solid rgba(255,193,7,0.35);border-left:3px solid #FFD700;
                     border-radius:10px;padding:14px 18px;margin:4px 0 12px 0;">
                    <span style="color:#FFD700;font-weight:800;font-size:.95rem;">💡 {_aide_titre}</span>
                    <pre style="color:rgba(255,255,255,0.85);font-family:inherit;font-size:.88rem;
                         margin:8px 0 0 0;white-space:pre-wrap;line-height:1.6;">{_aide_texte}</pre>
                </div>
                """, unsafe_allow_html=True)

        if service != st.session_state["last_service_seen"]:
            st.session_state["last_service_seen"] = service
            st.session_state["warning_triggered"] = False
            st.session_state["show_service_warning"] = False
            if service != SERVICE_SAISIE and service in SERVICE_PREREQUIS:
                st.session_state["show_service_warning"] = True

        if st.session_state["show_service_warning"] and service in SERVICE_PREREQUIS and service != SERVICE_SAISIE:
            info = SERVICE_PREREQUIS[service]

            SERVICE_AUDIO = {
                "📄 Création Word (depuis zéro)": "prerequis_word.mp3",
                "📋 Rapport de Stage IA":          "prerequis_rapport_stage.mp3",
                "📚 Affiches & Reçus":             "prerequis_affiches.mp3",
                "👔 CV & Lettre de Motivation":    "prerequis_cv.mp3",
            }

            st.info(f"""
**{info["icone"]} {info["titre"]} — Informations requises**

{info["intro"]}

{"".join(f"- {icone} {texte}\n" for icone, texte in info["items"])}
💡 *{info["note"]}*
""")
            audio_file = SERVICE_AUDIO.get(service) if service in SERVICE_AUDIO else None
            if audio_file and os.path.exists(audio_file):
                with open(audio_file, "rb") as f:
                    b64 = __import__('base64').b64encode(f.read()).decode()
                components.html(f"""
                    <script>
                    (function() {{
                        var binary = atob("{b64}");
                        var bytes = new Uint8Array(binary.length);
                        for (var i = 0; i < binary.length; i++) bytes[i] = binary.charCodeAt(i);
                        var blob = new Blob([bytes], {{type: "audio/mpeg"}});
                        var audio = new Audio(URL.createObjectURL(blob));
                        audio.volume = 1;
                        audio.play().catch(function(e) {{ console.log(e); }});
                    }})();
                    </script>
                """, height=1)
            col_mid = st.columns([1, 2, 1])[1]
            with col_mid:
                if st.button("✅ J'ai compris, je continue ma demande", key="close_service_warning"):
                    st.session_state["show_service_warning"] = False
                    # speechSynthesis supprimé
                    st.rerun()

        # ── SÉLECTION DU TYPE DE SUJET (uniquement pour le service Sujets/Examens) ──
        type_sujet_selectionne = None
        if "Sujets" in service or "Examens" in service:
            _show_splash("Sujets")
            st.markdown("#### 🎯 Type de sujet")
            TYPES_SUJETS = {
                "🔵 QCM — Questions à Choix Multiple": "QCM",
                "✅ Vrai ou Faux (avec justification)": "VRAI_FAUX",
                "🔤 Texte à Trous (lacunaire)": "TEXTE_TROU",
                "✍️ Questions Ouvertes (rédigées)": "QUESTIONS_OUVERTES",
                "🔀 Mixte (QCM + Vrai/Faux + Question ouverte)": "MIXTE",
                "📋 Cas Pratique / Étude de Cas": "CAS_PRATIQUE",
                "📐 Exercices de Calcul / Problèmes": "CALCUL",
                "🗺️ Étude de Document (texte, tableau, carte)": "ETUDE_DOCUMENT",
                "🔬 Schéma à Légender / Identification": "SCHEMA",
                "📝 Composition / Dissertation guidée": "DISSERTATION",
                "📄 Devoir Complet (comme sur les images)": "DEVOIR_COMPLET",
            }
            type_sujet_label = st.selectbox(
                "Choisissez le type d'exercice que vous voulez dans votre sujet",
                list(TYPES_SUJETS.keys()),
                help="Sélectionnez précisément le type de sujet souhaité. Nova Platform adaptera 100% du contenu à ce format."
            )
            type_sujet_selectionne = TYPES_SUJETS[type_sujet_label]

            TYPE_SUJET_DESCRIPTIONS = {
                "QCM": "**QCM sélectionné** — Nova Platform générera des questions à 4 choix (A/B/C/D) avec cases □ à cocher, distracteurs réalistes et corrigé si demandé.",
                "VRAI_FAUX": "**Vrai ou Faux sélectionné** — Nova Platform générera des affirmations à évaluer (V/F) avec lignes de justification pour les fausses réponses.",
                "TEXTE_TROU": "**Texte à trous sélectionné** — Nova Platform rédigera un texte cohérent avec des blancs à remplir et une liste de mots fournie.",
                "QUESTIONS_OUVERTES": "**Questions ouvertes sélectionnées** — Nova Platform formulera des questions de réflexion avec lignes de réponse proportionnelles au barème.",
                "MIXTE": "**Format Mixte sélectionné** — Nova Platform combinera QCM (Partie 1) + Vrai/Faux (Partie 2) + Question rédigée (Partie 3), barème équilibré.",
                "CAS_PRATIQUE": "**Cas Pratique sélectionné** — Nova Platform rédigera un texte/document contextualisé (Côte d'Ivoire) + questions d'analyse progressives.",
                "CALCUL": "**Exercices de Calcul sélectionnés** — Nova Platform rédigera des problèmes chiffrés contextualisés avec démarche guidée, formules rappelées et données réelles ivoiriennes.",
                "ETUDE_DOCUMENT": "**Étude de Document sélectionnée** — Nova Platform créera un document support (texte, tableau ou description de carte) + questions d'identification, analyse et interprétation.",
                "SCHEMA": "**Schéma à légender sélectionné** — Nova Platform décrira textuellement un schéma numéroté avec la liste des termes à placer et un corrigé de légendes.",
                "DISSERTATION": "**Dissertation guidée sélectionnée** — Nova Platform formulera un sujet de composition, fournira des consignes de méthode et proposera un plan détaillé guidé.",
                "DEVOIR_COMPLET": "**Devoir Complet sélectionné** — Nova Platform générera un vrai devoir ivoirien complet avec en-tête officiel + exercices variés progressifs (QCM → mise en situation → problème complexe) adaptés exactement au niveau et à la matière.",
            }
            st.info(TYPE_SUJET_DESCRIPTIONS.get(type_sujet_selectionne, ""))

        st.markdown("#### 📝 Spécifications de la mission")

        # Initialisations (re-évaluées à chaque run Streamlit)
        _niveau_val      = ""
        _matiere_val     = ""
        _fc_niveau_val   = ""
        _fc_matiere_val  = ""
        mf_fichier       = None
        mf_instructions  = ""
        conv_fichier     = None

        # ── FORMULAIRE STRUCTURÉ POUR SUJETS & EXAMENS ────────────────────────
        if "Sujets" in service or "Examens" in service:

            # ── TOGGLE FICHIER EN PREMIER ──────────────────────────────────
            st.markdown("""
            <style>
            @keyframes glowPulse2 {
                0%   { box-shadow: 0 0 6px 2px rgba(46,204,113,0.4), 0 0 12px 4px rgba(46,204,113,0.2); }
                50%  { box-shadow: 0 0 18px 6px rgba(46,204,113,0.9), 0 0 35px 12px rgba(46,204,113,0.4); }
                100% { box-shadow: 0 0 6px 2px rgba(46,204,113,0.4), 0 0 12px 4px rgba(46,204,113,0.2); }
            }
            @keyframes textPulse2 {
                0%   { opacity: 1; }
                50%  { opacity: 0.75; }
                100% { opacity: 1; }
            }
            .fichier-toggle-card2 {
                background: linear-gradient(135deg, rgba(46,204,113,0.12), rgba(39,174,96,0.08));
                border: 2px solid rgba(46,204,113,0.7);
                border-radius: 14px;
                padding: 16px 20px;
                margin: 0 0 6px 0;
                animation: glowPulse2 2s ease-in-out infinite;
            }
            .fichier-toggle-title2 {
                color: #2ecc71;
                font-weight: 800;
                font-size: 1.05rem;
                animation: textPulse2 2s ease-in-out infinite;
                display: block;
            }
            .fichier-toggle-sub2 {
                color: rgba(255,255,255,0.55);
                font-size: 0.8rem;
                margin-top: 4px;
                display: block;
            }
            </style>
            <div class="fichier-toggle-card2">
                <span class="fichier-toggle-title2">📂 ✨ Créer le sujet à partir d'un de mes fichiers</span>
                <span class="fichier-toggle-sub2">Importez votre cours ou leçon — Nova Platform génère un sujet basé uniquement sur votre document</span>
            </div>
            """, unsafe_allow_html=True)
            use_fichier_source = st.toggle(
                "✅ Activer cette option",
                key="use_fichier_source",
                help="Importez un fichier Word, PDF ou TXT — Nova Platform analysera son contenu."
            )

            if not use_fichier_source:
                st.markdown("""
                <div style="background:rgba(66,133,244,0.08);border:1px solid rgba(66,133,244,0.3);
                     border-radius:12px;padding:14px 18px;margin:14px 0;">
                    <span style="color:#4285f4;font-weight:700;">📋 Remplissez les champs ci-dessous — Nova s'appuie sur ces informations précises pour générer votre sujet</span>
                </div>
                """, unsafe_allow_html=True)
                col_a, col_b = st.columns(2)
            if not use_fichier_source:
                col_a, col_b = st.columns(2)
                with col_a:
                    exam_niveau = st.selectbox(
                        "🎓 Niveau scolaire *",
                        [
                            "── PRIMAIRE ──",
                            "CP1", "CP2", "CE1", "CE2", "CM1", "CM2 / CEPE",
                            "── COLLÈGE ──",
                            "6ème", "5ème", "4ème", "3ème / BEPC",
                            "── LYCÉE ──",
                            "2nde", "1ère - Série A1", "1ère - Série A2", "1ère - Série B",
                            "1ère - Série C", "1ère - Série D", "1ère - Série E",
                            "Terminale - Série A1", "Terminale - Série A2", "Terminale - Série B",
                            "Terminale - Série C", "Terminale - Série D", "Terminale - Série E",
                            "Terminale - Série F", "Terminale - Série G1", "Terminale - Série G2",
                            "Terminale - Série G3", "Terminale - Série H",
                            "── UNIVERSITÉ ──",
                            "Licence 1 (L1)", "Licence 2 (L2)", "Licence 3 (L3)",
                            "Master 1 (M1)", "Master 2 (M2)", "Doctorat",
                            "── CONCOURS ──",
                            "Concours ENS", "Concours CAFOP", "Concours INJS",
                            "Concours Fonction Publique", "Concours Douane / Police / Armée",
                            "Autre concours professionnel",
                        ],
                        index=0
                    )
                    exam_matiere = st.selectbox(
                        "📚 Matière / Discipline *",
                        [
                            "── TOUTES MATIÈRES ──",
                            "Français / Lettres", "Mathématiques",
                            "Sciences Physiques (PC)", "SVT / Biologie",
                            "Histoire-Géographie", "Économie / Gestion",
                            "Comptabilité", "Philosophie",
                            "EDHC / Éducation Civique",
                            "Anglais (LV1)", "Espagnol (LV2)", "Allemand (LV2)",
                            "Informatique / TIC",
                            "Technologie industrielle",
                            "EPS (Éducation Physique)",
                            "Arts Plastiques",
                            "Agronomie / Agriculture",
                            "Droit", "Économie politique",
                            "── PRIMAIRE ──",
                            "Lecture / Écriture (primaire)", "Calcul (primaire)",
                            "Sciences d'Éveil (primaire)",
                            "Histoire-Géo (primaire)", "ECM (primaire)",
                            "Autre matière (préciser dans les notes)",
                        ],
                        index=0
                    )
                with col_b:
                    exam_type_epreuve = st.selectbox(
                        "🎯 Type d'épreuve *",
                        [
                            "Devoir Surveillé (DS)", "Interrogation Écrite (IE)",
                            "Devoir de Maison (DM)", "Devoir du 1er Trimestre",
                            "Devoir du 2ème Trimestre", "Devoir du 3ème Trimestre",
                            "Examen Blanc / Blanc CEPE", "Examen Blanc / Brevet Blanc (BEPC)",
                            "Examen Blanc / BAC Blanc", "Épreuve de Rattrapage",
                            "Sujet de Concours", "Épreuve de Passage",
                            "Exercice de classe (rapide)",
                        ],
                        index=0
                    )
                    exam_duree = st.selectbox(
                        "⏱️ Durée prévue *",
                        [
                            "15 minutes", "30 minutes", "1 heure",
                            "1 heure 30", "2 heures", "2 heures 30",
                            "3 heures", "3 heures 30", "4 heures",
                            "Durée libre (DM / à domicile)",
                        ],
                        index=2
                    )
                col_c, col_d = st.columns(2)
                with col_c:
                    exam_nb_exercices = st.selectbox(
                        "📏 Nombre d'exercices / questions *",
                        [
                            "1 exercice", "2 exercices", "3 exercices",
                            "4 exercices", "5 exercices",
                            "10 questions", "15 questions", "20 questions",
                            "25 questions", "30 questions",
                            "Adapté automatiquement au niveau et à la durée",
                        ],
                        index=10
                    )
                with col_d:
                    exam_coefficient = st.selectbox(
                        "🔢 Coefficient",
                        ["1", "2", "3", "4", "5", "6", "7", "8"],
                        index=1
                    )
                col_e, col_f = st.columns(2)
                with col_e:
                    exam_etablissement = st.text_input(
                        "🏢 Établissement / Institution",
                        placeholder="Ex: Lycée Moderne de Cocody, CEG Treichville, UFHB..."
                    )
                with col_f:
                    exam_annee = st.text_input(
                        "📅 Année scolaire",
                        placeholder="Ex: 2024-2025",
                        value="2024-2025"
                    )
                exam_chapitre = st.text_input(
                    "📖 Chapitre / Notion spécifique (optionnel)",
                    placeholder="Ex: Les fractions, La cellule, La colonisation, La dérivation, La loi d'Ohm..."
                )
                exam_notes = st.text_area(
                    "💬 Informations complémentaires (optionnel)",
                    height=80,
                    placeholder="Ex: Avec corrigé, thème ivoirien, niveau difficile, chapitres 1 et 2, 4 QCM + 2 ouvertes..."
                )
            else:
                # Réinitialiser le contenu fichier si toggle désactivé
                st.session_state["contenu_fichier_source"] = ""
                contenu_fichier_source = ""
                # Valeurs par défaut quand fichier actif
                exam_niveau = "Non précisé"
                exam_matiere = "── TOUTES MATIÈRES ──"
                exam_type_epreuve = "Devoir Surveillé (DS)"
                exam_duree = "1 heure"
                exam_nb_exercices = "Adapté automatiquement au niveau et à la durée"
                exam_coefficient = "2"
                exam_etablissement = ""
                exam_annee = "2024-2025"
                exam_chapitre = ""
                exam_notes = ""

            contenu_fichier_source = ""
            if use_fichier_source:
                st.markdown('''
                <div style="background:rgba(46,204,113,0.08);border:1px solid rgba(46,204,113,0.3);
                     border-radius:10px;padding:12px 16px;margin:8px 0;">
                    <span style="color:#2ecc71;font-weight:700;">📄 Nova Platform va scanner votre document et créer le sujet uniquement à partir de son contenu</span>
                    <span style="color:rgba(255,255,255,.5);font-size:.82rem;display:block;margin-top:3px;">
                        Formats acceptés : PDF, Word (.docx), Texte (.txt)
                    </span>
                </div>
                ''', unsafe_allow_html=True)

                fichier_source = st.file_uploader(
                    "📂 Importer votre fichier *",
                    type=["pdf", "docx", "txt"],
                    key="fichier_source_exam"
                )

                if fichier_source:
                    with st.spinner("🔍 Lecture du fichier en cours..."):
                        try:
                            import io
                            fichier_bytes = fichier_source.read()

                            if fichier_source.name.endswith(".txt"):
                                contenu_fichier_source = fichier_bytes.decode("utf-8", errors="ignore")

                            elif fichier_source.name.endswith(".docx"):
                                from docx import Document as DocxDoc
                                doc_tmp = DocxDoc(io.BytesIO(fichier_bytes))
                                contenu_fichier_source = "\n".join([p.text for p in doc_tmp.paragraphs if p.text.strip()])

                            elif fichier_source.name.endswith(".pdf"):
                                try:
                                    import fitz  # PyMuPDF
                                    pdf_doc = fitz.open(stream=fichier_bytes, filetype="pdf")
                                    contenu_fichier_source = "\n".join([page.get_text() for page in pdf_doc])
                                except ImportError:
                                    import subprocess
                                    subprocess.run(["pip", "install", "PyMuPDF", "--break-system-packages", "-q"])
                                    import fitz
                                    pdf_doc = fitz.open(stream=fichier_bytes, filetype="pdf")
                                    contenu_fichier_source = "\n".join([page.get_text() for page in pdf_doc])

                            # Tronquer si trop long (max ~6000 caractères pour le prompt)
                            if len(contenu_fichier_source) > 6000:
                                contenu_fichier_source = contenu_fichier_source[:6000] + "\n...[document tronqué]"

                            if contenu_fichier_source.strip():
                                st.session_state["contenu_fichier_source"] = contenu_fichier_source
                                st.success(f"✅ **{fichier_source.name}** lu avec succès — {len(contenu_fichier_source)} caractères extraits")
                            else:
                                st.warning("⚠️ Le fichier semble vide ou illisible")
                                contenu_fichier_source = ""
                                st.session_state["contenu_fichier_source"] = ""

                        except Exception as e_fic:
                            st.error(f"❌ Erreur lecture fichier : {e_fic}")
                            contenu_fichier_source = ""
                            st.session_state["contenu_fichier_source"] = ""
                else:
                    # Récupérer depuis session_state si déjà chargé
                    if st.session_state["contenu_fichier_source"]:
                        contenu_fichier_source = st.session_state["contenu_fichier_source"]
                        st.success(f"✅ Fichier en mémoire — {len(contenu_fichier_source)} caractères")
                    else:
                        st.info("⬆️ Importez votre fichier pour que Nova génère le sujet à partir de son contenu")

                # ── CAHIER DES CHARGES SPÉCIFIQUE FICHIER ─────────────────────
                st.markdown("""
                <div style="background:rgba(255,165,0,0.08);border:1px solid rgba(255,165,0,0.4);
                     border-radius:10px;padding:12px 16px;margin:10px 0 4px 0;">
                    <span style="color:#FFA500;font-weight:700;">📋 Instructions pour le sujet</span>
                    <span style="color:rgba(255,255,255,.5);font-size:.8rem;display:block;margin-top:2px;">
                        Nova Platform se base <b>uniquement sur votre fichier</b> — précisez ici ce que vous voulez comme sujet
                    </span>
                </div>
                """, unsafe_allow_html=True)
                fichier_instructions = st.text_area(
                    "✍️ Décrivez le sujet que vous voulez *",
                    height=110,
                    placeholder="Ex: Fais un QCM de 10 questions sur les définitions...\nEx: Crée un devoir avec 2 exercices de calcul basés sur les formules du document...\nEx: Génère une interrogation sur la partie 2 du cours uniquement...",
                    key="fichier_instructions"
                )

            # ── CONSTRUCTION AUTOMATIQUE DU PROMPT STRUCTURÉ ──────────────────
            _niveau_val = exam_niveau if not exam_niveau.startswith("──") else ""
            _matiere_val = exam_matiere if not exam_matiere.startswith("──") else ""

            _etab_val = exam_etablissement.strip() if exam_etablissement.strip() else "Établissement non précisé"
            _annee_val = exam_annee.strip() if exam_annee.strip() else "2024-2025"
            _fichier_instr = fichier_instructions.strip() if use_fichier_source and "fichier_instructions" in dir() else ""

            prompt = f"""FICHE DE COMMANDE NOVA EXAM — INFORMATIONS STRUCTURÉES :

🎓 NIVEAU SCOLAIRE     : {_niveau_val if _niveau_val else "Non précisé"}
📚 MATIÈRE             : {_matiere_val if _matiere_val else "Non précisée"}
🎯 TYPE D'ÉPREUVE      : {exam_type_epreuve}
📏 EXERCICES/QUESTIONS : {exam_nb_exercices}
⏱️ DURÉE               : {exam_duree}
🔢 COEFFICIENT         : {exam_coefficient}
🏢 ÉTABLISSEMENT       : {_etab_val}
📅 ANNÉE SCOLAIRE      : {_annee_val}
📖 CHAPITRE/NOTION     : {"BASÉ SUR LE FICHIER FOURNI — voir document source ci-dessous" if use_fichier_source else (exam_chapitre if exam_chapitre.strip() else "Choisir un chapitre cohérent avec le programme officiel du niveau")}
💬 INSTRUCTIONS CLIENT : {"" + _fichier_instr if use_fichier_source else (exam_notes.strip() if exam_notes.strip() else "Aucune")}

INSTRUCTIONS NOVA EXAM :
- Respecte EXACTEMENT le niveau "{_niveau_val}" — applique le programme officiel MENET-FP de cette classe
- Génère UNIQUEMENT des notions au programme de ce niveau — rien hors-programme
- Adapte le vocabulaire, la complexité et la longueur à l'âge de l'élève de ce niveau
- Si un chapitre/notion est précisé, le sujet porte EXCLUSIVEMENT sur ce chapitre
- Si "avec corrigé" dans les notes, inclure le corrigé complet après ---SAUT_DE_PAGE---

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
SECTION CRITIQUE — FORMAT AUTHENTIQUE DES DEVOIRS IVOIRIENS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

① EN-TÊTE OBLIGATOIRE — Reproduis ce modèle EXACTEMENT :

{_etab_val}* {_etab_val}*
Année Scolaire : {_annee_val}                    Coefficient : {exam_coefficient}
Classe : {_niveau_val if _niveau_val else "___"}                              Durée : {exam_duree}
[ligne vide]
┌─────────────────────────────────────────────┐
│    DEVOIR DE {(_matiere_val.upper() if _matiere_val else "___")}    │
└─────────────────────────────────────────────┘


RÈGLES EN-TÊTE :
- Nom établissement répété 2 fois sur la même ligne séparé par *
- Année + Coefficient sur la MÊME ligne (gauche / droite)
- Classe + Durée sur la MÊME ligne (gauche / droite)
- Titre DEVOIR encadré ou en gras centré
- NE JAMAIS ajouter REPUBLIQUE, MINISTERE, DIRECTION non demandés

② STRUCTURE DES EXERCICES — ADAPTE-TOI INTELLIGEMMENT AU NIVEAU DÉTECTÉ :

Le niveau demandé est : {_niveau_val}

Tu es un expert du système éducatif ivoirien (MENET-FP). Tu connais parfaitement comment les devoirs sont structurés à CHAQUE niveau en Côte d'Ivoire. Adapte la structure, la complexité, le vocabulaire et les types d'exercices exactement comme le ferait un vrai professeur ivoirien de ce niveau.

▶ PRIMAIRE (CP1 → CM2) :
- Exercices courts, simples, directs
- Pas de mise en situation complexe
- Dictée, calcul mental, lecture, problèmes simples du quotidien ivoirien
- Consignes en langage simple accessible à l'enfant
- 3 à 4 exercices maximum, jamais de sous-questions complexes

▶ COLLÈGE — 6ème / 5ème :
- Exercice 1 : QCM ou Vrai/Faux ou tableau à compléter (facile)
- Exercice 2 : Définitions, propriétés, applications directes du cours
- Exercice 3 : Problème avec mise en situation simple, questions progressives 1- 2- 3-
- Exercice 4 (si demandé) : Problème de la vie courante ivoirienne, niveau accessible
- Langage simple, données concrètes, pas de démonstrations formelles

▶ COLLÈGE — 4ème / 3ème (BEPC) :
- Exercice 1 : QCM / Vrai-Faux / affirmations V ou F (mise en route)
- Exercice 2 : Questions courtes directes (définitions, calculs simples, compléter)
- Exercice 3 : TEXTE DE MISE EN SITUATION + 4-5 questions progressives avec sous-questions (1.1 / 1.2 / 1.3)
- Exercice 4 : PROBLÈME COMPLEXE ancré dans la réalité ivoirienne (PME, plantation, lycée CI) — démonstrations, calculs multi-étapes
- Barème : ex1-2 = 30%, ex3-4 = 70% des points

▶ LYCÉE — 2nde / 1ère :
- Structure plus libre, moins de QCM, plus de rédaction et démonstration
- Exercice 1 : Questions de cours ou vérification des connaissances (définitions, théorèmes)
- Exercice 2 : Application directe, calculs guidés
- Exercice 3 : Problème ouvert avec situation complexe, plusieurs sous-parties (A, B, C)
- Exercice 4 : Problème de synthèse liant plusieurs notions du programme
- Attendu : justifications rigoureuses, raisonnement logique développé

▶ LYCÉE — TERMINALE / BAC :
- Format proche des épreuves officielles du BAC ivoirien
- 3 à 4 exercices longs avec parties A / B / C
- Chaque partie autonome mais liée au thème général
- Calculs complexes, démonstrations formelles, analyse critique
- Problèmes souvent en contexte scientifique ou socio-économique ivoirien
- Barème précis sur /20, questions numérotées I- II- III- ou 1) 2) 3)

▶ POST-BAC (BTS, Licence, Master) :
- Format académique supérieur : énoncé dense, problème unique divisé en parties
- Partie I / Partie II / Partie III avec sous-questions a) b) c)
- Niveau de rigueur élevé, démonstrations formelles exigées
- Mise en contexte professionnel ou scientifique sérieux
- Pas de QCM — uniquement questions ouvertes et problèmes

RÈGLES GÉNÉRALES POUR TOUS LES NIVEAUX :
- Chaque exercice : ## EXERCICE N :
- Données toujours précises avec vrais chiffres
- Contextes ivoiriens authentiques : noms ivoiriens, villes CI (Abidjan, Bouaké, Yamoussoukro...), produits locaux (cacao, café, anacarde...), monnaie FCFA
- NE JAMAIS inventer une structure générique — colle au vrai format du niveau
"""
            # Injecter le contenu du fichier source si fourni
            if contenu_fichier_source.strip():
                prompt += """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
🎯 DIRECTIVE PRINCIPALE — CE QUE LE CLIENT VEUT EXACTEMENT :
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
""" + (_fichier_instr if _fichier_instr else "Crée un sujet adapté au niveau et à la durée indiqués, basé sur le document ci-dessous.") + """

⚠️ RÈGLES ABSOLUES — RESPECTER JUSQU'À LA FIN :
1. Exécute EXACTEMENT la demande du client ci-dessus, du début jusqu'à la toute dernière ligne.
2. Si le client a précisé un format, un nombre de questions, un style ou une structure — respecte-le à 100% sans t'en écarter.
3. Ne t'arrête pas avant d'avoir TOUT produit tel que demandé. Si le client veut 20 QCM, génère 20 QCM complets.
4. Le document ci-dessous est ta SEULE source de contenu — puise les notions, définitions et formules UNIQUEMENT dans ce texte.
5. N'invente AUCUNE notion absente du document. Ne complète pas avec tes propres connaissances.
6. Respecte l'intégralité du cahier des charges fourni, sans ignorer aucun détail.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📄 DOCUMENT SOURCE (contenu à utiliser) :
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
""" + contenu_fichier_source + """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
            # Afficher un résumé de la commande
            if use_fichier_source:
                # Mode fichier : validation basée sur le fichier + instructions
                if contenu_fichier_source.strip() and _fichier_instr.strip():
                    st.success(f"✅ Commande prête — sujet généré à partir de votre fichier")
                elif not contenu_fichier_source.strip():
                    st.warning("⚠️ Importez votre fichier pour continuer")
                else:
                    st.warning("⚠️ Décrivez le sujet que vous voulez dans le cahier des charges")
            else:
                if _niveau_val and _matiere_val and not _niveau_val.startswith("──") and not _matiere_val.startswith("──"):
                    st.success(f"✅ Commande prête : **{_matiere_val}** · **{_niveau_val}** · **{exam_type_epreuve}** · **{exam_duree}**")
                else:
                    if not _niveau_val or _niveau_val.startswith("──"):
                        st.warning("⚠️ Sélectionnez un niveau scolaire précis (pas le titre de catégorie)")
                    if not _matiere_val or _matiere_val.startswith("──"):
                        st.warning("⚠️ Sélectionnez une matière précise (pas le titre de catégorie)")

        elif "Exposé" in service:
            _show_splash("Expose")
            st.markdown('''
            <div style="background:rgba(255,165,0,0.08);border:1px solid rgba(255,165,0,0.35);
                 border-radius:12px;padding:14px 18px;margin-bottom:14px;">
                <span style="font-weight:700;color:#FFA500;">📝 Remplissez les champs — Nova génère votre exposé complet, structuré et prêt à présenter</span>
            </div>
            ''', unsafe_allow_html=True)

            col_a, col_b = st.columns(2)
            with col_a:
                exp_niveau = st.selectbox("🎓 Niveau scolaire *", [
                    "── PRIMAIRE ──","CP1","CP2","CE1","CE2","CM1","CM2",
                    "── COLLÈGE ──","6ème","5ème","4ème","3ème",
                    "── LYCÉE ──","2nde",
                    "1ère - Série A1","1ère - Série A2","1ère - Série B",
                    "1ère - Série C","1ère - Série D","1ère - Série E",
                    "Terminale - Série A1","Terminale - Série A2","Terminale - Série B",
                    "Terminale - Série C","Terminale - Série D","Terminale - Série E",
                    "── UNIVERSITÉ / BTS ──",
                    "Licence 1 (L1)","Licence 2 (L2)","Licence 3 (L3)",
                    "Master 1 (M1)","Master 2 (M2)",
                    "BTS 1ère année","BTS 2ème année",
                ], index=0, key="exp_niveau")
                exp_matiere = st.selectbox("📚 Matière / Discipline *", [
                    "── TOUTES MATIÈRES ──",
                    "Français / Lettres","Mathématiques","Sciences Physiques (PC)","SVT / Biologie",
                    "Histoire-Géographie","Économie / Gestion","Comptabilité","Philosophie",
                    "EDHC / Éducation Civique","Anglais (LV1)","Espagnol (LV2)","Allemand (LV2)",
                    "Informatique / TIC","Technologie industrielle","EPS","Arts Plastiques",
                    "Agronomie / Agriculture","Droit","Économie politique",
                    "Sciences de la Vie et de la Terre","Chimie","Physique",
                    "Autre (préciser dans les notes)",
                ], index=0, key="exp_matiere")
            with col_b:
                exp_pages = st.selectbox("📏 Nombre de pages souhaité *", [
                    "6 pages",
                    "7 pages",
                    "8 pages",
                    "9 pages",
                ], index=0, key="exp_pages")
                exp_langue = st.selectbox("🌍 Langue de rédaction *", [
                    "Français","Anglais","Français + résumé en anglais",
                ], index=0, key="exp_langue")

            exp_sujet = st.text_input("🎯 Thème / Sujet de l'exposé *",
                placeholder="Ex: Le changement climatique en Afrique, La Révolution française, L'ADN et la génétique...",
                key="exp_sujet")

            col_c, col_d = st.columns(2)
            with col_c:
                exp_type = st.selectbox("📄 Type d'exposé *", [
                    "Exposé scolaire classique (Introduction + Développement + Conclusion)",
                    "Rapport de recherche structuré (Parties I, II, III)",
                    "Exposé oral (plan + fiches de présentation)",
                    "Compte-rendu de TP / Expérience scientifique",
                    "Commentaire de texte / document",
                    "Dissertation guidée",
                ], index=0, key="exp_type")
            with col_d:
                exp_etablissement = st.text_input("🏢 Établissement (optionnel)",
                    placeholder="Ex: Lycée Moderne de Cocody, Université FHB...",
                    key="exp_etablissement")

            exp_notes = st.text_area("💬 Instructions complémentaires (optionnel)", height=70,
                placeholder="Ex: Insister sur le rôle de la Côte d'Ivoire, Niveau très basique, Inclure des statistiques récentes...",
                key="exp_notes")

            # ── PAGE DE GARDE ──────────────────────────────────────────
            with st.expander("🎨 Page de garde personnalisée (optionnel)", expanded=True):
                st.markdown("<small style='color:#aaa'>Ces informations apparaîtront directement sur la page de garde de votre exposé.</small>", unsafe_allow_html=True)

                col_pg1, col_pg2 = st.columns(2)
                with col_pg1:
                    exp_filiere  = st.text_input("🏛️ Filière", placeholder="Ex: Génie Civil – Option Bâtiment", key="exp_filiere")
                    exp_annee_pg = st.text_input("📅 Année scolaire", value="2025-2026", key="exp_annee_pg")
                with col_pg2:
                    exp_logo_ecole = st.file_uploader("🏫 Logo de votre école (PNG/JPG)", type=["png","jpg","jpeg"], key="exp_logo_ecole")
                    if exp_logo_ecole:
                        import tempfile, os
                        _logo_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                        _logo_tmp.write(exp_logo_ecole.read())
                        _logo_tmp.close()
                        st.session_state["logo_ecole_path"] = _logo_tmp.name
                        st.success("✅ Logo chargé")
                    elif "logo_ecole_path" not in st.session_state:
                        st.session_state["logo_ecole_path"] = None

                # ── MEMBRES DU GROUPE — 8 SLOTS COMPACTS AVEC RÔLES ─
                st.markdown("**👥 Présenté par**")
                _ROLES_PDG = [
                    "Chef de groupe", "Rapporteur", "Recherche",
                    "Mise en page", "Présentation orale", "Secrétaire",
                    "Illustrations", "Correction"
                ]
                # Initialiser 8 slots vides
                if "exp_membres" not in st.session_state or len(st.session_state["exp_membres"]) != 8:
                    _old = st.session_state.get("exp_membres", [])
                    st.session_state["exp_membres"] = (_old + [""] * 8)[:8]

                _membres = st.session_state["exp_membres"]
                # Affichage en grille 2 colonnes
                _col_g1, _col_g2 = st.columns(2)
                for _mi in range(8):
                    _col = _col_g1 if _mi % 2 == 0 else _col_g2
                    with _col:
                        _membres[_mi] = st.text_input(
                            f"{_mi+1:02d} — {_ROLES_PDG[_mi]}",
                            value=_membres[_mi],
                            placeholder="Nom Prénom",
                            key=f"exp_membre_{_mi}",
                            label_visibility="visible",
                        )
                st.session_state["exp_membres"] = _membres

            # Construire le prompt
            _exp_niveau_val  = exp_niveau  if not exp_niveau.startswith("──")  else ""
            _exp_matiere_val = exp_matiere if not exp_matiere.startswith("──") else ""
            _exp_membres_list = [m.strip() for m in st.session_state.get("exp_membres", []) if m.strip()]
            _exp_noms_val    = " ; ".join(_exp_membres_list)
            _exp_filiere_val = st.session_state.get("exp_filiere", "") if "exp_filiere" in st.session_state else ""
            _exp_annee_val   = st.session_state.get("exp_annee_pg", "2025-2026") if "exp_annee_pg" in st.session_state else "2025-2026"
            prompt = f"""FICHE DE COMMANDE NOVA EXPOSÉ :
🎯 SUJET            : {exp_sujet.strip() or "Non précisé"}
🎓 NIVEAU           : {_exp_niveau_val or "Non précisé"}
📚 MATIÈRE          : {_exp_matiere_val or "Non précisée"}
📄 TYPE             : {exp_type}
📏 PAGES            : {exp_pages}
🌍 LANGUE           : {exp_langue}
🏢 ÉTABLISSEMENT    : {exp_etablissement.strip() or "Non précisé"}
🏛️ FILIÈRE          : {_exp_filiere_val or "Non précisée"}
📅 ANNÉE SCOLAIRE   : {_exp_annee_val}
👥 NOMS EXPOSANTS   : {_exp_noms_val.replace(chr(10), " ; ") if _exp_noms_val else "Non précisés"}
💬 INSTRUCTIONS     : {exp_notes.strip() or "Aucune"}
"""
            if _exp_niveau_val and _exp_matiere_val and exp_sujet.strip():
                st.success(f"✅ Commande prête : **{exp_sujet.strip()[:45]}** · **{_exp_matiere_val}** · **{_exp_niveau_val}**")
            elif exp_sujet.strip():
                st.info("💡 Précisez le niveau et la matière pour un meilleur résultat")
            else:
                if not exp_sujet.strip(): st.warning("⚠️ Entrez le thème / sujet de l'exposé")
                if not _exp_niveau_val:   st.warning("⚠️ Sélectionnez un niveau scolaire précis")
                if not _exp_matiere_val:  st.warning("⚠️ Sélectionnez une matière précise")

        elif "Fiche de Cours" in service:
            _show_splash("Fiche")
            st.markdown('''
            <div style="background:rgba(155,89,182,0.08);border:1px solid rgba(155,89,182,0.35);
                 border-radius:12px;padding:14px 18px;margin-bottom:14px;">
                <span style="color:#9b59b6;font-weight:700;">📖 Remplissez les champs — Nova génère une fiche de cours complète utilisable directement en classe</span>
            </div>
            ''', unsafe_allow_html=True)
            col_a, col_b = st.columns(2)
            with col_a:
                fc_niveau = st.selectbox("🎓 Niveau de la classe *", [
                    "── PRIMAIRE ──","CP1","CP2","CE1","CE2","CM1","CM2",
                    "── COLLÈGE ──","6ème","5ème","4ème","3ème",
                    "── LYCÉE ──","2nde",
                    "1ère - Série A1","1ère - Série A2","1ère - Série B",
                    "1ère - Série C","1ère - Série D","1ère - Série E",
                    "Terminale - Série A1","Terminale - Série A2","Terminale - Série B",
                    "Terminale - Série C","Terminale - Série D","Terminale - Série E",
                    "── UNIVERSITÉ / BTS ──",
                    "Licence 1 (L1)","Licence 2 (L2)","Licence 3 (L3)",
                    "Master 1 (M1)","Master 2 (M2)",
                    "BTS 1ère année","BTS 2ème année",
                ], index=0, key="fc_niveau")
                fc_matiere = st.selectbox("📚 Matière *", [
                    "── TOUTES MATIÈRES ──",
                    "Français / Lettres","Mathématiques","Sciences Physiques (PC)","SVT / Biologie",
                    "Histoire-Géographie","Économie / Gestion","Comptabilité","Philosophie",
                    "EDHC / Éducation Civique","Anglais (LV1)","Espagnol (LV2)","Allemand (LV2)",
                    "Informatique / TIC","Technologie industrielle","EPS","Arts Plastiques",
                    "Agronomie / Agriculture","Droit","Économie politique",
                    "── PRIMAIRE ──",
                    "Lecture / Écriture (primaire)","Calcul (primaire)",
                    "Sciences d'Éveil (primaire)","Histoire-Géo (primaire)","ECM (primaire)",
                    "Autre matière (préciser dans les notes)",
                ], index=0, key="fc_matiere")
            with col_b:
                fc_duree = st.selectbox("⏱️ Durée de la séance *", [
                    "30 minutes","45 minutes","1 heure",
                    "1 heure 30","2 heures","3 heures",
                    "Cours complet (plusieurs séances)",
                ], index=2, key="fc_duree")
                fc_type = st.selectbox("📄 Type de document *", [
                    "Fiche de cours complète (objectifs + développement + exercices)",
                    "Cours magistral (contenu seul, très détaillé)",
                    "Plan de leçon (structure + timing + activités)",
                    "Fiche de révision élève (synthèse dense)",
                    "Cours + Exercices corrigés",
                    "Fiche méthode (procédure étape par étape)",
                ], index=0, key="fc_type")
            col_c, col_d = st.columns(2)
            with col_c:
                fc_chapitre = st.text_input("📖 Titre du chapitre / Notion *",
                    placeholder="Ex: Les fractions, La photosynthèse, La colonisation...",
                    key="fc_chapitre")
            with col_d:
                fc_etablissement = st.text_input("🏢 Établissement (optionnel)",
                    placeholder="Ex: Lycée Moderne de Cocody...", key="fc_etablissement")
            fc_inclure = st.multiselect("✨ Éléments à inclure", [
                "Objectifs pédagogiques détaillés","Prérequis des élèves",
                "Exercices d'application","Corrigé complet des exercices",
                "Évaluation formative (fin de séance)","Résumé / Synthèse à retenir",
                "Exemples contextualisés (Côte d'Ivoire)","Définitions des termes clés",
                "Référence au manuel officiel MENET-FP","Tableau récapitulatif",
            ], default=[
                "Objectifs pédagogiques détaillés","Exercices d'application",
                "Corrigé complet des exercices","Exemples contextualisés (Côte d'Ivoire)",
                "Définitions des termes clés","Résumé / Synthèse à retenir",
            ], key="fc_inclure")
            fc_notes = st.text_area("💬 Instructions supplémentaires (optionnel)", height=70,
                placeholder="Ex: Niveau très faible — simplifier, Inclure exemples en FCFA...",
                key="fc_notes")
            _fc_niveau_val  = fc_niveau  if not fc_niveau.startswith("──")  else ""
            _fc_matiere_val = fc_matiere if not fc_matiere.startswith("──") else ""
            _fc_inclure_str = ", ".join(fc_inclure) if fc_inclure else "Éléments standards"
            prompt = f"""FICHE DE COMMANDE NOVA COURS :
🎓 NIVEAU          : {_fc_niveau_val or "Non précisé"}
📚 MATIÈRE         : {_fc_matiere_val or "Non précisée"}
📖 CHAPITRE        : {fc_chapitre.strip() or "Choisir un chapitre cohérent"}
📄 TYPE DOCUMENT   : {fc_type}
⏱️ DURÉE           : {fc_duree}
🏢 ÉTABLISSEMENT   : {fc_etablissement.strip() or "Non précisé"}
✨ ÉLÉMENTS        : {_fc_inclure_str}
💬 INSTRUCTIONS    : {fc_notes.strip() or "Aucune"}
"""
            if _fc_niveau_val and _fc_matiere_val and fc_chapitre.strip():
                st.success(f"✅ Fiche prête : **{fc_chapitre.strip()[:40]}** · **{_fc_matiere_val}** · **{_fc_niveau_val}**")
            elif _fc_niveau_val and _fc_matiere_val:
                st.info("💡 Précisez le titre du chapitre pour un meilleur résultat")
            else:
                if not _fc_niveau_val: st.warning("⚠️ Sélectionnez un niveau scolaire précis")
                if not _fc_matiere_val: st.warning("⚠️ Sélectionnez une matière précise")

        elif "Modifier" in service and "Fichier" in service:
            _show_splash("Modifier")
            st.markdown('''
            <div style="background:rgba(0,210,255,0.07);border:1px solid rgba(0,210,255,0.3);
                 border-radius:12px;padding:14px 18px;margin-bottom:14px;">
                <span style="color:#00d2ff;font-weight:700;">📎 Importez votre fichier et décrivez vos modifications — notre équipe s'en charge</span>
            </div>
            ''', unsafe_allow_html=True)
            mf_fichier = st.file_uploader("📂 Importer votre fichier *",
                type=["docx","xlsx","xls","pptx","ppt","doc","pdf","csv","txt"],
                help="Formats : Word, Excel, PowerPoint, PDF, CSV, TXT",
                key="mf_fichier")
            if mf_fichier:
                st.success(f"✅ **{mf_fichier.name}** · {round(mf_fichier.size/1024,1)} Ko")
            mf_type_modif = st.selectbox("🛠️ Type de modification *", [
                "Correction et amélioration du contenu",
                "Mise en forme / Design professionnel",
                "Ajout de données / contenu supplémentaire",
                "Restructuration complète du document",
                "Traduction (Français ↔ Anglais)",
                "Fusion de plusieurs documents",
                "Extraction et réorganisation des données",
                "Création de graphiques / tableaux depuis les données",
                "Correction orthographique et grammaticale",
                "Adapter à un nouveau modèle / template",
                "Autre modification (préciser dans les instructions)",
            ], key="mf_type_modif")
            mf_instructions = st.text_area("✍️ Décrivez précisément ce que vous voulez *",
                height=130,
                placeholder="Ex: Mettre le tableau en page 3 en format auto, ajouter colonne Total avec formule, corriger les fautes...",
                key="mf_instructions")
            mf_urgence = st.selectbox("⏱️ Délai souhaité", [
                "Dès que possible (standard)",
                "Urgent — dans les 2 heures",
                "Très urgent — dans 1 heure",
            ], key="mf_urgence")
            _mf_nom = mf_fichier.name if mf_fichier else "Aucun fichier"
            _ext_fichier = _mf_nom.rsplit(".", 1)[-1].upper() if "." in _mf_nom else "INCONNU"
            prompt = f"""Tu es NOVA MODIFIER — l'expert Nova Platform chargé d'analyser la demande du client et de produire un document Word restructuré et amélioré, prêt à livrer.

════════════════════════════════════════
DEMANDE CLIENT
════════════════════════════════════════
📎 FICHIER SOURCE   : {_mf_nom} ({_ext_fichier})
🛠️ TYPE DE MODIF    : {mf_type_modif}
✍️ INSTRUCTIONS     :
{mf_instructions.strip() or "(aucune instruction spécifique — applique les meilleures pratiques)"}

════════════════════════════════════════
TON RÔLE — LIRE ATTENTIVEMENT
════════════════════════════════════════
Tu es l'intermédiaire intelligent entre le client et Python (le moteur Nova).
Python est capable de créer des documents Word parfaits — titres, tableaux, listes, mise en page,
polices, couleurs, sauts de page, etc. — mais il a besoin que TU lui fourniSSES le contenu
sous forme structurée avec les bonnes balises.

TON TRAVAIL :
1. Comprendre ce que le client veut faire sur son fichier
2. Produire le CONTENU COMPLET du document modifié, structuré avec les balises Nova
3. Python appliquera automatiquement toute la mise en page

TU NE GÈRES PAS : polices, couleurs, marges, tailles — Python s'en charge.
TU GÈRES : contenu, structure, logique, corrections, ajouts, restructuration.

════════════════════════════════════════
RÈGLES DE FORMATAGE NOVA — OBLIGATOIRES
════════════════════════════════════════
Ces balises sont converties automatiquement en vrai formatage Word par Python :

TITRES :
  # Titre principal      → H1 Word (une seule fois, en tête)
  ## Titre de section    → H2 Word
  ### Sous-titre         → H3 Word
  #### Micro-titre       → H4 Word

TEXTE :
  Paragraphe normal      → texte justifié 11pt
  **mot**                → gras dans Word
  - item                 → puce Word
  1. item                → liste numérotée Word
  ---SAUT_DE_PAGE---     → vrai saut de page Word

TABLEAUX :
  | Col1 | Col2 | Col3 |
  |------|------|------|
  | val  | val  | val  |
  → Tableau Word avec en-tête coloré et lignes alternées automatiquement

SÉPARATEURS :
  ════════════════════  → trait épais entre grandes parties
  ────────────────────  → trait fin entre sous-sections

INTERDIT ABSOLU :
  ✗ LaTeX ($formule$, \\frac, \\omega...)
  ✗ HTML (<br>, <b>, <div>...)
  ✗ Sections vides ou placeholders "[À compléter]"
  ✗ Commentaires méta comme "Note : voir page 3" ou "Section à personnaliser"

════════════════════════════════════════
RÈGLES PAR TYPE DE MODIFICATION
════════════════════════════════════════

① CORRECTION / AMÉLIORATION DU CONTENU :
  → Corrige les fautes, améliore les tournures, renforce les arguments
  → Conserve EXACTEMENT la structure originale (mêmes titres, même ordre)
  → Ne supprime aucune section sans raison explicite du client
  → Signale les passages reformulés avec **[amélioré]** après le titre de section

② MISE EN FORME / DESIGN PROFESSIONNEL :
  → Restructure avec les balises # ## ### pour une hiérarchie claire
  → Ajoute des séparateurs ════ entre les grandes parties
  → Transforme les listes en tirets en "- item" pour puces Word propres
  → Si le client a des données brutes → les mettre en tableau

③ AJOUT DE CONTENU :
  → Intègre les nouvelles données demandées naturellement dans le flux existant
  → Indique clairement avec **[AJOUTÉ]** au titre de la section nouvelle
  → Ne modifie pas les sections existantes sauf si demandé

④ RESTRUCTURATION COMPLÈTE :
  → Réorganise logiquement : Introduction → Développement → Conclusion
  → Crée une hiérarchie claire avec # ## ###
  → Conserve tout le contenu original, rien ne disparaît

⑤ CORRECTION ORTHOGRAPHIQUE / GRAMMATICALE :
  → Corrige uniquement les fautes, sans changer le fond
  → Conserve le style et le ton du client
  → Ne reformule pas si la phrase est correcte

⑥ TRADUCTION :
  → Traduis fidèlement sans adapter ni résumer
  → Conserve EXACTEMENT la structure, les titres, les tableaux
  → Garde les noms propres, termes techniques et chiffres intacts

⑦ EXTRACTION / RÉORGANISATION :
  → Si le client veut un tableau depuis des données → crée le tableau exact
  → Si le client veut une liste → formate en "- item" propre
  → Ne paraphrase pas, reproduis les données exactes

⑧ AUTRE MODIFICATION :
  → Analyse les instructions du client et applique le bon traitement
  → En cas de doute, applique la règle la plus conservative (ne rien supprimer)

════════════════════════════════════════
INSTRUCTION FINALE
════════════════════════════════════════
Produis le document COMPLET du début à la fin.
Ne laisse aucune section vide. Ne dis pas "le reste du document reste inchangé".
Si le fichier source n'est pas lisible directement, base-toi sur les instructions du client
pour reconstituer et améliorer le document tel qu'il te l'a décrit.
Le document livré doit être utilisable immédiatement, sans rien à compléter."""
            if mf_fichier and mf_instructions.strip():
                st.success("✅ Demande complète — génération IA en cours")
            else:
                if not mf_fichier: st.warning("⚠️ Importez votre fichier")
                if not mf_instructions.strip(): st.warning("⚠️ Décrivez les modifications souhaitées")

        elif "Conversion" in service:
            # ════════════════════════════════════════════════════════════════
            # SERVICE CONVERSION — 100% Python, pas de Gemini, pas de WhatsApp
            # ════════════════════════════════════════════════════════════════
            _show_splash("Conversion")
            st.markdown('''
            <div style="background:rgba(46,204,113,0.08);border:1px solid rgba(46,204,113,0.35);
                 border-radius:12px;padding:14px 18px;margin-bottom:14px;">
                <span style="font-weight:700;color:#2ecc71;">🔄 Conversion instantanée — 100% automatique, aucune attente</span>
                <span style="color:rgba(255,255,255,.5);font-size:.82rem;display:block;margin-top:4px;">
                    Importez votre fichier ci-dessous, choisissez le format de sortie — votre fichier est prêt en quelques secondes.
                </span>
            </div>
            ''', unsafe_allow_html=True)

            conv_type = st.selectbox("🔄 Type de conversion *", [
                "📝 Word (.docx) → PDF",
                "📊 Excel (.xlsx) → PDF",
                "📊 CSV → Excel (.xlsx)",
                "📊 Excel (.xlsx) → CSV",
                "🎞️ PowerPoint (.pptx) → PDF",
                "📄 PDF → Word (.docx)",
                "📄 PDF → PowerPoint (.pptx)",
            ], key="conv_type")

            # Définir les types acceptés selon la conversion choisie
            if "Word" in conv_type and "→ PDF" in conv_type:
                types_acceptes = ["docx", "doc"]
                label_upload   = "📂 Importer votre fichier Word (.docx / .doc)"
                suffix_in      = None  # déterminé dynamiquement depuis le nom du fichier
            elif "Excel" in conv_type and "→ PDF" in conv_type:
                types_acceptes = ["xlsx", "xls"]
                label_upload   = "📂 Importer votre fichier Excel (.xlsx)"
                suffix_in      = ".xlsx"
            elif "CSV → Excel" in conv_type:
                types_acceptes = ["csv"]
                label_upload   = "📂 Importer votre fichier CSV"
                suffix_in      = ".csv"
            elif "Excel" in conv_type and "→ CSV" in conv_type:
                types_acceptes = ["xlsx", "xls"]
                label_upload   = "📂 Importer votre fichier Excel (.xlsx)"
                suffix_in      = ".xlsx"
            elif "PowerPoint" in conv_type and "→ PDF" in conv_type:
                types_acceptes = ["pptx", "ppt"]
                label_upload   = "📂 Importer votre fichier PowerPoint (.pptx)"
                suffix_in      = None  # déterminé dynamiquement depuis le nom du fichier
            elif "PDF → Word" in conv_type:
                types_acceptes = ["pdf"]
                label_upload   = "📂 Importer votre fichier PDF"
                suffix_in      = ".pdf"
            else:  # PDF → PPT
                types_acceptes = ["pdf"]
                label_upload   = "📂 Importer votre fichier PDF"
                suffix_in      = ".pdf"

            conv_fichier = st.file_uploader(label_upload, type=types_acceptes, key="conv_fichier")

            if conv_fichier:
                # ── Mise en cache immédiate pour éviter buffer vide au clic ──
                cache_key = f"conv_bytes_{conv_fichier.name}_{conv_fichier.size}"
                if st.session_state.get("conv_cache_key") != cache_key:
                    st.session_state["conv_cache_key"]  = cache_key
                    st.session_state["conv_fichier_bytes"] = conv_fichier.read()
                    st.session_state["conv_fichier_name"]  = conv_fichier.name

                st.success(f"✅ **{conv_fichier.name}** · {round(conv_fichier.size/1024, 1)} Ko importé")

                if st.button("⚡ CONVERTIR MAINTENANT", use_container_width=True, key="btn_convertir"):
                    with st.spinner("🔄 Conversion en cours..."):
                        try:
                            import tempfile, os as _os, subprocess as _sub
                            fichier_bytes = st.session_state.get("conv_fichier_bytes", b"")
                            # suffix_in dynamique selon l'extension réelle du fichier
                            if suffix_in is None:
                                ext_reelle = "." + st.session_state["conv_fichier_name"].rsplit(".", 1)[-1].lower()
                                suffix_in_reel = ext_reelle
                            else:
                                suffix_in_reel = suffix_in
                            if not fichier_bytes:
                                st.error("❌ Fichier introuvable, veuillez le ré-importer.")
                                st.stop()

                            # ── Word / Excel / PPT → PDF via LibreOffice ─────
                            if "→ PDF" in conv_type:
                                # ── Détection du format réel du fichier ──────
                                magic = fichier_bytes[:8]
                                est_html_utf16 = fichier_bytes[:2] in (b'\xff\xfe', b'\xfe\xff')
                                est_html_utf8  = fichier_bytes[:5].lower().startswith(b'<html') or fichier_bytes[:14].lower().replace(b'\xef\xbb\xbf', b'').startswith(b'<html')
                                est_html = est_html_utf16 or est_html_utf8

                                if est_html:
                                    # ── Fichier HTML/Word HTML → reconstruire un vrai .docx ──
                                    import re as _re
                                    try:
                                        from bs4 import BeautifulSoup
                                    except ImportError:
                                        _sub.run(["pip", "install", "beautifulsoup4",
                                                  "--break-system-packages", "-q"], check=True)
                                        from bs4 import BeautifulSoup
                                    from docx import Document as _Doc
                                    from docx.shared import Pt as _Pt, RGBColor as _RGB, Cm as _Cm
                                    from io import BytesIO as _BIO

                                    if est_html_utf16:
                                        html_text = fichier_bytes.decode('utf-16', errors='ignore')
                                    else:
                                        html_text = fichier_bytes.decode('utf-8', errors='ignore')

                                    soup = BeautifulSoup(html_text, 'html.parser')
                                    # Supprimer VML, images, scripts, styles
                                    for tag in soup(['script', 'style', 'head', 'img',
                                                     'v:shape', 'v:shapetype', 'o:p']):
                                        tag.decompose()

                                    new_doc = _Doc()
                                    for sec in new_doc.sections:
                                        sec.top_margin = _Cm(2); sec.bottom_margin = _Cm(2)
                                        sec.left_margin = _Cm(2.5); sec.right_margin = _Cm(2.5)

                                    def get_font_size(el):
                                        """Récupère la font-size max dans les spans d'un paragraphe."""
                                        sizes = []
                                        for sp in el.find_all('span', style=True):
                                            m = _re.search(r'font-size\s*:\s*([\d.]+)pt', sp.get('style',''), _re.I)
                                            if m: sizes.append(float(m.group(1)))
                                        return max(sizes) if sizes else None

                                    def get_clean_text(el):
                                        """Texte propre — on prend le texte direct de l'élément sans récursion dans les sous-spans."""
                                        # Remplacer les sous-éléments par leur texte pour éviter les doublons
                                        return _re.sub(r'\s+', ' ', el.get_text(' ', strip=True)).strip()

                                    seen_texts = set()
                                    _roman_re  = _re.compile(r'^(I{1,3}|IV|V|VI{1,3}|IX|X)[\.\s]')
                                    for el in soup.find_all(['p','h1','h2','h3','h4','li']):
                                        texte = _re.sub(r'\s+', ' ', el.get_text(' ', strip=True)).strip()
                                        if not texte or texte in ('\xa0', '&nbsp;'): continue
                                        if texte in seen_texts: continue
                                        seen_texts.add(texte)
                                        tag      = el.name
                                        fs       = get_font_size(el)
                                        is_bold  = bool(el.find('b') or el.find('strong'))
                                        is_roman = bool(_roman_re.match(texte))

                                        if tag == 'h1' or (fs and fs >= 28):
                                            new_doc.add_heading(texte, level=1)
                                        elif tag == 'h2' or (fs and 20 <= fs < 28):
                                            new_doc.add_heading(texte, level=2)
                                        elif tag == 'h3' or (fs and 16 <= fs < 20) or is_roman:
                                            new_doc.add_heading(texte, level=3)
                                        elif tag == 'h4' or (fs and 13 <= fs < 16):
                                            new_doc.add_heading(texte, level=4)
                                        else:
                                            para = new_doc.add_paragraph()
                                            run  = para.add_run(texte)
                                            run.bold = is_bold
                                            if fs: run.font.size = _Pt(min(fs, 24))

                                    buf_docx = _BIO()
                                    new_doc.save(buf_docx)
                                    buf_docx.seek(0)
                                    suffix_in_reel = ".docx"
                                    fichier_bytes  = buf_docx.read()

                                with tempfile.NamedTemporaryFile(suffix=suffix_in_reel, delete=False) as tmp_in:
                                    tmp_in.write(fichier_bytes)
                                    tmp_in_path = tmp_in.name
                                tmp_out_dir = tempfile.mkdtemp()
                                res = _sub.run(
                                    ["libreoffice", "--headless", "--convert-to", "pdf",
                                     "--outdir", tmp_out_dir, tmp_in_path],
                                    capture_output=True, text=True, timeout=90
                                )
                                nom_base = _os.path.splitext(_os.path.basename(tmp_in_path))[0]
                                pdf_path = _os.path.join(tmp_out_dir, nom_base + ".pdf")
                                if _os.path.exists(pdf_path):
                                    with open(pdf_path, "rb") as f:
                                        buf_out = f.read()
                                    nom_sortie = st.session_state["conv_fichier_name"].rsplit(".", 1)[0] + ".pdf"
                                    st.success("✅ Conversion réussie !")
                                    st.download_button("📥 TÉLÉCHARGER LE PDF", data=buf_out,
                                        file_name=nom_sortie, mime="application/pdf",
                                        use_container_width=True, key="dl_pdf_out")
                                else:
                                    st.error(f"❌ Erreur LibreOffice : {res.stderr or res.stdout}")
                                try: _os.unlink(tmp_in_path)
                                except: pass
                                try: import shutil; shutil.rmtree(tmp_out_dir, ignore_errors=True)
                                except: pass

                            # ── CSV → Excel ──────────────────────────────────
                            elif "CSV → Excel" in conv_type:
                                import pandas as pd
                                from io import BytesIO as _BytesIO
                                df = pd.read_csv(_BytesIO(fichier_bytes))
                                buf_out = _BytesIO()
                                df.to_excel(buf_out, index=False, engine="openpyxl")
                                buf_out.seek(0)
                                nom_sortie = st.session_state["conv_fichier_name"].replace(".csv", "") + ".xlsx"
                                st.success("✅ Conversion réussie !")
                                st.download_button("📥 TÉLÉCHARGER L'EXCEL", data=buf_out.read(),
                                    file_name=nom_sortie,
                                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                    use_container_width=True, key="dl_xlsx_out")

                            # ── Excel → CSV ──────────────────────────────────
                            elif "Excel" in conv_type and "→ CSV" in conv_type:
                                import pandas as pd
                                from io import BytesIO as _BytesIO
                                df = pd.read_excel(_BytesIO(fichier_bytes))
                                csv_str = df.to_csv(index=False)
                                nom_sortie = st.session_state["conv_fichier_name"].rsplit(".", 1)[0] + ".csv"
                                st.success("✅ Conversion réussie !")
                                st.download_button("📥 TÉLÉCHARGER LE CSV", data=csv_str.encode("utf-8-sig"),
                                    file_name=nom_sortie, mime="text/csv",
                                    use_container_width=True, key="dl_csv_out")

                            # ── PDF → Word via pdf2docx ───────────────────────
                            elif "PDF → Word" in conv_type:
                                try:
                                    from pdf2docx import Converter as PdfConverter
                                except ImportError:
                                    _sub.run(["pip", "install", "pdf2docx",
                                              "--break-system-packages", "-q"], check=True)
                                    from pdf2docx import Converter as PdfConverter
                                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_in:
                                    tmp_in.write(fichier_bytes)
                                    tmp_in_path = tmp_in.name
                                tmp_out_path = tmp_in_path[:-4] + ".docx"
                                cv = PdfConverter(tmp_in_path)
                                cv.convert(tmp_out_path, start=0, end=None)
                                cv.close()
                                with open(tmp_out_path, "rb") as f:
                                    buf_out = f.read()
                                nom_sortie = st.session_state["conv_fichier_name"].replace(".pdf", "") + ".docx"
                                st.success("✅ Conversion réussie !")
                                st.download_button("📥 TÉLÉCHARGER LE WORD", data=buf_out,
                                    file_name=nom_sortie,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True, key="dl_word_out")
                                try: _os.unlink(tmp_in_path); _os.unlink(tmp_out_path)
                                except: pass

                            # ── PDF → PPT via LibreOffice Impress ────────────
                            else:
                                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_in:
                                    tmp_in.write(fichier_bytes)
                                    tmp_in_path = tmp_in.name
                                tmp_out_dir = tempfile.mkdtemp()
                                res = _sub.run(
                                    ["libreoffice", "--headless", "--convert-to", "pptx",
                                     "--outdir", tmp_out_dir, tmp_in_path],
                                    capture_output=True, text=True, timeout=90
                                )
                                nom_base  = _os.path.splitext(_os.path.basename(tmp_in_path))[0]
                                pptx_path = _os.path.join(tmp_out_dir, nom_base + ".pptx")
                                if _os.path.exists(pptx_path):
                                    with open(pptx_path, "rb") as f:
                                        buf_out = f.read()
                                    nom_sortie = st.session_state["conv_fichier_name"].replace(".pdf", "") + ".pptx"
                                    st.success("✅ Conversion réussie !")
                                    st.download_button("📥 TÉLÉCHARGER LE POWERPOINT", data=buf_out,
                                        file_name=nom_sortie,
                                        mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                                        use_container_width=True, key="dl_pptx_out")
                                else:
                                    st.error(f"❌ Erreur conversion PDF→PPT : {res.stderr or res.stdout}")
                                try: _os.unlink(tmp_in_path)
                                except: pass
                                try: import shutil; shutil.rmtree(tmp_out_dir, ignore_errors=True)
                                except: pass

                        except Exception as e:
                            st.error(f"❌ Erreur lors de la conversion : {e}")
            else:
                st.info("⬆️ Importez votre fichier pour démarrer la conversion")

            prompt = "CONVERSION_AUTO"  # Pas de prompt Gemini ni WhatsApp

        elif "OCR" in service or "Numérisation" in service:
            # ================================================================
            # SERVICE OCR - Numérisation de document scanné -> .docx éditable
            # 100% Python, pas de Gemini, pas de WhatsApp
            # ================================================================
            _show_splash("OCR")
            st.markdown("""
            <div style="background:rgba(180,100,255,0.08);border:1px solid rgba(180,100,255,0.35);
                 border-radius:12px;padding:14px 18px;margin-bottom:14px;">
                <span style="font-weight:700;color:#b464ff;">🔍 OCR — Extraction de texte automatique</span>
                <span style="color:rgba(255,255,255,.5);font-size:.82rem;display:block;margin-top:4px;">
                    Importez un PDF scanné, une image ou un document Word/Excel dont le contenu est une image —
                    Nova extrait le texte et vous livre un fichier <strong style="color:#b464ff;">.docx éditable</strong>.
                </span>
            </div>
            """, unsafe_allow_html=True)

            ocr_fichier = st.file_uploader(
                "📂 Importer votre document scanné",
                type=["pdf", "png", "jpg", "jpeg", "bmp", "tiff", "tif", "webp", "docx", "xlsx"],
                key="ocr_fichier_upload"
            )

            if ocr_fichier:
                ext_ocr = ocr_fichier.name.rsplit(".", 1)[-1].lower()
                st.info(f"📄 Fichier détecté : **{ocr_fichier.name}** ({ext_ocr.upper()})")

                # ── Diagnostic poppler ──────────────────────────────────────
                import shutil as _shutil
                _pdftoppm = _shutil.which("pdftoppm")
                _pdfinfo  = _shutil.which("pdfinfo")
                st.info(f"🔧 Diagnostic poppler — pdftoppm: `{_pdftoppm}` | pdfinfo: `{_pdfinfo}`")

                if st.button("🔍 LANCER L'OCR ET GÉNÉRER LE .DOCX", type="primary", use_container_width=True):
                    try:
                        import pytesseract
                        from PIL import Image
                        import io as _io
                        import shutil as _shutil2

                        texte_extrait = ""
                        images_ocr = []

                        if ext_ocr == "pdf":
                            try:
                                import fitz  # pymupdf - pas besoin de poppler
                                with st.spinner("📄 Conversion PDF → images..."):
                                    _pdf_bytes = ocr_fichier.read()
                                    _pdf_doc = fitz.open(stream=_pdf_bytes, filetype="pdf")
                                    for _page in _pdf_doc:
                                        _mat = fitz.Matrix(300 / 72, 300 / 72)  # 300 DPI
                                        _pix = _page.get_pixmap(matrix=_mat)
                                        _img_bytes = _pix.tobytes("png")
                                        from PIL import Image as _PILImage
                                        images_ocr.append(_PILImage.open(_io.BytesIO(_img_bytes)))
                                    _pdf_doc.close()
                            except ImportError:
                                st.error("❌ pymupdf non installé. Ajoutez 'pymupdf' dans requirements.txt")
                                st.stop()

                        elif ext_ocr in ["png", "jpg", "jpeg", "bmp", "tiff", "tif", "webp"]:
                            img = Image.open(_io.BytesIO(ocr_fichier.read()))
                            images_ocr = [img]

                        elif ext_ocr == "docx":
                            try:
                                import zipfile
                                docx_bytes = ocr_fichier.read()
                                with zipfile.ZipFile(_io.BytesIO(docx_bytes)) as zf:
                                    img_files = [f for f in zf.namelist() if f.startswith("word/media/")]
                                    if not img_files:
                                        st.error("❌ Aucune image trouvée dans ce Word. Utilisez le service Conversion à la place.")
                                        st.stop()
                                    for img_f in img_files:
                                        try:
                                            images_ocr.append(Image.open(_io.BytesIO(zf.read(img_f))))
                                        except Exception:
                                            pass
                            except Exception as e_docx:
                                st.error(f"❌ Erreur lecture Word : {e_docx}")
                                st.stop()

                        elif ext_ocr == "xlsx":
                            try:
                                import zipfile
                                xlsx_bytes = ocr_fichier.read()
                                with zipfile.ZipFile(_io.BytesIO(xlsx_bytes)) as zf:
                                    img_files = [f for f in zf.namelist() if f.startswith("xl/media/")]
                                    if not img_files:
                                        st.error("❌ Aucune image trouvée dans ce fichier Excel.")
                                        st.stop()
                                    for img_f in img_files:
                                        try:
                                            images_ocr.append(Image.open(_io.BytesIO(zf.read(img_f))))
                                        except Exception:
                                            pass
                            except Exception as e_xlsx:
                                st.error(f"❌ Erreur lecture Excel : {e_xlsx}")
                                st.stop()

                        if not images_ocr:
                            st.error("❌ Impossible d'extraire des images de ce fichier.")
                            st.stop()

                        bar_ocr = st.progress(0)
                        status_ocr = st.empty()
                        for _idx_img, _img in enumerate(images_ocr):
                            status_ocr.markdown(
                                f"<p style='color:#b464ff;font-weight:700;'>🔍 Analyse page {_idx_img+1}/{len(images_ocr)}...</p>",
                                unsafe_allow_html=True
                            )
                            _txt = pytesseract.image_to_string(_img, lang="fra", config="--psm 6")
                            texte_extrait += _txt + "\n\n"
                            bar_ocr.progress(int((_idx_img + 1) / len(images_ocr) * 100))

                        bar_ocr.empty()
                        status_ocr.empty()

                        if not texte_extrait.strip():
                            st.warning("⚠️ Aucun texte détecté. Vérifiez la qualité du scan (300 DPI minimum recommandé).")
                            st.stop()

                        from docx import Document as _DocxDoc
                        from docx.shared import Pt, Cm, RGBColor as _RGBColor
                        from docx.enum.text import WD_ALIGN_PARAGRAPH

                        doc_ocr = _DocxDoc()
                        for section in doc_ocr.sections:
                            section.top_margin    = Cm(2.5)
                            section.bottom_margin = Cm(2.5)
                            section.left_margin   = Cm(2.5)
                            section.right_margin  = Cm(2.5)

                        _style_normal = doc_ocr.styles["Normal"]
                        _style_normal.font.name = "Calibri"
                        _style_normal.font.size = Pt(11)

                        _titre_para = doc_ocr.add_heading("Document extrait par OCR — Nova Platform", level=1)
                        _titre_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        _sub = doc_ocr.add_paragraph(f"Source : {ocr_fichier.name}")
                        _sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _sub.runs[0].font.size = Pt(9)
                        _sub.runs[0].font.color.rgb = _RGBColor(0x88, 0x88, 0x88)
                        doc_ocr.add_paragraph("")

                        for _ligne in texte_extrait.split("\n"):
                            _ligne_strip = _ligne.strip()
                            if _ligne_strip:
                                _p = doc_ocr.add_paragraph(_ligne_strip)
                                _p.paragraph_format.space_after = Pt(4)
                            else:
                                doc_ocr.add_paragraph("")

                        _buf_ocr = _io.BytesIO()
                        doc_ocr.save(_buf_ocr)
                        _buf_ocr.seek(0)

                        nom_sortie_ocr = ocr_fichier.name.rsplit(".", 1)[0] + "_OCR_Nova.docx"
                        st.success(f"✅ OCR terminé — {len(images_ocr)} page(s) · {len(texte_extrait.split())} mots extraits")
                        st.download_button(
                            label="📥 TÉLÉCHARGER LE .DOCX ÉDITABLE",
                            data=_buf_ocr,
                            file_name=nom_sortie_ocr,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                        with st.expander("👁️ Aperçu du texte extrait (300 premiers caractères)"):
                            st.text(texte_extrait[:300] + ("..." if len(texte_extrait) > 300 else ""))

                    except ImportError as _ie:
                        st.error(f"❌ Module manquant : {_ie}. Installez : pip install pytesseract Pillow pdf2image")
                    except Exception as _e_ocr:
                        st.error(f"❌ Erreur OCR : {_e_ocr}")
            else:
                st.info("⬆️ Importez votre fichier scanné pour démarrer l'extraction OCR")

            prompt = "OCR_AUTO"  # Pas de prompt Gemini ni WhatsApp


        else:
            # ── Splash screen adapté au service ───────────────────────────────
            if "CV" in service or "Lettre" in service:
                _show_splash("CV")
            elif "Word" in service:
                _show_splash("Word")
            elif "Rapport de Stage" in service:
                _show_splash("RapportStage")
            # ── CHAMP TEXTE LIBRE POUR LES AUTRES SERVICES ────────────────────
            prompt = st.text_area("Cahier des charges Nova", height=150, placeholder="Détaillez votre projet pour une exécution parfaite...")

        if service == SERVICE_SAISIE and service in SERVICE_PREREQUIS:
            if prompt and not st.session_state["warning_triggered"]:
                st.session_state["warning_triggered"] = True
                st.session_state["show_service_warning"] = True
                st.rerun()

            if st.session_state["show_service_warning"]:
                info = SERVICE_PREREQUIS[service]
                st.info(f"""
**{info["icone"]} {info["titre"]} — Informations requises**

{info["intro"]}

{"".join(f"- {icone} {texte}\n" for icone, texte in info["items"])}
💡 *{info["note"]}*
""")

                col_mid = st.columns([1, 2, 1])[1]
                with col_mid:
                    if st.button("✅ J'ai compris, je continue ma demande", key="close_service_warning"):
                        st.session_state["show_service_warning"] = False
                        # speechSynthesis supprimé
                        st.rerun()
        
        st.markdown("""
        <div class="logo-container">
            <svg class="logo-item" viewBox="0 0 24 24" fill="#217346"><path d="M16.2 21H2.8c-.4 0-.8-.4-.8-.8V3.8c0-.4.4-.8.8-.8h13.4c.4 0 .8.4.8.8v16.4c0 .4-.4.8-.8.8z"/><path d="M14.7 15.3l-2.2-3.3 2.2-3.3h-1.6l-1.4 2.2-1.4-2.2H8.7l2.2 3.3-2.2 3.3h1.6l1.4-2.2 1.4 2.2z" fill="white"/></svg>
            <svg class="logo-item" viewBox="0 0 24 24" fill="#2b579a"><path d="M16.2 21H2.8c-.4 0-.8-.4-.8-.8V3.8c0-.4.4-.8.8-.8h13.4c.4 0 .8.4.8.8v16.4c0 .4-.4.8-.8.8z"/><path d="M11.5 15.3V8.7h1.4c.8 0 1.4.3 1.8.8.4.5.6 1.1.6 1.8s-.2 1.3-.6 1.8c-.4.5-1 .8-1.8.8h-1.4z" fill="white"/></svg>
            <svg class="logo-item" viewBox="0 0 24 24" fill="#3776ab"><path d="M12 2C6.47 2 2 6.47 2 12s4.47 10 10 10 10-4.47 10-10S17.53 2 12 2zm-1 14.5h-1v-5h1v5zm0-6.5h-1V9h1v1z"/></svg>
            <svg class="logo-item" viewBox="0 0 24 24" fill="#d24726"><path d="M16.2 21H2.8c-.4 0-.8-.4-.8-.8V3.8c0-.4.4-.8.8-.8h13.4c.4 0 .8.4.8.8v16.4c0 .4-.4.8-.8.8z"/><path d="M8.7 8.7h1.5v5.1h2.5v1.5H8.7V8.7z" fill="white"/></svg>
            <svg class="logo-item" viewBox="0 0 24 24" fill="#ff9900"><path d="M12 2L2 7l10 5 10-5-10-5zM2 17l10 5 10-5M2 12l10 5 10-5"/></svg>
        </div>
        <p style="text-align:center; color:rgba(255,255,255,0.4); font-size:0.8rem; margin-top:5px;">Data • Dev • Design • Expertise • Rapidité</p>
        """, unsafe_allow_html=True)

        champs_manquants = []
        if not wa_display:
            champs_manquants.append("WhatsApp de contact")
        # Pour Sujets & Examens : le prompt est auto-construit via les selectbox
        # On ne vérifie PAS "Cahier des charges" mais les champs structurés
        if "Sujets" in service or "Examens" in service:
            if use_fichier_source:
                # Mode fichier : vérifier fichier + instructions
                if not contenu_fichier_source.strip():
                    champs_manquants.append("Fichier source")
                if not _fichier_instr.strip():
                    champs_manquants.append("Description du sujet souhaité")
            else:
                if not _niveau_val or _niveau_val.startswith("──"):
                    champs_manquants.append("Niveau scolaire")
                if not _matiere_val or _matiere_val.startswith("──"):
                    champs_manquants.append("Matière")
        elif "Fiche de Cours" in service:
            if not _fc_niveau_val or fc_niveau.startswith("──"):
                champs_manquants.append("Niveau scolaire")
            if not _fc_matiere_val or fc_matiere.startswith("──"):
                champs_manquants.append("Matière")
        elif "Modifier" in service and "Fichier" in service:
            if not mf_fichier:
                champs_manquants.append("Fichier à modifier")
            if not mf_instructions.strip():
                champs_manquants.append("Instructions de modification")
        else:
            if not prompt:
                champs_manquants.append("Cahier des charges")
        if champs_manquants:
            st.markdown(f"""
            <div style="
                background: rgba(241,196,15,0.08);
                border: 1px dashed rgba(241,196,15,0.4);
                border-radius: 10px;
                padding: 10px 16px;
                margin-top: 8px;
                color: rgba(241,196,15,0.85);
                font-size: 0.85rem;
            ">
                👇 Décrivez votre besoin ci-dessus puis appuyez sur le bouton du bas pour lancer la génération.
            </div>
            """, unsafe_allow_html=True)

        if premium_actif and service in SERVICES_GEMINI:
            _udata_q = st.session_state["db"]["users"].get(user, {})
            _restant  = quota_restant(_udata_q)
            _plan_q   = _udata_q.get("premium_plan", "")
            _quota_q  = PLANS_PREMIUM.get(_plan_q, {}).get("generations", 0)
            _used_q, _ = get_gen_quota(_udata_q)
            _is_illimite = _quota_q >= 999
            _couleur_quota = "#FFD700" if _is_illimite else ("#2ecc71" if _restant > 1 else ("#FFD700" if _restant == 1 else "#e74c3c"))
            if _is_illimite:
                _quota_txt = "♾️ Illimité"
            elif _restant > 0:
                _quota_txt = f"{_used_q}/{_quota_q} utilisées — ✅ {_restant} restante(s)"
            else:
                _quota_txt = f"{_used_q}/{_quota_q} utilisées — 🚫 Quota atteint"
            st.markdown(f"""
            <div style="background:linear-gradient(135deg,rgba(255,215,0,.1),rgba(255,140,0,.06));
                 border:1px solid rgba(255,215,0,.5);border-radius:12px;padding:12px 18px;margin:10px 0;">
                <span style="color:#FFD700;font-weight:800;">⚡ PREMIUM — Génération IA automatique activée</span>
                <span style="color:rgba(255,255,255,.5);font-size:.8rem;display:block;margin-top:3px;">
                    Votre document sera généré et livré en moins d'1 minute.
                </span>
                <span style="color:{_couleur_quota};font-size:.85rem;font-weight:700;display:block;margin-top:6px;">
                    📊 Générations aujourd'hui : {_quota_txt}
                </span>
            </div>""", unsafe_allow_html=True)

        label_btn = "⚡ GÉNÉRER MAINTENANT AVEC L'IA NOVA" if (premium_actif and service in SERVICES_GEMINI) else "🚀 LANCER LA GÉNÉRATION"
        if "Conversion" not in service and "OCR_AUTO" not in (prompt or "") and "OCR" not in service and "Numérisation" not in service and st.button(label_btn):
            if not user:
                st.session_state["view"] = "auth"
                st.rerun()

            elif champs_manquants and "Cahier des charges" in champs_manquants:
                # Description vide → rediriger vers Nova IA support client
                st.session_state["view"] = "nova_support_ia"
                st.rerun()

            elif premium_actif and service in SERVICES_GEMINI and not champs_manquants:
                import threading

                # ── VÉRIFICATION DU QUOTA DE GÉNÉRATIONS ──────────────────
                user_data_frais = st.session_state["db"]["users"].get(user, {})
                restant = quota_restant(user_data_frais)
                plan_actuel = user_data_frais.get("premium_plan", "")
                quota_max = PLANS_PREMIUM.get(plan_actuel, {}).get("generations", 0)
                used_auj, _ = get_gen_quota(user_data_frais)

                if restant <= 0 and quota_max < 999:
                    st.error(f"🚫 Limite de générations atteinte pour aujourd'hui ({used_auj}/{quota_max} utilisées).")
                    st.info("💡 Votre quota se renouvelle demain, ou contactez Nova pour upgrader votre plan.")
                    # Basculer en mode demande manuelle
                    st.session_state["is_glowing"] = True
                    st.rerun()
                else:
                    processing_box = st.empty()
                    processing_box.markdown(f"""
                    <div class="nova-processing">
                        <div class="nova-processing-title">⚡ GÉNÉRATION EN COURS</div>
                        <div class="nova-processing-sub">Génération automatique · {'♾️ Illimité' if quota_max >= 999 else f'Quota restant après cette génération : {restant - 1}/{quota_max}'}</div>
                    </div>""", unsafe_allow_html=True)

                    barre = st.progress(0)
                    label_prog = st.empty()
                    t_start = time.time()
                    result_holder = {}

                    def generer():
                        try:
                            # Enrichir le prompt avec le type de sujet sélectionné (pour Sujets/Examens)
                            prompt_enrichi = prompt
                            if type_sujet_selectionne and ("Sujets" in service or "Examens" in service):
                                TYPE_SUJET_LABELS_FR = {
                                    "QCM": "QCM (Questions à Choix Multiple — 4 options A/B/C/D, cases □, UN SEUL TYPE)",
                                    "VRAI_FAUX": "VRAI ou FAUX UNIQUEMENT (V/F + justification si faux, UN SEUL TYPE)",
                                    "TEXTE_TROU": "TEXTE À TROUS UNIQUEMENT (lacunaire + liste de mots, UN SEUL TYPE)",
                                    "QUESTIONS_OUVERTES": "QUESTIONS OUVERTES UNIQUEMENT (rédigées + lignes de réponse, UN SEUL TYPE)",
                                    "MIXTE": "FORMAT MIXTE (Partie 1 QCM + Partie 2 Vrai/Faux + Partie 3 Question ouverte)",
                                    "CAS_PRATIQUE": "CAS PRATIQUE / ÉTUDE DE CAS (texte CI contextualisé + questions d'analyse)",
                                    "CALCUL": "EXERCICES DE CALCUL / PROBLÈMES (chiffrés, contextualisés CI, formules rappelées)",
                                    "ETUDE_DOCUMENT": "ÉTUDE DE DOCUMENT (document support + questions d'exploitation)",
                                    "SCHEMA": "SCHÉMA À LÉGENDER (description numérotée + termes à placer + corrigé)",
                                    "DISSERTATION": "COMPOSITION / DISSERTATION GUIDÉE (sujet + méthode + plan guidé)",
                                    "DEVOIR_COMPLET": "DEVOIR COMPLET AUTHENTIQUE IVOIRIEN",
                                }
                                label_fr = TYPE_SUJET_LABELS_FR.get(type_sujet_selectionne, type_sujet_selectionne)
                                prompt_enrichi = f"""{prompt}

⚠️ TYPE DE SUJET IMPOSÉ PAR L'UTILISATEUR — RESPECTER ABSOLUMENT :
TYPE UNIQUE SÉLECTIONNÉ : {label_fr}

RÈGLE ABSOLUE : Génère UNIQUEMENT ce type d'exercice. Ne pas mélanger avec d'autres types sauf si MIXTE ou DEVOIR_COMPLET est sélectionné.
Si QCM → QCM seulement. Si VRAI_FAUX → Vrai/Faux seulement. Si TEXTE_TROU → Texte à trous seulement.
Si QUESTIONS_OUVERTES → Questions ouvertes seulement. Si CALCUL → Calculs seulement.
Si ETUDE_DOCUMENT → Étude de document seulement. Si SCHEMA → Schéma à légender seulement.
Si DISSERTATION → Composition guidée seulement. Si CAS_PRATIQUE → Cas pratique seulement.
Si DEVOIR_COMPLET → Vrai devoir ivoirien COMPLET : applique EXACTEMENT la Section Critique (en-tête officiel + structure progressive adaptée au niveau détecté). Mélange intelligent de tous les types adaptés au niveau."""
                            contenu = generer_avec_gemini(service, prompt_enrichi, user, is_premium=True, gen_used=used_auj, _plan_for_model=plan_actuel)
                            if contenu.startswith("❌"):
                                result_holder["erreur"] = contenu
                                return
                            if service == "📊 Data & Excel Analytics":
                                buf  = creer_xlsx(prompt_enrichi, user)
                                nom  = f"{user}_{service[:20].strip()}.xlsx".replace(" ", "_").replace("/", "-")
                                mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                            else:
                                buf  = creer_docx(contenu, service, user)
                                # ── PAGE DE GARDE EXPOSÉ ─────────────────────────────
                                if "Expos" in service or "exposé" in service.lower():
                                    try:
                                        from docx import Document as _DocPDG
                                        from io import BytesIO as _BytesPDG
                                        import re as _re_pdg
                                        # Extraire les métadonnées depuis la description
                                        # ── EXTRACTION DIRECTE DEPUIS LE PROMPT NOVA ──────
                                        # Le prompt est structuré avec des labels fixes
                                        _titre = ""
                                        _mat   = ""
                                        _annee = "2025-2026"
                                        _fil   = ""
                                        _niv   = ""
                                        _etab  = ""
                                        _noms  = []
                                        for _line in prompt.split("\n"):
                                            _l = _line.strip()
                                            _ll = _l.lower()
                                            def _val(line):
                                                return line.split(":", 1)[-1].strip() if ":" in line else ""
                                            if "🎯 sujet" in _ll or "sujet" in _ll and ":" in _ll:
                                                _v = _val(_l)
                                                if _v and _v != "Non précisé": _titre = _v
                                            elif "📚 matière" in _ll or ("matière" in _ll and ":" in _ll):
                                                _v = _val(_l)
                                                if _v and _v != "Non précisée": _mat = _v
                                            elif "📅 année" in _ll or ("année scolaire" in _ll and ":" in _ll):
                                                _v = _val(_l)
                                                if _v: _annee = _v
                                            elif "🏛️ filière" in _ll or ("filière" in _ll and ":" in _ll):
                                                _v = _val(_l)
                                                if _v and _v != "Non précisée": _fil = _v
                                            elif "🎓 niveau" in _ll or ("niveau" in _ll and ":" in _ll):
                                                _v = _val(_l)
                                                if _v and _v != "Non précisé": _niv = _v
                                            elif "🏢 établissement" in _ll or ("établissement" in _ll and ":" in _ll):
                                                _v = _val(_l)
                                                if _v and _v != "Non précisé": _etab = _v
                                            elif "👥 noms exposants" in _ll or ("noms exposants" in _ll and ":" in _ll):
                                                _noms_raw = _val(_l)
                                                _noms = [n.strip() for n in _noms_raw.replace(" ; ", ";").replace(",", ";").split(";") if n.strip() and n.strip() != "Non précisés"]
                                        # Priorité : lire directement depuis session_state exp_membres
                                        _membres_session = [m.strip() for m in st.session_state.get("exp_membres", []) if m.strip()]
                                        if _membres_session:
                                            _noms = _membres_session
                                        # Fallback titre : première ligne H1 du contenu Gemini
                                        if not _titre and result_holder.get("contenu"):
                                            for _gl in result_holder["contenu"].split("\n"):
                                                _gl = _gl.strip()
                                                if _gl.startswith("# "):
                                                    _t = _gl.lstrip("# ").strip()
                                                    if _t and "exposé" not in _t.lower():
                                                        _titre = _t; break
                                        if not _titre:
                                            _titre = description[:80] if len(description) < 80 else description[:80] + "..."
                                        # Logo école uploadé ?
                                        _logo_ecole = None
                                        if hasattr(st.session_state, "logo_ecole_path") and st.session_state.logo_ecole_path:
                                            _logo_ecole = st.session_state.logo_ecole_path
                                        # Créer la page de garde depuis le template Word original
                                        import zipfile as _zf, shutil as _sh, os as _os
                                        from lxml import etree as _etree
                                        from io import BytesIO as _BytesPDG2

                                        _TEMPLATE_PATH = "page_de_Garde_public.docx"
                                        _pdg_buf = _BytesPDG2()

                                        # Lire le template
                                        with _zf.ZipFile(_TEMPLATE_PATH, 'r') as _zin:
                                            _xml_bytes = _zin.read('word/document.xml')

                                        _pdg_tree = _etree.fromstring(_xml_bytes)
                                        _W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

                                        def _pdg_get_text(elem):
                                            return ''.join(t.text or '' for t in elem.iter(f'{{{_W}}}t'))

                                        # Remplacer les données dans les zones de texte
                                        for _txbx in _pdg_tree.findall(f'.//{{{_W}}}txbxContent'):
                                            _txt = _pdg_get_text(_txbx).strip()

                                            # Zone THEME :
                                            if 'THEME' in _txt and ':' in _txt:
                                                _paras = _txbx.findall(f'.//{{{_W}}}p')
                                                if len(_paras) >= 2:
                                                    _p1 = _paras[1]
                                                    _runs = _p1.findall(f'{{{_W}}}r')
                                                    if _runs:
                                                        _t = _runs[-1].find(f'{{{_W}}}t')
                                                        if _t is not None: _t.text = _titre
                                                    else:
                                                        _r = _etree.SubElement(_p1, f'{{{_W}}}r')
                                                        _t = _etree.SubElement(_r, f'{{{_W}}}t')
                                                        _t.text = _titre

                                            # Zone Matière + Année scolaire
                                            elif 'Matière' in _txt and 'Année scolaire' in _txt:
                                                import re as _re_pdg2
                                                for _p in _txbx.findall(f'.//{{{_W}}}p'):
                                                    _pt = _pdg_get_text(_p)
                                                    _runs = _p.findall(f'{{{_W}}}r')
                                                    if 'Matière' in _pt and len(_runs) >= 2:
                                                        _t = _runs[-1].find(f'{{{_W}}}t')
                                                        if _t is not None: _t.text = _mat or "—"
                                                    elif 'Année scolaire' in _pt and len(_runs) >= 2:
                                                        _t = _runs[-1].find(f'{{{_W}}}t')
                                                        if _t is not None: _t.text = _annee

                                            # Zone liste noms 1- 2- ... 7-
                                            elif '1-' in _txt and '2-' in _txt:
                                                import re as _re_pdg3
                                                for _p in _txbx.findall(f'.//{{{_W}}}p'):
                                                    _pt = _pdg_get_text(_p).strip()
                                                    _m = _re_pdg3.match(r'^(\d+)-', _pt)
                                                    if _m:
                                                        _num = int(_m.group(1))
                                                        _nom = _noms[_num - 1] if _num - 1 < len(_noms) else ''
                                                        _runs = _p.findall(f'{{{_W}}}r')
                                                        if _runs:
                                                            _t = _runs[-1].find(f'{{{_W}}}t')
                                                            if _t is not None:
                                                                _t.text = f'{_num}- {_nom}'
                                                                if _t.text[-1] == ' ':
                                                                    _t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

                                        # Écrire le XML modifié dans un BytesIO zip
                                        _xml_out = _etree.tostring(_pdg_tree, xml_declaration=True, encoding='UTF-8', standalone=True)
                                        with _zf.ZipFile(_TEMPLATE_PATH, 'r') as _zin:
                                            with _zf.ZipFile(_pdg_buf, 'w', _zf.ZIP_DEFLATED) as _zout:
                                                for _item in _zin.infolist():
                                                    if _item.filename == 'word/document.xml':
                                                        _zout.writestr(_item, _xml_out)
                                                    else:
                                                        _zout.writestr(_item, _zin.read(_item.filename))
                                        _pdg_buf.seek(0)

                                        # Charger le doc page de garde
                                        _doc_pdg = _DocPDG(_pdg_buf)
                                        # Ajouter saut de page puis le contenu principal
                                        from docx.oxml import OxmlElement as _OEpdg
                                        from docx.oxml.ns import qn as _qnpdg
                                        _p_br = _doc_pdg.add_paragraph()
                                        _r_br = _p_br.add_run()
                                        _br_el = _OEpdg("w:br")
                                        _br_el.set(_qnpdg("w:type"), "page")
                                        _r_br._r.append(_br_el)
                                        # Ajouter le contenu de l'exposé depuis buf
                                        # ── Supprimer tout ce qui précède le SOMMAIRE ──
                                        # Le doc Gemini contient : header service/client
                                        # + ligne séparatrice + page de garde Gemini
                                        # + saut de page + SOMMAIRE. On saute tout ça.
                                        buf.seek(0)
                                        _doc_contenu = _DocPDG(buf)
                                        from copy import deepcopy as _dc
                                        _elems = list(_doc_contenu.element.body)
                                        _start_idx = 0

                                        def _get_txt(el):
                                            return "".join(
                                                t.text or "" for t in el.iter()
                                                if t.tag.endswith("}t")
                                            ).strip().upper()

                                        # Priorité 1 : trouver "SOMMAIRE"
                                        for _ei, _el in enumerate(_elems):
                                            if "SOMMAIRE" in _get_txt(_el):
                                                _start_idx = _ei
                                                break
                                        else:
                                            # Priorité 2 : premier Heading 1 = INTRODUCTION
                                            for _ei, _el in enumerate(_elems):
                                                _tag = _el.tag.split("}")[-1] if "}" in _el.tag else _el.tag
                                                _style = _el.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId", "")
                                                # Chercher pStyle dans pPr
                                                _pPr = _el.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
                                                _pStyle = _pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle") if _pPr is not None else None
                                                _style_val = _pStyle.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "") if _pStyle is not None else ""
                                                if "Heading" in _style_val or "heading" in _style_val.lower():
                                                    _start_idx = _ei
                                                    break
                                            else:
                                                # Fallback : 2e saut de page
                                                _nb_sauts = 0
                                                for _ei, _el in enumerate(_elems):
                                                    import lxml.etree as _etree
                                                    _xml = _etree.tostring(_el, encoding="unicode")
                                                    if "w:br" in _xml and "page" in _xml:
                                                        _nb_sauts += 1
                                                        if _nb_sauts >= 2:
                                                            _start_idx = _ei + 1
                                                            break

                                        for _elem in _elems[_start_idx:]:
                                            _doc_pdg.element.body.append(_dc(_elem))
                                        # Sauvegarder
                                        _buf_pdg = _BytesPDG()
                                        _doc_pdg.save(_buf_pdg)
                                        _buf_pdg.seek(0)
                                        buf = _buf_pdg
                                    except Exception as _e_garde:
                                        pass  # Si erreur page garde → garder le buf original
                                type_suffix = f"_{type_sujet_selectionne}" if type_sujet_selectionne else ""
                                nom  = f"{user}_{service[:20].strip()}{type_suffix}.docx".replace(" ", "_").replace("/", "-")
                                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            result_holder["buf"]  = buf
                            result_holder["nom"]  = nom
                            result_holder["mime"] = mime
                        except Exception as e:
                            result_holder["erreur"] = f"❌ Erreur : {e}"

                    thread = threading.Thread(target=generer)
                    thread.start()

                    pct = 0
                    while thread.is_alive():
                        elapsed = time.time() - t_start
                        pct = min(int(elapsed / 60 * 90), 90)
                        barre.progress(pct)
                        label_prog.markdown(f"<p style='text-align:center;color:#FFD700;font-weight:bold;'>⚡ Génération en cours... {pct}%</p>", unsafe_allow_html=True)
                        time.sleep(0.5)
                    thread.join()

                    barre.progress(100)
                    label_prog.markdown("<p style='text-align:center;color:#2ecc71;font-weight:bold;'>✅ Document généré !</p>", unsafe_allow_html=True)
                    time.sleep(0.8)
                    barre.empty(); label_prog.empty(); processing_box.empty()

                    duree = int(time.time() - t_start)

                    if "erreur" in result_holder:
                        st.error(result_holder["erreur"])
                        st.info("💡 Votre demande a été transmise à l'équipe Nova pour traitement manuel.")
                        new_req = {
                            "id": hashlib.md5(str(datetime.now()).encode()).hexdigest()[:8],
                            "user": user, "service": service,
                            "desc": prompt, "whatsapp": normalize_wa(wa_display),
                            "status": "Traitement Nova en cours...", "incomplet": False,
                            "champs_manquants": [], "timestamp": str(datetime.now()),
                            "modele_utilise": st.session_state.get("_last_modele_gemini", "—"),
                        }
                        st.session_state["db"]["demandes"].append(new_req)
                        save_demande(new_req)
                    else:
                        # Incrémenter le compteur de générations
                        incrementer_gen(user)
                        # Upload vers Supabase Storage pour accès permanent
                        _req_id_local = hashlib.md5(str(datetime.now()).encode()).hexdigest()[:8]
                        with st.spinner("⬆️ Sauvegarde du fichier..."):
                            _url_local = upload_fichier_client(user, _req_id_local, result_holder["buf"], result_holder["nom"])
                        if _url_local and not _url_local.startswith("ERREUR"):
                            save_lien(user, service, _url_local, datetime.now().strftime("%d/%m/%Y"))
                        else:
                            save_lien(user, service, f"__local__{result_holder['nom']}", datetime.now().strftime("%d/%m/%Y"))
                        # Notifier admin + client (livraison Gemini premium)
                        wa_display_local = st.session_state["db"]["users"].get(user, {}).get("whatsapp", "—")
                        _email_client = st.session_state["db"]["users"].get(user, {}).get("email", "")
                        notifier_livraison_gemini(user, wa_display_local, _email_client, service, result_holder["nom"], demande_complete=prompt)
                        st.session_state["premium_livrable"] = {
                            "buf":     result_holder["buf"],
                            "nom":     result_holder["nom"],
                            "mime":    result_holder["mime"],
                            "service": service,
                            "duree":   duree,
                        }
                        st.session_state["db"] = load_db()
                        st.rerun()

            else:
                st.session_state["is_glowing"] = True
                st.rerun()

        if st.session_state["is_glowing"]:
            # ── VÉRIFICATION LIMITE GRATUIT (7 demandes/jour) ─────────────
            if user and not premium_actif:
                _nb_dem_today = get_demandes_gratuit_today(user)
                if _nb_dem_today >= MAX_DEMANDES_GRATUIT_PAR_JOUR:
                    st.session_state["is_glowing"] = False
                    st.error(f"🚫 Limite atteinte : vous avez soumis {_nb_dem_today}/{MAX_DEMANDES_GRATUIT_PAR_JOUR} demandes aujourd'hui (plan gratuit).")
                    st.info("💡 Votre quota se renouvelle demain à minuit. Passez Premium pour des générations illimitées ! 👑")
                    st.stop()
            # ──────────────────────────────────────────────────────────────
            progress_placeholder = st.empty()
            status_text = st.empty()
            bar = progress_placeholder.progress(0)
            for percent_complete in range(100):
                time.sleep(0.02)
                bar.progress(percent_complete + 1)
                status_text.markdown(f"<p style='text-align:center; color:#00d2ff; font-size:1.2rem; font-weight:bold;'>NOVA PROCESSING : {percent_complete + 1}%</p>", unsafe_allow_html=True)

            statut = "En attente de vérification (informations incomplètes)" if champs_manquants else "Traitement Nova en cours..."

            req_id_new = hashlib.md5(str(datetime.now()).encode()).hexdigest()[:8]
            fichier_info = ""
            if "Modifier" in service and "Fichier" in service and mf_fichier:
                with st.spinner("📤 Upload du fichier..."):
                    url_f = upload_fichier_client(
                        user if user else "guest", req_id_new,
                        mf_fichier.getvalue(), mf_fichier.name)
                if url_f.startswith("ERREUR"):
                    st.warning(f"⚠️ Fichier non uploadé ({url_f})")
                else:
                    fichier_info = f"\n📎 FICHIER CLIENT : {url_f}"
            desc_finale = (prompt if prompt else "(aucune description fournie)") + fichier_info
            new_req = {
                "id": req_id_new,
                "user": user if user else "guest",
                "service": service,
                "desc": desc_finale,
                "whatsapp": normalize_wa(wa_display) if wa_display else "(non renseigné)",
                "status": statut,
                "incomplet": bool(champs_manquants),
                "champs_manquants": champs_manquants,
                "timestamp": str(datetime.now()),
                "modele_utilise": st.session_state.get("_last_modele_gemini", "—"),
            }
            st.session_state["db"]["demandes"].append(new_req)
            save_demande(new_req)
            envoyer_notification(
                client_nom  = user if user else "Visiteur",
                client_wa   = normalize_wa(wa_display) if wa_display else "(non renseigné)",
                service     = service,
                description = desc_finale
            )
            st.session_state["db"] = load_db()
            st.session_state["is_glowing"] = False
            progress_placeholder.empty()
            status_text.empty()
            if user:
                st.success("✅ Mission enregistrée ! L'équipe Nova examinera votre demande.")

                audio_path_confirm = "confirmation.mp3"
                if os.path.exists(audio_path_confirm):
                    with open(audio_path_confirm, "rb") as f:
                        audio_b64_confirm = __import__('base64').b64encode(f.read()).decode()
                    components.html(f"""
                        <script>
                        (function() {{
                            var b64 = "{audio_b64_confirm}";
                            var binary = atob(b64);
                            var bytes = new Uint8Array(binary.length);
                            for (var i = 0; i < binary.length; i++) {{
                                bytes[i] = binary.charCodeAt(i);
                            }}
                            var blob = new Blob([bytes], {{type: "audio/mpeg"}});
                            var url = URL.createObjectURL(blob);
                            var audio = new Audio(url);
                            audio.volume = 1;
                            audio.play().catch(function(e) {{ console.log("Autoplay bloqué:", e); }});
                        }})();
                        </script>
                    """, height=1)

                st.rerun()
            else:
                st.session_state["view"] = "auth"
                st.rerun()

    with tab2:
        if not user:
            st.warning("🔒 Authentification requise pour accéder au Cloud Nova.")
        else:
            fresh_db = load_db()
            user_links = fresh_db["liens"].get(user, [])
            user_reqs = [r for r in fresh_db["demandes"] if r["user"] == user]
            
            st.markdown("""
                <div style="background: rgba(46, 204, 113, 0.1); padding: 15px; border-radius: 10px; border: 1px dashed #2ecc71; margin-bottom: 20px; text-align: center;">
                    <h2 style="color: #2ecc71; margin: 0;">📥 HUB DE TÉLÉCHARGEMENT NOVA</h2>
                    <p style="color: white; font-size: 0.9rem;">Accédez à vos actifs numériques terminés.</p>
                </div>
            """, unsafe_allow_html=True)

            if st.session_state["premium_livrable"]:
                lv = st.session_state["premium_livrable"]
                st.markdown(f"""
                <div class="livrable-auto">
                    <div class="livrable-auto-title">⚡ Livrable IA Premium</div>
                    <div style="color:rgba(255,255,255,.7);margin-top:4px;">Généré en {lv['duree']}s · {lv['service']}</div>
                </div>""", unsafe_allow_html=True)
                st.download_button(
                    "📥 TÉLÉCHARGER MON DOCUMENT",
                    data=lv["buf"], file_name=lv["nom"], mime=lv["mime"],
                    use_container_width=True
                )
                st.markdown(f"""
                <div style="background:rgba(37,211,102,0.1);border:1px solid rgba(37,211,102,0.4);border-radius:10px;padding:12px 16px;margin-top:10px;text-align:center;">
                    <div style="font-size:1.05rem;font-weight:700;color:#25D366;">📲 Un fichier encore mieux vous attend !</div>
                    <div style="color:rgba(255,255,255,0.75);font-size:0.88rem;margin-top:5px;">
                        Notre équipe retravaille votre document et vous envoie une version améliorée et mise en page sur WhatsApp sous peu.
                    </div>
                </div>
                """, unsafe_allow_html=True)
                st.divider()

            if user_links:
                col_titre, col_vider = st.columns([3, 1])
                with col_titre:
                    st.markdown("#### 📁 Mes livrables")
                with col_vider:
                    if st.button("🗑️ Vider l'historique", key="vider_historique", use_container_width=True):
                        delete_all_liens(user)
                        st.success("Historique vidé !")
                        st.rerun()

                for link in user_links:
                    if link["url"].startswith("__refus__"):
                        msg_refus = link["url"].replace("__refus__", "")
                        st.markdown(f"""
                        <div class="file-card" style="border-color:rgba(231,76,60,0.5); background:rgba(231,76,60,0.05);">
                            <div style="display:flex;justify-content:space-between;align-items:center;">
                                <div>
                                    <h3 style="color:#e74c3c;margin:0;">❌ Mission refusée — {link['name']}</h3>
                                    <p style="color:#aaa;font-size:.85rem;margin:5px 0;">Le {link.get('date','')}</p>
                                    <p style="color:#eee;font-size:.9rem;margin-top:8px;">{msg_refus}</p>
                                </div>
                            </div>
                        </div>""", unsafe_allow_html=True)
                    elif link["url"].startswith("__local__"):
                        nom_fichier_local = link["url"].replace("__local__", "")
                        st.markdown(f"""
                        <div class="file-card" style="border-color:rgba(255,215,0,.5);">
                            <div style="display:flex;justify-content:space-between;align-items:center;">
                                <div>
                                    <h3 style="color:#FFD700;margin:0;">⭐ {link['name']}</h3>
                                    <p style="color:#aaa;font-size:.85rem;margin:5px 0;">Généré le {link.get('date',"Aujourd'hui")}</p>
                                </div>
                                <span class="badge-premium">IA AUTO</span>
                            </div>
                        </div>""", unsafe_allow_html=True)
                        # Bouton retélécharger si le fichier est encore en mémoire
                        lv = st.session_state.get("premium_livrable")
                        if lv and lv.get("nom") == nom_fichier_local and lv.get("buf"):
                            st.download_button(
                                label="📥 Retélécharger le fichier",
                                data=lv["buf"].getvalue() if hasattr(lv["buf"], "getvalue") else lv["buf"],
                                file_name=lv["nom"],
                                mime=lv.get("mime", "application/octet-stream"),
                                key=f"dl_local_{nom_fichier_local}",
                                use_container_width=True,
                            )
                        else:
                            st.info("⚠️ Fichier non disponible — reconnectez-vous ou contactez Nova pour le récupérer.")
                    else:
                        st.markdown(f"""
                        <div class="file-card">
                            <div style="display: flex; justify-content: space-between; align-items: center;">
                                <div>
                                    <h3 style="color:#00d2ff; margin:0;">💎 {link['name']}</h3>
                                    <p style="color:#aaa; font-size:0.85rem; margin: 5px 0;">Finalisé le {link.get('date', "Aujourd'hui")}</p>
                                </div>
                                <a href="{link['url']}" target="_blank" style="text-decoration:none;">
                                    <button style="padding:10px 25px; background:#2ecc71; color:white; border:none; border-radius:30px; font-weight:bold; cursor:pointer; box-shadow: 0 4px 10px rgba(46,204,113,0.3);">
                                        📥 TÉLÉCHARGER
                                    </button>
                                </a>
                            </div>
                        </div>
                        <div style="background:rgba(37,211,102,0.1);border:1px solid rgba(37,211,102,0.4);border-radius:10px;padding:12px 16px;margin-top:8px;margin-bottom:4px;text-align:center;">
                            <div style="font-size:1.0rem;font-weight:700;color:#25D366;">📲 Un fichier encore mieux vous attend !</div>
                            <div style="color:rgba(255,255,255,0.75);font-size:0.85rem;margin-top:4px;">
                                Notre équipe retravaille votre document et vous envoie une version améliorée et mise en page sur WhatsApp sous peu.
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
            
            if user_reqs:
                st.markdown("#### ⏳ Missions Nova en préparation")
                for r in user_reqs:
                    st.markdown(f"""
                        <div class="file-card" style="border-left: 5px solid #f1c40f; border-color: rgba(241, 196, 15, 0.3);">
                            <div style="display: flex; justify-content: space-between; align-items: center;">
                                <div>
                                    <strong style="color: #f1c40f;">{r['service']}</strong><br>
                                    <span style="color:#eee; font-size: 0.9rem;">Status: {r['status']}</span>
                                </div>
                                <div class="spinner" style="width: 20px; height: 20px; border: 3px solid rgba(255,255,255,0.1); border-top: 3px solid #f1c40f; border-radius: 50%; animation: spin 1s linear infinite;"></div>
                            </div>
                        </div>
                        <style>@keyframes spin {{ 0% {{ transform: rotate(0deg); }} 100% {{ transform: rotate(360deg); }} }}</style>
                    """, unsafe_allow_html=True)
            
            if not user_links and not user_reqs:
                st.info("Votre espace Nova est vide. Déployez votre première tâche !")
            
            st.write("---")
            st.markdown("### 🆘 Support Nova Direct")

            # ── CHAT SUPPORT IA ────────────────────────────────────────────
            st.markdown("""
            <div style="background:linear-gradient(135deg,rgba(66,133,244,0.1),rgba(66,133,244,0.05));
                 border:1px solid rgba(66,133,244,0.4);border-radius:14px;padding:14px 18px;margin-bottom:12px;">
                <span style="color:#4285f4;font-weight:800;font-size:1rem;">🤖 Assistant Support Nova</span>
                <span style="color:rgba(255,255,255,.5);font-size:.82rem;display:block;margin-top:3px;">
                    Décrivez votre problème — Nova Platform vous aide et envoie un résumé à Nova
                </span>
            </div>
            """, unsafe_allow_html=True)

            # Initialiser historique chat support
            if "support_chat" not in st.session_state:
                st.session_state["support_chat"] = [{
                    "role": "assistant",
                    "content": f"Salut, moi c'est Nova IA, ton assistant Nova ! 👋 Comment puis-je t'aider aujourd'hui, {user} ?"
                }]
            if "support_resolu" not in st.session_state:
                st.session_state["support_resolu"] = False

            # Afficher l'historique
            for msg in st.session_state["support_chat"]:
                role_icon = "🧑" if msg["role"] == "user" else "🤖"
                role_color = "#eee" if msg["role"] == "user" else "#4285f4"
                align = "flex-end" if msg["role"] == "user" else "flex-start"
                bg = "rgba(255,255,255,0.07)" if msg["role"] == "user" else "rgba(66,133,244,0.12)"
                st.markdown(f"""
                <div style="display:flex;justify-content:{align};margin:6px 0;">
                    <div style="background:{bg};border-radius:12px;padding:10px 14px;max-width:80%;">
                        <span style="color:{role_color};font-size:.82rem;font-weight:700;">{role_icon} {"Vous" if msg["role"]=="user" else "Nova Platform"}</span>
                        <p style="color:#eee;margin:4px 0 0 0;font-size:.9rem;">{msg["content"]}</p>
                    </div>
                </div>""", unsafe_allow_html=True)

            # Zone de saisie
            if not st.session_state["support_resolu"]:
                with st.form("support_chat_form", clear_on_submit=True):
                    msg_client = st.text_input(
                        "Votre message",
                        placeholder="Ex: Mon fichier n'a pas été livré, je ne peux pas me connecter...",
                        label_visibility="collapsed"
                    )
                    col_send, col_fin = st.columns([3, 1])
                    with col_send:
                        envoyer = st.form_submit_button("📨 Envoyer", use_container_width=True)
                    with col_fin:
                        terminer = st.form_submit_button("✅ Terminer", use_container_width=True)

                if envoyer and msg_client.strip():
                    # Ajouter message client
                    st.session_state["support_chat"].append({"role": "user", "content": msg_client.strip()})

                    # Construire historique pour Gemini
                    historique_txt = "\n".join([
                        f"{'Client' if m['role']=='user' else 'Assistant'}: {m['content']}"
                        for m in st.session_state["support_chat"]
                    ])

                    prompt_support = f"""Tu es NOVA IA, l'assistant support officiel de Nova Platform.
Tu t'appelles Nova IA. Tu ne t'appelles pas Gemini, pas ChatGPT, pas Claude. Tu es Nova IA.
Tu parles toujours en français, avec bienveillance, de façon claire et directe.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📚 CE QUE TU SAIS SUR NOVA PLATFORM :
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

IDENTITÉ :
- Nova Platform est une plateforme de services IA basée en Côte d'Ivoire
- Le WhatsApp Nova pour les commandes et le support humain : {WHATSAPP_NUMBER}
- Tu es l'assistant virtuel intégré dans l'application

SERVICES DISPONIBLES :
- 📊 Data & Excel Analytics : tableaux de bord, graphiques, analyse de données
- 📖 Fiche de Cours Professeur IA : fiches pédagogiques complètes pour enseignants
- 📎 Modifier mon Fichier : modification de fichiers Word, Excel, PowerPoint
- 📝 Exposé scolaire complet IA : exposés structurés du CP au Master (PREMIUM uniquement)
- 📝 Création de Sujets & Examens : devoirs, contrôles, QCM, examens (PREMIUM auto-généré)
- 📄 Création Word (depuis zéro) : le client décrit son document, Nova le crée complet et prêt à l'emploi
- 📋 Rapport de Stage IA : rapport complet structuré (BTS, Licence, Master) · Introduction · Missions · Analyse · Conclusion
- 📚 Affiches & Reçus : supports visuels pour entreprises et associations
- 👔 CV & Lettre de Motivation : CV et lettres percutants
- 🔄 Conversion & Fichier PDF : conversion entre formats (Word↔PDF, Excel↔CSV, etc.)

PLANS PREMIUM :
- 🌅 Journalier (1 jour) : 600 FC → 2 générations IA automatiques
- 🔟 10 Jours : 1000 FC → 9 générations IA automatiques
- 👑 30 Jours : 2500 FC → Générations ILLIMITÉES
- Pour s'abonner : contacter Nova sur WhatsApp {WHATSAPP_NUMBER}

FONCTIONNEMENT :
- Compte GRATUIT : le client soumet sa demande → Nova traite manuellement → livraison par lien
- Compte PREMIUM : génération automatique par Nova Platform en moins d'1 minute → disponible dans "Mes Livrables"
- Délai moyen pour les gratuits : quelques heures selon la charge
- Les livrables sont disponibles dans l'onglet "📂 MES LIVRABLES (CLOUD)"

CONNEXION / COMPTE :
- L'identifiant c'est le nom choisi à l'inscription
- Le mot de passe c'est le numéro WhatsApp (format : 225XXXXXXXX)
- Si le client oublie son WhatsApp d'inscription → contacter Nova pour récupération
- Si le client ne peut pas se connecter → vérifier que le numéro est exact avec le préfixe 225

PROBLÈMES FRÉQUENTS ET SOLUTIONS :
- "Je ne reçois pas mon fichier" → vérifier l'onglet "Mes Livrables", rafraîchir la page, ou attendre si demande récente
- "Je ne peux pas me connecter" → vérifier identifiant exact + numéro WhatsApp avec 225 au début
- "Mon quota est épuisé" → le quota se renouvelle chaque jour à minuit
- "Le fichier généré ne correspond pas" → soumettre à nouveau avec plus de détails dans le cahier des charges
- "Je veux upgrader mon plan" → contacter Nova sur WhatsApp {WHATSAPP_NUMBER}

RÈGLES QUE TU DOIS RESPECTER :
- Ne jamais promettre un remboursement sans l'accord de Nova
- Ne jamais promettre un délai précis que tu ne peux pas garantir
- Si le problème nécessite une intervention humaine, dire au client de contacter Nova sur WhatsApp {WHATSAPP_NUMBER}
- Toujours rester positif et rassurant

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
CLIENT ACTUEL : {user}
Premium actif : {"OUI" if premium_actif else "NON — compte gratuit"}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Historique de la conversation :
{historique_txt}

Réponds UNIQUEMENT au dernier message du client. Sois concis (3-5 phrases max). Tu es Nova IA, pas Gemini.
RÈGLE ESCALADE OBLIGATOIRE :
Si le client exprime un problème grave (paiement, fichier perdu, compte bloqué, bug critique, plainte urgente), tu DOIS lui proposer cette phrase exacte à la fin de ta réponse :
"👉 Veux-tu que je transmette ton problème directement au service client Nova ? Réponds juste OUI et je m'en occupe immédiatement."
Si dans l'historique le client répond OUI ou "oui" ou "ok" ou "ouais" à cette proposition, réponds UNIQUEMENT ce texte exact sans rien d'autre :
__ESCALADE_CONFIRMEE__"""

                    with st.spinner("🤖 Nova Platform répond..."):
                        reponse_ia = generer_avec_gemini("Support", prompt_support, user)

                    if reponse_ia.startswith("❌"):
                        reponse_ia = "Désolé, je rencontre une difficulté technique. Contactez Nova directement via WhatsApp."
                    if "__ESCALADE_CONFIRMEE__" in reponse_ia:
                        _wa_sup = st.session_state["db"]["users"].get(user, {}).get("whatsapp", "—")
                        ok = envoyer_escalade_support(user, _wa_sup, st.session_state["support_chat"], "Support Nova")
                        if ok:
                            reponse_ia = "✅ C'est fait ! Ton problème a été transmis au service client Nova. Nous te recontactons très bientôt. 🙏"
                            st.session_state["support_resolu"] = True
                        else:
                            reponse_ia = f"Désolé, l'envoi a échoué. Contacte Nova directement : {WHATSAPP_NUMBER} 📲"
                    st.session_state["support_chat"].append({"role": "assistant", "content": reponse_ia})
                    st.rerun()

                if terminer and st.session_state["support_chat"]:
                    # Envoyer résumé par email
                    try:
                        import resend
                        resend.api_key = st.secrets["RESEND_API_KEY"]
                        historique_email = "\n".join([
                            f"{'🧑 Client' if m['role']=='user' else '🤖 Nova Platform'} : {m['content']}"
                            for m in st.session_state["support_chat"]
                        ])
                        resend.Emails.send({
                            "from": "Nova Platform <onboarding@resend.dev>",
                            "to": [st.secrets["EMAIL_RECEIVER"]],
                            "subject": f"🆘 Support Nova — {user}",
                            "text": f"""RÉSUMÉ CONVERSATION SUPPORT
━━━━━━━━━━━━━━━━━━━━━━━━━━━
👤 Client   : {user}
📱 WhatsApp : {db["users"].get(user, {}).get("whatsapp", "—")}
⏰ Date     : {datetime.now().strftime("%d/%m/%Y à %H:%M")}
━━━━━━━━━━━━━━━━━━━━━━━━━━━

{historique_email}

━━━━━━━━━━━━━━━━━━━━━━━━━━━
Action requise si le problème n'est pas résolu.
"""
                        })
                        st.session_state["support_resolu"] = True
                        st.success("✅ Conversation terminée — un résumé a été envoyé à Nova. Nous revenons vers vous rapidement !")
                        st.rerun()
                    except Exception as e_sup:
                        st.error(f"❌ Erreur envoi résumé : {e_sup}")
            else:
                st.success("✅ Votre demande a été transmise à Nova. Nous vous répondrons bientôt.")
                if st.button("🔄 Nouveau ticket support", key="reset_support"):
                    st.session_state["support_chat"] = []
                    st.session_state["support_resolu"] = False
                    st.rerun()

            st.divider()
            col_rel, col_sup = st.columns(2)
            with col_rel:
                relance_msg = f"Bonjour, je souhaite un status sur ma mission Nova (ID: {user})."
                wa_relance = f"https://wa.me/{WHATSAPP_NUMBER}?text={relance_msg.replace(' ', '%20')}"
                st.markdown(f'<a href="{wa_relance}" target="_blank" class="support-btn" style="border-color:#f1c40f; color:#f1c40f !important;">🔔 Relancer Nova</a>', unsafe_allow_html=True)
            with col_sup:
                if st.button("🙋 Nova Platform", key="btn_nova_support_page", use_container_width=True):
                    st.session_state["view"] = "nova_support_ia"
                    st.rerun()

        st.divider()
        if st.button("🤖 Nova IA — Créer un document", key="btn_nova_ia_page", use_container_width=True):
            st.session_state.pop("nova_ia_chat", None)
            st.session_state.pop("nova_ia_phase", None)
            st.session_state.pop("nova_ia_service_detecte", None)
            st.session_state.pop("nova_ia_prompt_final", None)
            st.session_state["view"] = "nova_ia"
            st.rerun()

    with st.expander("🛠 Console Admin Nova"):
        _saisie_key = st.text_input("Clé d'accès", type="password", key="console_key_input")
        _est_admin  = (_saisie_key == ADMIN_CODE)
        _est_collab = (_saisie_key == COLLAB_CODE)

        if not _est_admin and not _est_collab and _saisie_key:
            st.error("❌ Clé incorrecte.")

        if _est_admin or _est_collab:
            current_db = st.session_state["db"]

            # ── Onglets selon le rôle ─────────────────────────────────
            if _est_admin:
                admin_tab1, admin_tab2, admin_tab3, admin_tab4 = st.tabs([
                    "📋 MISSIONS", "👑 GESTION MEMBRES", "🗑️ STORAGE", "🔧 MIGRATION UIDs"
                ])
            else:
                # Collaborateur : missions uniquement
                st.markdown("""
                <div style="background:rgba(0,210,255,0.08);border:1px solid rgba(0,210,255,0.3);
                border-radius:10px;padding:10px 16px;margin-bottom:12px;">
                <span style="color:#00d2ff;font-weight:700;">👷 Accès Collaborateur — Vue Missions uniquement</span>
                </div>""", unsafe_allow_html=True)
                admin_tab1 = st.container()
                admin_tab2 = admin_tab3 = admin_tab4 = None

            with admin_tab1:
                st.markdown("### 🛡️ Panneau de contrôle Nova")

                # ── DEBUG EMAIL CLIENT ─────────────────────────────────────
                try:
                    _err_row = supabase.table("config").select("value").eq("key", "email_client_last_error").execute().data
                    if _err_row:
                        st.error(f"📧 Dernière erreur email client : {_err_row[0]['value']}")
                        if st.button("🗑️ Effacer l'erreur email", key="clear_email_err"):
                            supabase.table("config").delete().eq("key", "email_client_last_error").execute()
                            st.rerun()
                except:
                    pass
                # ──────────────────────────────────────────────────────────

                # ── TOGGLE RÉPONSE AUTOMATIQUE PLAN GRATUIT ───────────────
                col_toggle_l, col_toggle_r = st.columns([3, 1])
                with col_toggle_l:
                    _auto_status = "🟢 ACTIVÉ" if st.session_state["auto_reply_gratuit"] else "🔴 DÉSACTIVÉ"
                    st.markdown(f"""
                    <div style="background:rgba(255,255,255,0.05);border:1px solid rgba(255,255,255,0.15);
                         border-radius:10px;padding:12px 16px;margin-bottom:12px;">
                        <span style="font-weight:700;color:#FFD700;">🤖 Réponse automatique — Plan Gratuit</span>
                        <span style="color:rgba(255,255,255,0.5);font-size:0.82rem;display:block;margin-top:3px;">
                            Si activé : Nova Platform répond après <b>6, 8 ou 10 min</b> (variable) pour : Sujets & Examens, Fiche de Cours, CV & Lettre.
                        </span>
                        <span style="font-weight:800;font-size:0.95rem;margin-top:6px;display:block;">
                            Statut actuel : {_auto_status}
                        </span>
                    </div>
                    """, unsafe_allow_html=True)
                with col_toggle_r:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.session_state["auto_reply_gratuit"]:
                        if st.button("🔴 Désactiver", key="btn_toggle_auto", use_container_width=True):
                            st.session_state["auto_reply_gratuit"] = False
                            set_auto_reply_setting(False)
                            st.success("✅ Réponse automatique désactivée")
                            st.rerun()
                    else:
                        if st.button("🟢 Activer", key="btn_toggle_auto", use_container_width=True):
                            st.session_state["auto_reply_gratuit"] = True
                            set_auto_reply_setting(True)
                            st.success("✅ Réponse automatique activée — Nova Platform répondra après 6, 8 ou 10 min (variable)")
                            st.rerun()

                st.divider()

                # ── DEBUG AUTO-REPLY ──────────────────────────────────────
                with st.expander("🔍 Diagnostic auto-reply (debug)"):
                    try:
                        _err = supabase.table("config").select("key,value").like("key", "auto%error%").execute().data
                        if _err:
                            for _e in _err:
                                st.error(f"**{_e['key']}** → {_e['value']}")
                            if st.button("🗑️ Effacer les erreurs", key="clear_auto_errors"):
                                for _e in _err:
                                    supabase.table("config").delete().eq("key", _e["key"]).execute()
                                st.rerun()
                        else:
                            st.success("✅ Aucune erreur auto-reply enregistrée")
                    except Exception as _de:
                        st.warning(f"Impossible de lire les erreurs : {_de}")

                st.divider()

                if not current_db["demandes"]:
                    st.info("✅ Aucune mission en attente.")

                def wa_url(numero, texte):
                    encoded = texte.replace(" ", "%20").replace("'", "%27").replace("\n", "%0A")
                    return f"https://wa.me/{numero}?text={encoded}"

                for i, req in enumerate(current_db["demandes"]):
                    client_wa_raw    = req.get("whatsapp", "(non renseigné)")
                    client_wa        = normalize_wa(client_wa_raw)
                    client_nom       = req.get("user", "Inconnu")
                    service          = req.get("service", "—")
                    description      = req.get("desc", "(aucune description)")
                    req_id           = req.get("id", f"{i+1}")
                    timestamp        = req.get("timestamp", "")[:16] if req.get("timestamp") else "—"
                    est_incomplet    = req.get("incomplet", False)
                    champs_manquants = req.get("champs_manquants", [])
                    client_premium   = is_premium_actif(current_db["users"].get(client_nom, {}))

                    if i > 0:
                        st.divider()

                    col_info, col_badge = st.columns([4, 1])
                    with col_info:
                        st.markdown(f"**Mission `#{req_id}`** · {timestamp}" + (" — ⚠️ *Incomplet : " + ", ".join(champs_manquants) + "*" if est_incomplet else ""))
                        st.markdown(f"👤 **Client :** {client_nom}")
                        st.markdown(f"📱 **WhatsApp :** {client_wa}")
                        st.markdown(f"🛠️ **Service demandé :** {service}")
                        _modele_log = req.get("modele_utilise", "—")
                        if _modele_log and _modele_log != "—":
                            _color = "#FFD700" if "pro" in _modele_log else ("#00aaff" if "flash" in _modele_log and "lite" not in _modele_log else "#aaaaaa")
                            st.markdown(f"🧠 **Modèle IA :** <span style='color:{_color};font-weight:700;font-family:monospace;'>{_modele_log}</span>", unsafe_allow_html=True)
                        st.markdown(f"📝 **Détails de la demande :** {description}")
                        if "Modifier" in service and "Fichier" in service:
                            _url_dl = None
                            _nom_dl = "fichier_client"
                            for _l in description.split("\n"):
                                if "📎 FICHIER CLIENT :" in _l:
                                    _url_dl = _l.replace("📎 FICHIER CLIENT :", "").strip()
                                if "FICHIER" in _l and ":" in _l and "MODIF" not in _l and "NOTE" not in _l and "📎 FICHIER CLIENT" not in _l:
                                    _nom_dl = _l.split(":")[-1].strip()
                            if _url_dl and _url_dl.startswith("http"):
                                st.markdown(f'''<a href="{_url_dl}" target="_blank"
                                   style="display:inline-flex;align-items:center;gap:8px;
                                   background:rgba(0,210,255,0.12);border:1px solid rgba(0,210,255,0.5);
                                   border-radius:10px;padding:10px 20px;color:#00d2ff;
                                   font-weight:700;text-decoration:none;margin-top:6px;">
                                   📥 TÉLÉCHARGER LE FICHIER — {_nom_dl}</a>''',
                                   unsafe_allow_html=True)
                            else:
                                st.info("📎 Fichier non disponible")
                    with col_badge:
                        if client_premium:
                            st.markdown('<span class="badge-premium">⭐ PREMIUM</span>', unsafe_allow_html=True)
                        else:
                            st.markdown('<span class="badge-free">🔓 Gratuit</span>', unsafe_allow_html=True)

                    if est_incomplet and champs_manquants:
                        champs_str = ", ".join(champs_manquants)
                        msg_rejet = (f"Bonjour {client_nom}, nous avons reçu votre demande Nova Platform "
                                     f"concernant : {service}. Cependant, nous ne pouvons pas la traiter "
                                     f"car les informations suivantes sont manquantes : {champs_str}. "
                                     f"Merci de soumettre à nouveau votre demande en complétant tous les champs. "
                                     f"— Équipe Nova Platform ⚡")
                    else:
                        msg_rejet = (f"Bonjour {client_nom}, nous avons bien reçu votre demande Nova Platform "
                                     f"concernant : {service}. Malheureusement, nous ne sommes pas en mesure "
                                     f"de traiter cette mission pour le moment. Merci de nous recontacter. "
                                     f"— Équipe Nova Platform ⚡")

                    msg_recu = (
                        f"📬 Bonjour {client_nom} !\n\n"
                        f"Nova Platform a bien reçu votre demande de *{service}*.\n"
                        f"Votre mission est en cours de traitement par notre équipe.\n\n"
                        f"Vous serez notifié dès que votre livrable sera prêt.\n"
                        f"Merci de votre confiance ! 🙏\n\n"
                        f"— Équipe Nova Platform ⚡"
                    )

                    msg_succes = (
                        f"✅ Bonjour {client_nom} !\n\n"
                        f"Votre *{service}* est terminé et disponible !\n\n"
                        f"👉 Connectez-vous à votre espace Nova et rendez-vous dans *Mes Livrables* pour télécharger votre document.\n\n"
                        f"Merci de votre confiance — à très bientôt ! 😊\n\n"
                        f"— Équipe Nova Platform ⚡"
                    )

                    # Lien WhatsApp direct (sans message prédéfini)
                    wa_direct = f"https://wa.me/{client_wa}" if client_wa and client_wa != "(non renseigné)" else None

                    col_rejet, col_contact, col_recu, col_succes = st.columns(4)
                    with col_rejet:
                        if st.button("❌ Rejeter", key=f"rej_{i}", use_container_width=True):
                            save_refus(client_nom, service, msg_rejet)
                            delete_demande(req["id"])
                            st.session_state["db"] = load_db()
                            st.success(f"Mission refusée — notification envoyée dans les livrables de {client_nom}.")
                            st.rerun()
                    with col_contact:
                        if wa_direct:
                            st.markdown(f'<a href="{wa_direct}" target="_blank" style="display:block; text-align:center; padding:10px; border-radius:10px; background:linear-gradient(135deg,rgba(37,211,102,0.2),rgba(18,140,126,0.15)); border:2px solid #25d366; color:#25d366; font-weight:800; text-decoration:none; font-size:0.85rem;">📞 Contacter</a>', unsafe_allow_html=True)
                        else:
                            st.markdown('<div style="text-align:center;padding:10px;color:rgba(255,255,255,0.3);font-size:0.8rem;">Pas de WA</div>', unsafe_allow_html=True)
                    with col_recu:
                        st.markdown(f'<a href="{wa_url(client_wa, msg_recu)}" target="_blank" style="display:block; text-align:center; padding:10px; border-radius:10px; background:rgba(255,215,0,0.1); border:1px solid rgba(255,215,0,0.4); color:#FFD700; font-weight:700; text-decoration:none;">📬 Reçu</a>', unsafe_allow_html=True)
                    with col_succes:
                        st.markdown(f'<a href="{wa_url(client_wa, msg_succes)}" target="_blank" style="display:block; text-align:center; padding:10px; border-radius:10px; background:rgba(46,204,113,0.15); border:1px solid rgba(46,204,113,0.5); color:#2ecc71; font-weight:700; text-decoration:none;">✅ Succès</a>', unsafe_allow_html=True)

                    # ── Vérifie si le service est éligible à la génération Gemini ──
                    # (compatible avec les noms avec ou sans emoji — ex. demandes du chat Nova IA)
                    import re as _re
                    def _strip_emoji(s):
                        return _re.sub(r'[^\w\s\-&/]', '', s).strip().lower()
                    _svc_norm = _strip_emoji(service)
                    _gemini_eligible = service in SERVICES_GEMINI or any(
                        _svc_norm in _strip_emoji(s) or _strip_emoji(s) in _svc_norm
                        for s in SERVICES_GEMINI
                    ) or (
                        # Demandes du chat Nova IA : service détecté libre mais description riche
                        bool(description.strip()) and len(description.strip()) > 30
                        and service not in ("—", "", "Demande Nova IA")
                    )

                    if _gemini_eligible:
                        st.markdown("<br>", unsafe_allow_html=True)
                        st.markdown(f"""
                        <div class="gemini-card">
                            <div class="gemini-title">🤖 Nova Platform — GÉNÉRATION AUTOMATIQUE DISPONIBLE</div>
                            <div class="gemini-sub">Génère le document complet en .docx en 30-60 secondes</div>
                        </div>
                        """, unsafe_allow_html=True)

                        if st.button(f"🔍 Voir modèles disponibles", key=f"diag_{req_id}"):
                            with st.spinner("Interrogation de l'API Nova Platform..."):
                                modeles_dispo = get_modeles_disponibles(st.secrets["GEMINI_API_KEY"])
                            if modeles_dispo:
                                st.success(f"✅ {len(modeles_dispo)} modèles trouvés :")
                                for m in modeles_dispo:
                                    st.code(m)
                            else:
                                st.error("❌ Aucun modèle disponible — vérifiez votre clé API.")

                        if st.button(f"⚡ APPROUVER & GÉNÉRER AVEC NOVA IA", key=f"gemini_{req_id}", use_container_width=True):
                            with st.spinner("🔍 Détection automatique du meilleur modèle disponible..."):
                                modeles_dispo = get_modeles_disponibles(st.secrets["GEMINI_API_KEY"])
                                if modeles_dispo:
                                    st.info(f"✅ Modèle sélectionné : **{modeles_dispo[0]}**")
                                else:
                                    st.error("❌ Aucun modèle Gemini disponible pour cette clé API.")
                            with st.spinner("🤖 Nova Platform génère le document... (30-60 secondes)"):
                                contenu = generer_avec_gemini(service, description, client_nom)

                            if contenu.startswith("❌"):
                                st.error(contenu)
                            else:
                                st.session_state["gemini_results"][req_id] = {
                                    "contenu": contenu,
                                    "service": service,
                                    "client": client_nom
                                }
                                st.success("✅ Document généré avec succès !")
                                st.rerun()

                        if req_id in st.session_state["gemini_results"]:
                            result = st.session_state["gemini_results"][req_id]

                            with st.expander("👁️ Aperçu du contenu généré", expanded=False):
                                st.markdown(result["contenu"])

                            if st.button("🚀 LIVRER DIRECTEMENT AU CLIENT", key=f"livrer_auto_{req_id}", use_container_width=True):
                                try:
                                    SERVICE_EXCEL = "📊 Data & Excel Analytics"
                                    if result["service"] == SERVICE_EXCEL:
                                        buf = creer_xlsx(result.get("desc", ""), result["client"])
                                        nom_fichier = f"{client_nom}_Suivi_Depenses.xlsx".replace(" ", "_")
                                        mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                                    else:
                                        buf = creer_docx(result["contenu"], result["service"], result["client"])
                                        nom_fichier = f"{client_nom}_{result['service'][:20].strip()}.docx".replace(" ", "_").replace("/", "-")
                                        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                                    with st.spinner("⬆️ Upload du fichier vers Supabase..."):
                                        url_fichier = upload_fichier_client(client_nom, req_id, buf, nom_fichier)

                                    if url_fichier.startswith("ERREUR"):
                                        st.error(f"❌ Upload échoué : {url_fichier}")
                                    else:
                                        save_lien(client_nom, service, url_fichier, datetime.now().strftime("%d/%m/%Y"))
                                        delete_demande(req["id"])
                                        if req_id in st.session_state["gemini_results"]:
                                            del st.session_state["gemini_results"][req_id]
                                        # Notifier admin + client
                                        _email_c = st.session_state["db"]["users"].get(client_nom, {}).get("email", "")
                                        _wa_c = st.session_state["db"]["users"].get(client_nom, {}).get("whatsapp", "—")
                                        notifier_livraison_gemini(client_nom, _wa_c, _email_c, service, nom_fichier)
                                        st.session_state["db"] = load_db()
                                        st.success(f"✅ Fichier livré directement dans les livrables de {client_nom} !")
                                        st.rerun()
                                except Exception as e:
                                    st.error(f"Erreur livraison : {e}")

                    st.markdown("<br>", unsafe_allow_html=True)
                    url_dl = st.text_input("🔗 Lien de livraison manuel (optionnel)", key=f"url_{i}", placeholder="https://drive.google.com/...")
                    if st.button("📦 LIVRER AVEC LIEN MANUEL", key=f"btn_{i}", use_container_width=True):
                        if url_dl:
                            save_lien(req['user'], req['service'], url_dl, datetime.now().strftime("%d/%m/%Y"))
                            delete_demande(req['id'])
                            if req_id in st.session_state["gemini_results"]:
                                del st.session_state["gemini_results"][req_id]
                            # Notifier admin + client
                            _email_c = st.session_state["db"]["users"].get(req['user'], {}).get("email", "")
                            _wa_c = st.session_state["db"]["users"].get(req['user'], {}).get("whatsapp", "—")
                            notifier_livraison_gemini(req['user'], _wa_c, _email_c, req['service'], "Votre fichier")
                            st.session_state["db"] = load_db()
                            st.success(f"✅ Mission livrée à {client_nom} !")
                            st.rerun()

            if admin_tab2 is not None:
             with admin_tab2:
                st.markdown("### 👑 Gestion des membres")
                total  = len(current_db["users"])
                prems  = [u for u, d in current_db["users"].items() if is_premium_actif(d)]
                c1, c2, c3 = st.columns(3)
                c1.metric("👥 Total membres", total)
                c2.metric("⭐ Premium actifs", len(prems))
                c3.metric("🔓 Gratuits", total - len(prems))
                st.divider()

                # ══ SECTION 1 : ACTIVER PREMIUM ══════════════════════════
                st.markdown("#### ➕ Activer un Premium")
                recherche = st.text_input(
                    "🔍 Rechercher un membre par numéro WhatsApp",
                    placeholder="Ex: 2250707...  ou  0707...",
                    key="admin_search_premium"
                )
                tous_users = current_db["users"]
                if recherche.strip():
                    r = recherche.strip().lower().replace(" ", "").replace("+", "")
                    users_filtres = {
                        uid: d for uid, d in tous_users.items()
                        if r in str(d.get("whatsapp", "")).lower() or r in uid.lower()
                    }
                else:
                    users_filtres = tous_users

                if not users_filtres:
                    st.warning("❌ Aucun membre trouvé pour ce numéro.")
                else:
                    options_keys = list(users_filtres.keys())
                    co1, co2, co3 = st.columns([3, 2, 1])
                    with co1:
                        uid_target = st.selectbox(
                            "Membre",
                            options=options_keys,
                            format_func=lambda u: (
                                f"{'⭐' if is_premium_actif(tous_users[u]) else '🔓'}  "
                                f"📱 {tous_users[u].get('whatsapp', u)}"
                                f"  · Inscrit {str(tous_users[u].get('joined',''))[:10]}"
                            ),
                            key="admin_select_uid"
                        )
                    with co2:
                        plan_ch = st.selectbox(
                            "Plan",
                            list(PLANS_PREMIUM.keys()),
                            format_func=lambda p: f"{PLANS_PREMIUM[p]['emoji']} {p} — {PLANS_PREMIUM[p]['prix']}",
                            key="admin_select_plan"
                        )
                    with co3:
                        st.markdown("<br>", unsafe_allow_html=True)
                        if st.button("⚡ ACTIVER", key="btn_act_global"):
                            activer_premium(uid_target, plan_ch)
                            st.session_state["db"] = load_db()
                            wa_cible = tous_users[uid_target].get("whatsapp", uid_target)
                            st.success(f"✅ Premium **{plan_ch}** activé pour 📱 **{wa_cible}** !")
                            st.rerun()

                st.divider()

                # ══ SECTION 2 : BONUS DE GÉNÉRATIONS GRATUITES ═══════════
                st.markdown("#### 🎁 Bonus de générations gratuites")
                st.caption("Donne des générations gratuites valables aujourd'hui (1 jour) à n'importe quel membre.")
                if users_filtres:
                    bon1, bon2, bon3 = st.columns([3, 1, 1])
                    with bon1:
                        uid_bonus = st.selectbox(
                            "Membre à créditer",
                            options=list(tous_users.keys()),
                            format_func=lambda u: (
                                f"{'⭐' if is_premium_actif(tous_users[u]) else '🔓'}  "
                                f"📱 {tous_users[u].get('whatsapp', u)}"
                            ),
                            key="admin_select_bonus"
                        )
                    with bon2:
                        nb_bonus = st.selectbox(
                            "Nb générations",
                            options=[1, 2, 3, 4, 5, 6, 7, 8, 10],
                            index=1,
                            key="admin_nb_bonus"
                        )
                    with bon3:
                        st.markdown("<br>", unsafe_allow_html=True)
                        if st.button("🎁 DONNER", key="btn_donner_bonus"):
                            ok = donner_bonus_gen(uid_bonus, nb_bonus)
                            wa_b = tous_users[uid_bonus].get("whatsapp", uid_bonus)
                            if ok:
                                st.session_state["db"] = load_db()
                                st.success(f"✅ **{nb_bonus}** génération(s) bonus données à 📱 **{wa_b}** — valables aujourd'hui !")
                            else:
                                st.error("❌ Erreur lors de l'attribution du bonus.")

                st.divider()

                # ══ SECTION 3 : LISTE DES MEMBRES ════════════════════════
                st.markdown("#### 👥 Liste des membres")
                filtre = st.radio("Afficher", ["Tous", "Premium uniquement", "Gratuits uniquement"], horizontal=True)

                for uid_m, udata in current_db["users"].items():
                    p_actif = is_premium_actif(udata)
                    p_info  = get_premium_info(udata)
                    if filtre == "Premium uniquement" and not p_actif: continue
                    if filtre == "Gratuits uniquement" and p_actif:    continue

                    wa_affiche = udata.get("whatsapp", uid_m)
                    # Afficher le bonus gratuit si applicable
                    bonus_q, bonus_u = get_bonus_gen_gratuit(uid_m)
                    bonus_txt = (f"<br><small style='color:#2ecc71;'>🎁 Bonus aujourd'hui : "
                                 f"{bonus_q - bonus_u}/{bonus_q} restant(s)</small>") if bonus_q > 0 else ""

                    col_m, col_a = st.columns([3, 2])
                    with col_m:
                        badge = f'<span class="badge-premium">⭐ {udata.get("premium_plan","—")}</span>' if p_actif else '<span class="badge-free">🔓 Gratuit</span>'
                        exp_txt = f"<br><small style='color:rgba(255,215,0,.6);'>Expire : {p_info['expiry']} ({p_info['jours_restants']}j)</small>" if p_actif and p_info else ""
                        st.markdown(f"""<div class="admin-premium-row">
                            <div>
                                <div class="admin-user-name">📱 {wa_affiche}</div>
                                <div class="admin-user-meta">Inscrit le {str(udata.get('joined',''))[:10]}</div>
                                {exp_txt}{bonus_txt}
                            </div>
                            <div>{badge}</div>
                        </div>""", unsafe_allow_html=True)
                    with col_a:
                        if p_actif:
                            cp1, cp2, cp3 = st.columns(3)
                            with cp1:
                                ext_p = st.selectbox("", list(PLANS_PREMIUM.keys()), key=f"ext_{uid_m}",
                                    format_func=lambda p: f"{PLANS_PREMIUM[p]['emoji']} {p}")
                                if st.button("➕ Prolonger", key=f"pro_{uid_m}"):
                                    curr_exp = datetime.fromisoformat(udata.get("premium_expiry", datetime.now().isoformat()))
                                    new_exp  = max(curr_exp, datetime.now()) + timedelta(days=PLANS_PREMIUM[ext_p]["jours"])
                                    update_premium_status(uid_m, True, ext_p, new_exp.isoformat())
                                    st.session_state["db"] = load_db()
                                    st.success(f"✅ Prolongé jusqu'au {new_exp.strftime('%d/%m/%Y')} !")
                                    st.rerun()
                            with cp2:
                                st.markdown("<br>", unsafe_allow_html=True)
                                if st.button("🚫 Révoquer", key=f"rev_{uid_m}"):
                                    desactiver_premium(uid_m)
                                    st.session_state["db"] = load_db()
                                    st.warning(f"Premium révoqué pour 📱 {wa_affiche}.")
                                    st.rerun()
                            with cp3:
                                st.markdown("<br>", unsafe_allow_html=True)
                                if st.button("🗑️ Suppr.", key=f"del_{uid_m}"):
                                    st.session_state[f"confirm_del_{uid_m}"] = True
                        else:
                            cp1, cp2 = st.columns(2)
                            with cp1:
                                ap = st.selectbox("", list(PLANS_PREMIUM.keys()), key=f"act_{uid_m}",
                                    format_func=lambda p: f"{PLANS_PREMIUM[p]['emoji']} {p}")
                                if st.button("⚡ Activer", key=f"actbtn_{uid_m}"):
                                    activer_premium(uid_m, ap)
                                    st.session_state["db"] = load_db()
                                    st.success(f"✅ Premium activé pour 📱 {wa_affiche} !")
                                    st.rerun()
                            with cp2:
                                st.markdown("<br>", unsafe_allow_html=True)
                                if st.button("🗑️ Supprimer", key=f"del_{uid_m}"):
                                    st.session_state[f"confirm_del_{uid_m}"] = True

                        # ── Confirmation suppression ──────────────────
                        if st.session_state.get(f"confirm_del_{uid_m}"):
                            st.error(f"⚠️ Supprimer définitivement **📱 {wa_affiche}** ? (données + livrables + historique)")
                            cc1, cc2 = st.columns(2)
                            with cc1:
                                if st.button("✅ OUI, supprimer", key=f"confirm_yes_{uid_m}", type="primary"):
                                    delete_user(uid_m)
                                    st.session_state.pop(f"confirm_del_{uid_m}", None)
                                    st.session_state["db"] = load_db()
                                    st.success(f"✅ Membre 📱 {wa_affiche} supprimé définitivement.")
                                    st.rerun()
                            with cc2:
                                if st.button("❌ Annuler", key=f"confirm_no_{uid_m}"):
                                    st.session_state.pop(f"confirm_del_{uid_m}", None)
                                    st.rerun()

            if admin_tab3 is not None:
             with admin_tab3:
                st.markdown("### 🗑️ Nettoyage Storage — Fichiers +20 jours")
                st.info("Supprime du Storage Supabase et de la table `liens` tous les fichiers livrés depuis plus de 20 jours.")

                col_dry, col_purge = st.columns(2)

                with col_dry:
                    if st.button("🔍 Simuler (sans supprimer)", key="btn_dry_run", use_container_width=True):
                        with st.spinner("Analyse en cours..."):
                            res = purger_fichiers_anciens(jours=20, dry_run=True)
                        if res["supprimes"]:
                            st.warning(f"⚠️ {len(res['supprimes'])} fichier(s) seraient supprimés :")
                            for f in res["supprimes"]:
                                st.markdown(f"- `{f['uid']}` · {f['date']} · `{f['chemin'].split('/')[-1]}`")
                        else:
                            st.success("✅ Aucun fichier à supprimer pour l'instant.")
                        st.caption(f"{res['ignores']} entrée(s) ignorée(s) (refus, liens spéciaux, non expirés)")

                with col_purge:
                    if st.button("🗑️ Lancer le nettoyage", key="btn_purge", use_container_width=True, type="primary"):
                        with st.spinner("Nettoyage en cours..."):
                            res = purger_fichiers_anciens(jours=20, dry_run=False)
                        if res["supprimes"]:
                            st.success(f"✅ {len(res['supprimes'])} fichier(s) supprimé(s) avec succès.")
                            for f in res["supprimes"]:
                                st.markdown(f"- `{f['uid']}` · {f['date']} · `{f['chemin'].split('/')[-1]}`")
                        else:
                            st.info("Aucun fichier à supprimer.")
                        if res["erreurs"]:
                            st.error(f"{len(res['erreurs'])} erreur(s) :")
                            for e in res["erreurs"]:
                                st.caption(f"❌ {e}")
                        st.caption(f"{res['ignores']} entrée(s) ignorée(s)")
                        st.session_state["db"] = load_db()


            if admin_tab4 is not None:
             with admin_tab4:
                st.markdown("### 🔧 Migration — Corriger les anciens UIDs")
                st.info(
                    "Certains anciens comptes ont un uid aléatoire (ex: `8a387850d656f49f`) au lieu "
                    "de leur numéro WhatsApp. Cet outil les corrige en Supabase automatiquement."
                )

                tous = current_db["users"]
                # Détecter les comptes dont l'uid != whatsapp normalisé
                a_migrer = {
                    uid: d for uid, d in tous.items()
                    if uid != d.get("whatsapp", "")
                }

                if not a_migrer:
                    st.success("✅ Tous les comptes ont déjà leur numéro WhatsApp comme identifiant.")
                else:
                    st.warning(f"⚠️ **{len(a_migrer)} compte(s)** avec un ancien uid détecté(s) :")
                    for uid_old, d in a_migrer.items():
                        wa = d.get("whatsapp", "—")
                        st.markdown(
                            f"- `{uid_old}` → sera renommé en `{wa}`"
                            + (f" ⭐ {d.get('premium_plan')}" if d.get('premium') else "")
                        )

                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("🔄 Lancer la migration", key="btn_migrate_uids", type="primary", use_container_width=True):
                        ok, ko = 0, []
                        for uid_old, d in a_migrer.items():
                            wa = d.get("whatsapp", "")
                            if not wa:
                                ko.append(f"{uid_old} — pas de numéro WhatsApp")
                                continue
                            try:
                                # 1. Créer le nouveau compte avec le bon uid
                                supabase.table("users").upsert({
                                    "uid": wa,
                                    "whatsapp": wa,
                                    "email": d.get("email", "Non renseigné"),
                                    "joined": d.get("joined", str(datetime.now())),
                                    "premium": d.get("premium", False),
                                    "premium_plan": d.get("premium_plan", None),
                                    "premium_expiry": d.get("premium_expiry", None),
                                    "gen_used": d.get("gen_used", 0),
                                    "gen_date": d.get("gen_date", None),
                                }).execute()
                                # 2. Migrer les demandes liées
                                supabase.table("demandes").update({"uid": wa}).eq("uid", uid_old).execute()
                                # 3. Migrer les liens livrables
                                supabase.table("liens").update({"uid": wa}).eq("uid", uid_old).execute()
                                # 4. Supprimer l'ancien uid
                                supabase.table("users").delete().eq("uid", uid_old).execute()
                                ok += 1
                            except Exception as e_mig:
                                ko.append(f"{uid_old} → {wa} : {e_mig}")

                        st.session_state["db"] = load_db()
                        if ok:
                            st.success(f"✅ {ok} compte(s) migré(s) avec succès. Recharge la page.")
                        if ko:
                            st.error(f"❌ {len(ko)} erreur(s) :")
                            for err in ko:
                                st.caption(err)
                        st.rerun()


inject_custom_css()

# localStorage géré en haut du fichier (nova_user_id + nova_user_ts)

# Masquer l'iframe vide créée par components.html
st.markdown("""
    <style>
    iframe[title="components.v1.html"] { display: none !important; height: 0 !important; }
    </style>
""", unsafe_allow_html=True)

def show_nova_support_ia_page():
    user = st.session_state.get("current_user", "")
    st.markdown("""
    <style>
    @keyframes goldGlow {
        0%   { box-shadow: 0 0 12px 4px rgba(255,193,7,0.3), 0 0 30px 8px rgba(255,140,0,0.15); }
        50%  { box-shadow: 0 0 30px 10px rgba(255,215,0,0.7), 0 0 60px 20px rgba(255,140,0,0.35); }
        100% { box-shadow: 0 0 12px 4px rgba(255,193,7,0.3), 0 0 30px 8px rgba(255,140,0,0.15); }
    }
    @keyframes goldShimmer {
        0%   { background-position: -200% center; }
        100% { background-position: 200% center; }
    }
    @keyframes floatBot {
        0%, 100% { transform: translateY(0px); }
        50%       { transform: translateY(-8px); }
    }
    @keyframes goldPulseText {
        0%, 100% { opacity: 1; text-shadow: 0 0 8px rgba(255,215,0,0.5); }
        50%       { opacity: 0.85; text-shadow: 0 0 20px rgba(255,215,0,0.9); }
    }
    @keyframes starSpin {
        0%   { transform: rotate(0deg) scale(1); }
        50%  { transform: rotate(180deg) scale(1.2); }
        100% { transform: rotate(360deg) scale(1); }
    }

    .nova-assistant-page {
        background: linear-gradient(160deg, rgba(20,15,5,0.98), rgba(30,22,5,0.95));
        min-height: 100vh;
    }
    .arsene-header-gold {
        background: linear-gradient(135deg,
            rgba(255,215,0,0.12) 0%,
            rgba(255,140,0,0.08) 40%,
            rgba(255,215,0,0.06) 100%);
        border: 2px solid;
        border-image: linear-gradient(135deg, #FFD700, #FFA500, #FFD700) 1;
        border-radius: 22px;
        padding: 36px 28px 28px 28px;
        text-align: center;
        animation: goldGlow 2.5s ease-in-out infinite;
        margin-bottom: 24px;
        position: relative;
        overflow: hidden;
    }
    .arsene-header-gold::before {
        content: "";
        position: absolute;
        top: 0; left: -100%; right: 0; bottom: 0;
        background: linear-gradient(90deg,
            transparent 0%,
            rgba(255,215,0,0.08) 40%,
            rgba(255,255,255,0.12) 50%,
            rgba(255,215,0,0.08) 60%,
            transparent 100%);
        background-size: 200% 100%;
        animation: goldShimmer 3s linear infinite;
    }
    .arsene-bot-icon {
        display: none;  /* supprimé — remplacé par le nom en grand */
    }
    @keyframes arsene-name-glow {
        0%,100% { filter: drop-shadow(0 0 14px rgba(255,215,0,0.5)); }
        50%      { filter: drop-shadow(0 0 32px rgba(255,215,0,0.95)); }
    }
    @keyframes arsene-cursor-blink {
        0%,100% { opacity:1; }
        50%      { opacity:0; }
    }
    @keyframes arsene-slogan-in {
        0%   { opacity:0; transform:translateX(-16px); }
        100% { opacity:1; transform:translateX(0); }
    }
    .arsene-ia-name {
        background: linear-gradient(135deg, #FFD700 0%, #FFA500 30%, #FFD700 55%, #FFFACD 70%, #FFD700 100%);
        background-size: 250% auto;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-size: 3.6rem;
        font-weight: 900;
        letter-spacing: 6px;
        text-transform: uppercase;
        display: block;
        line-height: 1.1;
        animation: goldShimmer 3s linear infinite, arsene-name-glow 2.5s ease-in-out infinite;
    }
    .arsene-ia-cursor {
        display: inline-block;
        width: 3px; height: 3.4rem;
        background: #FFD700;
        vertical-align: middle;
        margin-left: 4px;
        border-radius: 2px;
        animation: arsene-cursor-blink 1s step-end infinite;
        box-shadow: 0 0 10px rgba(255,215,0,0.8);
    }
    .arsene-ia-slogan {
        display: block;
        margin-top: 14px;
        font-size: 1.05rem;
        font-weight: 600;
        color: rgba(255,240,180,0.82);
        letter-spacing: 0.5px;
        animation: arsene-slogan-in 1s ease-out 0.4s both;
    }
    .arsene-ia-slogan em {
        color: #FFD700;
        font-style: normal;
        font-weight: 700;
    }
    .arsene-title-gold {
        display: none;  /* remplacé par arsene-ia-name */
    }
    .arsene-badge {
        display: inline-block;
        background: linear-gradient(135deg, rgba(255,215,0,0.2), rgba(255,140,0,0.15));
        border: 1px solid rgba(255,215,0,0.5);
        border-radius: 30px;
        padding: 5px 16px;
        margin-top: 10px;
        font-size: 0.82rem;
        color: #FFD700;
        font-weight: 700;
        letter-spacing: 0.5px;
    }
    .arsene-stars {
        font-size: 1.1rem;
        display: block;
        margin-top: 8px;
        animation: starSpin 4s linear infinite;
        display: inline-block;
    }
    .arsene-divider {
        height: 2px;
        background: linear-gradient(90deg, transparent, #FFD700, #FFA500, #FFD700, transparent);
        border: none;
        margin: 16px 0;
        border-radius: 2px;
    }
    .msg-assistant-gold {
        background: linear-gradient(135deg, rgba(255,193,7,0.1), rgba(255,140,0,0.06));
        border: 1px solid rgba(255,215,0,0.25);
        border-left: 3px solid #FFD700;
        border-radius: 0 14px 14px 14px;
    }
    .msg-user-gold {
        background: rgba(255,255,255,0.06);
        border: 1px solid rgba(255,255,255,0.1);
        border-radius: 14px 0 14px 14px;
    }
    .online-dot {
        display: inline-block;
        width: 9px; height: 9px;
        background: #2ecc71;
        border-radius: 50%;
        margin-right: 6px;
        box-shadow: 0 0 6px #2ecc71;
        animation: goldPulseText 1.5s ease-in-out infinite;
    }
    </style>

    <div class="arsene-header-gold">
        <span class="arsene-ia-name">NOVA IA<span class="arsene-ia-cursor"></span></span>
        <span class="arsene-ia-slogan">Votre IA qui génère des <em>Word</em>, <em>Excel</em> &amp; <em>PDF</em> par vos mots</span>
        <div style="margin-top:14px;">
            <span class="arsene-badge">
                <span class="online-dot"></span>En ligne · Propulsé par Nova Platform
            </span>
        </div>
    </div>
    <hr class="arsene-divider">
    """, unsafe_allow_html=True)

    if st.button("← Retour au dashboard", key="retour_nova_support"):
        st.session_state["view"] = "home"
        st.rerun()

    st.markdown('<hr class="arsene-divider">', unsafe_allow_html=True)

    # Initialiser chat
    if "nova_support_chat" not in st.session_state:
        st.session_state["nova_support_chat"] = [{
            "role": "assistant",
            "content": f"Salut, moi c'est Nova IA, ton assistant Nova 24/7 ! ✨ Je connais tous les services, les abonnements et je suis là pour toi. Comment puis-je t'aider aujourd'hui{', ' + user if user else ''} ?"
        }]
    if "nova_support_resolu" not in st.session_state:
        st.session_state["nova_support_resolu"] = False

    # Afficher historique
    for msg in st.session_state["nova_support_chat"]:
        align = "flex-end" if msg["role"] == "user" else "flex-start"
        css_class = "msg-user-gold" if msg["role"] == "user" else "msg-assistant-gold"
        color = "rgba(255,255,255,0.85)" if msg["role"] == "user" else "#FFD700"
        icon = "🧑" if msg["role"] == "user" else "🤖"
        label = "Vous" if msg["role"] == "user" else "Assistant Nova 24/7"
        st.markdown(f"""
        <div style="display:flex;justify-content:{align};margin:8px 0;">
            <div class="{css_class}" style="padding:12px 16px;max-width:82%;">
                <span style="color:{color};font-size:.8rem;font-weight:800;">{icon} {label}</span>
                <p style="color:#eee;margin:5px 0 0 0;font-size:.92rem;line-height:1.6;">{msg["content"]}</p>
            </div>
        </div>""", unsafe_allow_html=True)

    if not st.session_state["nova_support_resolu"]:
        with st.form("nova_support_form", clear_on_submit=True):
            msg_user = st.text_input(
                "Message",
                placeholder="Pose ta question à Nova Platform...",
                label_visibility="collapsed"
            )
            col_s, col_t = st.columns([4, 1])
            with col_s:
                envoyer = st.form_submit_button("📨 Envoyer", use_container_width=True)
            with col_t:
                terminer = st.form_submit_button("✅ Fin", use_container_width=True)

        if envoyer and msg_user.strip():
            st.session_state["nova_support_chat"].append({"role": "user", "content": msg_user.strip()})
            db = st.session_state["db"]
            user_data = db["users"].get(user, {}) if user else {}
            premium_actif = is_premium_actif(user_data)
            historique_txt = "\n".join([
                f"{'Client' if m['role']=='user' else 'Nova IA'}: {m['content']}"
                for m in st.session_state["nova_support_chat"]
            ])
            prompt_nova_support = f"""Tu es NOVA IA, l'assistant officiel de Nova Platform.
Tu t'appelles Nova IA — jamais Gemini, jamais ChatGPT, jamais Claude.
Tu parles toujours en français, avec bienveillance et clarté.

TOUT CE QUE TU SAIS SUR NOVA PLATFORM :

SERVICES :
- 📊 Data & Excel Analytics : tableaux de bord, graphiques, analyse de données
- 📖 Fiche de Cours Professeur IA : fiches pédagogiques pour enseignants
- 📎 Modifier mon Fichier : modification Word, Excel, PowerPoint
- 📝 Exposé scolaire complet IA : exposés du CP au Master (PREMIUM)
- 📝 Création de Sujets & Examens : devoirs, QCM, contrôles (PREMIUM = auto)
- 📄 Création Word (depuis zéro) : le client décrit son document Word, Nova le génère complet
- 📋 Rapport de Stage IA : rapport académique complet (BTS, Licence, Master) · Missions · Analyse · Conclusion
- 📚 Affiches & Reçus : supports visuels entreprises
- 👔 CV & Lettre de Motivation : CV et lettres percutants
- 🔄 Conversion & Fichier PDF : conversion entre formats

PLANS PREMIUM :
- 🌅 Journalier (1 jour) : 600 FC → 2 générations IA automatiques
- 🔟 10 Jours : 1000 FC → 9 générations IA automatiques
- 👑 30 Jours : 2500 FC → Générations ILLIMITÉES
- Abonnement via WhatsApp : {WHATSAPP_NUMBER}

CONNEXION :
- Identifiant = nom choisi à l'inscription
- Mot de passe = numéro WhatsApp (ex: 2250XXXXXXXX)
- Problème connexion → vérifier le numéro avec 225 au début

FONCTIONNEMENT :
- Gratuit : demande soumise → Nova traite → livraison par lien (quelques heures)
- Premium : génération automatique en moins d'1 minute → onglet "Mes Livrables"
- Quota premium se renouvelle chaque jour à minuit

SUPPORT HUMAIN : WhatsApp {WHATSAPP_NUMBER}

RÈGLE IMPORTANTE : Si le problème est grave (paiement, compte bloqué, fichier perdu, bug critique) ou si tu n'arrives pas à le résoudre toi-même, propose TOUJOURS au client de contacter Nova directement sur WhatsApp : {WHATSAPP_NUMBER}. Formule-le ainsi : "Pour ce problème, je te recommande de contacter Nova directement sur WhatsApp : {WHATSAPP_NUMBER} 📲"

CLIENT : {user if user else "visiteur"} | Premium : {"OUI" if premium_actif else "NON"}

Historique :
{historique_txt}

Réponds UNIQUEMENT au dernier message. 3-5 phrases max. Tu es Nova IA.
RÈGLE IMPORTANTE : Ne commence JAMAIS tes réponses par "Bonjour [nom]" ou en te présentant à nouveau. Tu l'as déjà fait au début. Reste naturel, direct, comme dans une vraie conversation. Réponds directement à ce que le client dit.
RÈGLE ESCALADE OBLIGATOIRE :
Si le client exprime un problème grave (paiement, fichier perdu, compte bloqué, bug critique, plainte urgente), tu DOIS lui proposer cette phrase exacte à la fin de ta réponse :
"👉 Veux-tu que je transmette ton problème directement au service client Nova ? Réponds juste OUI et je m'en occupe immédiatement."
Si dans l'historique le client répond OUI ou "oui" ou "ok" ou "ouais" à cette proposition, réponds UNIQUEMENT ce texte exact sans rien d'autre :
__ESCALADE_CONFIRMEE__"""

            with st.spinner("🤖 Nova Platform réfléchit..."):
                reponse = generer_avec_gemini("Support Nova IA", prompt_nova_support, user or "visiteur")
            if reponse.startswith("❌"):
                reponse = f"Désolé, je rencontre une difficulté. Contacte Nova directement sur WhatsApp : {WHATSAPP_NUMBER}"
            if "__ESCALADE_CONFIRMEE__" in reponse:
                _db_ars = st.session_state["db"]
                _wa_ars = _db_ars["users"].get(user, {}).get("whatsapp", "—") if user else "—"
                ok = envoyer_escalade_support(user or "visiteur", _wa_ars, st.session_state["nova_support_chat"], "Assistant Nova 24/7")
                if ok:
                    reponse = "✅ C'est fait ! Ton problème a été transmis au service client Nova. Nous te recontactons très bientôt. 🙏"
                    st.session_state["nova_support_resolu"] = True
                else:
                    reponse = f"Désolé, l'envoi a échoué. Contacte Nova directement : {WHATSAPP_NUMBER} 📲"
            st.session_state["nova_support_chat"].append({"role": "assistant", "content": reponse})
            st.rerun()

        if terminer and len(st.session_state["nova_support_chat"]) > 1:
            try:
                import resend
                resend.api_key = st.secrets["RESEND_API_KEY"]
                db = st.session_state["db"]
                historique_email = "\n".join([
                    f"{'🧑 Client' if m['role']=='user' else '🤖 Nova Platform'} : {m['content']}"
                    for m in st.session_state["nova_support_chat"]
                ])
                resend.Emails.send({
                    "from": "Nova Platform <onboarding@resend.dev>",
                    "to": [st.secrets["EMAIL_RECEIVER"]],
                    "subject": f"🤖 Nova Platform — Conversation {user or 'visiteur'}",
                    "text": f"""RÉSUMÉ CONVERSATION NOVA IA
━━━━━━━━━━━━━━━━━━━━━━━━━━━
👤 Client   : {user or "visiteur"}
📱 WhatsApp : {db["users"].get(user, {}).get("whatsapp", "—") if user else "—"}
⏰ Date     : {datetime.now().strftime("%d/%m/%Y à %H:%M")}
━━━━━━━━━━━━━━━━━━━━━━━━━━━

{historique_email}

━━━━━━━━━━━━━━━━━━━━━━━━━━━
Intervenir si problème non résolu.
"""
                })
                st.session_state["nova_support_resolu"] = True
                st.success("✅ Conversation envoyée à Nova. Nous revenons vers toi rapidement !")
                st.rerun()
            except Exception as e:
                st.error(f"Erreur envoi : {e}")
    else:
        st.success("✅ Ta demande a été transmise à Nova. Nous te répondrons bientôt.")
        if st.button("🔄 Nouvelle conversation", key="reset_nova_support"):
            st.session_state["nova_support_chat"] = [{
                "role": "assistant",
                "content": f"Salut, moi c'est Nova IA ! 👋 Comment puis-je t'aider ?"
            }]
            st.session_state["nova_support_resolu"] = False
            st.rerun()


# ══════════════════════════════════════════════════════════════════
# PAGE NOVA IA — Chat intelligent dédié, génération ou demande
# ══════════════════════════════════════════════════════════════════
def show_nova_ia_page():
    user = st.session_state.get("current_user", "")
    db   = st.session_state.get("db", {"users": {}, "demandes": [], "liens": {}})
    user_data     = db["users"].get(user, {}) if user else {}
    premium_actif = is_premium_actif(user_data)
    wa_user       = user_data.get("whatsapp", "—")

    # ── STYLES ────────────────────────────────────────────────────
    st.markdown("""
    <style>
    @keyframes novaGlow {
        0%,100%{box-shadow:0 0 14px 4px rgba(0,200,255,0.25),0 0 35px 8px rgba(0,150,255,0.12);}
        50%     {box-shadow:0 0 32px 12px rgba(0,220,255,0.55),0 0 65px 22px rgba(0,180,255,0.3);}
    }
    @keyframes novaShimmer {
        0%  {background-position:-200% center;}
        100%{background-position: 200% center;}
    }
    @keyframes floatBot {
        0%,100%{transform:translateY(0);}
        50%    {transform:translateY(-8px);}
    }
    @keyframes pulseDot {
        0%,100%{opacity:1;transform:scale(1);}
        50%    {opacity:.4;transform:scale(.6);}
    }
    .nova-ia-header{
        background:linear-gradient(135deg,rgba(0,180,255,.13) 0%,rgba(0,100,200,.08) 50%,rgba(0,180,255,.07) 100%);
        border:2px solid rgba(0,200,255,.4);
        border-radius:22px;padding:22px 14px 18px;text-align:center;
        animation:novaGlow 2.8s ease-in-out infinite;margin-bottom:14px;
        position:relative;overflow:hidden;
    }
    .nova-ia-header::before{
        content:"";position:absolute;top:0;left:-100%;right:0;bottom:0;
        background:linear-gradient(90deg,transparent,rgba(0,200,255,.08),rgba(255,255,255,.1),rgba(0,200,255,.08),transparent);
        background-size:200% 100%;animation:novaShimmer 3.5s linear infinite;
    }
    @keyframes nova-name-appear {
        0%   { opacity:0; transform:translateY(12px) scale(.96); }
        100% { opacity:1; transform:translateY(0) scale(1); }
    }
    @keyframes nova-slogan-slide {
        0%   { opacity:0; transform:translateX(-18px); }
        100% { opacity:1; transform:translateX(0); }
    }
    @keyframes nova-cursor-blink {
        0%,100% { opacity:1; }
        50%      { opacity:0; }
    }
    .nova-ia-name {
        background: linear-gradient(90deg, #00c6ff 0%, #7df9ff 30%, #ffffff 50%, #7df9ff 70%, #00c6ff 100%);
        background-size: 250% auto;
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-size: clamp(1.8rem, 8vw, 3.6rem);
        font-weight: 900;
        letter-spacing: 4px;
        text-transform: uppercase;
        display: block;
        line-height: 1.1;
        animation: novaShimmer 3s linear infinite, nova-name-appear 0.8s ease-out both;
        filter: drop-shadow(0 0 22px rgba(0,220,255,0.55));
    }
    .nova-ia-cursor {
        display: inline-block;
        width: 3px; height: 3.4rem;
        background: #00d2ff;
        vertical-align: middle;
        margin-left: 4px;
        border-radius: 2px;
        animation: nova-cursor-blink 1s step-end infinite;
        box-shadow: 0 0 10px rgba(0,210,255,0.8);
    }
    .nova-ia-slogan {
        display: block;
        margin-top: 10px;
        font-size: clamp(0.78rem, 3.5vw, 1.05rem);
        font-weight: 600;
        color: rgba(200,240,255,0.82);
        letter-spacing: 0.3px;
        animation: nova-slogan-slide 1s ease-out 0.4s both;
    }
    .nova-ia-slogan em {
        color: #7df9ff;
        font-style: normal;
        font-weight: 700;
    }
    .nova-badge{
        display:inline-block;
        background:linear-gradient(135deg,rgba(0,200,255,.18),rgba(0,120,200,.12));
        border:1px solid rgba(0,200,255,.5);border-radius:30px;
        padding:5px 16px;margin-top:14px;font-size:.82rem;
        color:#00d2ff;font-weight:700;letter-spacing:.5px;
    }
    .online-dot-nova{
        display:inline-block;width:9px;height:9px;
        background:#2ecc71;border-radius:50%;margin-right:6px;
        box-shadow:0 0 6px #2ecc71;animation:pulseDot 1.5s ease-in-out infinite;
    }
    .nova-divider{height:2px;
        background:linear-gradient(90deg,transparent,#00d2ff,#7df9ff,#00d2ff,transparent);
        border:none;margin:14px 0;border-radius:2px;}
    .msg-nova{
        background:linear-gradient(135deg,rgba(0,200,255,.1),rgba(0,120,200,.06));
        border:1px solid rgba(0,200,255,.22);border-left:3px solid #00d2ff;
        border-radius:0 14px 14px 14px;
        word-break:break-word;
    }
    .msg-user-nova{
        background:rgba(255,255,255,.06);border:1px solid rgba(255,255,255,.1);
        border-radius:14px 0 14px 14px;
        word-break:break-word;
    }
    .nova-service-chip{
        display:inline-block;background:rgba(0,200,255,.1);
        border:1px solid rgba(0,200,255,.3);border-radius:20px;
        padding:4px 12px;margin:3px;font-size:.8rem;color:#7df9ff;cursor:pointer;
    }
    .nova-livraison-badge{
        background:linear-gradient(135deg,rgba(46,204,113,.18),rgba(0,200,100,.1));
        border:1px solid rgba(46,204,113,.4);border-radius:12px;
        padding:10px 16px;margin-top:10px;font-size:.85rem;color:#2ecc71;
    }
    .nova-attente-badge{
        background:linear-gradient(135deg,rgba(255,193,7,.15),rgba(255,150,0,.08));
        border:1px solid rgba(255,193,7,.4);border-radius:12px;
        padding:10px 16px;margin-top:10px;font-size:.85rem;color:#FFD700;
    }
    </style>
    """, unsafe_allow_html=True)

    # ── EN-TÊTE ───────────────────────────────────────────────────
    statut_label = "⭐ Membre Premium" if premium_actif else "👤 Compte Gratuit"
    st.markdown(f"""
    <div class="nova-ia-header">
        <span class="nova-ia-name">NOVA IA<span class="nova-ia-cursor"></span></span>
        <span class="nova-ia-slogan">Votre IA qui génère des <em>Word</em>, <em>Excel</em> &amp; <em>PDF</em> par vos mots</span>
        <div style="margin-top:14px;">
            <span class="nova-badge">
                <span class="online-dot-nova"></span>En ligne · {statut_label}
            </span>
        </div>
    </div>
    <hr class="nova-divider">
    """, unsafe_allow_html=True)

    if st.button("← Retour au dashboard", key="retour_nova_ia"):
        st.session_state["view"] = "home"
        st.rerun()

    st.markdown('<hr class="nova-divider">', unsafe_allow_html=True)

    # ── INIT SESSION ──────────────────────────────────────────────
    if "nova_ia_chat" not in st.session_state:
        intro = (
            f"Salut {user} ! 👋 Moi c'est **Nova IA**, ton assistant intelligent Nova Platform. "
            if user else
            "Salut ! 👋 Moi c'est **Nova IA**, ton assistant intelligent Nova Platform. "
        )
        if premium_actif:
            intro += "✨ Tu es **Premium** — je peux générer tes documents instantanément ! Dis-moi ce dont tu as besoin."
        else:
            intro += "Je peux soumettre une demande pour toi, elle sera traitée par l'équipe Nova. Dis-moi ce que tu veux créer !"
        st.session_state["nova_ia_chat"] = [{"role": "assistant", "content": intro}]

    if "nova_ia_phase" not in st.session_state:
        st.session_state["nova_ia_phase"] = "dialogue"   # dialogue | confirmation | traitement | termine

    if "nova_ia_service_detecte" not in st.session_state:
        st.session_state["nova_ia_service_detecte"] = None

    if "nova_ia_prompt_final" not in st.session_state:
        st.session_state["nova_ia_prompt_final"] = ""

    if "nova_ia_livrable" not in st.session_state:
        st.session_state["nova_ia_livrable"] = None

    if "nova_ia_fichier" not in st.session_state:
        st.session_state["nova_ia_fichier"] = None  # {"nom": str, "type": str, "contenu": str|bytes, "est_image": bool}

    if "nova_ia_pending_image" not in st.session_state:
        st.session_state["nova_ia_pending_image"] = None  # {"b64": str, "mime": str} — image sauvegardée entre dialogue et traitement

    # ── AFFICHAGE HISTORIQUE ──────────────────────────────────────
    for msg in st.session_state["nova_ia_chat"]:
        align = "flex-end" if msg["role"] == "user" else "flex-start"
        css   = "msg-user-nova" if msg["role"] == "user" else "msg-nova"
        color = "rgba(255,255,255,.85)" if msg["role"] == "user" else "#00d2ff"
        icon  = "🧑" if msg["role"] == "user" else "🤖"
        label = "Vous" if msg["role"] == "user" else "Nova IA"
        # Remplace **texte** par <strong>texte</strong> pour le gras en HTML
        contenu = msg["content"].replace("**", "<strong>", 1)
        while "**" in contenu:
            contenu = contenu.replace("**", "</strong>", 1).replace("**", "<strong>", 1) if "**" in contenu else contenu.replace("**", "</strong>", 1)
        st.markdown(f"""
        <div style="display:flex;justify-content:{align};margin:6px 0;">
            <div class="{css}" style="padding:10px 14px;max-width:92%;min-width:0;">
                <span style="color:{color};font-size:.75rem;font-weight:800;">{icon} {label}</span>
                <p style="color:#eee;margin:4px 0 0 0;font-size:clamp(.85rem,3.8vw,.95rem);line-height:1.65;word-break:break-word;">{contenu}</p>
            </div>
        </div>""", unsafe_allow_html=True)

    # ── PHASE : DIALOGUE / COLLECTE ───────────────────────────────
    if st.session_state["nova_ia_phase"] in ("dialogue",):

        # ── FILE UPLOADER hors form (Streamlit ne permet pas uploader dans form) ──
        st.markdown("""
        <style>
        /* Zone principale du file uploader Nova IA */
        [data-testid="stFileUploader"] section {
            background: rgba(0, 180, 255, 0.06) !important;
            border: 1.5px dashed rgba(0, 180, 255, 0.45) !important;
            border-radius: 10px !important;
            padding: 8px 14px !important;
        }
        [data-testid="stFileUploader"] section:hover {
            background: rgba(0, 180, 255, 0.12) !important;
            border-color: rgba(0, 180, 255, 0.75) !important;
        }
        /* Texte "Drag and drop" et formats */
        [data-testid="stFileUploader"] section span,
        [data-testid="stFileUploader"] section small,
        [data-testid="stFileUploader"] section p {
            color: rgba(0, 200, 255, 0.85) !important;
            font-size: 0.82rem !important;
        }
        /* Bouton Browse files */
        [data-testid="stFileUploader"] section button {
            background: rgba(0, 180, 255, 0.15) !important;
            color: #00d2ff !important;
            border: 1px solid rgba(0, 180, 255, 0.4) !important;
            border-radius: 6px !important;
            font-size: 0.8rem !important;
            font-weight: 600 !important;
        }
        [data-testid="stFileUploader"] section button:hover {
            background: rgba(0, 180, 255, 0.28) !important;
        }
        /* Fichier uploadé — nom du fichier */
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFile"] {
            background: rgba(0, 180, 255, 0.1) !important;
            border: 1px solid rgba(0, 180, 255, 0.3) !important;
            border-radius: 8px !important;
            color: #00d2ff !important;
        }
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileName"] {
            color: #00d2ff !important;
            font-weight: 600 !important;
        }
        [data-testid="stFileUploader"] [data-testid="stFileUploaderFileSize"] {
            color: rgba(0, 200, 255, 0.6) !important;
        }
        </style>
        """, unsafe_allow_html=True)
        st.markdown("<div style='color:rgba(0,200,255,0.7);font-size:0.82rem;margin-bottom:4px;'>📎 Joindre un fichier à ta demande (photo, Word, PDF, txt) — optionnel</div>", unsafe_allow_html=True)
        fichier_joint = st.file_uploader(
            "Joindre un fichier à ta demande",
            type=["png", "jpg", "jpeg", "webp", "pdf", "docx", "txt"],
            label_visibility="collapsed",
            key="nova_ia_uploader"
        )
        if fichier_joint is not None:
            import base64 as _b64
            _nom = fichier_joint.name
            _ext = _nom.rsplit(".", 1)[-1].lower()
            _est_image = _ext in ("png", "jpg", "jpeg", "webp")

            if _est_image:
                # Image → base64 pour Gemini vision
                _bytes = fichier_joint.getvalue()
                _mime_img = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "webp": "image/webp"}.get(_ext, "image/png")
                st.session_state["nova_ia_fichier"] = {
                    "nom": _nom, "type": _mime_img,
                    "contenu": _b64.b64encode(_bytes).decode("utf-8"),
                    "est_image": True
                }
                st.success(f"📸 Image jointe : **{_nom}** — Nova va l'analyser avec ta demande.")

            elif _ext == "txt":
                _texte = fichier_joint.getvalue().decode("utf-8", errors="ignore")
                st.session_state["nova_ia_fichier"] = {
                    "nom": _nom, "type": "text/plain",
                    "contenu": _texte, "est_image": False
                }
                st.success(f"📄 Fichier texte joint : **{_nom}** ({len(_texte)} caractères)")

            elif _ext == "docx":
                try:
                    from docx import Document as _Doc
                    import io as _io
                    _raw_bytes = fichier_joint.getvalue()  # getvalue() plus fiable que read() avec Streamlit
                    _doc = _Doc(_io.BytesIO(_raw_bytes))
                    # Extraire paragraphes + cellules de tableaux
                    _lignes = []
                    for p in _doc.paragraphs:
                        if p.text.strip():
                            _lignes.append(p.text.strip())
                    for tbl in _doc.tables:
                        for row in tbl.rows:
                            _row_txt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                            if _row_txt:
                                _lignes.append(_row_txt)
                    _texte = "\n".join(_lignes)
                    st.session_state["nova_ia_fichier"] = {
                        "nom": _nom, "type": "application/docx",
                        "contenu": _texte, "est_image": False
                    }
                    st.success(f"📝 Document Word joint : **{_nom}** ({len(_texte)} caractères extraits)")
                except Exception as _e:
                    st.warning(f"⚠️ Impossible de lire ce fichier Word : {_e}")

            elif _ext == "pdf":
                try:
                    import io as _io
                    _pdf_bytes = fichier_joint.getvalue()
                    # Tentative extraction texte PDF via pdfminer si dispo
                    try:
                        from pdfminer.high_level import extract_text as _pdf_extract
                        _texte = _pdf_extract(_io.BytesIO(_pdf_bytes))
                    except ImportError:
                        # Fallback : envoyer comme image la première page n'est pas possible sans lib
                        # On stocke le PDF en base64 et on prévient
                        _texte = f"[PDF joint : {_nom} — contenu non extractible sans pdfminer]"
                    st.session_state["nova_ia_fichier"] = {
                        "nom": _nom, "type": "application/pdf",
                        "contenu": _texte, "est_image": False
                    }
                    st.success(f"📋 PDF joint : **{_nom}** ({len(_texte)} caractères extraits)")
                except Exception as _e:
                    st.warning(f"⚠️ Impossible de lire ce PDF : {_e}")

        # Afficher le fichier actuellement en mémoire si déjà uploadé avant
        elif st.session_state.get("nova_ia_fichier"):
            _f = st.session_state["nova_ia_fichier"]
            st.info(f"📎 Fichier en mémoire : **{_f['nom']}** — il sera utilisé avec ta prochaine demande. | [Supprimer](javascript:void(0))")

        with st.form("nova_ia_form", clear_on_submit=True):
            msg_input = st.text_input(
                "Message",
                placeholder="Ex: Je veux un exposé sur la photosynthèse pour le lycée...",
                label_visibility="collapsed"
            )
            col_env, col_reinit = st.columns([5, 1])
            with col_env:
                envoyer = st.form_submit_button("📨 Envoyer", use_container_width=True)
            with col_reinit:
                reinit = st.form_submit_button("🔄", use_container_width=True)

        if reinit:
            st.session_state.pop("nova_ia_chat", None)
            st.session_state.pop("nova_ia_phase", None)
            st.session_state.pop("nova_ia_service_detecte", None)
            st.session_state.pop("nova_ia_prompt_final", None)
            st.rerun()

        if envoyer and msg_input.strip():
            _fichier = st.session_state.get("nova_ia_fichier")
            _msg_affiche = msg_input.strip()
            if _fichier:
                if _fichier["est_image"]:
                    _msg_affiche += f"\n📎 *[Image jointe : {_fichier['nom']}]*"
                else:
                    _msg_affiche += f"\n📎 *[Fichier joint : {_fichier['nom']}]*"
            st.session_state["nova_ia_chat"].append({"role": "user", "content": _msg_affiche})

            historique_txt = "\n".join([
                f"{'Client' if m['role']=='user' else 'Nova IA'}: {m['content']}"
                for m in st.session_state["nova_ia_chat"]
            ])

            # ── Injection contenu fichier dans le prompt si présent ──
            _fichier = st.session_state.get("nova_ia_fichier")
            _bloc_fichier = ""
            _image_b64  = None
            _image_mime = None
            if _fichier:
                if _fichier["est_image"]:
                    _image_b64  = _fichier["contenu"]
                    _image_mime = _fichier["type"]
                    _bloc_fichier = f"\n\n════════════════════════════════════════\nFICHIER IMAGE JOINT PAR LE CLIENT : {_fichier['nom']}\n════════════════════════════════════════\nINSTRUCTIONS STRICTES :\n- Si le client demande d'EXTRAIRE du texte ou des données → extrais exactement ce qui est visible, RIEN D'AUTRE.\n- Si le client demande de reproduire un tableau → extrais toutes les colonnes et données visibles en markdown, sans commentaire.\n- Si le client veut UTILISER cette image pour créer un document → collecte les infos manquantes puis confirme normalement.\n- INTERDIT : ajouter une introduction, un développement, une conclusion, une mise en contexte non demandés.\n"
                else:
                    _contenu_txt = _fichier["contenu"][:8000]  # sécurité tokens
                    _bloc_fichier = f"\n\n════════════════════════════════════════\nFICHIER JOINT PAR LE CLIENT : {_fichier['nom']}\n════════════════════════════════════════\n{_contenu_txt}\n════════════════════════════════════════\nINSTRUCTIONS STRICTES SUR CE FICHIER :\n- Si le client demande d'EXTRAIRE du texte → copie exactement les lignes demandées, RIEN D'AUTRE. Zéro intro, zéro conclusion, zéro commentaire.\n- Si le client demande de RÉSUMER → résume en 2-3 phrases max, pas de structure académique.\n- Si le client veut UTILISER ce contenu pour créer un document → collecte les infos manquantes puis confirme normalement.\n- INTERDIT : ajouter une introduction, un développement, une conclusion, une mise en contexte non demandés.\n"

            prompt_nova = f"""Tu es NOVA IA, l'assistante intelligente de Nova Platform (Côte d'Ivoire).
Tu t'appelles Nova IA — jamais Gemini, jamais ChatGPT, jamais Claude, jamais le support IA.
Tu parles toujours en français, avec bienveillance et précision.
TON RÔLE : comprendre ce que le client veut créer, collecter les informations nécessaires, puis confirmer.
{_bloc_fichier}
════════════════════════════════════════
CATALOGUE COMPLET DES SERVICES NOVA PLATFORM
════════════════════════════════════════

1. 📝 EXPOSÉ SCOLAIRE COMPLET IA
   → Infos ESSENTIELLES (dans cet ordre si manquantes) :
     1. Le sujet exact de l'exposé → question : "C'est sur quel sujet ?"
     2. Le niveau scolaire → question : "Tu es en quelle classe ?"
     3. La matière → question : "C'est pour quelle matière ?"
   → Dès que sujet + niveau sont connus : RÉCAPITULATIF immédiat, matière optionnelle

2. 📖 FICHE DE COURS PROFESSEUR IA
   → Infos ESSENTIELLES (dans cet ordre si manquantes) :
     1. La matière → question : "C'est pour quelle matière ?"
     2. Le chapitre ou sujet précis → question : "C'est sur quel chapitre ou notion ?"
     3. Le niveau des élèves → question : "Tes élèves sont en quelle classe ?"
   → Dès que matière + sujet sont connus : RÉCAPITULATIF immédiat

3. 📝 CRÉATION DE SUJETS & EXAMENS
   → Infos ESSENTIELLES (dans cet ordre si manquantes) :
     1. La matière → question : "C'est pour quelle matière ?"
     2. Le niveau → question : "C'est pour quelle classe ?"
     3. Le type d'évaluation → question : "Tu veux quel type de sujet ? (QCM, devoir, dissertation, cas pratique...)"
   → Dès que matière + niveau sont connus : RÉCAPITULATIF immédiat

4. 📄 CRÉATION DOCUMENT WORD (depuis zéro)
   → Ce service couvre TOUT ce que Gemini peut produire en Word : article de journal, tableau, contrat, lettre, rapport, discours, affiche, note de service, fiche technique, formulaire, planning, facture, règlement, etc.
   → Infos ESSENTIELLES (dans cet ordre si manquantes) :
     1. Le type exact de document → question : "C'est quel type de document ?"
     2. Le sujet ou contexte → question : "C'est sur quel sujet ou dans quel contexte ?"
   → RÈGLE COLLECTE INTELLIGENTE : selon le type détecté, pose UNE question supplémentaire si critique :
     - TABLEAU → si nombre de colonnes/lignes non précisé : "Combien de colonnes et de lignes tu veux ?"
     - ARTICLE DE JOURNAL → si rubrique non précisée : "C'est pour quelle rubrique ? (actualité, sport, culture, économie...)"
     - CONTRAT / LETTRE OFFICIELLE → si parties non précisées : "C'est entre qui et qui ?"
     - PLANNING / PROGRAMME → si période non précisée : "C'est sur quelle période ?"
     - FACTURE / DEVIS → si produits non précisés : "Cite les produits ou services avec les montants."
     - Autres types simples (discours, affiche, note de service...) → RÉCAPITULATIF immédiat dès type + sujet connus
   → RÈGLE RESPECT CLIENT : si le client donne toutes les infos dans son premier message → RÉCAPITULATIF IMMÉDIAT, zéro question
   → Dès que les infos critiques sont connues : RÉCAPITULATIF immédiat

5. 👔 CV & LETTRE DE MOTIVATION
   → Infos ESSENTIELLES (dans cet ordre si manquantes) :
     1. Prénom et nom complet → question : "C'est pour quel nom complet ?"
     2. Le poste visé → question : "Tu postules pour quel poste ?"
     3. Le secteur d'activité → question : "C'est dans quel secteur ?"
     4. La formation / diplôme(s) → question : "Tu as quel diplôme et dans quel établissement ?"
     5. Les expériences professionnelles ou stages → question : "Tu as des expériences ou stages ? (poste, entreprise, durée, tâches principales)"
     6. Les compétences principales → question : "Cite tes principales compétences (logiciels, outils, savoir-faire)."
     7. Les langues parlées → question : "Tu parles quelles langues et à quel niveau ?"
     8. CV seul, lettre seule ou les deux ? → question : "Tu veux le CV seul, la lettre seule ou les deux ?"
   → RÈGLE SPÉCIALE CV — COLLECTE SÉQUENTIELLE (un champ à la fois, dans cet ordre) :
     Si le client n'a pas encore donné sa formation → pose UNIQUEMENT : "Ta formation ? (ex : BAC, BTS, Licence — précise l'établissement et l'année)"
     Si formation connue mais pas les expériences → pose UNIQUEMENT : "Tu as des expériences ou stages ? (poste, entreprise, durée et 2-3 tâches) — réponds 'aucun' si non"
     Si expériences connues mais pas les compétences → pose UNIQUEMENT : "Tes compétences principales ? (logiciels, outils, savoir-faire)"
     Si compétences connues mais pas les langues → pose UNIQUEMENT : "Les langues que tu parles et ton niveau ? (ex : Français courant, Anglais débutant)"
     Si tout est connu → RÉCAPITULATIF immédiat
   → RÈGLE D'OR : une seule question à la fois, jamais deux champs dans le même message
   → RÈGLE RESPECT DU CLIENT : si le client répond "aucun", "pas encore", "je sais pas", "rien", "N/A" ou équivalent pour un champ → accepte sa réponse, note l'absence, passe au champ suivant. Ne repose JAMAIS la même question.
   → Dès que tous les champs ont été abordés (même si certains sont vides par choix du client) : RÉCAPITULATIF immédiat

6. 📋 RAPPORT DE STAGE IA
   → Infos ESSENTIELLES (dans cet ordre si manquantes) :
     1. Prénom et nom complet du stagiaire → question : "C'est au nom de qui ? (prénom et nom)"
     2. Nom complet de l'entreprise d'accueil → question : "C'est dans quelle entreprise ?"
     3. Ville / localisation de l'entreprise → question : "L'entreprise est dans quelle ville ?"
     4. Secteur d'activité de l'entreprise → question : "C'est quoi le secteur de l'entreprise ? (banque, BTP, informatique, santé...)"
     5. La filière ou spécialité → question : "Tu es en quelle filière ?"
     6. Le niveau d'études → question : "Tu es en BTS, Licence ou Master ?"
     7. La période exacte du stage → question : "Le stage s'est déroulé du ... au ... ?"
     8. Les missions effectuées → question : "Tu as fait quoi comme tâches pendant le stage ? (même en quelques mots)"
   → Dès que nom + entreprise + ville + secteur + filière + niveau + période + missions sont connus : RÉCAPITULATIF immédiat

7. 📊 DATA & EXCEL ANALYTICS
   → Infos ESSENTIELLES (dans cet ordre si manquantes) :
     1. Le type d'analyse → question : "Tu veux quoi exactement ? (tableau, graphique, rapport, dashboard...)"
     2. Description des données → question : "Décris tes données ou colle-les directement ici."
   → Dès que type + données sont connus : RÉCAPITULATIF immédiat

8. 📎 MODIFIER MON FICHIER (Word, Excel, PowerPoint)
   → Infos ESSENTIELLES (dans cet ordre si manquantes) :
     1. Le type de fichier → question : "C'est un fichier Word, Excel ou PowerPoint ?"
     2. Les modifications souhaitées → question : "Tu veux qu'on change quoi exactement ?"
   → Dès que type + modifications sont connus : RÉCAPITULATIF immédiat

9. 🔄 CONVERSION & FICHIER PDF
   → Infos ESSENTIELLES (dans cet ordre si manquantes) :
     1. Le format source et cible → question : "Tu veux convertir de quel format vers quel format ?"
   → Dès que les formats sont connus : RÉCAPITULATIF immédiat

10. 🔍 NUMÉRISATION OCR
    → Infos ESSENTIELLES (dans cet ordre si manquantes) :
      1. Le type de document → question : "C'est une image ou un PDF scanné ?"
      2. La langue du texte → question : "Le texte est en quelle langue ?"
    → Dès que type + langue sont connus : RÉCAPITULATIF immédiat

════════════════════════════════════════
PLANS & FONCTIONNEMENT
════════════════════════════════════════
- PREMIUM ({premium_actif}) : génération AUTOMATIQUE IA en moins d'1 min → livraison directe dans "Mes Livrables"
  · Journalier : 600 FC → 2 générations
  · 10 Jours : 1000 FC → 9 générations
  · 30 Jours : 2500 FC → Illimité
- GRATUIT : demande soumise → traitée par l'équipe Nova sous quelques heures → livraison par lien WhatsApp

════════════════════════════════════════
INSTRUCTIONS COMPORTEMENT
════════════════════════════════════════
- Si le client exprime son besoin → identifie le service correspondant parmi les 10 ci-dessus
- RÈGLE D'OR : si le client a donné les infos essentielles (même en un seul message) → RÉCAPITULATIF IMMÉDIAT, zéro question
- Si UNE SEULE info essentielle manque → pose UNE question, maximum 8 mots, pas de choix, pas d'options
- INTERDIT ABSOLU : poser 2 questions ou plus dans un même message
- INTERDIT ABSOLU : proposer des "options", "axes de réflexion", "précisions", "personnalisations" — le client n'a pas demandé ça
- INTERDIT ABSOLU : demander si le client veut ajouter quelque chose — Nova sait faire, pas besoin de demander
- Les infos FACULTATIVES n'existent pas pour toi — ignore-les complètement, lance le travail tel quel
- Quand tu as toutes les infos ESSENTIELLES → propose un récapitulatif et demande confirmation avec la phrase EXACTE :
  "✅ J'ai tout ce qu'il me faut ! Voici le récapitulatif : [récap] — Tu confirmes ? Réponds OUI pour lancer."
- Si le client répond OUI / oui / ok / ouais / confirme → réponds UNIQUEMENT : __NOVA_CONFIRME__|SERVICE:[nom_service]|DESC:[description_complète]
- Ne génère JAMAIS le document toi-même dans le chat
- Ne commence JAMAIS par "Bonjour [nom]" après le premier message
- Problème grave / technique → propose WhatsApp : {WHATSAPP_NUMBER}

CLIENT : {user if user else "visiteur"} | Premium : {"OUI" if premium_actif else "NON"}

Historique :
{historique_txt}

Réponds UNIQUEMENT au dernier message du client. 2-4 phrases max sauf pour le récapitulatif."""

            with st.spinner("🤖 Nova IA réfléchit..."):
                reponse = generer_avec_gemini(
                    "Nova IA Chat", prompt_nova, user or "visiteur",
                    _image_b64=_image_b64, _image_mime=_image_mime
                )

            if reponse.startswith("❌"):
                reponse = f"Désolé, une erreur s'est produite. Contacte Nova directement sur WhatsApp : {WHATSAPP_NUMBER} 📲"

            # Détection de la confirmation
            if "__NOVA_CONFIRME__" in reponse:
                # ── Sauvegarder l'image en session_state AVANT de vider le fichier ──
                # Sans ça, Gemini reçoit la commande "reproduis le tableau" sans avoir l'image → hallucinations
                if _image_b64 and _image_mime:
                    st.session_state["nova_ia_pending_image"] = {
                        "b64": _image_b64,
                        "mime": _image_mime
                    }
                else:
                    st.session_state["nova_ia_pending_image"] = None

                # Vider le fichier seulement maintenant, après avoir sauvegardé l'image
                st.session_state["nova_ia_fichier"] = None

                # Parser service et description
                try:
                    partie = reponse.split("__NOVA_CONFIRME__|")[1]
                    service_part = partie.split("|DESC:")[0].replace("SERVICE:", "").strip()
                    desc_part    = partie.split("|DESC:")[1].strip() if "|DESC:" in partie else msg_input.strip()
                except Exception:
                    service_part = "Demande Nova IA"
                    desc_part    = msg_input.strip()

                st.session_state["nova_ia_service_detecte"] = service_part
                st.session_state["nova_ia_prompt_final"]    = desc_part
                st.session_state["nova_ia_phase"]           = "traitement"
                reponse_affichee = f"✅ Parfait ! Je lance maintenant {'la génération automatique ⚡' if premium_actif else 'ta demande 📋'}..."
                st.session_state["nova_ia_chat"].append({"role": "assistant", "content": reponse_affichee})
                st.rerun()
            else:
                st.session_state["nova_ia_chat"].append({"role": "assistant", "content": reponse})
                st.rerun()

    # ── PHASE : TRAITEMENT ────────────────────────────────────────
    elif st.session_state["nova_ia_phase"] == "traitement":
        service_final = st.session_state.get("nova_ia_service_detecte", "Demande Nova IA")
        desc_finale   = st.session_state.get("nova_ia_prompt_final", "")

        # ── Normalisation du service pour matcher la console admin ──
        _NOVA_SERVICE_MAP = [
            (["exposé", "expose", "exposi"],        "📝 Exposé scolaire complet IA"),
            (["fiche", "cours", "professeur"],       "📖 Fiche de Cours Professeur IA"),
            (["sujet", "examen", "qcm", "évaluat"],  "📝 Création de Sujets & Examens"),
            (["rapport", "stage"],                   "📋 Rapport de Stage IA"),
            (["cv", "lettre", "motivation", "curriculum"], "👔 CV & Lettre de Motivation"),
            (["word", "document", "rédact", "contrat", "rapport profess"], "📄 Création Word (depuis zéro)"),
            (["data", "excel", "tableau", "analyt", "graphique"], "📊 Data & Excel Analytics"),
            (["modif", "correc", "retouche"],        "📎 Modifier mon Fichier (Word / Excel / PPT)"),
            (["conver", "pdf", "format"],            "📄 Conversion & Fichier PDF"),
            (["ocr", "numér", "scan"],               "🔍 OCR — Numérisation de Document"),
        ]
        _svc_low = service_final.lower()
        for _keywords, _canonical in _NOVA_SERVICE_MAP:
            if any(k in _svc_low for k in _keywords):
                service_final = _canonical
                st.session_state["nova_ia_service_detecte"] = _canonical
                break

        if premium_actif:
            # ── CAS PREMIUM : génération Gemini immédiate ────────
            user_data_frais = db["users"].get(user, {})
            restant = quota_restant(user_data_frais)
            plan_actuel = user_data_frais.get("premium_plan", "")
            quota_max = PLANS_PREMIUM.get(plan_actuel, {}).get("generations", 0)
            used_auj, _ = get_gen_quota(user_data_frais)

            if restant <= 0 and quota_max < 999:
                msg_quota = (
                    f"🚫 Tu as atteint ta limite de générations pour aujourd'hui "
                    f"({used_auj}/{quota_max}). Ton quota se renouvelle demain à minuit, "
                    f"ou contacte Nova pour étendre ton abonnement : {WHATSAPP_NUMBER} 📲"
                )
                st.session_state["nova_ia_chat"].append({"role": "assistant", "content": msg_quota})
                st.session_state["nova_ia_phase"] = "dialogue"
                st.rerun()
            else:
                # ── Récupérer l'image sauvegardée depuis la phase dialogue ──────
                _pending_img = st.session_state.get("nova_ia_pending_image")
                _trait_img_b64  = _pending_img["b64"]  if _pending_img else None
                _trait_img_mime = _pending_img["mime"] if _pending_img else None

                # ── Enrichir le prompt si image présente ────────────────────────
                desc_finale_trait = desc_finale
                if _trait_img_b64:
                    desc_finale_trait = (
                        f"{desc_finale}\n\n"
                        f"⚠️ INSTRUCTION ABSOLUE — REPRODUCTION FIDÈLE DE L'IMAGE :\n"
                        f"Une image est jointe. Tu DOIS utiliser EXCLUSIVEMENT les données visibles sur cette image.\n"
                        f"INTERDIT : inventer, halluciner, compléter ou modifier quoi que ce soit.\n"
                        f"Chaque chiffre, chaque nom, chaque cellule doit être recopié exactement tel quel depuis l'image.\n"
                        f"Si une cellule est vide dans l'image → elle reste vide. Aucune donnée fictive tolérée."
                    )

                # ── Anti-rédaction : si service Word + données collées ───────────
                _mots_repro = ["reproduis", "recopie", "mets en word", "extrais", "liste des",
                               "voici les", "voici la liste", "convertis", "mets en forme",
                               "liste :", "tableau depuis", "les données", "les noms", "les membres"]
                _desc_lower = desc_finale_trait.lower()
                if "word" in (service_final or "").lower() and any(m in _desc_lower for m in _mots_repro):
                    desc_finale_trait = (
                        f"{desc_finale_trait}\n\n"
                        f"⚠️ RÈGLE ANTI-RÉDACTION ABSOLUE :\n"
                        f"Le client t'a fourni des données brutes à mettre en forme. "
                        f"INTERDIT d'écrire un texte d'introduction, d'analyse ou de contexte. "
                        f"INTERDIT d'ajouter des paragraphes qui ne sont pas dans les données fournies. "
                        f"Tu dois UNIQUEMENT reproduire et structurer les données exactes du client. "
                        f"Commence directement par le titre et les données. Aucune phrase introductive."
                    )

                with st.spinner("⚡ Génération en cours... Cela prend moins d'1 minute."):
                    resultat = generer_avec_gemini(
                        service_final,
                        desc_finale_trait,
                        user,
                        is_premium=True,
                        gen_used=used_auj,
                        _image_b64=_trait_img_b64,
                        _image_mime=_trait_img_mime
                    )

                # ── Vider l'image pending APRÈS génération réussie ──────────────
                st.session_state["nova_ia_pending_image"] = None

                if resultat.startswith("❌"):
                    msg_err = f"❌ Une erreur est survenue lors de la génération. Contacte Nova : {WHATSAPP_NUMBER} 📲"
                    st.session_state["nova_ia_chat"].append({"role": "assistant", "content": msg_err})
                    st.session_state["nova_ia_phase"] = "dialogue"
                    st.rerun()
                else:
                    # ── Générer le vrai fichier .docx ─────────────────────
                    import uuid as _uuid_chat
                    _req_id_chat = str(_uuid_chat.uuid4())[:8]
                    _nom_fichier = f"Nova_IA_{service_final.replace(' ', '_')[:30]}_{datetime.now().strftime('%d%m%Y_%H%M')}.docx"
                    _buf_chat = creer_docx(resultat, service_final, user)

                    # ── Upload vers Supabase Storage ──────────────────────
                    _url_chat = upload_fichier_client(user, _req_id_chat, _buf_chat, _nom_fichier)

                    # ── Incrémenter quota + notifier ──────────────────────
                    incrementer_gen(user)
                    if _url_chat and not _url_chat.startswith("ERREUR"):
                        save_lien(user, f"✨ {service_final}", _url_chat, datetime.now().strftime("%d/%m/%Y"))
                    else:
                        save_lien(user, f"✨ {service_final}", f"__nova_ia__{resultat[:2000]}", datetime.now().strftime("%d/%m/%Y"))
                    envoyer_notification_gemini_ok(user, wa_user, service_final, _nom_fichier, demande_complete=desc_finale)

                    # ── Stocker le buf pour le bouton download dans le chat ──
                    _buf_chat.seek(0)
                    st.session_state["nova_ia_livrable"] = {
                        "buf": _buf_chat.read(),
                        "nom": _nom_fichier,
                        "service": service_final,
                        "url": _url_chat if (_url_chat and not _url_chat.startswith("ERREUR")) else None,
                    }

                    msg_ok = (
                        f"✅ Ton document **{service_final}** est prêt ! "
                        f"Tu peux le télécharger directement ici ou le retrouver dans **Mes Livrables**.\n\n"
                        f"🔧 En plus, nos ingénieurs vont retravailler et peaufiner ton fichier — tu recevras une version améliorée directement sur **WhatsApp** dans peu de temps.\n\n"
                        f"Il te reste **{quota_restant(db['users'].get(user, {})) - 1}** génération(s) aujourd'hui."
                    )
                    st.session_state["nova_ia_chat"].append({"role": "assistant", "content": msg_ok})
                    st.session_state["nova_ia_phase"] = "termine"
                    st.rerun()

        else:
            # ── CAS GRATUIT : save_demande comme les autres services ──
            import uuid
            req_id = str(uuid.uuid4())[:8]
            nouvelle_demande = {
                "id": req_id,
                "user": user or "visiteur",
                "service": service_final,
                "desc": desc_finale,
                "whatsapp": wa_user,
                "status": "en_attente",
                "incomplet": False,
                "champs_manquants": [],
                "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M")
            }
            save_demande(nouvelle_demande)
            db["demandes"].append(nouvelle_demande)
            st.session_state["db"] = db
            # Vider l'image pending même en cas gratuit
            st.session_state["nova_ia_pending_image"] = None

            # ── Notifier l'admin par email (comme le mode standard) ──
            try:
                envoyer_notification(
                    user or "visiteur",
                    wa_user,
                    f"[Nova IA Chat] {service_final}",
                    f"Réf : {req_id}\n\n{desc_finale}"
                )
            except Exception:
                pass

            msg_attente = (
                f"✅ Parfait ! Ta demande **{service_final}** a bien été reçue par l'équipe Nova.\n\n"
                f"🔧 Nos ingénieurs vont retravailler ton fichier et te l'envoyer directement sur **WhatsApp ({wa_user})** dans peu de temps.\n\n"
                f"📌 Référence : `{req_id}`\n\n"
                f"Pour recevoir tes documents en moins d'1 minute automatiquement, passe en **Premium** 👉 {WHATSAPP_NUMBER} 📲"
            )
            st.session_state["nova_ia_chat"].append({"role": "assistant", "content": msg_attente})
            st.session_state["nova_ia_phase"] = "termine"
            st.rerun()

    # ── PHASE : TERMINÉ ───────────────────────────────────────────
    elif st.session_state["nova_ia_phase"] == "termine":
        if premium_actif:
            st.markdown('<div class="nova-livraison-badge">✅ Livraison instantanée effectuée</div>', unsafe_allow_html=True)
            # ── Bouton de téléchargement direct dans le chat ──────
            lv_chat = st.session_state.get("nova_ia_livrable")
            if lv_chat:
                st.download_button(
                    label=f"📥 TÉLÉCHARGER — {lv_chat['nom']}",
                    data=lv_chat["buf"],
                    file_name=lv_chat["nom"],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="dl_nova_ia_chat"
                )
                st.info("💡 Ce document est aussi disponible dans **📂 Mes Livrables** de ton tableau de bord.")
        else:
            st.markdown('<div class="nova-attente-badge">📋 Demande soumise · L\'équipe Nova traite ta demande · Livraison par WhatsApp</div>', unsafe_allow_html=True)

        col_new, col_dash = st.columns(2)
        with col_new:
            if st.button("✨ Nouvelle demande", key="nova_ia_new", use_container_width=True):
                st.session_state.pop("nova_ia_chat", None)
                st.session_state.pop("nova_ia_phase", None)
                st.session_state.pop("nova_ia_service_detecte", None)
                st.session_state.pop("nova_ia_prompt_final", None)
                st.session_state.pop("nova_ia_livrable", None)
                st.rerun()
        with col_dash:
            if st.button("🏠 Tableau de bord", key="nova_ia_dash", use_container_width=True):
                st.session_state["view"] = "home"
                st.rerun()


if st.session_state["view"] == "auth" and st.session_state["current_user"] is None:
    show_auth_page()
elif st.session_state["view"] == "nova_support_ia":
    show_nova_support_ia_page()
elif st.session_state["view"] == "nova_ia":
    show_nova_ia_page()
else:
    main_dashboard()
